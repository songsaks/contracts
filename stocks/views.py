# ====== views.py - View หลักของระบบวิเคราะห์หุ้น AI ======
# ทุก view ต้องผ่านการ login (@login_required)
# ใช้ yfinance, yahooquery ดึงข้อมูลตลาด และ Gemini AI วิเคราะห์

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
from django.contrib import messages
from django.conf import settings
from google import genai
from .models import (
    Watchlist, AnalysisCache, AssetCategory, MarketType, Portfolio,
    MomentumCandidate, ScannableSymbol, MultiFactorCandidate, SoldStock,
    TitheRecord, ValueScanCandidate, PrecisionScanCandidate,
    InvestmentDashboardInsight,
)
from .forms import AddPortfolioForm, SellStockForm, AddWatchlistForm
from .utils import (
    get_stock_data, analyze_with_ai, calculate_trailing_stop,
    refresh_all_thai_symbols, find_supply_demand_zones, find_supply_demand_zones_v2,
    detect_price_pattern, detect_vcp_pattern, _is_commodity, _fetch_commodity_macro, _score_commodity_signal,
    analyze_momentum_technical_v2
)
from .crew_analysis import MomentumCrew
from decimal import Decimal
from yahooquery import Ticker as YQTicker
import requests
import yfinance as yf
import pandas as pd
import json
import os
import traceback
import pandas_ta as ta

# Removed custom session for yfinance to let it handle curl_cffi internally

# ====== ฟังก์ชันตรวจสอบสิทธิ์ผู้ใช้ ======

def admin_only(user):
    # อนุญาตให้ผู้ใช้ทุกคนที่ login แล้วเข้าถึงข้อมูลของตัวเองได้
    return user.is_authenticated # Changed to allow any user to see their own data

# ====== _build_us_symbol_set - สร้าง set ของสัญลักษณ์หุ้น US สำหรับ user ======
def _build_us_symbol_set(user):
    """
    Return a set of stock symbols (uppercase, no .BK) that are US stocks for this user.
    Uses MomentumCandidate(market='US') as the source of truth.
    """
    from .models import MomentumCandidate
    return set(
        MomentumCandidate.objects.filter(user=user, market='US')
        .values_list('symbol', flat=True)
    )


def _is_us_symbol(symbol, us_set):
    """True only if symbol is a known US stock. Thai stocks without .BK default to SET."""
    if '.BK' in symbol:
        return False
    return symbol.split('.')[0].upper() in us_set


# ====== _get_usd_thb - ดึงอัตราแลกเปลี่ยน USD/THB (cache 15 นาที) ======
def _get_usd_thb():
    """Return current USD/THB rate, cached 15 min. Fallback = 33.5."""
    from django.core.cache import cache
    rate = cache.get('usd_thb_rate')
    if rate:
        return float(rate)
    try:
        t = yf.Ticker('USDTHB=X')
        hist = t.history(period='2d')
        rate = float(hist['Close'].iloc[-1]) if not hist.empty else 0
        if not rate:
            rate = float(t.info.get('regularMarketPrice') or t.info.get('currentPrice') or 33.5)
    except Exception:
        rate = 33.5
    if rate < 25 or rate > 50:
        rate = 33.5
    cache.set('usd_thb_rate', rate, 900)
    return float(rate)


# ====== _compute_signals - คำนวณ BUY/SELL Score + Exit Signal จาก PrecisionScanCandidate ======
def _compute_signals(prec, current_price=None, is_turtle=False, turtle_stop=None):
    """Reusable scorer v3 - ใช้ใน Portfolio, Watchlist, และ Precision Scanner."""
    price  = float(current_price or getattr(prec, 'price', 0) or 0)
    dz_s   = getattr(prec, 'demand_zone_start', None)
    dz_e   = getattr(prec, 'demand_zone_end', None)
    prox   = getattr(prec, 'zone_proximity', 999)
    rvol_b = getattr(prec, 'rvol_bullish', True)
    rvol   = getattr(prec, 'rvol', 0)
    rr     = getattr(prec, 'risk_reward_ratio', 0) or 0
    adx    = getattr(prec, 'adx', 0)
    erc    = getattr(prec, 'erc_volume_confirmed', False)
    rsi    = getattr(prec, 'rsi', 0)
    pat    = getattr(prec, 'price_pattern_score', 0)
    rel3   = getattr(prec, 'rel_momentum_3m', 0.0)
    rel1   = getattr(prec, 'rel_momentum_1m', 0.0)
    sz_s   = getattr(prec, 'supply_zone_start', None)
    yr_h   = getattr(prec, 'year_high', 0)
    # v3 new fields (graceful fallback for old records)
    macd_hist  = getattr(prec, 'macd_histogram', None) or 0.0
    macd_cross = getattr(prec, 'macd_crossover', False)
    bb_sq      = getattr(prec, 'bb_squeeze', False)
    ema20_aln  = getattr(prec, 'ema20_aligned', False)
    rs_rat     = getattr(prec, 'rs_rating', 0)
    eps        = getattr(prec, 'eps_growth', 0)
    rev        = getattr(prec, 'rev_growth', 0)
    # v4 trend following fields
    ema20_rising = getattr(prec, 'ema20_rising', False)
    hh_hl        = getattr(prec, 'hh_hl_structure', False)
    ema20_slope  = getattr(prec, 'ema20_slope', 0.0)
    # v7 money flow & breakout
    cmf          = getattr(prec, 'cmf', None)
    is_52w_bo    = getattr(prec, 'is_52w_breakout', False)
    stage2       = getattr(prec, 'stage2', False)

    # ── BUY SCORE ─────────────────────────────────────────────────────
    # base technical: technical_score คือผลลัพธ์จาก tech analyzer (max 100)
    # เราลดน้ำหนักลงเหลือ 0.25 (max 25 pts) เพื่อเปิดทางให้ signals อื่นๆ
    buy = int(getattr(prec, 'technical_score', 0) * 0.25) 

    # Zone proximity (max 25)
    in_zone = dz_s and dz_e and price <= dz_s and price >= dz_e
    if in_zone:         buy += 25
    elif prox <= 10:    buy += 15
    elif prox <= 30:    buy += 8
    elif prox <= 60:    buy += 3

    # RVOL direction-aware (max 22)
    if rvol_b and rvol >= 2.0:   buy += 22
    elif rvol_b and rvol >= 1.5: buy += 17
    elif rvol_b and rvol >= 1.0: buy += 12
    elif rvol_b and rvol >= 0.7: buy += 4

    # R/R ratio (max 15)
    if rr >= 3:     buy += 15
    elif rr >= 2:   buy += 10
    elif rr >= 1.5: buy += 5

    # ADX (max 8) - ปรับ threshold ให้เหมาะกับ SET (liquidity ต่ำกว่า US → ADX มักอยู่ 15-25)
    if adx >= 25:   buy += 8
    elif adx >= 20: buy += 5
    elif adx >= 15: buy += 2

    # ERC confirmed (max 5)
    if erc: buy += 5

    # RSI - v3: optimal zone ขยับเป็น 65-80 (max 8)
    if 65 <= rsi <= 80:  buy += 8
    elif 55 <= rsi < 65: buy += 4
    elif rsi > 80:       buy += 2   # overbought แต่ trend ยังแรง

    # Price pattern (max +10 / min -10)
    buy += pat

    # Relative momentum (max +12 / min -8)
    rel = rel3 if rel3 != 0.0 else rel1
    if rel >= 15:   buy += 12
    elif rel >= 8:  buy += 9
    elif rel >= 3:  buy += 6
    elif rel >= 0:  buy += 3
    elif rel >= -5: buy += 0
    else:           buy -= 8

    # ── v3 new signals ────────────────────────────────────────────────
    # MACD bullish crossover (max 12)
    if macd_cross:          buy += 12
    elif macd_hist > 0:     buy += 8   # histogram positive (buying pressure building)

    # Bollinger Band Squeeze - pending breakout (max 6)
    if bb_sq: buy += 6

    # EMA20 > EMA50 > EMA200 full 3-layer alignment (max 5)
    if ema20_aln: buy += 5

    # ── v3 Relative Strength Rating (0-99 percentile) ───────────────────
    if rs_rat >= 85:   buy += 15
    elif rs_rat >= 70: buy += 8
    elif rs_rat >= 50: buy += 3

    # ── v4 Trend Following quality (max 10 pts) ───────────────────────
    # EMA20 rising = momentum มีโครงสร้างรองรับ ไม่ใช่แค่ spike สั้น
    if ema20_rising and hh_hl:  buy += 10  # ทั้ง EMA rising + HH/HL = trend สมบูรณ์
    elif ema20_rising:           buy += 5   # EMA rising เพียงอย่างเดียว
    elif hh_hl:                  buy += 4   # HH/HL โดยที่ EMA ยังไม่ขึ้นชัด

    # ── v7 CMF buy bonus (max 6 pts) ─────────────────────────────────
    if cmf is not None:
        if cmf >= 0.1:    buy += 6   # สถาบันสะสมชัดเจน
        elif cmf >= 0.05: buy += 3   # มีแรงซื้อสุทธิ

    # ── v8 SET-specific signals ───────────────────────────────────────
    # 52-week breakout (max 15) - signal สำคัญมากใน SET (Minervini/O'Neil)
    if is_52w_bo: buy += 15

    # Stage 2 Weinstein: price > SMA150 AND SMA150 rising (max 8)
    if stage2: buy += 8

    buy_score = max(0, min(100, buy))

    # ── SELL SCORE ────────────────────────────────────────────────────
    sell = 0
    if is_turtle:
        # Turtles only exit on 10D/20D Low (turtle_stop)
        if turtle_stop and price <= turtle_stop:
            sell = 100
        else:
            sell = 0
    else:
        if sz_s and price >= sz_s:              sell += 45
        # ยกเลิกเงื่อนไขขายเมื่อใกล้วน 52w High เพราะเบรกเอาต์คือสัญญาณโมเมนตัมที่ดี
        if rsi > 78:    sell += 20
        elif rsi > 72:  sell += 12
        elif rsi > 68:  sell += 5
        # Volume bearish penalty - ผ่อนลงสำหรับ SET เพราะหลายวัน volume เป็น neutral ไม่ใช่ bearish จริง
        if not rvol_b and rvol >= 2.0:  sell += 18   # volume สูงมากและเป็น bearish = แรงขายจริง
        elif not rvol_b and rvol >= 1.5: sell += 10  # volume สูงปานกลางและ bearish
        # rvol_b = False แต่ volume ปกติ → ไม่ penalty
        if rel1 < -5:   sell += 12
        elif rel1 < 0:  sell += 6
        if pat < -5:    sell += 10
        elif pat < 0:   sell += 5
        if adx < 15:    sell += 8
        elif adx < 20:  sell += 4
        # v3: MACD bearish (histogram negative + no crossover)
        if not macd_cross and macd_hist < 0 and abs(macd_hist) > 0.01:
            sell += 8
        # v7: CMF distribution - เงินไหลออกสุทธิ
        if cmf is not None:
            if cmf < -0.1:    sell += 10  # Distribution ชัดเจน
            elif cmf < -0.05: sell += 5   # เริ่มมีแรงขายสุทธิ
    sell_score = min(100, sell)

    if sell_score >= 70:   exit_signal = 'STRONG EXIT'
    elif sell_score >= 50: exit_signal = 'EXIT'
    elif sell_score >= 30: exit_signal = 'WATCH'
    else:                  exit_signal = ''

    return {'buy_score': buy_score, 'sell_score': sell_score, 'exit_signal': exit_signal}


# ====== Dashboard - หน้าแสดง Watchlist พร้อมราคาและ RSI แบบ Real-time ======

@login_required
def dashboard(request):
    """
    แสดงรายการ Watchlist ของผู้ใช้พร้อมราคาปัจจุบัน, % เปลี่ยนแปลง
    และค่า RSI 14 วัน คำนวณแบบ Real-time ผ่าน yfinance
    """
    watchlist = Watchlist.objects.filter(user=request.user)
    items = []
    import pandas_ta as ta
    from .utils import analyze_momentum_technical
    for item in watchlist:
        try:
            t = yf.Ticker(item.symbol)
            hist = t.history(period="1y")

            # Fallback: try alternate symbol if empty
            if hist.empty:
                alt_sym = f"{item.symbol}.BK" if ".BK" not in item.symbol else item.symbol.replace(".BK", "")
                t = yf.Ticker(alt_sym)
                hist = t.history(period="1y")

            current = None
            change = 0
            if not hist.empty:
                if isinstance(hist.columns, pd.MultiIndex):
                    hist.columns = [col[0] for col in hist.columns]
                hist = hist.loc[:, ~hist.columns.duplicated()]
                current = float(hist['Close'].iloc[-1])
                if len(hist) >= 2:
                    prev = float(hist['Close'].iloc[-2])
                    change = ((current - prev) / prev * 100) if prev else 0
            if not current:
                try:
                    info = t.info
                    if isinstance(info, dict):
                        current = info.get('currentPrice') or info.get('regularMarketPrice') or info.get('previousClose')
                        change = info.get('regularMarketChangePercent', 0)
                except:
                    pass

            rsi_val = None
            rsi_status = "Neutral"
            if not hist.empty and len(hist) >= 14:
                rsi_series = ta.rsi(hist['Close'], length=14)
                if not rsi_series.empty:
                    rsi_val = rsi_series.iloc[-1]
                    if rsi_val < 30: rsi_status = "Oversold"
                    elif rsi_val > 70: rsi_status = "Overbought"

            # ดึงข้อมูลจาก PrecisionScanCandidate ก่อน (ตรงกับ Scanner ทุกค่า)
            clean_symbol = item.symbol.split('.')[0].upper()
            from .models import PrecisionScanCandidate
            prec_data = (PrecisionScanCandidate.objects
                         .filter(user=request.user, symbol=clean_symbol)
                         .order_by('-scan_run').first())

            if prec_data:
                class QuickMom: pass
                mom_data = QuickMom()
                mom_data.technical_score      = prec_data.technical_score
                mom_data.rvol                 = prec_data.rvol
                mom_data.rvol_bullish         = prec_data.rvol_bullish
                mom_data.adx                  = prec_data.adx
                mom_data.rsi                  = prec_data.rsi
                mom_data.erc_volume_confirmed = prec_data.erc_volume_confirmed
                mom_data.risk_reward_ratio    = prec_data.risk_reward_ratio
                mom_data.demand_zone_start    = prec_data.demand_zone_start
                mom_data.demand_zone_end      = prec_data.demand_zone_end
                mom_data.supply_zone_start    = prec_data.supply_zone_start
                mom_data.stop_loss            = prec_data.stop_loss
                mom_data.zone_proximity       = prec_data.zone_proximity
                mom_data.year_high            = prec_data.year_high
                mom_data.price_pattern        = prec_data.price_pattern
                mom_data.price_pattern_score  = prec_data.price_pattern_score
                mom_data.rel_momentum_1m      = prec_data.rel_momentum_1m
                mom_data.rel_momentum_3m      = prec_data.rel_momentum_3m
            else:
                # Fallback: คำนวณ on-the-fly ด้วย v2
                from .utils import analyze_momentum_technical_v2
                mom_data = None
                if not hist.empty:
                    tech_analysis = analyze_momentum_technical_v2(hist)
                    if tech_analysis and tech_analysis.get('score', 0) > 0:
                        class QuickMom: pass
                        mom_data = QuickMom()
                        mom_data.technical_score      = tech_analysis['score']
                        mom_data.rvol                 = tech_analysis.get('rvol', 1.0)
                        mom_data.rvol_bullish         = tech_analysis.get('rvol_bullish', True)
                        mom_data.adx                  = 0
                        mom_data.rsi                  = float(rsi_val or 0)
                        mom_data.erc_volume_confirmed = False
                        mom_data.year_high            = 0
                        mom_data.price_pattern        = ''
                        mom_data.price_pattern_score  = 0
                        mom_data.rel_momentum_1m      = 0.0
                        mom_data.rel_momentum_3m      = 0.0
                        sd = tech_analysis.get('sd_zone')
                        if sd and sd.get('start') and sd['start'] > 0:
                            mom_data.risk_reward_ratio = sd.get('rr_ratio', 0)
                            mom_data.demand_zone_start = sd['start']
                            mom_data.demand_zone_end   = sd.get('end', 0)
                            mom_data.supply_zone_start = sd.get('target', 0)
                            mom_data.stop_loss         = sd.get('stop_loss', None)
                            mom_data.zone_proximity    = 0 if (current and current <= sd['start']) else ((float(current or 0) - sd['start']) / sd['start']) * 100
                        else:
                            mom_data.risk_reward_ratio = 0
                            mom_data.demand_zone_start = 0
                            mom_data.demand_zone_end   = 0
                            mom_data.supply_zone_start = 0
                            mom_data.stop_loss         = None
                            mom_data.zone_proximity    = 999

            signals = _compute_signals(mom_data, current) if mom_data else {'buy_score': 0, 'sell_score': 0, 'exit_signal': ''}

            # Heuristic market detection
            mkt = 'SET'
            if '.BK' in item.symbol: mkt = 'SET'
            elif item.category == AssetCategory.CRYPTO: mkt = 'CRYPTO'
            elif '-' in item.symbol and item.category != AssetCategory.CRYPTO: mkt = 'US'
            elif '.' not in item.symbol and '=' not in item.symbol and '-' not in item.symbol:
                mkt = 'US'
            elif '=' in item.symbol: mkt = 'OTHER' # Commodities usually use =F

            items.append({
                'obj': item,
                'price': current,
                'change': change,
                'rsi': rsi_val,
                'rsi_status': rsi_status,
                'mom_data': mom_data,
                'buy_score': signals['buy_score'],
                'sell_score': signals['sell_score'],
                'exit_signal': signals['exit_signal'],
                'market': mkt,
            })

        except:
            items.append({'obj': item, 'price': 'Error', 'change': 0, 'rsi': None, 'rsi_status': 'Error', 'mom_data': None})

    return render(request, 'stocks/dashboard.html', {
        'items': items,
        'categories': AssetCategory.choices,
        'market_types': MarketType.choices,
    })

# ====== Analyze - วิเคราะห์หุ้นรายตัวด้วย AI (Gemini) ======

@login_required
def analyze(request, symbol):
    """
    ดึงข้อมูลหุ้นจาก yfinance + yahooquery และส่งให้ Gemini AI วิเคราะห์
    ผลการวิเคราะห์จะถูกแคชไว้ใน AnalysisCache เพื่อใช้ซ้ำได้
    แสดงกราฟราคา 90 วัน, ข่าวล่าสุด, และข้อมูลพื้นฐาน
    """
    try:
        # Stop passing custom session, let yfinance handle its internal logic
        ticker = yf.Ticker(symbol)
        # ดึงข้อมูลหุ้นทั้งหมดผ่าน utility function
        data = get_stock_data(symbol)
        # ====== Fetch Extra Context from Cache (Value Stock data) ======
        extra_ctx = ""
        cached_lists = ['THAI_REC_ALL', 'US_REC_ALL']
        for list_sym in cached_lists:
            c = AnalysisCache.objects.filter(user=request.user, symbol=list_sym).first()
            if c:
                try:
                    s_list = json.loads(c.analysis_data)
                    match = next((s for s in s_list if s['symbol'] == symbol), None)
                    if match:
                        extra_ctx = f"Legendary Score: {match.get('value_score')}\n"
                        extra_ctx += f"Pillars: {match.get('legendary')}\n"
                        extra_ctx += f"Fair Value (Estimated): {match.get('fair_value')}\n"
                        extra_ctx += f"Upside (%): {match.get('upside')}\n"
                        break
                except: pass

        # ====== Fetch Extra Context from Momentum Scanner (Technical data) ======
        mom = MomentumCandidate.objects.filter(user=request.user, symbol_bk=symbol).first()
        if not mom:
            # ลองค้นหาด้วย symbol แบบไม่มี .BK
            clean_sym = symbol.replace('.BK', '')
            mom = MomentumCandidate.objects.filter(user=request.user, symbol=clean_sym).first()
        
        if mom:
            extra_ctx += f"\n[Technical Momentum Analysis]:\n"
            extra_ctx += f"Momentum Score: {mom.technical_score}/100\n"
            extra_ctx += f"RVOL: {mom.rvol}x, ADX: {mom.adx}, MFI: {mom.mfi}\n"
            if mom.entry_strategy:
                extra_ctx += f"Entry Strategy: {mom.entry_strategy}\n"
                extra_ctx += f"Demand Zone: {mom.demand_zone_start} - {mom.demand_zone_end}\n"
                extra_ctx += f"Target (Supply): {mom.supply_zone_start}\n"
                extra_ctx += f"Stop Loss: {mom.stop_loss}\n"
                extra_ctx += f"RR Ratio: {mom.risk_reward_ratio}\n"

        # ====== Pre-fetch macro data & compute signal for commodities ======
        # Done BEFORE AI call so we can pass signal into the prompt
        history = data.get('history', pd.DataFrame())
        is_commodity = _is_commodity(symbol)
        macro_data   = {}
        macro_signal = None
        if is_commodity:
            macro_data = _fetch_commodity_macro()
            _ema200 = float(history['EMA_200'].iloc[-1]) if 'EMA_200' in history.columns and not history.empty else None
            _rsi    = history['RSI'].iloc[-1]             if 'RSI'     in history.columns and not history.empty else None
            _price  = float(history['Close'].iloc[-1])    if not history.empty else 0
            macro_signal = _score_commodity_signal(symbol, _price, _ema200, _rsi, macro_data)
            # Stash raw macro inside signal so AI function can re-use without double-fetching
            macro_signal['_raw_macro'] = macro_data

        # ส่งข้อมูลให้ AI วิเคราะห์และรับผลเป็น Markdown
        analysis_text = analyze_with_ai(symbol, data, extra_context=extra_ctx, macro_signal=macro_signal)

        # ====== เตรียมข้อมูลกราฟราคาและวอลลุ่ม ======
        # Prepare Chart Data (Price & Volume)
        chart_labels = []
        chart_values = []
        chart_volumes = []
        if not history.empty:
            # แสดงเฉพาะ 90 วันล่าสุด
            history_subset = history.tail(90)
            chart_labels = [d.strftime('%Y-%m-%d') for d in history_subset.index]
            chart_values = [round(float(v), 2) for v in history_subset['Close'].values]
            chart_volumes = [int(v) for v in history_subset['Volume'].values]

        # ====== เตรียมข้อมูลข่าว - แปลง timestamp ให้อ่านได้ ======
        # Prepare News Data (Convert timestamp to readable)
        from datetime import datetime
        news_list = data.get('news', [])
        for n in news_list:
            if 'providerPublishTime' in n:
                try:
                    # รองรับทั้ง string (ISO format) และ int (Unix timestamp)
                    if isinstance(n['providerPublishTime'], str):
                        n['display_time'] = datetime.fromisoformat(n['providerPublishTime'].replace('Z', '+00:00'))
                    else:
                        n['display_time'] = datetime.fromtimestamp(n['providerPublishTime'])
                except Exception:
                    n['display_time'] = n['providerPublishTime']

        # ====== บันทึกผลวิเคราะห์ลงใน cache ของแต่ละ user ======
        # Cache it per user
        AnalysisCache.objects.update_or_create(
            user=request.user,
            symbol=symbol,
            defaults={'analysis_data': analysis_text}
        )

        # ====== คำนวณค่า RSI ล่าสุดเพื่อแสดงใน header ======
        # Prepare RSI
        current_rsi = history['RSI'].iloc[-1] if 'RSI' in history.columns and not history.empty else None
        rsi_status = "Neutral"
        if current_rsi:
            if current_rsi < 30: rsi_status = "Oversold"
            elif current_rsi > 70: rsi_status = "Overbought"

        # ====== เตรียมข้อมูลพื้นฐานการเงิน ======
        info = data.get('info') or {}
        if not isinstance(info, dict):
            info = {}

        # ทำ copy เพื่อไม่แก้ไข dict ต้นฉบับ
        info = info.copy() # Safe copy
        # แปลงค่าสัดส่วนพื้นฐาน (ROE, Dividend Yield, NPM) จากทศนิยมเป็นเปอร์เซ็นต์ (เฉพาะถ้าเป็นทศนิยม < 1)
        for key in ['returnOnEquity', 'dividendYield', 'profitMargins']:
            val = info.get(key)
            if isinstance(val, (int, float)):
                if abs(val) < 1.0:
                    info[key] = val * 100
                else:
                    info[key] = val

        # ====== คำนวณ D/E Ratio จาก Balance Sheet ======
        bs = data.get('balance_sheet')
        de_calculated = None
        if bs is not None and not bs.empty:
            try:
                col = bs.columns[0]
                # พยายามดึง Total Liabilities จากหลายชื่อ field ที่ yfinance อาจใช้
                tot_liab = bs.loc['Total Liabilities Net Minority Interest', col] if 'Total Liabilities Net Minority Interest' in bs.index else bs.loc['Total Liabilities', col]
                tot_eq = bs.loc['Stockholders Equity', col] if 'Stockholders Equity' in bs.index else bs.loc['Total Equity Gross Minority Interest', col]
                de_calculated = tot_liab / tot_eq
            except Exception:
                pass

        # อัปเดต D/E ใน info dict (ใช้ค่าจาก balance sheet ถ้ามี หรือ fallback เป็นค่าจาก yfinance)
        if de_calculated is not None:
            info['debtToEquity'] = de_calculated
        elif isinstance(info.get('debtToEquity'), (int, float)) if isinstance(info, dict) else False:
            # yfinance ส่งค่า D/E เป็น % (คูณ 100 มาแล้ว) ต้องหาร 100 กลับ
            info['debtToEquity'] = info['debtToEquity'] / 100

        # ====== คำนวณ Breakout / Resistance Levels ======
        fifty_two_week_high = history['High'].max() if not history.empty and 'High' in history.columns else None
        # แนวต้านย่อยในช่วง 20 วันล่าสุด
        recent_resistance = history['High'].tail(20).max() if not history.empty and 'High' in history.columns else None
        # แนวรับ/จุดตัดขาดทุนในช่วง 20 วันล่าสุด
        recent_support = history['Low'].tail(20).min() if not history.empty and 'Low' in history.columns else None
        curr_price = history['Close'].iloc[-1] if not history.empty and 'Close' in history.columns else info.get('currentPrice', 0)
        # ตรวจสอบว่าราคาปัจจุบันทะลุ 52-Week High หรือไม่ (Breakout Signal)
        is_breakout = (curr_price >= fifty_two_week_high) if (fifty_two_week_high and curr_price) else False

        context = {
            'symbol': symbol,
            'info': info,
            'analysis': analysis_text,
            'chart_labels': chart_labels,
            'chart_values': chart_values,
            'chart_volumes': chart_volumes,
            'fifty_two_week_high': fifty_two_week_high,
            'recent_resistance': recent_resistance,
            'recent_support': recent_support,
            'is_breakout': is_breakout,
            'current_rsi': current_rsi,
            'rsi_status': rsi_status,
            'news': news_list,
            'title': f"AI Analysis: {symbol}",
            'is_commodity': is_commodity,
            'macro_data': macro_data,
            'macro_signal': macro_signal,
        }
        return render(request, 'stocks/analysis.html', context)
    except Exception as e:
        messages.error(request, f"Error analyzing {symbol}: {str(e)}")
        return redirect('stocks:dashboard')

@login_required
def crew_analyze(request, symbol):
    """
    CrewAI Multi-Agent Deep Analysis - runs in background thread,
    progress polled via AJAX so the browser never times out.

    Optional GET params (from Portfolio page):
      ?entry_price=2.40&quantity=10000&gain_loss_pct=5.2&gain_loss=1200&market_value=25200
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'crew_analysis_{user_id}_{symbol}'

    # ── Portfolio context from GET params ────────────────────────────
    def _safe_float(val, default=0.0):
        try:
            return float(val) if val else default
        except (ValueError, TypeError):
            return default

    portfolio_context = {}
    if request.GET.get('entry_price'):
        portfolio_context = {
            'entry_price':   _safe_float(request.GET.get('entry_price')),
            'quantity':      _safe_float(request.GET.get('quantity')),
            'gain_loss_pct': _safe_float(request.GET.get('gain_loss_pct')),
            'gain_loss':     _safe_float(request.GET.get('gain_loss')),
            'market_value':  _safe_float(request.GET.get('market_value')),
        }
        ep_key = str(int(_safe_float(request.GET.get('entry_price')) * 100))
        cache_key = f'crew_analysis_{user_id}_{symbol}_p{ep_key}'

    # ── AJAX status poll - use ck param to find correct cache key ────
    if request.GET.get('crew_status') == '1':
        ck = request.GET.get('ck')
        poll_key = ck if ck else cache_key
        st = _cp.get(poll_key, {'state': 'idle'})
        return _JR(st)

    strategy_param = request.GET.get('strategy')
    market_param = request.GET.get('market', 'SET')

    # ── Result ready (page reload after done) ────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'done':
        result = cached.get('result', '')
        _cp.delete(cache_key)
        data = get_stock_data(symbol)
        return render(request, 'stocks/crew_result.html', {
            'symbol':            symbol,
            'crew_result':       result,
            'info':              data.get('info', {}),
            'title':             f'CrewAI Deep Analysis: {symbol}',
            'loading':           False,
            'portfolio_context': portfolio_context,
        })

    # ── Background worker ────────────────────────────────────────────
    def _run_crew_bg(ckey, sym, pctx, strat, mkt):
        import concurrent.futures as _cf
        from django.core.cache import cache as _c
        from .crew_analysis import MomentumCrew as _MC

        try:
            phase = 'วิเคราะห์ Portfolio + Technical…' if pctx else 'กำลังวิเคราะห์…'
            _c.set(ckey, {'state': 'running', 'phase': phase}, timeout=600)

            mc = _MC(sym, portfolio_context=pctx, strategy=strat, market=mkt)

            # Hard timeout: kill entire analysis after 90 seconds
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(mc.run_analysis)
                try:
                    result = future.result(timeout=120)
                except _cf.TimeoutError:
                    result = '## วิเคราะห์ไม่สำเร็จ\n\nหมดเวลา 90 วินาที กรุณาลองใหม่อีกครั้ง'

            _c.set(ckey, {'state': 'done', 'result': result}, timeout=600)
        except Exception as exc:
            _c.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=600)

    # Start only if not already running
    if not cached or cached.get('state') == 'idle':
        _cp.set(cache_key, {'state': 'running', 'phase': 'เริ่มต้น Multi-Agent…'}, timeout=600)
        _th.Thread(target=_run_crew_bg, args=(cache_key, symbol, portfolio_context, strategy_param, market_param), daemon=True).start()

    # ── Show loading page (fast - no heavy data fetch) ───────────────
    try:
        import yfinance as _yf
        _t = _yf.Ticker(symbol)
        _fi = _t.fast_info
        info = {
            'longName': getattr(_fi, 'display_name', None) or symbol,
            'currentPrice': getattr(_fi, 'last_price', None),
        }
    except Exception:
        info = {}

    return render(request, 'stocks/crew_result.html', {
        'symbol':            symbol,
        'info':              info,
        'title':             f'CrewAI Deep Analysis: {symbol}',
        'loading':           True,
        'portfolio_context': portfolio_context,
        'cache_key':         cache_key,   # pass to template for correct poll URL
    })

@login_required
def core_analyze(request, symbol):
    """
    "The Core Project" - Renaissance-inspired Multi-Agent Deep Analysis.
    Uses TheCoreCrew with Anomaly Hunter, Backtest Engineer, and Execution Decider.
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'core_analysis_{user_id}_{symbol}'

    # ── AJAX status poll ─────────────────────────────────────────────
    if request.GET.get('core_status') == '1':
        ck = request.GET.get('ck')
        poll_key = ck if ck else cache_key
        st = _cp.get(poll_key, {'state': 'idle'})
        return _JR(st)

    strategy_param = request.GET.get('strategy')
    market_param = request.GET.get('market', 'SET')

    # ── Result ready ─────────────────────────────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'done':
        result = cached.get('result', '')
        _cp.delete(cache_key)
        data = get_stock_data(symbol)
        return render(request, 'stocks/crew_result.html', {
            'symbol':      symbol,
            'crew_result': result,
            'info':        data.get('info', {}),
            'title':       f'The Core Project: {symbol}',
            'loading':     False,
        })

    # ── Background worker ────────────────────────────────────────────
    def _run_core_bg(ckey, sym, mkt):
        import concurrent.futures as _cf
        from django.core.cache import cache as _c
        from .crew_analysis import TheCoreCrew as _TCC

        try:
            _c.set(ckey, {'state': 'running', 'phase': 'Anomaly Hunter กำลังสแกน WACC/PEGY…'}, timeout=600)
            
            # รันการวิเคราะห์เชิงลึก
            crew = _TCC(sym, market=mkt)
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(crew.run_analysis)
                try:
                    # Anomaly Hunter + Backtest ใช้เวลาพอสมควร (timeout 2 นาที)
                    result = future.result(timeout=120) 
                except _cf.TimeoutError:
                    result = '## วิเคราะห์ "The Core" ไม่สำเร็จ\n\nการวิเคราะห์เชิงลึกใช้เวลาเกินกำหนด กรุณาลองใหม่อีกครั้ง'

            _c.set(ckey, {'state': 'done', 'result': result}, timeout=600)
        except Exception as exc:
            _c.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=600)

    # Start Worker
    if not cached or cached.get('state') == 'idle':
        _cp.set(cache_key, {'state': 'running', 'phase': 'เริ่มต้นโครงวิเคราะห์ The Core…'}, timeout=600)
        _th.Thread(target=_run_core_bg, args=(cache_key, symbol, market_param), daemon=True).start()

    # Show loading
    return render(request, 'stocks/crew_result.html', {
        'symbol':    symbol,
        'title':     f'The Core Analysis: {symbol}',
        'loading':   True,
        'cache_key': cache_key,
        # Special polling parameter for the client JS to use Correct URL
        'is_core':   True, 
    })

# ====== Momentum Quick CrewAI Analysis (AJAX modal) ======

@login_required
def momentum_quick_analysis(request, symbol):
    """
    Short-term CrewAI multi-agent analysis for a momentum candidate.
    Returns JSON - designed to be called from a modal (no page reload).

    Flow:
      1. POST/GET → start background analysis → return {'state': 'running'}
      2. Poll ?mq_status=1 until {'state': 'done', 'result': '...'}
      3. Render markdown in modal via marked.js
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'mq_analysis_{user_id}_{symbol}'

    # ── Poll ─────────────────────────────────────────────────────────
    if request.GET.get('mq_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        return _JR(st)

    # ── Already running ───────────────────────────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'running':
        return _JR({'state': 'running'})

    # ── Return cached done result ─────────────────────────────────────
    if cached and cached.get('state') == 'done':
        _cp.delete(cache_key)
        return _JR({'state': 'done', 'result': cached.get('result', '')})

    # ── Collect scan data from MomentumCandidate model ───────────────
    scan_data = {}
    try:
        from .models import MomentumCandidate as _MCM

        def _sf(val):
            try:
                return float(val) if val is not None else None
            except (TypeError, ValueError):
                return None

        cand = _MCM.objects.filter(user=request.user, symbol=symbol).first()
        if cand:
            scan_data = {
                'price':             _sf(cand.price),
                'technical_score':   cand.technical_score,
                'rsi':               _sf(cand.rsi),
                'adx':               _sf(cand.adx),
                'mfi':               _sf(cand.mfi),
                'rvol':              _sf(cand.rvol),
                'demand_zone_start': _sf(cand.demand_zone_start),
                'demand_zone_end':   _sf(cand.demand_zone_end),
                'supply_zone_start': _sf(cand.supply_zone_start),
                'supply_zone_end':   _sf(cand.supply_zone_end),
                'risk_reward_ratio': _sf(cand.risk_reward_ratio),
                'zone_proximity':    _sf(cand.zone_proximity),
                'eps_growth':        _sf(cand.eps_growth) or 0,
                'rev_growth':        _sf(cand.rev_growth) or 0,
                'sector':            cand.sector or 'N/A',
                'year_high':         _sf(cand.year_high),
                'upside_to_high':    _sf(cand.upside_to_high),
            }
    except Exception:
        pass

    # ── Background worker ─────────────────────────────────────────────
    def _run_bg(ckey, sym, sd, mkt):
        from django.core.cache import cache as _c
        try:
            _c.set(ckey, {'state': 'running', 'phase': 'กำลังวิเคราะห์ด้วย 3 Expert Agents…'}, timeout=600)
            from .crew_analysis import MomentumShortTermCrew as _STC
            import concurrent.futures as _cf
            crew = _STC(sym, scan_data=sd, market=mkt)
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(crew.run_analysis)
                try:
                    result = future.result(timeout=180)
                except _cf.TimeoutError:
                    result = '## หมดเวลาวิเคราะห์\n\nกรุณาลองใหม่อีกครั้ง'
            _c.set(ckey, {'state': 'done', 'result': result}, timeout=900)
        except Exception as exc:
            from django.core.cache import cache as _c2
            _c2.set(ckey, {'state': 'done', 'result': f'## เกิดข้อผิดพลาด\n\n{exc}'}, timeout=60)

    market = cand.market if cand else 'SET'
    _cp.set(cache_key, {'state': 'running'}, timeout=600)
    _th.Thread(target=_run_bg, args=(cache_key, symbol, scan_data, market), daemon=True).start()
    return _JR({'state': 'running', 'cache_key': cache_key})


# ====== CrewAI Export - Word / PDF ======

@login_required
def crew_export_docx(request, symbol):
    """Export CrewAI analysis as a formatted Word document (.docx)"""
    if request.method != 'POST':
        from django.shortcuts import redirect
        return redirect('stocks:crew_analyze', symbol=symbol)

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re
    from io import BytesIO

    markdown_text = request.POST.get('markdown_content', '')
    company_name  = request.POST.get('company_name', symbol)

    # ── Document setup ──────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Theme colors ────────────────────────────────────────────────
    COLOR_TITLE   = RGBColor(0x0f, 0x17, 0x2a)   # navy
    COLOR_H1      = RGBColor(0x1e, 0x29, 0x3b)
    COLOR_H2      = RGBColor(0x1d, 0x40, 0xaf)   # blue-800
    COLOR_H3      = RGBColor(0x07, 0x89, 0x16)   # green
    COLOR_BODY    = RGBColor(0x1e, 0x29, 0x3b)
    COLOR_MUTED   = RGBColor(0x64, 0x74, 0x8b)

    def _set_font(run, size, bold=False, color=None, italic=False):
        run.font.name      = 'Sarabun'
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.italic    = italic
        if color:
            run.font.color.rgb = color

    def _para_spacing(para, before=0, after=6, line=None):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(before * 20))
        spacing.set(qn('w:after'),  str(after  * 20))
        if line:
            spacing.set(qn('w:line'),     str(line * 20))
            spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)

    def _add_border_bottom(para, color='1d40af', size=12):
        """Add bottom border to paragraph (used for h1/h2)"""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    str(size))
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _apply_inline(run_text, para, default_size=11, default_color=None):
        """Parse **bold**, *italic*, `code` inline markers into runs"""
        segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', run_text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                run = para.add_run(seg[2:-2])
                _set_font(run, default_size, bold=True, color=default_color or COLOR_BODY)
            elif seg.startswith('*') and seg.endswith('*'):
                run = para.add_run(seg[1:-1])
                _set_font(run, default_size, italic=True, color=default_color or COLOR_BODY)
            elif seg.startswith('`') and seg.endswith('`'):
                run = para.add_run(seg[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(default_size - 1)
                run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            else:
                run = para.add_run(seg)
                _set_font(run, default_size, color=default_color or COLOR_BODY)

    # ── Cover header ────────────────────────────────────────────────
    # Top accent bar (using shaded paragraph)
    accent = doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(0)
    run_accent = accent.add_run('▬' * 60)
    run_accent.font.color.rgb = COLOR_H2
    run_accent.font.size      = Pt(6)

    # Title
    title_p = doc.add_paragraph()
    _para_spacing(title_p, before=8, after=2)
    r1 = title_p.add_run(f'🤖  CrewAI Multi-Agent Analysis')
    _set_font(r1, 22, bold=True, color=COLOR_TITLE)

    # Symbol + company
    sym_p = doc.add_paragraph()
    _para_spacing(sym_p, before=0, after=2)
    r2 = sym_p.add_run(f'{symbol}')
    _set_font(r2, 28, bold=True, color=COLOR_H2)
    r3 = sym_p.add_run(f'  ·  {company_name}')
    _set_font(r3, 16, color=COLOR_MUTED)

    # Date line
    from django.utils import timezone as _tz
    date_p = doc.add_paragraph()
    _para_spacing(date_p, before=0, after=12)
    r4 = date_p.add_run(f'Generated: {_tz.now().strftime("%d %B %Y  %H:%M")}')
    _set_font(r4, 10, italic=True, color=COLOR_MUTED)

    # Divider
    div = doc.add_paragraph()
    _add_border_bottom(div, color='1d40af', size=16)
    _para_spacing(div, before=0, after=12)

    # ── Parse markdown ──────────────────────────────────────────────
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # H1
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            _add_border_bottom(p, color='0f172a', size=12)
            _para_spacing(p, before=14, after=4)
            run = p.add_run(text)
            _set_font(run, 18, bold=True, color=COLOR_H1)
            i += 1
            continue

        # H2
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            _add_border_bottom(p, color='1d40af', size=8)
            _para_spacing(p, before=12, after=3)
            run = p.add_run(text)
            _set_font(run, 14, bold=True, color=COLOR_H2)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=10, after=2)
            run = p.add_run(text)
            _set_font(run, 12, bold=True, color=COLOR_H3)
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=8, after=2)
            run = p.add_run(text)
            _set_font(run, 11, bold=True, color=COLOR_BODY)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', line.strip()):
            hr = doc.add_paragraph()
            _add_border_bottom(hr, color='cbd5e1', size=6)
            _para_spacing(hr, before=8, after=8)
            i += 1
            continue

        # Unordered list
        if re.match(r'^[-*+]\s', line):
            text = re.sub(r'^[-*+]\s', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            _para_spacing(p, before=1, after=1)
            _apply_inline(text, p)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            _para_spacing(p, before=1, after=1)
            _apply_inline(text, p)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            _para_spacing(p, before=4, after=4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left_bdr = OxmlElement('w:left')
            left_bdr.set(qn('w:val'),   'single')
            left_bdr.set(qn('w:sz'),    '16')
            left_bdr.set(qn('w:space'), '8')
            left_bdr.set(qn('w:color'), '7c3aed')
            pBdr.append(left_bdr)
            pPr.append(pBdr)
            run = p.add_run(text)
            _set_font(run, 11, italic=True, color=COLOR_MUTED)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _para_spacing(p, before=2, after=4)
        _apply_inline(line, p)
        i += 1

    # ── Footer disclaimer ────────────────────────────────────────────
    doc.add_paragraph()
    disc_outer = doc.add_paragraph()
    _add_border_bottom(disc_outer, color='f59e0b', size=16)
    _para_spacing(disc_outer, before=16, after=4)
    dr = disc_outer.add_run('⚠  หมายเหตุ / Disclaimer')
    _set_font(dr, 10, bold=True, color=RGBColor(0xd9, 0x77, 0x06))

    disc_p = doc.add_paragraph()
    _para_spacing(disc_p, before=2, after=2)
    disc_r = disc_p.add_run(
        'รายงานนี้สร้างโดย AI Multi-Agent (CrewAI + Gemini) อิงจากข้อมูลตลาดจริงและข่าวล่าสุด '
        'ใช้เป็นข้อมูลประกอบการตัดสินใจเท่านั้น - ไม่ใช่คำแนะนำการลงทุน'
    )
    _set_font(disc_r, 9, italic=True, color=COLOR_MUTED)

    # ── Return file ──────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from django.http import HttpResponse
    resp = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    safe_symbol = re.sub(r'[^\w\-]', '_', symbol)
    resp['Content-Disposition'] = f'attachment; filename="CrewAI_{safe_symbol}_Analysis.docx"'
    return resp


# ====== Watchlist Management - เพิ่ม/ลบ รายการ Watchlist ======

@login_required
def add_to_watchlist(request):
    """รับ POST form เพิ่ม symbol เข้า Watchlist ของ user ปัจจุบัน"""
    if request.method == 'POST':
        form = AddWatchlistForm(request.POST)
        if form.is_valid():
            symbol = form.cleaned_data['symbol']
            Watchlist.objects.get_or_create(
                user=request.user,
                symbol=symbol,
                defaults={
                    'name': form.cleaned_data['name'],
                    'category': form.cleaned_data['category'],
                }
            )
            messages.success(request, f"เพิ่ม {symbol} เข้าใน Watchlist แล้ว")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{error}")

    return redirect('stocks:dashboard')

@login_required
def delete_from_watchlist(request, pk):
    """ลบรายการ Watchlist ตาม pk (เฉพาะของ user ปัจจุบันเท่านั้น)"""
    item = get_object_or_404(Watchlist, pk=pk, user=request.user)
    symbol = item.symbol
    item.delete()
    messages.success(request, f"ลบ {symbol} ออกจาก Watchlist แล้ว")
    return redirect('stocks:dashboard')

# ====== Portfolio - แสดงพอร์ตการลงทุนพร้อมวิเคราะห์ AI ======

@login_required
def portfolio_list(request):
    """
    แสดงรายการสินทรัพย์ในพอร์ต พร้อม:
    - ราคาปัจจุบัน, P/L, Market Value
    - RSI 14 วัน
    - Trailing Stop (3% จาก High)
    - Supply & Demand Zone
    - AI Portfolio Analysis (PyPortfolioOpt + Gemini)
      เมื่อผู้ใช้กดปุ่ม Analyze (?analyze=true)
    """
    portfolio_items = Portfolio.objects.filter(user=request.user)
    items = []
    total_market_value = 0
    total_gain_loss = 0
    total_set_value = 0
    total_set_cost = 0
    total_set_pl = 0
    total_us_value = 0
    total_crypto_value = 0
    total_crypto_cost = 0
    total_crypto_pl = 0
    total_us_cost = 0
    total_us_pl = 0
    print(f"DEBUG: Portfolio Scan Started for {getattr(request.user, 'username', 'Anonymous')}")
    _portfolio_us_set = _build_us_symbol_set(request.user)

    for item in portfolio_items:
        try:
            symbol = item.symbol
            print(f"DEBUG: Processing {symbol}")

            # ====== ดึงข้อมูลราคาจาก yfinance ======
            # Determine correct symbol string for yfinance based on database market field
            fetch_symbol = symbol
            if item.market == MarketType.SET and not symbol.endswith('.BK'):
                fetch_symbol = f"{symbol}.BK"
            elif item.market == MarketType.CRYPTO and '-' not in symbol:
                fetch_symbol = f"{symbol}-USD"
            
            t = yf.Ticker(fetch_symbol)
            hist = t.history(period="1y")

            # Fallback if empty (for robustness with manually entered symbols)
            used_symbol = fetch_symbol
            if hist.empty and fetch_symbol == symbol:
                alt_sym = f"{symbol}.BK" if ".BK" not in symbol else symbol.replace(".BK", "")
                print(f"DEBUG: {symbol} empty, trying {alt_sym}")
                t = yf.Ticker(alt_sym)
                hist = t.history(period="1y")
                if not hist.empty:
                    used_symbol = alt_sym

            current_price = 0
            rsi_val = None

            if not hist.empty:
                # จัดการ MultiIndex columns ที่อาจเกิดขึ้นเมื่อ yfinance ส่งข้อมูลหลาย ticker
                if isinstance(hist.columns, pd.MultiIndex):
                    hist.columns = [col[0] for col in hist.columns]
                hist = hist.loc[:, ~hist.columns.duplicated()]

                current_price = float(hist['Close'].iloc[-1])

                # ตรวจสอบซ้ำอีกรอบเพื่อกรณีที่ Close เป็น NaN
                # Double check price
                if not current_price or pd.isna(current_price):
                    try:
                        info = t.info
                        if isinstance(info, dict):
                            current_price = info.get('currentPrice') or info.get('regularMarketPrice') or 0
                    except: pass

                current_price = float(current_price or 0)

                # คำนวณ RSI 14 วัน จากข้อมูลราคาปิด
                # RSI
                rsi_series = ta.rsi(hist['Close'], length=14)
                rsi_val = rsi_series.iloc[-1] if (rsi_series is not None and not rsi_series.empty) else None
                print(f"DEBUG: {symbol} Success Price={current_price}")
            else:
                print(f"DEBUG: {symbol} FAILED - No data")

            # คำนวณ % เปลี่ยนแปลงวันนี้ vs เมื่อวาน
            day_change = 0
            if not hist.empty and len(hist) >= 2:
                prev_close = float(hist['Close'].iloc[-2])
                day_change = ((current_price - prev_close) / prev_close * 100) if prev_close else 0

            # ====== คำนวณ P/L และ Market Value ======
            # Calculations
            market_value = float(item.quantity) * float(current_price)
            cost_basis = float(item.quantity) * float(item.entry_price or 0)
            gain_loss = market_value - cost_basis
            gain_loss_pct = (gain_loss / cost_basis * 100) if cost_basis > 0 else 0
            is_us = item.market == MarketType.US

            # ====== คำนวณ ATR Trailing Stop ======
            from .utils import calculate_atr_trailing_stop
            atr_ts = calculate_atr_trailing_stop(
                df=hist if not hist.empty else None,
                entry_price=float(item.entry_price or 0),
                highest_price_db=float(item.highest_price or 0),
                multiplier=float(item.trail_multiplier or 2.5),
            ) if current_price > 0 else None

            print(f"DEBUG: {symbol} Price={current_price}, DB_High={item.highest_price}, ATR_High={atr_ts['highest'] if atr_ts else 'N/A'}")
            
            # อัปเดต highest_price และ ATR ใน DB ถ้าสูงขึ้น
            if atr_ts and current_price > 0:
                new_high = atr_ts['highest']
                update_fields = []
                if float(item.highest_price or 0) < new_high:
                    item.highest_price = new_high
                    update_fields.append('highest_price')
                if abs(float(item.atr or 0) - atr_ts['atr']) > 0.0001:
                    item.atr = atr_ts['atr']
                    update_fields.append('atr')
                if update_fields:
                    item.save(update_fields=update_fields)

            # ====== Override Trailing Stop with Turtle Logic ======
            is_turtle = item.strategy and ('turtle' in item.strategy.lower() or '🐢' in item.strategy)
            if is_turtle and atr_ts and hist is not None and not hist.empty:
                is_s2 = 'S2' in item.strategy.upper() or '20' in item.strategy
                periods = 20 if is_s2 else 10
                nday_low = float(hist['Low'].tail(periods).min())
                initial_stop = float(item.entry_price or 0) - (2.0 * atr_ts['atr'])
                current_stop = max(initial_stop, nday_low) if float(item.entry_price or 0) > 0 else nday_low
                pyramid_price = current_price + (atr_ts['atr'] * 0.5) if current_price > 0 else 0
                
                dist_pct = ((current_price - current_stop) / current_price * 100) if current_price > 0 else 0
                
                if current_price <= current_stop:
                    status = 'EXIT HIT'
                    color = 'danger'
                elif dist_pct <= 3.0:
                    status = 'NEAR EXIT'
                    color = 'warning'
                else:
                    status = 'RIDE TREND 🚀'
                    color = 'success'
                    
                atr_ts.update({
                    'trailing_stop': current_stop,
                    'color': color,
                    'status': status,
                    'distance_pct': dist_pct,
                    'is_turtle': True,
                    'turtle_sys': 'S2 (20D Low)' if is_s2 else 'S1 (10D Low)',
                    'nday_low': nday_low,
                    'initial_stop': initial_stop,
                    'pyramid_price': pyramid_price
                })

            ts_data = atr_ts  # ยังคง key เดิมใน template

            # ====== ดึง/คำนวณ Zone Data - ใช้ PrecisionScanCandidate (v2) เสมอ ======
            clean_symbol = item.symbol.split('.')[0].upper()
            from .utils import analyze_momentum_technical_v2
            from .models import PrecisionScanCandidate

            # 1. ลองหาผล Precision Scan ล่าสุดก่อน (ตรงกับ Precision Scanner ทุกค่า)
            prec_data = (PrecisionScanCandidate.objects
                         .filter(user=request.user, symbol=clean_symbol)
                         .order_by('-scan_run').first())

            if prec_data and not request.GET.get('refresh') == 'true':
                # ใช้ข้อมูลจาก Precision Scanner โดยตรง
                class QuickMom: pass
                mom_data = QuickMom()
                mom_data.technical_score   = prec_data.technical_score
                mom_data.rvol              = prec_data.rvol
                mom_data.rvol_bullish      = prec_data.rvol_bullish
                mom_data.adx               = prec_data.adx
                mom_data.rsi               = prec_data.rsi
                mom_data.erc_volume_confirmed = prec_data.erc_volume_confirmed
                mom_data.risk_reward_ratio = prec_data.risk_reward_ratio
                mom_data.demand_zone_start = prec_data.demand_zone_start
                mom_data.demand_zone_end   = prec_data.demand_zone_end
                mom_data.supply_zone_start = pyramid_price if is_turtle else prec_data.supply_zone_start
                mom_data.stop_loss         = current_stop if is_turtle else prec_data.stop_loss
                mom_data.zone_proximity    = prec_data.zone_proximity
                mom_data.year_high         = prec_data.year_high
                mom_data.price_pattern     = prec_data.price_pattern
                mom_data.price_pattern_score = prec_data.price_pattern_score
                mom_data.rel_momentum_1m   = prec_data.rel_momentum_1m
                mom_data.rel_momentum_3m   = prec_data.rel_momentum_3m
                mom_data.macd_histogram    = prec_data.macd_histogram
                mom_data.macd_crossover    = prec_data.macd_crossover
                mom_data.bb_squeeze        = prec_data.bb_squeeze
                mom_data.ema20_aligned     = prec_data.ema20_aligned
                mom_data.rs_rating         = prec_data.rs_rating
                mom_data.ema20_rising      = prec_data.ema20_rising
                mom_data.hh_hl_structure   = prec_data.hh_hl_structure
                mom_data.cmf               = prec_data.cmf
                mom_data.is_52w_breakout   = prec_data.is_52w_breakout
            else:
                # 2. คำนวณ on-the-fly ด้วย v2 (ตรงกับ Precision Scanner)
                tech_analysis = analyze_momentum_technical_v2(hist) if not hist.empty else None
                class QuickMom: pass
                mom_data = QuickMom()
                if tech_analysis:
                    mom_data.technical_score = tech_analysis['score']
                    mom_data.rvol = tech_analysis['rvol']
                    sd = tech_analysis.get('sd_zone')
                    if sd and sd.get('start') and sd['start'] > 0:
                        mom_data.risk_reward_ratio = sd.get('rr_ratio', 0)
                        mom_data.demand_zone_start = sd['start']
                        mom_data.demand_zone_end = sd.get('end', 0)
                        mom_data.supply_zone_start = pyramid_price if is_turtle else sd.get('target', 0)
                        mom_data.stop_loss = current_stop if is_turtle else sd.get('stop_loss', None)
                        mom_data.zone_proximity = 0 if current_price <= sd['start'] else ((float(current_price) - sd['start']) / sd['start']) * 100
                    else:
                        mom_data.risk_reward_ratio = 0
                        mom_data.demand_zone_start = 0
                        mom_data.stop_loss = None
                        mom_data.zone_proximity = 999
                else:
                    mom_data.technical_score = 0
                    mom_data.rvol = 0
                    mom_data.risk_reward_ratio = 0
                    mom_data.demand_zone_start = 0
                    mom_data.stop_loss = None
                    mom_data.zone_proximity = 999

            total_market_value += market_value
            total_gain_loss += gain_loss
            _market = item.market
            if _market == MarketType.US:
                total_us_value += market_value
                total_us_cost += cost_basis
                total_us_pl += gain_loss
            elif _market == MarketType.CRYPTO:
                total_crypto_value += market_value
                total_crypto_cost += cost_basis
                total_crypto_pl += gain_loss
            else:
                total_set_value += market_value
                total_set_cost += cost_basis
                total_set_pl += gain_loss

            signals = _compute_signals(
                mom_data, 
                current_price, 
                is_turtle=is_turtle, 
                turtle_stop=(current_stop if is_turtle else None)
            ) if mom_data else {'buy_score': 0, 'sell_score': 0, 'exit_signal': ''}

            items.append({
                'obj': item,
                'current_price': current_price,
                'day_change': day_change,
                'market_value': market_value,
                'gain_loss': gain_loss,
                'gain_loss_pct': gain_loss_pct,
                'rsi': rsi_val,
                'trailing_stop_data': ts_data,
                'mom_data': mom_data,
                'buy_score': signals['buy_score'],
                'sell_score': signals['sell_score'],
                'exit_signal': signals['exit_signal'],
                'in_scan': prec_data is not None,
                'scan_score': prec_data.technical_score if prec_data else None,
                'is_us': is_us,
                'market': item.market,
            })
        except Exception as e:
            print(f"DEBUG: ERROR for {item.symbol}: {e}")
            traceback.print_exc()
            # ถ้า error ใส่ข้อมูลเปล่าเพื่อแสดง error state ใน template
            items.append({
                'obj': item, 'current_price': 0, 'day_change': 0, 'market_value': 0,
                'gain_loss': 0, 'gain_loss_pct': 0, 'rsi': None,
                'trailing_stop_data': None, 'mom_data': None,
                'is_us': item.market == MarketType.US,
                'market': item.market,
            })

    # ── USD/THB rate for combined totals ──
    usd_thb = _get_usd_thb() if any(it.get('market') in (MarketType.US, MarketType.CRYPTO) for it in items) else 1.0

    # ── Sort items: SET → US → CRYPTO → OTHER; mark group headers ──
    _market_order = {MarketType.SET: 0, MarketType.US: 1, MarketType.CRYPTO: 2, MarketType.OTHER: 3}
    items.sort(key=lambda x: (_market_order.get(x.get('market', MarketType.SET), 99), x['obj'].symbol))
    _prev_market = None
    for _it in items:
        _it['show_group_header'] = (_it.get('market') != _prev_market)
        _prev_market = _it.get('market')

    # ====== AI Portfolio Analysis ด้วย Gemini + PyPortfolioOpt ======
    ai_analysis = None
    if request.GET.get('analyze') == 'true' and items:
        # เลือก Gemini model ที่ดีที่สุดที่ตอบสนองได้
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        model_name_to_use = 'gemini-2.5-flash'

        # สร้าง string สรุปพอร์ตสำหรับส่งให้ AI
        port_data = []
        for it in items:
            mom_info = ""
            if it.get('mom_data'):
                m = it['mom_data']
                rvol = m.rvol if hasattr(m, 'rvol') else 0
                rs = getattr(m, 'rs_rating', 0)
                adx = m.adx if hasattr(m, 'adx') else 0
                score = m.technical_score if hasattr(m, 'technical_score') else 0
                buy_sc = it.get('buy_score', 0)
                exit_sig = it.get('exit_signal', '')
                mom_info = f", Score(0-100): {score}, BUY Score: {buy_sc}, RS(0-99): {rs}, RVOL: {rvol:.1f}x, ADX: {adx:.1f}, Exit Signal: '{exit_sig}'"
            port_data.append(f"{it['obj'].symbol}: {it['obj'].quantity} units @ {it['obj'].entry_price} (Current: {it['current_price']}, P/L: {it['gain_loss_pct']:.2f}%, RSI: {it['rsi']}{mom_info})")
        port_str = "\n".join(port_data)

        # ====== PyPortfolioOpt - คำนวณ Efficient Frontier / Max Sharpe ======
        # --- PyPortfolioOpt Integration ---
        # Get historical data for all symbols to calculate correlation & efficient frontier
        symbols = [it['obj'].symbol for it in items if it['obj'].quantity > 0]
        ppo_advice = ""
        if len(symbols) > 1:
            try:
                from pypfopt import expected_returns, risk_models
                from pypfopt.efficient_frontier import EfficientFrontier

                # ดึงราคาปิดย้อนหลัง 1 ปีสำหรับทุก symbol พร้อมกัน
                # Fetch 1 yr of closing prices for correlation
                data = yf.download(symbols, period="1y")

                # แปลง MultiIndex columns ให้เหลือแค่ 'Close' level
                if isinstance(data.columns, pd.MultiIndex):
                    data = data['Close']
                elif 'Close' in data:
                    data = data[['Close']]
                else:
                    data = pd.DataFrame() # Fallback

                # จัดการ missing values ด้วย forward-fill และ backward-fill
                # Deal with missing values
                data = data.dropna(how="all")
                data = data.ffill().bfill()

                # ตรวจสอบว่ามีข้อมูลเพียงพอสำหรับการคำนวณ
                # Make sure data is not flat (in case of single symbol fallback bug, though handled by len(symbols) > 1)
                if data.empty or len(data.columns) < 2:
                    raise ValueError("Not enough overlapping price data to calculate correlation.")

                # คำนวณ Expected Return และ Covariance Matrix
                mu = expected_returns.mean_historical_return(data)
                S = risk_models.sample_cov(data)

                ef = EfficientFrontier(mu, S)

                # หา Portfolio ที่ Max Sharpe Ratio (ผลตอบแทนดีที่สุดเมื่อเทียบกับความเสี่ยง)
                # Optimise for maximal Sharpe ratio
                try:
                    raw_weights = ef.max_sharpe()
                except Exception as ef_e:
                    # Fallback: กรณีที่ max_sharpe ไม่ converge ใช้ equal weight แทน
                    # Fallback to equal weighting if max_sharpe fails (e.g. non-convex/all negative returns)
                    raw_weights = {sym: 1.0/len(symbols) for sym in symbols}

                cleaned_weights = ef.clean_weights()

                # เปรียบเทียบน้ำหนักปัจจุบันกับน้ำหนักที่เหมาะสม
                # Compare current weights to optimal weights
                portfolio_total = sum(it['market_value'] for it in items)
                current_weights = {it['obj'].symbol: (it['market_value'] / portfolio_total if portfolio_total > 0 else 0) for it in items}

                ppo_advice += "\n[PyPortfolioOpt Portfolio Optimization (Max Sharpe)]\n"
                for sym in symbols:
                    c_weight = current_weights.get(sym, 0) * 100
                    o_weight = cleaned_weights.get(sym, 0) * 100
                    action = "Hold"
                    # คำแนะนำ: Buy เมื่อน้ำหนักปัจจุบันต่ำกว่า optimal > 5%
                    if o_weight > c_weight + 5: action = "Buy/Increase Weight"
                    elif o_weight < c_weight - 5: action = "Sell/Reduce Weight"
                    ppo_advice += f"- {sym}: Current Weight = {c_weight:.1f}%, Optimal Weight = {o_weight:.1f}% -> Model says: {action}\n"

                # แสดงผลการวิเคราะห์ประสิทธิภาพของ Portfolio ที่เหมาะสม
                try:
                    perf = ef.portfolio_performance(verbose=False)
                    ppo_advice += f"\nOptimal Expected Annual Return: {perf[0]*100:.2f}%\n"
                    ppo_advice += f"Optimal Annual Volatility: {perf[1]*100:.2f}%\n"
                    ppo_advice += f"Optimal Sharpe Ratio: {perf[2]:.2f}\n"
                except:
                    pass

            except ImportError:
                ppo_advice = f"\n[PyPortfolioOpt] Unable to optimize portfolio: PyPortfolioOpt is not installed.\n"
            except Exception as e:
                ppo_advice = f"\n[PyPortfolioOpt] Unable to optimize portfolio due to an error: {str(e)}\n"

        # ====== สร้าง Prompt สำหรับ AI วิเคราะห์พอร์ต ======
        prompt = f"""
        You are an expert Stock Portfolio Analyst specializing in "Precision Momentum Trading" (similar to Mark Minervini and CANSLIM).
        The user has the following assets in their portfolio (with Entry Price, Current Price, and Profit/Loss, along with Momentum metrics):
        {port_str}

        {ppo_advice}

        Please analyze this portfolio and provide:
        2. A brief analysis and clear recommendation for EACH individual asset (e.g., Hold, Buy More, Take Profit, Cut Loss).
           - CRITICAL RULE: In your analysis, you MUST heavily weigh the Momentum metrics (Score, BUY Score, RS, RVOL, ADX, and Exit Signal).
           - Do NOT immediately suggest cutting a loss JUST because P/L is slightly negative or RSI is > 70.
           - **Consider Entry Price:** If a stock's price is below its technical Stop Loss but still above the Entry Price (in profit), suggest "Locking Profit" (ล็อกกำไร) or "Trailing Exit" instead of a harsh "Cut Loss".
           - If a stock has High RS (e.g. > 75), strong RVOL (> 1.5x) or a high BUY Score/Score, it indicates it is a "Market Leader" and in a strong uptrend. In such cases, suggest "Holding" to ride the momentum as long as the Exit Signal is not triggered, even if P/L is temporarily negative.
           - Acknowledge the strength of the momentum. E.g., "แม้จะขาดทุน -2.35% แต่หุ้นมี RS แข็งแกร่งถึง 88 และ RVOL สูง บ่งบอกถึงแรงซื้อที่ยังมีอยู่ แนะนำให้ถือเพื่อรอจังหวะเด้งกลับ"
        3. Actionable strategic advice on what sectors or types of assets to consider adding next to balance the portfolio.

        Format your response beautifully in Markdown using Thai Language (Sarabun professional tone).
        IMPORTANT RULES:
        1. DO NOT include any conversational preamble or outro (e.g. "Here is the analysis...", "Explanation of Choices:").
        2. Output ONLY the raw markdown text.
        3. DO NOT wrap the output in ```markdown code blocks. Start immediately with the analysis headings.
        """
        try:
            response = client.models.generate_content(
                model=model_name_to_use,
                contents=prompt
            )
            ai_analysis = response.text

            # ลบ markdown code block wrapper ถ้า AI ไม่ปฏิบัติตาม prompt
            # Strip any residual markdown blocks if AI disobeys
            if ai_analysis.startswith("```markdown"):
                ai_analysis = ai_analysis[len("```markdown"):].strip()
            if ai_analysis.endswith("```"):
                ai_analysis = ai_analysis[:-3].strip()
        except Exception as e:
            ai_analysis = f"ไม่สามารถวิเคราะห์พอร์ตได้ในขณะนี้: {str(e)}"

    # ── Performance Chart data (All time) ──
    all_sold_stocks = SoldStock.objects.filter(user=request.user).order_by('sold_at')
    chart_labels = []
    chart_data = []
    running_pl = 0
    for s in all_sold_stocks:
        val = float(s.profit_loss)
        if s.market in (MarketType.US, MarketType.CRYPTO):
            val *= usd_thb
        running_pl += val
        chart_labels.append(s.sold_at.strftime('%Y-%m-%d %H:%M'))
        chart_data.append(running_pl)

    # ── Available Months for Filter ──
    # สร้าง list ของเดือน/ปีที่มีรายการขายจริง เพื่อให้ User เลือก
    available_months = []
    seen_months = set()
    for s in all_sold_stocks[::-1]:
        m_key = s.sold_at.strftime('%Y-%m')
        if m_key not in seen_months:
            available_months.append({
                'key': m_key,
                'name': s.sold_at.strftime('%B %Y')
            })
            seen_months.add(m_key)

    # ── Filter Transactions ──
    from django.utils import timezone
    now = timezone.now()
    default_month = now.strftime('%Y-%m')
    
    # ถ้าเดือนปัจจุบันไม่มีรายการขาย ให้เลือกเดือนล่าสุดที่มีรายการเป็นค่าเริ่มต้น
    if not all_sold_stocks.filter(sold_at__year=now.year, sold_at__month=now.month).exists() and available_months:
        default_month = available_months[0]['key']

    selected_month = request.GET.get('month', default_month)
    
    # ถ้าระบุ 'all' จะแสดงทั้งหมดเหมือนเดิม
    if selected_month == 'all':
        sold_stocks = all_sold_stocks[::-1]
    else:
        try:
            yr, mn = map(int, selected_month.split('-'))
            sold_stocks = all_sold_stocks.filter(sold_at__year=yr, sold_at__month=mn).order_by('-sold_at')
        except:
            sold_stocks = all_sold_stocks[::-1]

    # ── Monthly Summary (Table on the right) ──
    from collections import defaultdict
    monthly_summary_dict = defaultdict(lambda: {'items': [], 'total_pl': 0})
    for s in all_sold_stocks:
        month_key = s.sold_at.strftime('%B %Y')
        month_id = s.sold_at.strftime('%Y-%m')
        monthly_summary_dict[month_key]['items'].append(s)
        monthly_summary_dict[month_key]['month_id'] = month_id
        val = float(s.profit_loss)
        if s.market in (MarketType.US, MarketType.CRYPTO):
            val *= usd_thb
        monthly_summary_dict[month_key]['total_pl'] += val
    
    monthly_summary = []
    for m_name in [m['name'] for m in available_months]:
        monthly_summary.append({
            'month_name': m_name,
            'month_id': monthly_summary_dict[m_name]['month_id'],
            'items': monthly_summary_dict[m_name]['items'],
            'total_pl': monthly_summary_dict[m_name]['total_pl']
        })

    context = {
        'items': items,
        'total_market_value': total_market_value,
        'total_gain_loss': total_gain_loss,
        'total_set_value': total_set_value,
        'total_set_cost': total_set_cost,
        'total_set_pl': total_set_pl,
        'total_us_value': total_us_value,
        'total_us_cost': total_us_cost,
        'total_us_pl': total_us_pl,
        'total_crypto_value': total_crypto_value,
        'total_crypto_cost': total_crypto_cost,
        'total_crypto_pl': total_crypto_pl,
        'has_set': any(it.get('market') == MarketType.SET for it in items),
        'has_us': any(it.get('market') == MarketType.US for it in items),
        'has_crypto': any(it.get('market') == MarketType.CRYPTO for it in items),
        'usd_thb': round(usd_thb, 2),
        'total_combined_value': total_set_value + (total_us_value + total_crypto_value) * usd_thb,
        'total_combined_cost': total_set_cost + (total_us_cost + total_crypto_cost) * usd_thb,
        'total_combined_pl': total_set_pl + (total_us_pl + total_crypto_pl) * usd_thb,
        'categories': AssetCategory.choices,
        'market_types': MarketType.choices,
        'title': 'My Portfolio',
        'ai_analysis': ai_analysis,
        'sold_stocks': sold_stocks,
        'monthly_summary': monthly_summary,
        'available_months': available_months,
        'selected_month': selected_month,
        'chart_labels': json.dumps(chart_labels),
        'chart_data': json.dumps(chart_data),
    }
    return render(request, 'stocks/portfolio.html', context)


# ====== Portfolio Exit Plan - แผนออกหุ้นแต่ละตัว เรียงตามความเร่งด่วน ======

@login_required
def portfolio_exit_plan(request):
    """
    แสดงแผนออกจากหุ้นแต่ละตัวในพอร์ต พร้อม:
    - Progress bar: SL → Entry → Current → TP
    - Action recommendation ชัดเจน (ออกทันที / ทยอยขาย / เฝ้า / ถือต่อ)
    - สัญญาณออกที่ active อยู่
    - เรียงตาม SELL Score สูงสุดก่อน (urgent first)
    """
    portfolio_items = Portfolio.objects.filter(user=request.user)
    items = []

    for item in portfolio_items:
        try:
            symbol = item.symbol
            # Determine correct symbol string for yfinance based on database market field
            fetch_symbol = symbol
            if item.market == MarketType.SET and not symbol.endswith('.BK'):
                fetch_symbol = f"{symbol}.BK"
            elif item.market == MarketType.CRYPTO and '-' not in symbol:
                fetch_symbol = f"{symbol}-USD"
            
            t = yf.Ticker(fetch_symbol)
            hist = t.history(period="1y")

            # Fallback if empty (for robustness with manually entered symbols)
            if hist.empty and fetch_symbol == symbol:
                alt_sym = f"{symbol}.BK" if ".BK" not in symbol else symbol.replace(".BK", "")
                print(f"DEBUG: {symbol} empty, trying {alt_sym}")
                t = yf.Ticker(alt_sym)
                hist = t.history(period="1y")
            if isinstance(hist.columns, pd.MultiIndex):
                hist.columns = [col[0] for col in hist.columns]
            hist = hist.loc[:, ~hist.columns.duplicated()]

            current_price = float(hist['Close'].iloc[-1]) if not hist.empty else 0
            day_change = 0
            if not hist.empty and len(hist) >= 2:
                prev = float(hist['Close'].iloc[-2])
                day_change = ((current_price - prev) / prev * 100) if prev else 0

            # ====== Turtle Exit Logic (S1: 10D Low, S2: 20D Low) ======
            turtle_s1_exit = 0
            turtle_s2_exit = 0
            s1_hit = s2_hit = False
            
            if not hist.empty:
                # Use daily Low for Turtle exits
                hist_lows = hist['Low']
                if len(hist_lows) >= 10:
                    turtle_s1_exit = float(hist_lows.tail(10).min())
                if len(hist_lows) >= 20:
                    turtle_s2_exit = float(hist_lows.tail(20).min())
                
                s1_hit = current_price <= turtle_s1_exit if turtle_s1_exit > 0 else False
                s2_hit = current_price <= turtle_s2_exit if turtle_s2_exit > 0 else False

            entry_price  = float(item.entry_price or 0)
            quantity     = float(item.quantity or 0)
            gain_loss_pct = ((current_price - entry_price) / entry_price * 100) if entry_price else 0

            # days held
            from datetime import date
            days_held = (date.today() - item.added_at.date()).days if item.added_at else 0

            # ดึง PrecisionScanCandidate
            clean_symbol = symbol.split('.')[0].upper()
            from .models import PrecisionScanCandidate
            prec_data = (PrecisionScanCandidate.objects
                         .filter(user=request.user, symbol=clean_symbol)
                         .order_by('-scan_run').first())

            sl_price = tp_price = rsi_val = adx_val = None
            price_pattern = ''
            price_pattern_score = 0
            rel_1m = rel_3m = 0.0
            rvol_bullish = True
            rvol = 1.0
            supply_zone_start = year_high = 0
            cmf_val = None

            if prec_data:
                sl_price    = prec_data.stop_loss
                tp_price    = prec_data.supply_zone_start
                rsi_val     = prec_data.rsi
                adx_val     = prec_data.adx
                price_pattern       = prec_data.price_pattern
                price_pattern_score = prec_data.price_pattern_score
                rel_1m      = prec_data.rel_momentum_1m
                rel_3m      = prec_data.rel_momentum_3m
                rvol_bullish = prec_data.rvol_bullish
                rvol        = prec_data.rvol
                supply_zone_start = prec_data.supply_zone_start or 0
                year_high   = prec_data.year_high or 0
                cmf_val     = prec_data.cmf

            signals = _compute_signals(prec_data, current_price) if prec_data else {'buy_score': 0, 'sell_score': 0, 'exit_signal': ''}
            sell_score   = signals['sell_score']
            exit_signal  = signals['exit_signal']

            # Boost sell score if Turtle exits are hit
            if s1_hit: sell_score = max(sell_score, 75)
            if s2_hit: sell_score = max(sell_score, 90)
            
            if sell_score >= 70: exit_signal = 'STRONG EXIT'
            elif sell_score >= 50: exit_signal = 'EXIT'
            elif sell_score >= 30: exit_signal = 'WATCH'

            # ====== Progress Bar: SL → Entry → Current → TP ======
            progress_pct   = None
            current_pct    = None
            entry_pct      = None
            sl_hit         = False
            tp_hit         = False

            if sl_price and tp_price and tp_price > sl_price:
                total_range    = tp_price - sl_price
                sl_hit         = current_price <= sl_price
                tp_hit         = current_price >= tp_price
                current_pct    = min(100, max(0, (current_price - sl_price) / total_range * 100))
                entry_pct      = min(100, max(0, (entry_price - sl_price) / total_range * 100))

            # ====== Momentum Classification ======
            # หุ้นผู้นำ (Leader) = Momentum แข็งแกร่งกว่าตลาดมาก
            is_leader = (rel_3m and rel_3m > 10) or (rel_1m and rel_1m > 5) or (adx_val and adx_val > 30)
            # หุ้นล้าหลัง (Laggard) = Momentum อ่อนแอกว่าตลาด
            is_laggard = (rel_3m and rel_3m < -5) or (adx_val and adx_val < 18)

            # ====== Action Recommendation (Momentum-Aware) ======
            if exit_signal == 'STRONG EXIT':
                if is_leader:
                    action       = 'ทยอยขาย (Leader)'
                    action_style = 'warning'
                    action_detail = f"สัญญาณขายแรง แต่หุ้นยังเป็นผู้นำ (RS สูง) - ขาย 70% เก็บ 30% เผื่อเด้งแรง"
                else:
                    action       = 'ออกทันที'
                    action_style = 'danger'
                    action_detail = f"ขายทั้งหมด {quantity:.0f} หุ้น - สัญญาณเทคนิคขาลงชัดเจนและหุ้นเริ่มล้าหลัง"
            
            elif sl_hit:
                if entry_price > 0 and current_price < entry_price:
                    if is_leader:
                        action       = 'เฝ้าจุดเด้ง (Cut?)'
                        action_style = 'warning'
                        action_detail = "หลุด SL แต่เป็นหุ้นผู้นำ - รอดูการดึงกลับที่เส้นค่าเฉลี่ย ถ้าไม่เด้งต้องคัท"
                    else:
                        action       = 'ตัดขาดทุน (Cut Loss)'
                        action_style = 'danger'
                        action_detail = f"ราคาหลุด SL และหุ้นอ่อนแอกว่าตลาด - แนะนำขายทันทีเพื่อปกป้องเงินทุน"
                else:
                    action       = 'ล็อกกำไร (Trailing)'
                    action_style = 'warning'
                    action_detail = f"ราคาหลุดจุดเฝ้าระวัง (SL) - กำไรยังเหลือ {gain_loss_pct:.1f}% แนะนำขายล็อกกำไร"

            elif exit_signal == 'EXIT':
                if is_leader:
                    action       = 'ถือต่อ (Leader)'
                    action_style = 'success'
                    action_detail = "มีสัญญาณขายบ้าง แต่ momentum แข็งแกร่งมาก - ถือต่อเพื่อรันเทรน"
                else:
                    action       = 'ทยอยขาย 50%'
                    action_style = 'warning'
                    action_detail = f"หุ้นเริ่มหมดแรงและไม่ใช่ผู้นำ - ขายครึ่งหนึ่งเก็บกำไรไว้ก่อน"

            elif is_laggard and gain_loss_pct < 0:
                action       = 'พิจารณาเปลี่ยนตัว'
                action_style = 'warning-soft'
                action_detail = "หุ้นเคลื่อนไหวช้ากว่าตลาด (Laggard) - แนะนำพิจารณาเปลี่ยนไปถือหุ้นผู้นำตัวอื่น"

            elif tp_price and current_price >= tp_price * 0.95:
                action       = 'ใกล้ TP'
                action_style = 'info'
                action_detail = f"ราคาใกล้เป้าหมายแล้ว - เตรียมทยอยรับทรัพย์"

            else:
                action       = 'ถือต่อ'
                action_style = 'success'
                action_detail = "ยังไม่มีสัญญาณออก และโครงสร้างราคายังดี - ถือรันเทรนต่อไป"

            # ====== Active Exit Triggers ======
            triggers = []
            if sl_hit:
                triggers.append({'label': 'SL HIT - หลุด Stop Loss', 'level': 'danger'})
            if tp_hit:
                triggers.append({'label': f'TP Hit - ถึงเป้า ฿{tp_price:.2f}', 'level': 'danger'})
            if rsi_val and rsi_val > 78:
                triggers.append({'label': f'RSI {rsi_val:.0f} - overbought มาก', 'level': 'danger'})
            elif rsi_val and rsi_val > 72:
                triggers.append({'label': f'RSI {rsi_val:.0f} - เริ่ม overbought', 'level': 'warning'})
            if not rvol_bullish and rvol >= 1.5:
                triggers.append({'label': f'RVOL {rvol:.1f}x Bear - แรงขายเข้ามา', 'level': 'danger'})
            elif not rvol_bullish:
                triggers.append({'label': 'RVOL Bear - volume หันขาลง', 'level': 'warning'})
            if rel_1m < -5:
                triggers.append({'label': f'Rel Mom 1m {rel_1m:.1f}% - แพ้ SET มาก', 'level': 'warning'})
            elif rel_1m < 0:
                triggers.append({'label': f'Rel Mom 1m {rel_1m:.1f}% - เริ่มแพ้ SET', 'level': 'info'})
            if price_pattern_score < -5:
                triggers.append({'label': f'Pattern: {price_pattern} - สัญญาณขาย', 'level': 'danger'})
            elif price_pattern_score < 0:
                triggers.append({'label': f'Pattern: {price_pattern}', 'level': 'warning'})
            if adx_val and adx_val < 20:
                triggers.append({'label': f'ADX {adx_val:.0f} - เทรนด์อ่อนแรง', 'level': 'warning'})
            if cmf_val is not None:
                if cmf_val < -0.1:
                    triggers.append({'label': f'CMF {cmf_val:.2f} - Distribution ชัดเจน เงินไหลออก', 'level': 'danger'})
                elif cmf_val < -0.05:
                    triggers.append({'label': f'CMF {cmf_val:.2f} - เริ่มมีแรงขายสุทธิ', 'level': 'warning'})
            if not triggers and exit_signal == '':
                triggers.append({'label': 'ไม่มีสัญญาณออก - ถือต่อได้', 'level': 'success'})
            
            # Turtle-specific signals (always show if near or hit)
            if s1_hit:
                triggers.append({'label': f'TURTLE S1 EXIT - หลุด 10D Low ({turtle_s1_exit:.2f})', 'level': 'danger'})
            elif turtle_s1_exit > 0 and current_price <= turtle_s1_exit * 1.02:
                triggers.append({'label': f'TURTLE S1 NEAR - ใกล้ 10D Low ({turtle_s1_exit:.2f})', 'level': 'warning'})

            if s2_hit:
                triggers.append({'label': f'TURTLE S2 EXIT - หลุด 20D Low ({turtle_s2_exit:.2f})', 'level': 'danger'})
            elif turtle_s2_exit > 0 and current_price <= turtle_s2_exit * 1.02:
                triggers.append({'label': f'TURTLE S2 NEAR - ใกล้ 20D Low ({turtle_s2_exit:.2f})', 'level': 'warning'})

            items.append({
                'obj':          item,
                'current_price': current_price,
                'day_change':   day_change,
                'entry_price':  entry_price,
                'gain_loss_pct': gain_loss_pct,
                'quantity':     quantity,
                'days_held':    days_held,
                'sl_price':     sl_price,
                'tp_price':     tp_price,
                'turtle_s1':    turtle_s1_exit,
                'turtle_s2':    turtle_s2_exit,
                's1_hit':       s1_hit,
                's2_hit':       s2_hit,
                'rsi':          rsi_val,
                'adx':          adx_val,
                'price_pattern': price_pattern,
                'price_pattern_score': price_pattern_score,
                'rel_1m':       rel_1m,
                'rel_3m':       rel_3m,
                'rvol':         rvol,
                'rvol_bullish': rvol_bullish,
                'sell_score':   sell_score,
                'exit_signal':  exit_signal,
                'current_pct':  current_pct,
                'entry_pct':    entry_pct,
                'sl_hit':       sl_hit,
                'tp_hit':       tp_hit,
                'action':       action,
                'action_style': action_style,
                'action_detail': action_detail,
                'triggers':     triggers,
                'is_leader':    is_leader,
                'is_laggard':   is_laggard,
                'cmf':          cmf_val,
            })
        except Exception as e:
            print(f"[ExitPlan] Error {item.symbol}: {e}")
            continue

    # เรียงตาม SELL Score สูงสุดก่อน
    items.sort(key=lambda x: x['sell_score'], reverse=True)

    # ====== Portfolio Health Summary ======
    urgent_count   = sum(1 for i in items if i['exit_signal'] in ('STRONG EXIT',) or i['sl_hit'])
    warning_count  = sum(1 for i in items if i['exit_signal'] == 'EXIT')
    watch_count    = sum(1 for i in items if i['exit_signal'] == 'WATCH')
    healthy_count  = sum(1 for i in items if not i['exit_signal'] and not i['sl_hit'])
    total_count    = len(items)
    avg_sell_score = round(sum(i['sell_score'] for i in items) / total_count, 1) if total_count else 0

    # Market Condition
    market_condition = {'phase': 'UNKNOWN', 'label': 'ไม่มีข้อมูล', 'color': 'secondary', 'score': 0}
    try:
        from datetime import datetime as _mcdt, timedelta as _mctd
        import pytz as _mcpytz
        _mc_bkk   = _mcpytz.timezone('Asia/Bangkok')
        _mc_now   = _mcdt.now(_mc_bkk)
        _mc_end   = _mc_now.date().strftime('%Y-%m-%d')
        _mc_start = (_mc_now.date() - _mctd(days=430)).strftime('%Y-%m-%d')
        _mc_df = yf.download("^SET", start=_mc_start, end=_mc_end, interval="1d", progress=False)
        if _mc_df is not None and not _mc_df.empty:
            if isinstance(_mc_df.columns, pd.MultiIndex):
                _mc_df.columns = _mc_df.columns.droplevel(1)
            market_condition = _get_market_condition(_mc_df)
    except Exception:
        pass

    return render(request, 'stocks/portfolio_exit_plan.html', {
        'items': items,
        'urgent_count':  urgent_count,
        'warning_count': warning_count,
        'watch_count':   watch_count,
        'healthy_count': healthy_count,
        'total_count':   total_count,
        'avg_sell_score': avg_sell_score,
        'market_condition': market_condition,
    })


# ====== Portfolio Management - เพิ่ม/ลบ รายการพอร์ต ======

@login_required
def add_to_portfolio(request):
    """
    รับ POST form เพิ่มหรืออัปเดต position ในพอร์ต
    ใช้ update_or_create เพื่อรองรับการแก้ไขข้อมูล (เช่น เพิ่ม quantity)
    """
    if request.method == 'POST':
        symbol = request.POST.get('symbol').upper()
        name = request.POST.get('name', '')
        quantity = request.POST.get('quantity', 0)
        entry_price = request.POST.get('entry_price', 0)
        category = request.POST.get('category', AssetCategory.STOCK)

        form = AddPortfolioForm(request.POST)
        if form.is_valid():
            symbol = form.cleaned_data['symbol']
            market = form.cleaned_data['market']
            
            # Standardize SET symbols to have .BK suffix
            if market == MarketType.SET and not symbol.endswith('.BK'):
                symbol = f"{symbol}.BK"

            Portfolio.objects.update_or_create(
                user=request.user,
                symbol=symbol,
                defaults={
                    'name': form.cleaned_data['name'],
                    'quantity': form.cleaned_data['quantity'],
                    'entry_price': form.cleaned_data['entry_price'],
                    'category': form.cleaned_data['category'],
                    'market': market,
                    'strategy': form.cleaned_data.get('strategy', ''),
                }
            )
            messages.success(request, f"เพิ่ม {symbol} เข้าพอร์ตเรียบร้อยแล้ว")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{error}")
    return redirect('stocks:portfolio_list')

@login_required
def delete_from_portfolio(request, pk):
    """ลบรายการจากพอร์ต (เฉพาะ object ของ user ปัจจุบันเท่านั้น)"""
    item = get_object_or_404(Portfolio, pk=pk, user=request.user)
    symbol = item.symbol
    item.delete()
    messages.success(request, f"ลบ {symbol} ออกจากพอร์ตแล้ว")
    return redirect('stocks:portfolio_list')

@login_required
def sell_stock(request, pk):
    """
    จัดการการขายหุ้นพร้อมคำนวณกำไร/ขาดทุน
    """
    portfolio_item = get_object_or_404(Portfolio, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = SellStockForm(request.POST)
        if form.is_valid():
            sell_quantity = form.cleaned_data['quantity']
            sell_price = form.cleaned_data['sell_price']

            if sell_quantity > portfolio_item.quantity:
                messages.error(request, f"จำนวนเกินที่ถือครองอยู่ (มีอยู่ {portfolio_item.quantity} หุ้น)")
                return redirect('stocks:portfolio_list')

            # คำนวณกำไร/ขาดทุน
            cost_of_sold_shares = sell_quantity * portfolio_item.entry_price
            sell_revenue = sell_quantity * sell_price
            profit_loss = sell_revenue - cost_of_sold_shares
            profit_loss_pct = (profit_loss / cost_of_sold_shares * 100) if cost_of_sold_shares > 0 else 0

            # ── Currency Conversion (Tithe calculation requirement) ──
            # ดึงอัตราแลกเปลี่ยน ณ เดี๋ยวนี้ (ตอนขาย)
            fx_rate = 1.0
            if portfolio_item.market == MarketType.US:
                fx_rate = _get_usd_thb()
            
            # บันทึกประวัติการขาย พร้อม market จาก Portfolio
            SoldStock.objects.create(
                user=request.user,
                symbol=portfolio_item.symbol,
                quantity=sell_quantity,
                buy_price=portfolio_item.entry_price,
                bought_at=portfolio_item.added_at,
                sell_price=sell_price,
                profit_loss=profit_loss,
                profit_loss_pct=profit_loss_pct,
                market=portfolio_item.market,
                settlement_rate=fx_rate,
                profit_loss_thb=float(profit_loss) * fx_rate,
                sell_revenue_thb=float(sell_revenue) * fx_rate,
            )

            # อัปเดตพอร์ต
            portfolio_item.quantity -= sell_quantity
            if portfolio_item.quantity <= 0:
                portfolio_item.delete()
                messages.success(request, f"ขาย {portfolio_item.symbol} เรียบร้อยแล้ว (ปิดสถานะ)")
            else:
                portfolio_item.save()
                messages.success(request, f"ขาย {portfolio_item.symbol} จำนวน {sell_quantity} หุ้น เรียบร้อยแล้ว")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{error}")

    return redirect('stocks:portfolio_list')

# ====== Recommendations - คำแนะนำหุ้นรายวันจาก AI ======

@login_required
def recommendations(request):
    """
    Thai Stock Recommendations with Legendary 5-Pillar Scoring.
    Implements Manual Scan and Persistence (AnalysisCache).
    """
    import random
    import json
    from datetime import datetime
    import pandas as pd
    import pandas_ta as ta
    import yfinance as yf
    from concurrent.futures import ThreadPoolExecutor

    def process_single_stock(sym):
        try:
            t = yf.Ticker(sym)
            try:
                inf = t.info
                if not isinstance(inf, dict): inf = {}
            except:
                inf = {}
            
            hist_short = t.history(period="1y")
            
            # If sym without .BK fails, try with .BK
            if (not inf or hist_short.empty) and ".BK" not in sym:
                try:
                    alt_t = yf.Ticker(f"{sym}.BK")
                    alt_inf = alt_t.info
                    alt_hist = alt_t.history(period="1y")
                    if isinstance(alt_inf, dict) and alt_inf:
                        t = alt_t
                        inf = alt_inf
                        hist_short = alt_hist
                except: pass

            rsi_val = 'N/A'
            rvol = 1.0
            if not hist_short.empty:
                rsi_series = ta.rsi(hist_short['Close'], length=14)
                if rsi_series is not None and not rsi_series.empty:
                    rsi_val = float(rsi_series.iloc[-1])
                current_vol = float(hist_short['Volume'].iloc[-1])
                avg_vol_20 = float(hist_short['Volume'].tail(20).mean())
                rvol = current_vol / avg_vol_20 if avg_vol_20 > 0 else 1.0

            de = 'N/A'
            try:
                bs = t.quarterly_balance_sheet if not t.quarterly_balance_sheet.empty else t.balance_sheet
                if not bs.empty:
                    col = bs.columns[0]
                    tot_liab = bs.loc['Total Liabilities Net Minority Interest', col] if 'Total Liabilities Net Minority Interest' in bs.index else bs.loc['Total Liabilities', col]
                    tot_eq = bs.loc['Stockholders Equity', col] if 'Stockholders Equity' in bs.index else bs.loc['Total Equity Gross Minority Interest', col]
                    de = tot_liab / tot_eq
            except: pass

            if de == 'N/A' or pd.isna(de):
                de = inf.get('debtToEquity', 'N/A')
                if isinstance(de, (int, float)): de = de / 100

            def scale_to_percent(val):
                if not isinstance(val, (int, float)): return val
                if abs(val) < 1.0: return val * 100
                return val

            pe = inf.get('trailingPE')
            pb = inf.get('priceToBook')
            peg = inf.get('pegRatio')
            roe = scale_to_percent(inf.get('returnOnEquity'))
            dy = scale_to_percent(inf.get('dividendYield'))
            npm = scale_to_percent(inf.get('profitMargins'))
            bv = inf.get('bookValue')
            price = inf.get('currentPrice') or inf.get('regularMarketPrice')

            p_graham = 0; p_buffett = 0; p_lynch = 0; p_greenblatt = 0; p_templeton = 0
            if isinstance(pe, (int, float)) and pe < 15: p_graham += 7
            if isinstance(pb, (int, float)) and pb < 1.2: p_graham += 7
            if isinstance(de, (int, float)) and de < 1.0: p_graham += 6
            fcf = inf.get('freeCashflow')
            if isinstance(roe, (int, float)) and roe > 18: p_buffett += 8
            if isinstance(fcf, (int, float)) and fcf > 0: p_buffett += 7
            if isinstance(npm, (int, float)) and npm > 10: p_buffett += 5
            eg = inf.get('earningsGrowth')
            if isinstance(peg, (int, float)) and 0 < peg < 1.0: p_lynch += 12
            elif isinstance(peg, (int, float)) and 0 < peg < 1.5: p_lynch += 8
            if isinstance(eg, (int, float)) and eg > 0.20: p_lynch += 8
            ev = inf.get('enterpriseValue'); ebitda = inf.get('ebitda')
            if ev and ebitda:
                ey = ebitda / ev
                if ey > 0.12: p_greenblatt += 10
            if isinstance(roe, (int, float)) and roe > 15: p_greenblatt += 10
            if isinstance(pe, (int, float)) and pe < 10: p_templeton += 10

            final_score = (p_graham * 0.75) + (p_buffett * 0.75) + (p_lynch * 2.0) + (p_greenblatt * 1.0) + (p_templeton * 0.5)
            momentum_bonus = 0
            if isinstance(rsi_val, (int, float)) and 30 <= rsi_val <= 50: momentum_bonus += 5
            if rvol > 1.2: momentum_bonus += 5
            final_score = min(100, final_score + momentum_bonus)

            ev_spread = (roe - 10.0) if isinstance(roe, (int, float)) else None
            
            # ====== ENHANCED VALUATION FRAMEWORK (Thai Market) ======
            # วิธีที่ใช้: 3 วิธีผสมกันแบบ Weighted Average
            #   1. Graham Number       - พื้นฐานราคาตามทรัพย์สิน
            #   2. Graham Revised      - คำนึงถึง Growth + Bond Yield (สูตรทองของ Graham)
            #   3. DCF (FCF-based)     - มูลค่าจากกระแสเงินสดอิสระ (multiplier ปรับตาม Growth)
            v_graham = None; v_graham_rev = None; v_dcf = None
            fair_value = None; upside = None; mos_price = None

            # ใช้ Forward EPS ก่อน (มองไปข้างหน้า) ถ้าไม่มีให้ใช้ Trailing EPS
            eps = inf.get('forwardEps') or inf.get('trailingEps') or (price / pe if pe and pe > 0 and price else 0)
            bv = inf.get('bookValue') or 0
            # เฉลี่ย Earnings Growth + Revenue Growth เพื่อลด noise
            eg = inf.get('earningsGrowth') or 0
            rg = inf.get('revenueGrowth') or 0
            raw_growth = ((eg + rg) / 2) if eg and rg else (eg or rg or 0.10)
            fcf = inf.get('freeCashflow')
            shares = inf.get('sharesOutstanding')

            # Bond Yield ของไทย (ใช้ค่าประมาณ Thai 10-yr Govt Bond)
            THAI_BOND_YIELD = 3.0

            import math
            try:
                # 1. Graham Number: √(22.5 × EPS × Book Value)
                #    เหมาะกับหุ้นที่มีทรัพย์สินมาก (Banks, Property, Industrial)
                if isinstance(eps, (int, float)) and eps > 0 and isinstance(bv, (int, float)) and bv > 0:
                    v_graham = math.sqrt(22.5 * eps * bv)

                # 2. Graham Revised: EPS × (8.5 + 2g) × (4.4 / bond_yield)
                #    สูตรดั้งเดิมของ Graham ที่ปรับด้วย Bond Yield เพื่อสะท้อนสภาพดอกเบี้ย
                g_rate = min(25, max(3, raw_growth * 100 if raw_growth < 1 else raw_growth))
                bond_adj = 4.4 / THAI_BOND_YIELD  # ปรับตามดอกเบี้ย (สูงกว่า 1.0 เมื่อ yield ต่ำ)
                if isinstance(eps, (int, float)) and eps > 0:
                    v_graham_rev = eps * (8.5 + 2 * g_rate) * bond_adj

                # 3. FCF-based DCF: (FCF/Share) × Multiplier
                #    Multiplier ปรับตาม Growth Rate (growth สูง = มูลค่าในอนาคตสูงกว่า)
                if fcf and shares and shares > 0:
                    fcf_per_share = fcf / shares
                    if fcf_per_share > 0:
                        fcf_multiple = 20 if g_rate > 15 else (15 if g_rate >= 7 else 12)
                        v_dcf = fcf_per_share * fcf_multiple

                # รวม 3 วิธีด้วย Weighted Average (Graham Rev มีน้ำหนักมากสุด)
                weighted_sum = 0.0; total_weight = 0.0
                if v_graham and v_graham > 0:
                    weighted_sum += v_graham * 1.0; total_weight += 1.0
                if v_graham_rev and v_graham_rev > 0:
                    weighted_sum += v_graham_rev * 2.0; total_weight += 2.0
                if v_dcf and v_dcf > 0:
                    weighted_sum += v_dcf * 1.5; total_weight += 1.5

                if total_weight > 0:
                    fair_value = weighted_sum / total_weight
                elif bv > 0:
                    fair_value = bv * 0.9  # Fallback: 90% ของ Book Value

                if fair_value and price and price > 0:
                    fair_value = min(fair_value, price * 3.0)  # Cap ที่ 200% upside
                    upside = ((fair_value / price) - 1) * 100
                    mos_price = fair_value * 0.75  # ราคาที่ควรซื้อ = Fair Value - 25% Margin of Safety
            except: pass

            return {
                'symbol': sym, 'name': inf.get('longName', sym),
                'pe': pe, 'pb': pb, 'roe': roe, 'dy': dy, 'npm': npm, 'de': de,
                'rsi': rsi_val, 'peg': peg, 'price': price, 'rvol': round(rvol, 2),
                'ev_spread': ev_spread,
                'fair_value': round(fair_value, 2) if fair_value else None,
                'mos_price': round(mos_price, 2) if mos_price else None,
                'upside': round(upside, 1) if upside else None,
                'value_score': round(final_score, 1),
                'legendary': {'graham': p_graham, 'buffett': p_buffett, 'lynch': p_lynch, 'greenblatt': p_greenblatt, 'templeton': p_templeton}
            }
        except: return None

    cache_symbol = 'THAI_REC_ALL'
    cached_data = AnalysisCache.objects.filter(user=request.user, symbol=cache_symbol).first()
    
    stock_previews = []
    last_scanned = None
    if cached_data:
        try:
            stock_previews = json.loads(cached_data.analysis_data)
            last_scanned = cached_data.last_updated
        except: pass

    # If manual scan is requested
    if request.GET.get('scan') == 'true':
        set100_pool = [
            'ADVANC.BK', 'AOT.BK', 'AWC.BK', 'BANPU.BK', 'BBL.BK', 'BDMS.BK', 'BEM.BK', 'BGRIM.BK', 'BH.BK', 'BJC.BK',
            'BTS.BK', 'CBG.BK', 'CENTEL.BK', 'COM7.BK', 'CPALL.BK', 'CPAXT.BK', 'CPF.BK', 'CPN.BK', 'CRC.BK', 'DELTA.BK',
            'EA.BK', 'EGCO.BK', 'GLOBAL.BK', 'GPSC.BK', 'GULF.BK', 'HMPRO.BK', 'INTUCH.BK', 'IRPC.BK', 'IVL.BK', 'JMART.BK',
            'JMT.BK', 'KBANK.BK', 'KCE.BK', 'KKP.BK', 'KTB.BK', 'KTC.BK', 'LH.BK', 'MINT.BK', 'MTC.BK', 'OR.BK',
            'OSP.BK', 'PTT.BK', 'PTTEP.BK', 'PTTGC.BK', 'RATCH.BK', 'SCB.BK', 'SCC.BK', 'SCGP.BK', 'TISCO.BK', 'TOP.BK',
            'TRUE.BK', 'TTB.BK', 'TU.BK', 'WHA.BK', 'AMATA.BK', 'AP.BK', 'BAM.BK', 'BCH.BK', 'BCP.BK', 'BCPG.BK',
            'BLA.BK', 'BPP.BK', 'CHG.BK', 'CK.BK', 'CKP.BK', 'DOHOME.BK', 'ERW.BK', 'FORTH.BK', 'GUNKUL.BK', 'HANA.BK',
            'ICHI.BK', 'ITC.BK', 'M.BK', 'MBK.BK', 'MEGA.BK', 'ORI.BK', 'PLANB.BK', 'PRM.BK', 'PSL.BK', 'PTG.BK',
            'QH.BK', 'RCL.BK', 'ROJNA.BK', 'RS.BK', 'SABINA.BK', 'SAWAD.BK', 'SINGER.BK', 'SIRI.BK', 'SPALI.BK', 'SPRC.BK',
            'STA.BK', 'STEC.BK', 'STGT.BK', 'SUPER.BK', 'TASCO.BK', 'TCAP.BK', 'THANI.BK', 'THCOM.BK', 'THG.BK', 'TIDLOR.BK',
            'TKN.BK', 'TLI.BK', 'TOA.BK', 'TPIPL.BK', 'TPIPP.BK', 'TQM.BK', 'TTA.BK', 'VGI.BK', 'WHAUP.BK'
        ]
        mai_pool = [
            'SPA.BK', 'AU.BK', 'D.BK', 'CHAYO.BK', 'YGG.BK', 'BE8.BK', 'BBIK.BK', 'SNNP.BK', 'TNP.BK', 'TACC.BK',
            'SICT.BK', 'ADD.BK', 'ABM.BK', 'CHO.BK', 'PSTC.BK', 'TVDH.BK', 'NDR.BK', 'BOL.BK', 'IP.BK', 'PLANET.BK'
        ]
        full_pool = set100_pool + mai_pool
        candidate_symbols = random.sample(full_pool, min(100, len(full_pool))) # Limit to 100 for speed
        
        with ThreadPoolExecutor(max_workers=10) as executor:
            scanned_list = list(filter(None, executor.map(process_single_stock, candidate_symbols)))
        
        # Save to Cache
        scanned_list = sorted(scanned_list, key=lambda x: x['value_score'], reverse=True)
        AnalysisCache.objects.update_or_create(
            user=request.user, symbol=cache_symbol,
            defaults={'analysis_data': json.dumps(scanned_list)}
        )
        stock_previews = scanned_list
        last_scanned = datetime.now()


    # Generate Report
    report_text = None
    if request.GET.get('analyze') == 'true' and stock_previews:
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        data_str = "\n".join([f"{s['symbol']} Price:{s['price']} Score:{s['value_score']} PEG:{s['peg']}" for s in stock_previews[:30]])
        prompt = f"""คุณคือนักวิเคราะห์หุ้นที่เน้น Maximum Returns ในปี 2026 โดยใช้สไตล์ Peter Lynch (GARP) และ Greenblatt (Magic Formula) เป็นหลัก
        นี่คือข้อมูลหุ้น 30 อันดับแรก:
        {data_str}
        โปรดวิเคราะห์หุ้นที่น่าเข้าซื้อที่สุดสำหรับปี 2026 ภายใต้ธีม AI และ พลังงานสะอาด เขียนรายงานภาษาไทยแบบมืออาชีพ เจาะลึกรายตัวท็อป 10"""
        try:
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            report_text = response.text
            if report_text.startswith("```markdown"): report_text = report_text[11:].strip()
            if report_text.endswith("```"): report_text = report_text[:-3].strip()
        except Exception as e: report_text = f"Error: {e}"

    context = {
        'title': 'AI Thai Stock Recommendations',
        'report': report_text,
        'stocks': stock_previews,
        'market': 'thai',
        'last_scanned': last_scanned
    }
    return render(request, 'stocks/recommendations.html', context)


@login_required
def us_recommendations(request):
    """
    US Stock Recommendations with Manual Scan and Persistence.
    """
    import random
    import json
    from datetime import datetime
    import pandas as pd
    import pandas_ta as ta

    cache_symbol = 'US_REC_ALL'
    cached_data = AnalysisCache.objects.filter(user=request.user, symbol=cache_symbol).first()
    
    stock_previews = []
    last_scanned = None
    if cached_data:
        try:
            stock_previews = json.loads(cached_data.analysis_data)
            last_scanned = cached_data.last_updated
        except: pass

    if request.GET.get('scan') == 'true':
        full_pool = [
            'AAPL', 'MSFT', 'GOOGL', 'AMZN', 'NVDA', 'META', 'TSLA', 'BRK-B', 'UNH', 'JNJ',
            'XOM', 'JPM', 'V', 'PG', 'MA', 'CVX', 'HD', 'PFE', 'ABBV', 'KO',
            'PEP', 'COST', 'MCD', 'WMT', 'DIS', 'ADBE', 'CRM', 'NFLX', 'AMD', 'INTC',
            'QCOM', 'TXN', 'AMAT', 'MU', 'LRCX', 'AVGO', 'NKE', 'TMUS', 'SBUX', 'LOW',
            'UPS', 'CAT', 'GE', 'HON', 'DE', 'MMM', 'LMT', 'BA', 'RTX', 'T'
        ]
        candidate_symbols = random.sample(full_pool, min(50, len(full_pool)))
        
        scanned_list = []
        for sym in candidate_symbols:
            try:
                t = yf.Ticker(sym)
                inf = t.info
                hist = t.history(period="1y")
                if hist.empty: continue
                
                rsi_series = ta.rsi(hist['Close'], length=14)
                rsi_val = float(rsi_series.iloc[-1]) if rsi_series is not None and not rsi_series.empty else None
                current_vol = float(hist['Volume'].iloc[-1])
                avg_vol_20 = float(hist['Volume'].tail(20).mean())
                rvol = current_vol / avg_vol_20 if avg_vol_20 > 0 else 1.0

                def scale_to_percent(val):
                    if not isinstance(val, (int, float)): return val
                    if abs(val) < 1.0: return val * 100
                    return val

                pe = inf.get('trailingPE')
                pb = inf.get('priceToBook')
                peg = inf.get('pegRatio')
                roe = scale_to_percent(inf.get('returnOnEquity'))
                dy = scale_to_percent(inf.get('dividendYield'))
                npm = scale_to_percent(inf.get('profitMargins'))
                price = inf.get('currentPrice') or inf.get('regularMarketPrice')
                de = inf.get('debtToEquity')
                if isinstance(de, (int, float)) and de > 5: de = de / 100

                # 5 PILLARS v2 (US Focus)
                p_graham = 0; p_buffett = 0; p_lynch = 0; p_greenblatt = 0; p_templeton = 0
                
                if isinstance(pe, (int, float)) and pe < 18: p_graham += 7
                if isinstance(de, (int, float)) and de < 1.0: p_graham += 6
                
                fcf = inf.get('freeCashflow')
                if isinstance(roe, (int, float)) and roe > 20: p_buffett += 8
                if isinstance(fcf, (int, float)) and fcf > 0: p_buffett += 7
                
                eg = inf.get('earningsGrowth')
                if isinstance(peg, (int, float)) and 0 < peg < 1.0: p_lynch += 12
                if isinstance(eg, (int, float)) and eg > 0.15: p_lynch += 8
                
                ev = inf.get('enterpriseValue'); ebitda = inf.get('ebitda')
                if ev and ebitda:
                    ey = ebitda / ev
                    if ey > 0.10: p_greenblatt += 10
                if isinstance(roe, (int, float)) and roe > 15: p_greenblatt += 10
                
                if isinstance(pe, (int, float)) and pe < 12: p_templeton += 10
                if isinstance(pb, (int, float)) and pb < 1.2: p_templeton += 10

                final_score = (p_graham * 0.75) + (p_buffett * 0.75) + (p_lynch * 2.0) + (p_greenblatt * 1.0) + (p_templeton * 0.5)
                final_score = min(100, final_score + 5 if rvol > 1.2 else final_score)

                # Economic Profit (ROE - CoE 10%)
                ev_spread = None
                if isinstance(roe, (int, float)):
                    ev_spread = roe - 10.0

                # ====== ENHANCED VALUATION FRAMEWORK (US Market) ======
                # ใช้ 3 วิธีเหมือน Thai Framework แต่ Bond Yield อิง US 10-yr Treasury
                fair_value = None; upside = None; mos_price = None
                v_graham = None; v_graham_rev = None; v_dcf = None

                eps = inf.get('forwardEps') or inf.get('trailingEps') or (price / pe if pe and pe > 0 and price else 0)
                bv = inf.get('bookValue') or 0
                eg = inf.get('earningsGrowth') or 0
                rg = inf.get('revenueGrowth') or 0
                raw_growth = ((eg + rg) / 2) if eg and rg else (eg or rg or 0.10)
                fcf_us = inf.get('freeCashflow')
                shares_us = inf.get('sharesOutstanding')

                # Bond Yield ของสหรัฐฯ (US 10-yr Treasury - ประมาณ 4.4%)
                US_BOND_YIELD = 4.4

                if isinstance(eps, (int, float)):
                    import math
                    try:
                        g_rate = min(25, max(3, raw_growth * 100 if raw_growth < 1 else raw_growth))
                        bond_adj = 4.4 / US_BOND_YIELD  # ≈ 1.0 เมื่อ yield ปกติ

                        # 1. Graham Number
                        if eps > 0 and isinstance(bv, (int, float)) and bv > 0:
                            v_graham = math.sqrt(22.5 * eps * bv)

                        # 2. Graham Revised + Bond Yield adjustment
                        if eps > 0:
                            v_graham_rev = eps * (8.5 + 2 * g_rate) * bond_adj

                        # 3. DCF (FCF-based, dynamic multiplier)
                        if fcf_us and shares_us and shares_us > 0:
                            fcf_ps = fcf_us / shares_us
                            if fcf_ps > 0:
                                fcf_multiple = 20 if g_rate > 15 else (15 if g_rate >= 7 else 12)
                                v_dcf = fcf_ps * fcf_multiple

                        # Weighted Average (Graham Rev: น้ำหนัก 2, DCF: 1.5, Graham No.: 1)
                        ws = 0.0; tw = 0.0
                        if v_graham and v_graham > 0: ws += v_graham * 1.0; tw += 1.0
                        if v_graham_rev and v_graham_rev > 0: ws += v_graham_rev * 2.0; tw += 2.0
                        if v_dcf and v_dcf > 0: ws += v_dcf * 1.5; tw += 1.5

                        if tw > 0:
                            fair_value = ws / tw
                        elif bv > 0:
                            fair_value = bv * 0.7  # Fallback: 70% ของ Book Value

                        if fair_value and price and price > 0:
                            fair_value = min(fair_value, price * 3.0)  # Cap ที่ 200% upside
                            upside = ((fair_value / price) - 1) * 100
                            mos_price = fair_value * 0.75  # ราคาที่ควรซื้อ (25% Margin of Safety)
                    except: pass

                scanned_list.append({
                    'symbol': sym, 'name': inf.get('shortName', sym),
                    'pe': pe, 'pb': pb, 'roe': roe, 'dy': dy, 'npm': npm, 'de': de,
                    'rsi': rsi_val, 'peg': peg, 'price': price, 'rvol': round(rvol, 2),
                    'ev_spread': ev_spread,
                    'fair_value': round(fair_value, 2) if fair_value else None,
                    'mos_price': round(mos_price, 2) if mos_price else None,
                    'upside': round(upside, 1) if upside else None,
                    'value_score': round(final_score, 1),
                    'legendary': {'graham': p_graham, 'buffett': p_buffett, 'lynch': p_lynch, 'greenblatt': p_greenblatt, 'templeton': p_templeton}
                })
            except: continue
            
        scanned_list = sorted(scanned_list, key=lambda x: x['value_score'], reverse=True)
        AnalysisCache.objects.update_or_create(
            user=request.user, symbol=cache_symbol,
            defaults={'analysis_data': json.dumps(scanned_list)}
        )
        stock_previews = scanned_list
        last_scanned = datetime.now()

    report_text = None
    if request.GET.get('analyze') == 'true' and stock_previews:
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        data_str = "\n".join([f"{s['symbol']} Price:{s['price']} Score:{s['value_score']} PEG:{s['peg']}" for s in stock_previews[:20]])
        prompt = f"คุณคือนักวิเคราะห์หุ้นอเมริกัน เน้น Maximum Returns ปี 2026 โดยใช้ Lynch และ Greenblatt สแกนหุ้น {data_str} โปรดสรุปตัวท็อปในธีม AI/Semiconductor/Energy ภาษาไทย"
        try:
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            report_text = response.text
        except: report_text = "Analysis service unavailable."

    context = {
        'title': 'AI US Stock Recommendations',
        'report': report_text,
        'stocks': stock_previews,
        'market': 'us',
        'last_scanned': last_scanned
    }
    return render(request, 'stocks/recommendations.html', context)



# ====== Morning Briefing - รายงานสรุปประจำวัน ======

@login_required
def morning_briefing(request):
    """
    รายงานสรุปประจำวัน - กดปุ่มเดียว AI รวมข้อมูลทั้งหมด:
    Portfolio + Momentum SET/US + Precision + SEPA + Cup&Handle + Macro
    แล้วสร้างแผนซื้อ/ขายและภาพรวมเศรษฐกิจ
    """
    from .models import (
        MorningBriefing as _MB, Portfolio as _Port,
        MomentumCandidate as _MC, PrecisionScanCandidate as _PSC,
        CupHandleCandidate as _CHC, USSepaCandidate as _USC,
    )
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    cache_key = f'morning_briefing_{request.user.id}'

    # ── AJAX poll ────────────────────────────────────────────────
    if request.GET.get('mb_status') == '1':
        return _JR(_cp.get(cache_key, {'state': 'idle'}))

    # ── POST: trigger generation ──────────────────────────────────
    if request.method == 'POST':
        existing = _cp.get(cache_key, {})
        if existing.get('state') == 'running':
            return redirect('stocks:morning_briefing')

        _cp.set(cache_key, {'state': 'running', 'phase': 'กำลังรวบรวมข้อมูล…'}, timeout=600)

        def _run(uid, ckey):
            import django; django.setup()
            import yfinance as _yf
            from django.core.cache import cache as _c
            from django.contrib.auth import get_user_model
            from django.utils import timezone as tz
            from .models import (
                MorningBriefing as MB, Portfolio,
                MomentumCandidate, PrecisionScanCandidate,
                CupHandleCandidate, USSepaCandidate,
            )
            from django.conf import settings as _s
            from google import genai as _genai

            User = get_user_model()
            user = User.objects.get(pk=uid)

            try:
                # ── 1. Portfolio ────────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Portfolio…'}, timeout=600)
                portfolio = list(Portfolio.objects.filter(user=user))
                port_lines = []
                for p in portfolio:
                    try:
                        ticker_sym = p.symbol if not p.symbol.endswith('.BK') else p.symbol
                        hist = _yf.Ticker(ticker_sym).history(period='5d')
                        if hist.empty:
                            hist = _yf.Ticker(f'{p.symbol}.BK').history(period='5d')
                        if not hist.empty:
                            cur = float(hist['Close'].iloc[-1])
                            entry = float(p.entry_price)
                            pl_pct = (cur - entry) / entry * 100
                            port_lines.append(f"  - {p.symbol}: ราคาปัจจุบัน {cur:.2f} (ทุน {entry:.2f}, P/L {pl_pct:+.1f}%)")
                        else:
                            port_lines.append(f"  - {p.symbol}: ทุน {float(p.entry_price):.2f} (ไม่สามารถดึงราคาได้)")
                    except Exception:
                        port_lines.append(f"  - {p.symbol}: ทุน {float(p.entry_price):.2f}")

                # ── 2. Momentum SET ─────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Momentum SET…'}, timeout=600)
                mom_set = list(MomentumCandidate.objects.filter(user=user, market='SET').order_by('-technical_score')[:10])
                mom_set_lines = [f"  - {c.symbol}: Score={c.technical_score} RSI={c.rsi:.0f} RS={c.rs_rating} Price={c.price:.2f}" for c in mom_set]

                # ── 3. Momentum US ──────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Momentum US…'}, timeout=600)
                mom_us = list(MomentumCandidate.objects.filter(user=user, market='US').order_by('-technical_score')[:10])
                mom_us_lines = [f"  - {c.symbol}: Score={c.technical_score} RSI={c.rsi:.0f} RS={c.rs_rating} Price={c.price:.2f}" for c in mom_us]

                # ── 4. Precision SET (latest run) ───────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Precision…'}, timeout=600)
                prec_run = PrecisionScanCandidate.objects.filter(user=user, market='SET').values_list('scan_run', flat=True).order_by('-scan_run').first()
                prec_set = []
                if prec_run:
                    prec_set = list(PrecisionScanCandidate.objects.filter(user=user, market='SET', scan_run=prec_run).order_by('-technical_score')[:8])
                prec_lines = [f"  - {c.symbol}: Score={c.technical_score} RS={c.rs_rating} Stage2={'✓' if c.stage2 else '✗'} RR={c.risk_reward_ratio:.1f} Prox={c.zone_proximity:.1f}%" for c in prec_set]

                # ── 5. SEPA SET (from Precision) ────────────────
                sepa_set = [c for c in prec_set if c.stage2 and c.rs_rating >= 70]
                sepa_lines = [f"  - {c.symbol}: RS={c.rs_rating} VCP={'✓' if c.vcp_setup else '✗'} Score={c.technical_score}" for c in sepa_set]

                # ── 6. Cup & Handle ─────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Cup & Handle…'}, timeout=600)
                cup_run = CupHandleCandidate.objects.filter(user=user).values_list('scan_run', flat=True).order_by('-scan_run').first()
                cup_list = []
                if cup_run:
                    cup_list = list(CupHandleCandidate.objects.filter(user=user, scan_run=cup_run).order_by('-rs_rating')[:8])
                cup_lines = [f"  - {c.symbol}: Price={c.price:.2f} Breakout={c.breakout_price:.2f} Target={c.target_price:.2f} RS={c.rs_rating}" for c in cup_list]

                # ── 7. US SEPA ──────────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล US SEPA…'}, timeout=600)
                us_sepa_run = USSepaCandidate.objects.filter(user=user).values_list('scan_run', flat=True).order_by('-scan_run').first()
                us_sepa_list = []
                if us_sepa_run:
                    us_sepa_list = list(USSepaCandidate.objects.filter(user=user, scan_run=us_sepa_run, stage2=True).order_by('-rs_rating')[:8])
                us_sepa_lines = [f"  - {c.symbol}: RS={c.rs_rating} VCP={'✓' if c.vcp_setup else '✗'} ADX={c.adx:.0f} Price={c.price:.2f}" for c in us_sepa_list]

                # ── 8. Macro data ───────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'ดึงข้อมูล Macro…'}, timeout=600)
                macro_symbols = {
                    'SET Index': '^SET', 'S&P 500': '^GSPC', 'Nasdaq': '^IXIC',
                    'USD/THB': 'USDTHB=X', 'DXY': 'DX-Y.NYB', 'US 10Y Yield': '^TNX',
                    'Gold': 'GC=F', 'WTI Oil': 'CL=F', 'Bitcoin': 'BTC-USD',
                }
                macro_lines = []
                for name, sym in macro_symbols.items():
                    try:
                        h = _yf.Ticker(sym).history(period='5d')
                        if not h.empty and len(h) >= 2:
                            cur = float(h['Close'].iloc[-1])
                            prev = float(h['Close'].iloc[-2])
                            chg = (cur - prev) / prev * 100
                            macro_lines.append(f"  - {name}: {cur:.2f} ({chg:+.2f}%)")
                    except Exception:
                        pass

                # ── 9. Build prompt ─────────────────────────────
                _c.set(ckey, {'state': 'running', 'phase': 'AI กำลังวิเคราะห์และสร้างรายงาน…'}, timeout=600)

                _now = tz.now()
                _months_th = ['', 'มกราคม','กุมภาพันธ์','มีนาคม','เมษายน','พฤษภาคม','มิถุนายน',
                              'กรกฎาคม','สิงหาคม','กันยายน','ตุลาคม','พฤศจิกายน','ธันวาคม']
                today_str = f"{_now.day} {_months_th[_now.month]} {_now.year}"  # ปี ค.ศ. เสมอ

                prompt = f"""คุณคือ Senior Portfolio Manager และ Macro Strategist ระดับสถาบัน
วันที่: {today_str}

จงสร้าง **Morning Briefing Report** ภาษาไทยแบบครบถ้วน จากข้อมูลด้านล่าง:

---
## 📊 MACRO ECONOMY
{chr(10).join(macro_lines) if macro_lines else 'ไม่มีข้อมูล'}

## 💼 PORTFOLIO (หุ้นที่ถืออยู่)
{chr(10).join(port_lines) if port_lines else 'ไม่มีพอร์ต'}

## 🇹🇭 MOMENTUM SET (Top 10 by Score)
{chr(10).join(mom_set_lines) if mom_set_lines else 'ยังไม่ได้สแกน'}

## 🇺🇸 MOMENTUM US (Top 10 by Score)
{chr(10).join(mom_us_lines) if mom_us_lines else 'ยังไม่ได้สแกน'}

## 🎯 PRECISION SCAN - SET (Top 8)
{chr(10).join(prec_lines) if prec_lines else 'ยังไม่ได้สแกน'}

## 🦅 SEPA - SET Stage 2 + RS≥70
{chr(10).join(sepa_lines) if sepa_lines else 'ยังไม่ได้สแกน'}

## ☕ CUP & HANDLE - SET
{chr(10).join(cup_lines) if cup_lines else 'ยังไม่ได้สแกน'}

## 🦅 US SEPA - Stage 2 + VCP
{chr(10).join(us_sepa_lines) if us_sepa_lines else 'ยังไม่ได้สแกน'}

---
จงเขียนรายงานเป็นภาษาไทย **Markdown** โดยมีหัวข้อดังนี้:

## 1. 🌍 ภาพรวมเศรษฐกิจและตลาดโลกวันนี้
วิเคราะห์ Macro: Risk-on/Risk-off, Fund Flow, SET vs S&P500, ทิศทางดอกเบี้ย, Gold/BTC

## 2. 💼 สถานะพอร์ต - ควรทำอะไรวันนี้?
รายหุ้นใน Portfolio - แต่ละตัวควร: ✅ Hold | ➕ Add | ⚠️ Trail Stop | 🔴 ขาย
ระบุเหตุผลสั้น ๆ จาก P/L% และบรรยากาศตลาด

## 3. 🇹🇭 หุ้น SET น่าสนใจวันนี้
Top 3-5 จาก Momentum + Precision + SEPA + Cup&Handle รวมกัน
พร้อม Entry Zone, Stop Loss, Target และ Priority (🔥 สูง / ⚡ กลาง / 👀 เฝ้าดู)

## 4. 🇺🇸 หุ้น US น่าสนใจวันนี้
Top 3-5 จาก Momentum US + US SEPA
พร้อม Entry, Stop, Target และ Priority

## 5. ⚡ สรุปแผนปฏิบัติการวันนี้
ตารางสรุป: หุ้น | ตลาด | Action | ราคาเข้า | Stop | เหตุผล
เรียงตาม Priority สูงสุดก่อน

## 6. ⚠️ ความเสี่ยงที่ต้องระวังวันนี้
Macro risks, Earnings, การเมือง หรือสัญญาณที่น่าเป็นห่วง
"""

                client = _genai.Client(api_key=_s.GEMINI_API_KEY)
                resp = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt,
                )
                report_text = resp.text or '## ไม่สามารถสร้างรายงานได้'

                # ── 10. Save to DB ──────────────────────────────
                MB.objects.create(
                    user=user,
                    report_md=report_text,
                    portfolio_count=len(portfolio),
                    momentum_set_count=len(mom_set),
                    momentum_us_count=len(mom_us),
                    precision_count=len(prec_set),
                    sepa_count=len(sepa_set),
                    cup_handle_count=len(cup_list),
                )

                # Keep only 7 latest reports per user
                old_ids = list(MB.objects.filter(user=user).order_by('-created_at').values_list('id', flat=True)[7:])
                if old_ids:
                    MB.objects.filter(id__in=old_ids).delete()

                _c.set(ckey, {'state': 'done'}, timeout=300)

            except Exception as exc:
                import logging
                logging.getLogger('stocks').exception(f'[MorningBriefing] error: {exc}')
                from django.core.cache import cache as _c2
                _c2.set(ckey, {'state': 'done', 'error': str(exc)}, timeout=60)
                from django.core.cache import cache as _c2
                _c2.set(ckey, {'state': 'done', 'error': str(exc)}, timeout=60)

        import threading as _th
        _th.Thread(target=_run, args=(request.user.id, cache_key), daemon=True).start()
        return redirect('stocks:morning_briefing')

    # ── GET: display ─────────────────────────────────────────────
    import json as _json
    from .models import MorningBriefing as _MB
    is_generating = (_cp.get(cache_key, {}).get('state') == 'running')
    briefings = list(_MB.objects.filter(user=request.user)[:7])

    # Pass markdown as safe JSON so template JS can render without escape issues
    briefings_md_json = _json.dumps([b.report_md for b in briefings])

    return render(request, 'stocks/morning_briefing.html', {
        'briefings':       briefings,
        'latest':          briefings[0] if briefings else None,
        'is_generating':   is_generating,
        'briefings_md_json': briefings_md_json,
    })


# ====== Macro Economy - ภาพรวมเศรษฐกิจมหภาคและสินค้าโภคภัณฑ์ ======

@login_required
def macro_economy(request):
    """
    ดึงข้อมูล SET Index, USD/THB, ทองคำ, น้ำมัน WTI/Brent
    แสดงกราฟย้อนหลัง 3 เดือน และวิเคราะห์ด้วย AI ตามคำขอ
    """
    # รายการข้อมูลมหภาคที่ต้องดึงพร้อม symbol Yahoo Finance
    macro_items = [
        {'id': 'set', 'name': 'SET Index (ดัชนีหุ้นไทย)', 'symbol': '^SET', 'unit': 'Points', 'desc': 'ดัชนีตลาดหลักทรัพย์แห่งประเทศไทย บ่งบอกสภาวะตลาดโดยรวม'},
        {'id': 'spx', 'name': 'S&P 500 (ตลาดหุ้นสหรัฐฯ)', 'symbol': '^GSPC', 'unit': 'Points', 'desc': 'ดัชนีหุ้นใหญ่ 500 ตัวของสหรัฐฯ สะท้อนภาพรวมตลาดโลก'},
        {'id': 'usdthb', 'name': 'USD/THB (อัตราแลกเปลี่ยน)', 'symbol': 'USDTHB=X', 'unit': 'THB', 'desc': 'ค่าเงินบาทเทียบดอลลาร์สิงคโปร์/สหรัฐฯ'},
        {'id': 'dxy', 'name': 'Dollar Index (DXY)', 'symbol': 'DX-Y.NYB', 'unit': 'Points', 'desc': 'ดัชนีดอลลาร์สหรัฐ บ่งบอกถึงกระแสเงินทุน (Fund Flow)'},
        {'id': 'us10y', 'name': 'US 10Y Bond Yield', 'symbol': '^TNX', 'unit': '%', 'desc': 'อัตราดอกเบี้ยพันธบัตรฯ 10 ปี สหรัฐฯ'},
        {'id': 'gold', 'name': 'Gold (ทองคำโลก)', 'symbol': 'GC=F', 'unit': 'USD/oz', 'desc': 'สินทรัพย์ปลอดภัยและตัววัดเงินเฟ้อ'},
        {'id': 'btc', 'name': 'Bitcoin (BTC-USD)', 'symbol': 'BTC-USD', 'unit': 'USD', 'desc': 'สินทรัพย์ดิจิทัล (Crypto) สะท้อนความกล้าเสี่ยง (Risk-on) ของตลาด'},
        {'id': 'wti', 'name': 'WTI Crude Oil (น้ำมันดิบ)', 'symbol': 'CL=F', 'unit': 'USD/bbl', 'desc': 'ต้นทุนพลังงานและเศรษฐกิจโลก'}
    ]

    data = []
    charts = {}

    # วนดึงข้อมูลราคาปัจจุบันและ % เปลี่ยนแปลงของแต่ละตัวชี้วัด
    for item in macro_items:
        try:
            t = yf.Ticker(item['symbol'])
            hist = t.history(period='1y')
            if not hist.empty:
                current_price = hist['Close'].iloc[-1]
                prev_price = hist['Close'].iloc[-2]
                pct_change = ((current_price - prev_price) / prev_price) * 100

                # Chart data
                dates = [d.strftime('%Y-%m-%d') for d in hist.index]
                prices = [float(p) for p in hist['Close'].tolist()]

                charts[item['id']] = {
                    'labels': dates,
                    'values': prices
                }

                data.append({
                    'id': item['id'],
                    'name': item['name'],
                    'price': current_price,
                    'change': pct_change,
                    'is_up': pct_change >= 0,
                    'unit': item['unit'],
                    'desc': item['desc'],
                })
        except Exception:
            continue

    # ====== AI Macro Analysis - วิเคราะห์ภาพรวมเศรษฐกิจด้วย Gemini ======
    _MACRO_CACHE_KEY = 'MACRO_ECONOMY_V2'
    analysis_text = None
    analysis_last_updated = None

    # Load cached analysis on every page visit
    _cached = AnalysisCache.objects.filter(user=request.user, symbol=_MACRO_CACHE_KEY).first()
    if _cached:
        analysis_text = _cached.analysis_data
        analysis_last_updated = _cached.last_updated

    if request.GET.get('analyze') == 'true' and data:
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        model_name_to_use = 'gemini-2.5-flash'

        # สร้าง string สรุปข้อมูลมหภาคเพื่อส่งให้ AI
        data_str = "\n".join([f"{d['name']}: {d['price']:.2f} ({d['change']:+.2f}%) - {d['desc']}" for d in data])
        prompt = f"""
        คุณคือผู้เชี่ยวชาญด้านเศรษฐศาสตร์มหาภาค (Senior Economist) และนักกลยุทธ์การลงทุนระดับโลก (Global Investment Strategist) 
        จงวิเคราะห์ข้อมูลตลาดปัจจุบันด้านล่างนี้ และสรุปภาพรวมเศรษฐกิจและการลงทุนให้มีความลุ่มลึกระดับสถาบัน:
        
        [Market Data Summary]:
        {data_str}

        โปรดเขียนรายงาน "Global Investment Outlook & Asset Allocation" เป็นภาษาไทย โดยมีหัวข้อดังนี้:
        1. **Fund Flow & Risk Appetite**: วิเคราะห์ความสัมพันธ์ของ DXY, Bond Yield เเละ Bitcoin ตอนนี้ตลาดอยู่ในภาวะ Risk-on (กล้าเสี่ยง) หรือ Risk-off (กลัว) เงินกำลังไหลออกจากตลาดหุ้นไปสู่หลุมหลบภัย (Gold) หรือไหลเข้าสู่ระบบใหม่ (Crypto)?
        2. **US vs Thai Market Direction**: วิเคราะห์เปรียบเทียบตลาดหุ้นสหรัฐฯ (S&P 500) และไทย (SET) ทิศทางเป็นอย่างไร? มีปัจจัยอะไรที่สวนทางกันหรือไม่?
        3. **Asset Allocation (การจัดพอร์ตแนะนำ)**: แนะนำสัดส่วนการลงทุนที่เหมาะสมใน 'ตอนนี้' (เช่น หุ้นกี่ %, ทองคำกี่ %, คริปโตกี่ %, เงินสดกี่ %) โดยอ้างอิงจากความเสี่ยงเศรษฐกิจมหาภาค
        4. **Deep Dive: Gold & Crypto**: เจาะลึกทองคำและบิทคอยน์ในฐานะสินทรัพย์ทางเลือก (Alternative Assets) ในสภาวะปัจจุบันควรสะสม, ถือเฉยๆ หรือหาจังหวะขาย?
        5. **Sector Strategy (US & Thai)**: เจาะกลุ่มอุตสาหกรรมโดดเด่น:
           - **US Sectors**: แนะนำกลุ่มที่น่าสนใจในตลาดสหรัฐฯ (เช่น Tech, Healthcare, Energy) พร้อมตัวอย่างบริษัทยักษ์ใหญ่
           - **Thai Sectors**: แนะนำกลุ่ม Winner ในไทย (เช่น Tourism, Export, Banking) พร้อมระบุรายชื่อหุ้น 3-5 ตัว
        6. **Strategic Summary (1-3 Months Outlook)**: สรุปกลยุทธ์สั้นๆ 1-3 เดือนข้างหน้าว่าควรโฟกัสที่จุดใดมากที่สุด

        ข้อกำหนดการตอบ:
        - ใช้โทนเสียงระดับมืออาชีพ น่าเชื่อถือ (Institutional Tone)
        - ใช้ Markdown จัดรูปแบบให้สวยงาม (ใช้ตารางเเนะนำ Asset Allocation จะดีมาก)
        - ห้ามมีคำเกริ่นนำหรือคำส่งท้าย
        - ส่งออกมาเป็น Raw Markdown เท่านั้น
        """

        try:
            response = client.models.generate_content(
                model=model_name_to_use,
                contents=prompt
            )
            analysis_text = response.text or ""

            # ลบ markdown block wrapper ถ้า AI ไม่ปฏิบัติตาม prompt
            # Strip any residual markdown blocks if AI disobeys
            if analysis_text.startswith("```markdown"):
                analysis_text = analysis_text[len("```markdown"):].strip()
            if analysis_text.endswith("```"):
                analysis_text = analysis_text[:-3].strip()
            if not analysis_text:
                analysis_text = "ไม่สามารถรับผลจาก AI ได้ในขณะนี้ กรุณาลองใหม่อีกครั้ง"

            # Save to cache so user can read it on next visit
            obj, _ = AnalysisCache.objects.update_or_create(
                user=request.user,
                symbol=_MACRO_CACHE_KEY,
                defaults={'analysis_data': analysis_text}
            )
            analysis_last_updated = obj.last_updated

        except Exception as e:
            analysis_text = f"ไม่สามารถสร้างบทวิเคราะห์ได้ในขณะนี้: {str(e)}"

    context = {
        'title': 'Macro Economy & Commodities',
        'data': data,
        'analysis': analysis_text,
        'analysis_last_updated': analysis_last_updated,
        # ส่ง charts data เป็น JSON string สำหรับ JavaScript
        'charts_json': json.dumps(charts)
    }
    return render(request, 'stocks/macro.html', context)

# ====== Momentum Scanner - สแกนหาหุ้น SET100+MAI ตามเกณฑ์ Trend Template ======

@login_required
def momentum_scanner(request):
    """
    Globally scans SET100 roughly matching Mark Minervini Trend Template.
    Runs in a background thread to avoid 504 timeout.
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'momentum_scan_{user_id}'

    # ── AJAX status poll ──────────────────────────────────────────────
    if request.GET.get('scan_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        if st.get('state') == 'done':
            _cp.delete(cache_key)
        return _JR(st)

    # ── Trigger background scan ───────────────────────────────────────
    if request.GET.get('scan') == 'true' or request.method == 'POST':
        from .utils import refresh_all_thai_symbols, get_top_ranked_symbols
        # ใช้ Top 300 หุ้นใหญ่เท่านั้นเพื่อความเร็วและคุณภาพ
        scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)
        
        if not scan_symbols:
            refresh_all_thai_symbols()
            scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)

        already = _cp.get(cache_key, {})
        if already.get('state') != 'running':
            total_syms = len(scan_symbols)
            _cp.set(cache_key, {'state': 'running', 'progress': 0, 'total': total_syms, 'phase': 'เริ่มสแกน…'}, timeout=900)

            def _run_momentum_bg(uid, ckey, sym_list):
                try:
                    import pandas as _pd
                    import pandas_ta as _ta
                    import yfinance as _yf
                    import numpy as _np
                    from django.core.cache import cache as _c
                    from .models import MomentumCandidate as _MC
                    from .utils import analyze_momentum_technical, find_supply_demand_zones, get_top_ranked_symbols as _GTRS
                    from django.contrib.auth import get_user_model
                    User = get_user_model()
                    user = User.objects.get(pk=uid)
                    
                    sym_list = _GTRS(market='SET', limit=200, auto_refresh=True)
                    _MC.objects.filter(user=user).delete()
                    
                    # --- STAGE 1: Fast Screening (The Radar) ---
                    # Scan all 800+ symbols for basic liquidity and trend
                    total_syms = len(sym_list)
                    _c.set(ckey, {'state': 'running', 'progress': 5, 'total': total_syms, 'phase': f'Stage 1: สแกนด่วน {total_syms} ตัว...'}, timeout=900)
                    
                    # Align dates with Precision scanner for better data consistency
                    import pytz as _pytz
                    from datetime import datetime as _dt, timedelta as _td
                    _bkk_tz = _pytz.timezone('Asia/Bangkok')
                    _now_bkk = _dt.now(_bkk_tz)
                    scan_end_date  = _now_bkk.date() + _td(days=1)
                    scan_end_str   = scan_end_date.strftime('%Y-%m-%d')
                    scan_start_str = (_now_bkk.date() - _td(days=600)).strftime('%Y-%m-%d')
                    
                    # --- STAGE 1: Fast Screening (Turbo Chunks) ---
                    from yahooquery import Ticker as _TQ
                    _c.set(ckey, {'state': 'running', 'progress': 5, 'total': total_syms, 'phase': 'Stage 1: 🔎 Fast Screening...'}, timeout=900)
                    
                    candidates = []
                    chunk_size = 100
                    for i in range(0, total_syms, chunk_size):
                        chunk = sym_list[i : i + chunk_size]
                        chunk_bk = [f"{s}.BK" for s in chunk]
                        _c.set(ckey, {'state': 'running', 'progress': 5 + int((i/total_syms)*15), 'total': total_syms, 'phase': f'Phase 1: กรองราคาด่วน {i}/{total_syms}...'}, timeout=600)
                        try:
                            tq = _TQ(chunk_bk, timeout=60)
                            prices = tq.price
                            for symbol in chunk:
                                try:
                                    s_bk = f"{symbol}.BK"
                                    if not isinstance(prices, dict) or s_bk not in prices:
                                        candidates.append({'symbol': symbol}); continue
                                    
                                    p_data = prices.get(s_bk)
                                    if not isinstance(p_data, dict) or 'regularMarketPrice' not in p_data:
                                        candidates.append({'symbol': symbol}); continue
                                    
                                    curr_p = p_data.get('regularMarketPrice')
                                    avg_vol = p_data.get('averageDailyVolume3Month', 0)
                                    # Very loose liquidity filter to ensure we get results
                                    if (curr_p * avg_vol) < 150000: continue
                                    candidates.append({'symbol': symbol})
                                except Exception: candidates.append({'symbol': symbol})
                        except Exception:
                            for sym in chunk: candidates.append({'symbol': sym})

                    if len(candidates) < 20: # Emergency fallback
                        candidates = [{'symbol': s} for s in sym_list[:150]]
                        
                    # --- STAGE 2: Deep Technical Analysis (Multi-threaded Turbo) ---
                    total_cand = len(candidates)
                    pre_results = []
                    _c.set(ckey, {'state': 'running', 'progress': 20, 'total': total_cand, 'phase': f'Stage 2: Technical Scan ({total_cand})...'}, timeout=900)
                    
                    import concurrent.futures as _cf
                    def _analyze_one(symbol):
                        try:
                            s_bk = f"{symbol}.BK"
                            df = _yf.download(s_bk, period="1y", interval="1d", progress=False, timeout=20)
                            if df is None or df.empty or len(df) < 55: return None
                            if isinstance(df.columns, pd.MultiIndex): df.columns = df.columns.droplevel(1)
                            
                            df['EMA50'] = _ta.ema(df['Close'], length=50)
                            df['EMA200'] = _ta.ema(df['Close'], length=200)
                            df['RSI'] = _ta.rsi(df['Close'], length=14)
                            adx = _ta.adx(df['High'], df['Low'], df['Close'], length=14)
                            if adx is not None: df = pd.concat([df, adx], axis=1)
                            df['MFI'] = _ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                            
                            tech = analyze_momentum_technical_v2(df)
                            if tech.get('rsi', 0) < 35: return None 
                            
                            return {
                                'symbol': symbol, 'df': df, 'tech': tech,
                                'price': float(df['Close'].iloc[-1]),
                                'year_high': float(df['High'].tail(252).max()),
                                'sd_zone': find_supply_demand_zones(df)
                            }
                        except Exception: return None

                    # Run Stage 2 in parallel for speed
                    with _cf.ThreadPoolExecutor(max_workers=15) as executor:
                        futs = {executor.submit(_analyze_one, c['symbol']): c['symbol'] for c in candidates[:150]}
                        done = 0
                        for fut in _cf.as_completed(futs):
                            done += 1
                            if done % 10 == 0:
                                _c.set(ckey, {'state': 'running', 'progress': 20 + int((done/150)*65), 'phase': f'Analyzing {done}/150...'}, timeout=600)
                            
                            res = None
                            try:
                                res = fut.result(timeout=25)
                            except Exception: pass
                            
                            # FALLBACK: If Threaded yfinance fails, try ONE last Sync check with YahooQuery for this symbol
                            if not res:
                                symbol = futs[fut]
                                try:
                                    s_bk = f"{symbol}.BK"
                                    tq_single = _TQ(s_bk)
                                    h = tq_single.history(period='1y', interval='1d')
                                    if h is not None and not h.empty:
                                        if isinstance(h.index, pd.MultiIndex): h = h.xs(s_bk, level=0)
                                        # Minimal data requirement reduced to 40
                                        if len(h) >= 40:
                                            h.rename(columns={'adjclose': 'Close', 'high': 'High', 'low': 'Low', 'open': 'Open', 'volume': 'Volume'}, inplace=True)
                                            h['EMA50'] = _ta.ema(h['Close'], length=50)
                                            h['EMA200'] = _ta.ema(h['Close'], length=min(200, len(h)-1))
                                            h['RSI'] = _ta.rsi(h['Close'], length=14)
                                            tech = analyze_momentum_technical_v2(h)
                                            if tech.get('rsi', 0) > 30: 
                                                 res = {
                                                    'symbol': symbol, 'df': h, 'tech': tech,
                                                    'price': float(h['Close'].iloc[-1]),
                                                    'year_high': float(h['High'].tail(252).max()),
                                                    'sd_zone': find_supply_demand_zones(h)
                                                }
                                except Exception: pass
                                
                            if res: pre_results.append(res)

                    # --- STAGE 3: Bulk Fundamental ---
                    fund_data = {}
                    if pre_results:
                        _c.set(ckey, {'state': 'running', 'progress': 85, 'phase': 'Stage 3: Fundamental...'}, timeout=900)
                        from .utils import YQTicker
                        match_bk = [f"{r['symbol']}.BK" for r in pre_results]
                        try:
                            yq_all = YQTicker(match_bk)
                            modules = yq_all.get_modules('financialData summaryProfile')
                            for s_bk, val in modules.items():
                                if not isinstance(val, dict): continue
                                sym_clean = s_bk.replace('.BK','')
                                prof = val.get('summaryProfile', {})
                                fin  = val.get('financialData', {})
                                fund_data[sym_clean] = {
                                    'sector': prof.get('sector', 'Other'),
                                    'eps_growth': float(fin.get('earningsQuarterlyGrowth', 0) or 0) * 100,
                                    'rev_growth': float(fin.get('revenueGrowth', 0) or 0) * 100
                                }
                        except Exception: pass

                    # FINAL: Save to DB
                    _c.set(ckey, {'state': 'running', 'progress': 95, 'phase': 'Saving results...'}, timeout=600)
                    bulk_objs = []
                    for r in pre_results:
                        sym = r['symbol']
                        sd  = r['sd_zone']
                        tech = r['tech']
                        df = r['df']
                        f   = fund_data.get(sym, {'sector': 'N/A', 'eps_growth': 0.0, 'rev_growth': 0.0})
                        
                        dz_start = dz_end = sz_start = sz_end = sl_price = rr_val = None
                        entry_strat = ''
                        if sd:
                            entry_strat = sd['type']; dz_start = sd['start']; dz_end = sd['end']
                            sz_start = sd['target']; sz_end = sd['target'] * 1.02
                            sl_price = sd['stop_loss']; rr_val = sd['rr_ratio']

                        bulk_objs.append(_MC(
                            user=user, symbol=sym, symbol_bk=f"{sym}.BK", market='SET', price=r['price'],
                            rsi=tech.get('rsi', 0), 
                            adx=float(df['ADX_14'].iloc[-1]) if 'ADX_14' in df.columns else 0,
                            mfi=float(df['MFI'].iloc[-1]) if 'MFI' in df.columns else 0,
                            rvol=tech.get('rvol', 0), 
                            rvol_bullish=tech.get('rvol_bullish', False),
                            technical_score=tech.get('score', 0), 
                            rs_rating=0,
                            entry_strategy=entry_strat, 
                            demand_zone_start=dz_start, 
                            demand_zone_end=dz_end,
                            supply_zone_start=sz_start, 
                            supply_zone_end=sz_end, 
                            stop_loss=sl_price, 
                            risk_reward_ratio=rr_val,
                            year_high=r['year_high'], 
                            upside_to_high=((r['year_high'] - r['price'])/r['price'])*100 if r['price'] > 0 else 0,
                            sector=f['sector'], 
                            eps_growth=f['eps_growth'], 
                            rev_growth=f['rev_growth'],
                            stage2=r['price'] > float(df['EMA200'].iloc[-1]) if 'EMA200' in df.columns else False
                        ))
                    if bulk_objs:
                        _MC.objects.bulk_create(bulk_objs)
                except Exception as e:
                    import logging; logging.getLogger('stocks').error(f"Momentum Scan Error: {e}")
                finally:
                    _c.set(ckey, {'state': 'done', 'progress': 100, 'total': total_syms, 'phase': 'เสร็จสิ้น'}, timeout=120)

            # Start Worker
            _th.Thread(target=_run_momentum_bg, args=(user_id, cache_key, scan_symbols), daemon=True).start()

    # ====== Handle AI Summary Analysis (Optional) ======
    ai_analysis = ""
    if request.GET.get('analyze') == 'true':
        try:
            data_to_analyze = []
            top_best = MomentumCandidate.objects.filter(user=request.user).order_by('-technical_score')[:15]
            for t in top_best:
                data_to_analyze.append({
                    'symbol': t.symbol, 'score': t.technical_score, 'rvol': t.rvol, 'rsi': t.rsi,
                    'eps_growth': t.eps_growth, 'rev_growth': t.rev_growth, 'upside': t.upside_to_high
                })
            
            if data_to_analyze:
                prompt = f"หุ้นไทย Momentum แรงที่สุด 15 ตัวจากระบบสแกน: {json.dumps(data_to_analyze)}\nช่วยวิเคราะห์และคัดเลือก 3-5 ตัวที่น่าสนใจที่สุด พร้อมเหตุผลเชิงกลยุทธ์ตามสไตล์ Mark Minervini และบอกจุดระวัง"
                ai_analysis = analyze_with_ai(prompt)
        except Exception as e:
            ai_analysis = f"AI Analysis Error: {str(e)}"

    # ป้องกัน FieldError จาก ?sort= ค่าที่ไม่มีในฐานข้อมูล
    _MOMENTUM_SORT_MAP = {
        'score':          '-technical_score',
        'technical_score':'-technical_score',
        'price':          '-price',
        'rsi':            '-rsi',
        'rvol':           '-rvol',
        'rs':             '-rs_rating',
        'rs_rating':      '-rs_rating',
        'symbol':         'symbol',
        'adx':            '-adx',
        'upside':         '-upside_to_high',
    }
    raw_sort   = request.GET.get('sort', '-technical_score')
    sort_by    = _MOMENTUM_SORT_MAP.get(raw_sort, raw_sort if raw_sort.lstrip('-') in {
        'technical_score','price','rsi','rvol','rs_rating','symbol','adx','upside_to_high','mfi'
    } else '-technical_score')
    candidates = MomentumCandidate.objects.filter(user=request.user).order_by(sort_by)
    scanned_at = candidates.first().scanned_at if candidates.exists() else None

    # ตรวจว่ากำลังสแกนอยู่ - ถ้าใช่ ซ่อน results เพื่อไม่ให้กระพริบ
    _scan_state = _cp.get(cache_key, {})
    is_scanning = _scan_state.get('state') == 'running'

    candidate_list = list(candidates) if not is_scanning else []

    # ====== Live Price + Fresh Zone - recompute zone จาก historical data ใหม่ทุกครั้ง ======
    if candidate_list:
        try:
            import concurrent.futures as _mcf
            # คำนวณ end date เหมือน entry_finder - ห้ามรวม today's incomplete bar ตอนตลาดเปิด
            import pytz as _mpytz
            from datetime import datetime as _mdt, timedelta as _mtd, time as _mtime
            _mnow   = _mdt.now(_mpytz.timezone('Asia/Bangkok'))
            _mt     = _mnow.time()
            _market_open_now = (
                _mnow.weekday() < 5 and
                (
                    _mt < _mtime(10, 0) or
                    (_mtime(10, 0) <= _mt <= _mtime(12, 30)) or
                    (_mtime(12, 30) < _mt < _mtime(14, 30)) or
                    (_mtime(14, 30) <= _mt <= _mtime(16, 30))
                )
            )
            _mend_date  = (_mnow.date() - _mtd(days=1)) if _market_open_now else _mnow.date()
            _mend_str   = _mend_date.strftime('%Y-%m-%d')
            _mstart_str = (_mend_date - _mtd(days=600)).strftime('%Y-%m-%d')

            def _mom_live(sym):
                try:
                    full_sym = f"{sym}.BK"
                    fi = yf.Ticker(full_sym).fast_info
                    p = getattr(fi, 'last_price', None)
                    live_price = float(p) if p else None

                    # Recompute zone - ใช้ Ticker().history() (thread-safe), end date เหมือน entry_finder
                    _t = yf.Ticker(full_sym)
                    df = _t.history(start=_mstart_str, end=_mend_str, interval='1d')
                    fresh_zone = None
                    if df is not None and len(df) >= 50:
                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = df.columns.get_level_values(0)
                        fresh_zone = find_supply_demand_zones_v2(df)
                    return sym, live_price, fresh_zone
                except Exception:
                    return sym, None, None

            live_map = {}
            zone_map = {}   # fresh zones - keyed by symbol
            with _mcf.ThreadPoolExecutor(max_workers=6) as _mex:
                for _s, _p, _z in _mex.map(_mom_live, [c.symbol for c in candidate_list]):
                    if _p: live_map[_s] = _p
                    if _z: zone_map[_s] = _z
        except Exception:
            live_map = {}
            zone_map = {}

        for c in candidate_list:
            lp  = live_map.get(c.symbol)
            c.live_price = lp
            ref = lp if lp else float(c.price or 0)

            # ใช้ fresh zone ถ้ามี มิฉะนั้น fallback ไป DB zone
            fz = zone_map.get(c.symbol)
            if fz:
                dz_s = float(fz.get('start') or 0)
                dz_e = float(fz.get('end')   or 0)
                sz_s = float(fz.get('target') or 0)
                # อัปเดต zone ใน candidate object ด้วย เพื่อให้ template แสดงถูก
                c.demand_zone_start = dz_s or c.demand_zone_start
                c.demand_zone_end   = dz_e or c.demand_zone_end
                c.supply_zone_start = sz_s or c.supply_zone_start
            else:
                dz_s = float(c.demand_zone_start or 0)
                dz_e = float(c.demand_zone_end   or 0)
                sz_s = float(c.supply_zone_start or 0)

            # zone status ตาม live price + fresh zone
            c.live_in_zone    = dz_s > 0 and dz_e > 0 and dz_e <= ref <= dz_s
            c.live_broke_zone = dz_e > 0 and ref < dz_e
            c.live_above_tp   = sz_s > 0 and ref >= sz_s
            c.live_near_tp    = (not c.live_above_tp) and sz_s > 0 and dz_s > 0 and (sz_s - ref) / (sz_s - dz_s) * 100 <= 15 if (sz_s - dz_s) > 0 else False
            c.live_zone_prox  = 0.0 if ref <= dz_s else round((ref - dz_s) / dz_s * 100, 1) if dz_s > 0 else 999
            if lp and float(c.price or 0) > 0:
                c.live_change_pct = round((lp - float(c.price)) / float(c.price) * 100, 2)
            else:
                c.live_change_pct = None

    context = {
        'title': 'Global Momentum Scanner (CAN SLIM)',
        'candidates': candidate_list,
        'ai_analysis': ai_analysis,
        'scanned_at': scanned_at,
        'current_sort': sort_by,
        'is_scanning': is_scanning,
        'has_scanned': bool(candidate_list) or (candidates.exists() and not is_scanning),
    }
    return render(request, 'stocks/momentum.html', context)


# ====== Market Condition Analyzer - วิเคราะห์สภาวะตลาด SET Index ======

def _get_market_condition(set_df):
    """
    วิเคราะห์ SET Index เพื่อกำหนด Market Phase ปัจจุบัน
    คืน dict: phase, label, color, score, indicators
    """
    import pandas as pd
    if set_df is None or set_df.empty:
        return {'phase': 'UNKNOWN', 'label': 'ไม่มีข้อมูล SET', 'color': 'secondary', 'score': 0}

    close = set_df['Close'].dropna()
    if len(close) < 50:
        return {'phase': 'UNKNOWN', 'label': 'ข้อมูลไม่พอ', 'color': 'secondary', 'score': 0}

    curr = float(close.iloc[-1])
    ema50  = float(close.ewm(span=50, adjust=False).mean().iloc[-1])
    ema200 = float(close.ewm(span=200, adjust=False).mean().iloc[-1]) if len(close) >= 200 else None
    sma150 = float(close.rolling(150).mean().iloc[-1]) if len(close) >= 150 else None

    # SMA150 slope เทียบ 20 วันที่แล้ว
    sma150_slope = None
    if sma150 and len(close) >= 170:
        sma150_20d = float(close.rolling(150).mean().iloc[-21])
        sma150_slope = ((sma150 - sma150_20d) / sma150_20d * 100) if sma150_20d else None

    # 1m / 3m return ของ SET
    m1 = float((close.iloc[-1] - close.iloc[-22]) / close.iloc[-22] * 100) if len(close) >= 22 else 0.0
    m3 = float((close.iloc[-1] - close.iloc[-66]) / close.iloc[-66] * 100) if len(close) >= 66 else 0.0

    score = 0
    bullets = []  # เงื่อนไขที่ผ่าน

    # EMA200
    if ema200:
        if curr > ema200:
            score += 3
            bullets.append(f"SET ({curr:.0f}) > EMA200 ({ema200:.0f}) ✅")
        else:
            score -= 3
            bullets.append(f"SET ({curr:.0f}) < EMA200 ({ema200:.0f}) ❌")

    # SMA150
    if sma150:
        if curr > sma150:
            score += 2
            bullets.append(f"SET > SMA150 ({sma150:.0f}) ✅")
        else:
            score -= 1
            bullets.append(f"SET < SMA150 ({sma150:.0f}) ❌")

    # SMA150 slope
    if sma150_slope is not None:
        if sma150_slope > 0.2:
            score += 2
            bullets.append(f"SMA150 ขาขึ้น (+{sma150_slope:.2f}%) ✅")
        elif sma150_slope < -0.3:
            score -= 2
            bullets.append(f"SMA150 ขาลง ({sma150_slope:.2f}%) ❌")
        else:
            bullets.append(f"SMA150 sideways ({sma150_slope:.2f}%) ⚠️")

    # EMA50
    if curr > ema50:
        score += 1

    # Momentum 1m
    if m1 > 3:
        score += 2
        bullets.append(f"SET +{m1:.1f}% (1 เดือน) ✅")
    elif m1 > 0:
        score += 1
    elif m1 < -5:
        score -= 2
        bullets.append(f"SET {m1:.1f}% (1 เดือน) ❌")
    else:
        bullets.append(f"SET {m1:.1f}% (1 เดือน) ⚠️")

    # กำหนด Phase
    if score >= 7:
        phase, color, label = 'UPTREND',    'success', 'ตลาดขาขึ้น - เหมาะสำหรับ Swing Buy'
    elif score >= 4:
        phase, color, label = 'RECOVERY',   'info',    'ตลาดฟื้นตัว - คัดเฉพาะหุ้นแข็งแกร่ง'
    elif score >= 1:
        phase, color, label = 'MIXED',      'warning', 'ตลาดผสม - เน้น Watchlist และ Risk Management'
    elif score >= -2:
        phase, color, label = 'CORRECTION', 'warning', 'ตลาดพักฐาน - ระวังสูง ลดขนาด Position'
    else:
        phase, color, label = 'DOWNTREND',  'danger',  'ตลาดขาลง - หลีกเลี่ยงการซื้อใหม่'

    return {
        'phase':  phase,
        'color':  color,
        'label':  label,
        'score':  score,
        'curr':   round(curr, 2),
        'ema200': round(ema200, 2) if ema200 else None,
        'sma150': round(sma150, 2) if sma150 else None,
        'sma150_slope': round(sma150_slope, 2) if sma150_slope is not None else None,
        'm1':     round(m1, 2),
        'm3':     round(m3, 2),
        'above_ema200':   bool(ema200 and curr > ema200),
        'above_sma150':   bool(sma150 and curr > sma150),
        'sma150_rising':  bool(sma150_slope and sma150_slope > 0.2),
        'bullets': bullets,
    }


# ====== Scan Watchlist Views ======

@login_required
def watchlist_item_toggle(request):
    """AJAX POST - เพิ่ม/ลบหุ้นออกจาก ScanWatchlistItem และส่งไปที่ Market Watchlist (สำหรับรับ Alert เข้า Telegram)"""
    import json
    from django.http import JsonResponse
    from .models import ScanWatchlistItem, Watchlist
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        data = json.loads(request.body)
    except (ValueError, KeyError):
        return JsonResponse({'error': 'invalid JSON'}, status=400)
    symbol   = data.get('symbol', '').strip().upper()
    sector   = data.get('sector', 'Unknown')
    market   = data.get('market', 'SET')
    strategy = data.get('strategy', 'PRECISION')
    note     = data.get('note', '')

    if not symbol:
        return JsonResponse({'error': 'symbol required'}, status=400)
        
    obj, created = ScanWatchlistItem.objects.get_or_create(
        user=request.user, symbol=symbol, market=market,
        defaults={'sector': sector, 'strategy': strategy, 'note': note}
    )
    
    if not created:
        # หากมีอยู่แล้ว ให้ลบออก (Un-toggle)
        obj.delete()
        if market == 'SET':
            Watchlist.objects.filter(user=request.user, symbol=symbol).delete()
        return JsonResponse({'status': 'removed', 'symbol': symbol})
        
    # หากเพิ่มใหม่ ให้อัปเดตค่าหากมีการส่งมา (กรณี get_or_create ใช้ defaults แค่ตอนสร้าง)
    obj.strategy = strategy
    obj.note     = note
    obj.save()

    # สำหรับ SET สั่งให้เพิ่มเข้าไปที่ฝั่ง Market Watchlist ด้วย
    if market == 'SET':
        # เราเก็บข้อมูล Pattern/Strategy ลงในฟิลด์ strategy ของ Portfolio ได้ แต่ Watchlist ปกติไม่มี
        # ดังนั้นจะเน้นเก็บใน ScanWatchlistItem เป็นหลัก
        Watchlist.objects.get_or_create(user=request.user, symbol=symbol)
    
    return JsonResponse({'status': 'added', 'symbol': symbol})


@login_required
def scan_watchlist_view(request):
    """แสดง Scan Watchlist พร้อม score ปัจจุบัน / รอบก่อน / delta / alert"""
    from .models import ScanWatchlistItem, PrecisionScanCandidate
    
    market = request.GET.get('market', 'SET')
    items = ScanWatchlistItem.objects.filter(user=request.user, market=market)

    runs = list(
        PrecisionScanCandidate.objects
        .filter(user=request.user, market=market)
        .values_list('scan_run', flat=True)
        .order_by('-scan_run')
        .distinct()[:2]
    )
    latest_run = runs[0] if len(runs) >= 1 else None
    prev_run   = runs[1] if len(runs) >= 2 else None

    latest_map = {c.symbol: c for c in PrecisionScanCandidate.objects.filter(user=request.user, market=market, scan_run=latest_run)} if latest_run else {}
    prev_map   = {c.symbol: c for c in PrecisionScanCandidate.objects.filter(user=request.user, market=market, scan_run=prev_run)}   if prev_run   else {}

    enriched = []
    for item in items:
        latest = latest_map.get(item.symbol)
        prev   = prev_map.get(item.symbol)
        cur_score  = latest.technical_score if latest else None
        prev_score = prev.technical_score   if prev   else None
        delta = (cur_score - prev_score) if (cur_score is not None and prev_score is not None) else None
        enriched.append({
            'watchlist':   item,
            'scan_data':   latest,
            'delta':       delta,
            'triggered':   cur_score is not None and cur_score >= item.alert_threshold,
        })

    return render(request, 'stocks/scan_watchlist.html', {
        'items':       enriched,
        'latest_run':  latest_run,
        'market':      market,
    })


# ====== Precision Momentum Scanner - เวอร์ชันกรองคุณภาพสูง ======

@login_required
def vcp_manual(request):
    return render(request, 'stocks/vcp_manual.html')

@login_required
def minervini_sepa_scanner(request):
    """
    Minervini SEPA Scanner - ระบบสแกนเจาะจงเฉพาะตามตำรา Mark Minervini
    กรองเฉพาะหุ้นที่อยู่ใน Stage 2 และมีฟอร์ม VCP/VDU
    """
    # ดึงรายชื่อรอบการสแกน (ใช้จาก PrecisionScanCandidate)
    all_runs = list(
        PrecisionScanCandidate.objects
        .filter(user=request.user, market='SET')
        .values_list('scan_run', flat=True)
        .order_by('-scan_run')
        .distinct()
    )
    
    candidates = []
    last_updated = None
    selected_run_idx = int(request.GET.get('run_idx', 0))

    if all_runs:
        if selected_run_idx < len(all_runs):
            run_time = all_runs[selected_run_idx]
            # กรองเฉพาะ Stage 2 + RS Rating >= 70 (ตาม SEPA criteria)
            candidates = list(PrecisionScanCandidate.objects.filter(
                user=request.user,
                scan_run=run_time,
                stage2=True,
                rs_rating__gte=70,
            ).order_by('-vcp_setup', '-rs_rating'))
            last_updated = run_time

    # เพิ่มเติม: กรองเฉพาะตัวที่มีสัญญาณ VCP เพื่อความ Clean
    vcp_only = request.GET.get('vcp_only') == '1'
    if vcp_only:
        candidates = [c for c in candidates if c.vcp_setup]

    # hide_at_tp: ซ่อนหุ้นที่ราคาอยู่ใกล้/ถึง Target (upside_to_high < 5%)
    hide_at_tp = request.GET.get('hide_at_tp', '1') == '1'
    if hide_at_tp:
        candidates = [c for c in candidates if c.upside_to_high >= 5.0]

    # คำนวณ field เพิ่มเติมสำหรับแสดงผล + SEPA Score
    for c in candidates:
        # % ห่างจาก Pivot (52w High)
        c.dist_from_pivot = round(c.upside_to_high, 1)
        # สถานะ TP
        if c.upside_to_high < 5:
            c.tp_status = 'at_tp'
        elif c.upside_to_high < 10:
            c.tp_status = 'near_tp'
        else:
            c.tp_status = None
        # SEPA Score
        sc = 0
        if c.vcp_setup:
            sc += 30
            sc += int(max(0, (10 - min(c.vcp_tightness, 10)) * 2))
            sc += min(c.vcp_contractions, 5) * 3
        if c.vcp_vdu or c.vdu_near_zone:
            sc += 20
        if c.pocket_pivot:
            sc += 10
        sc += int(c.rs_rating * 0.7)
        if c.adx >= 25:
            sc += 10
        elif c.adx >= 15:
            sc += 5
        if c.vcp_setup:
            if c.dist_from_pivot <= 5:
                sc += 10
            elif c.dist_from_pivot <= 10:
                sc += 5
            elif c.dist_from_pivot > 15:
                sc -= 5
        c.sepa_score = sc

    # Sort by SEPA Score descending
    candidates.sort(key=lambda c: c.sepa_score, reverse=True)

    # Assign rank
    for i, c in enumerate(candidates, 1):
        c.sepa_rank = i

    context = {
        'candidates': candidates,
        'last_updated': last_updated,
        'all_runs': all_runs,
        'selected_run_idx': selected_run_idx,
        'vcp_only': vcp_only,
        'hide_at_tp': hide_at_tp,
    }
    return render(request, 'stocks/sepa_scanner.html', context)

@login_required
def precision_momentum_scanner(request):
    """
    Precision Momentum Scanner - กรองคุณภาพสูงกว่า momentum_scanner
    ปรับปรุงจาก momentum_scanner:
    1. ERC ต้องมี Body + Volume > 1.5x avg (ทั้งสองเงื่อนไข)
    2. ADX >= 20 (กรองเทรนด์แข็งแกร่งเท่านั้น)
    3. Liquidity filter: avg 20d volume >= 500,000 หุ้น
    4. Supply target = 52-week high เสมอ
    5. ATR-based stop loss
    6. Direction-aware RVOL scoring
    7. เก็บประวัติ scan 3 รอบล่าสุด
    8. is_new_entry flag (หุ้นใหม่ vs ยังอยู่จากรอบก่อน)
    """
    from .models import PrecisionScanCandidate
    from .utils import analyze_momentum_technical_v2, get_top_ranked_symbols
    from django.utils import timezone as tz
    from yahooquery import Ticker as YQTicker

    scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)
    if not scan_symbols:
        refresh_all_thai_symbols()
        scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)

    # ====== AJAX Status Poll ======
    if request.GET.get('scan_status') == '1':
        from django.core.cache import cache as _cp
        from django.http import JsonResponse as _JR
        _key = f'precision_scan_{request.user.id}'
        _st = _cp.get(_key, {'state': 'idle'})
        if _st.get('state') == 'done':
            _cp.delete(_key)
        return _JR(_st)

    if request.method == "POST" and request.POST.get('action') == 'scan':
        from django.core.cache import cache as _cache_bg
        import threading

        user_id   = request.user.id
        cache_key = f'precision_scan_{user_id}'
        
        # เก็บหน้าที่ต้องกลับไปหลังสแกนเสร็จ
        raw_next = request.POST.get('next_url')
        next_url = 'stocks:minervini_sepa_scanner' if raw_next == 'sepa' else 'stocks:precision_momentum_scanner'

        _cur = _cache_bg.get(cache_key, {})
        if _cur.get('state') == 'running':
            return redirect(next_url)

        _cache_bg.set(cache_key, {'state': 'running', 'progress': 0, 'total': 0, 'phase': 'เตรียมข้อมูล…'}, timeout=900)

        def _run_precision_bg(uid, ckey, sym_list):
            try:
                import django
                django.setup()
                import pandas_ta as ta
                from datetime import datetime as _dt, timedelta as _td, time as _dtime
                import pytz as _pytz
                import concurrent.futures
                from django.contrib.auth import get_user_model
                from django.core.cache import cache as _cache
                from django.utils import timezone as tz
                from yahooquery import Ticker as YQTicker
                from .models import PrecisionScanCandidate
                from .utils import analyze_momentum_technical_v2, get_top_ranked_symbols as _GTRS
                sym_list = _GTRS(market='SET', limit=200, auto_refresh=True)

                User = get_user_model()
                user = User.objects.get(pk=uid)
                scan_run_time = tz.now()

                # ====== Pin Scan Date ======
                _bkk_tz = _pytz.timezone('Asia/Bangkok')
                _now_bkk = _dt.now(_bkk_tz)
                _t = _now_bkk.time()
                _market_open = (
                    _now_bkk.weekday() < 5 and
                    (_dtime(10, 0) <= _t <= _dtime(12, 30) or
                     _dtime(12, 30) < _t < _dtime(16, 30))
                )
                # yfinance download end= is exclusive. To include today's data, use tomorrow.
                scan_end_date  = _now_bkk.date() + _td(days=1)
                scan_end_str   = scan_end_date.strftime('%Y-%m-%d')
                scan_start_str = (_now_bkk.date() - _td(days=600)).strftime('%Y-%m-%d')  # 600 วัน → ~430 trading days, EMA200 warm-up มีพอ
                set_start_str  = (_now_bkk.date() - _td(days=400)).strftime('%Y-%m-%d')

                prev_run = (
                    PrecisionScanCandidate.objects
                    .filter(user=user, market='SET')
                    .values_list('scan_run', flat=True)
                    .order_by('-scan_run')
                    .distinct()
                    .first()
                )
                prev_symbols = set()
                if prev_run:
                    prev_symbols = set(
                        PrecisionScanCandidate.objects
                        .filter(user=user, scan_run=prev_run)
                        .values_list('symbol', flat=True)
                    )

                # ====== SET Index ======
                set_1m_return = 0.0
                set_3m_return = 0.0
                try:
                    set_df = yf.download("^SET", start=set_start_str, end=scan_end_str, interval="1d", progress=False)
                    if set_df is not None and not set_df.empty:
                        if isinstance(set_df.columns, pd.MultiIndex):
                            set_df.columns = set_df.columns.droplevel(1)
                        set_close = set_df['Close'].dropna()
                        if len(set_close) >= 66:
                            set_1m_return = float((set_close.iloc[-1] - set_close.iloc[-22]) / set_close.iloc[-22] * 100)
                            set_3m_return = float((set_close.iloc[-1] - set_close.iloc[-66]) / set_close.iloc[-66] * 100)
                    import logging; logging.getLogger('stocks').info(f"[Precision] SET Index: 1m={set_1m_return:.2f}% 3m={set_3m_return:.2f}%")
                except Exception as e:
                    import logging; logging.getLogger('stocks').warning(f"[Precision] SET Index fetch failed: {e}")


                # ====== Phase 1: Fast Screening with YahooQuery (Institutional Speed) ======
                total_syms = len(sym_list)
                _cache.set(ckey, {'state': 'running', 'progress': 5, 'total': total_syms, 'phase': 'ดึงข้อมูล RS Rating...'}, timeout=900)
                
                from yahooquery import Ticker as _TQ
                rs_returns_all = {}
                
                # Fetch history in chunks of 80 to prevent hangs
                import logging; logger = logging.getLogger('stocks')
                chunk_size = 80
                for i in range(0, len(sym_list), chunk_size):
                    chunk = sym_list[i : i + chunk_size]
                    chunk_bk = [f"{s}.BK" for s in chunk]
                    _cache.set(ckey, {'state': 'running', 'progress': 5 + int((i/total_syms)*15), 'total': total_syms, 'phase': f'Phase 1: โหลดข้อมูลกลุ่ม {i//chunk_size + 1}...'}, timeout=900)
                    try:
                        tq = _TQ(chunk_bk)
                        tq_hist = tq.history(start=scan_start_str, end=scan_end_str, interval="1d")
                        if tq_hist is not None and not tq_hist.empty:
                            if isinstance(tq_hist.index, pd.MultiIndex):
                                for symbol in chunk:
                                    try:
                                        s_bk = f"{symbol}.BK"
                                        if s_bk in tq_hist.index.get_level_values(0):
                                            _close = tq_hist.loc[s_bk]['adjclose'].dropna()
                                            if len(_close) >= 66:
                                                ret = float((_close.iloc[-1] - _close.iloc[-66]) / abs(_close.iloc[-66]) * 100)
                                                rs_returns_all[symbol] = ret
                                    except Exception: continue
                    except Exception as e:
                        logger.error(f"RS Chunk Error at {i}: {e}")

                # FAILSAFE: If results are empty or too small, force evaluation of a subset
                if len(rs_returns_all) < 10:
                    import logging; logging.getLogger('stocks').warning(f"[Precision] Data recovery mode: Only {len(rs_returns_all)} found. Force fallback.")
                    # Use at least top 50 symbols to ensure some results
                    for s in sym_list[:100]:
                        if s not in rs_returns_all: rs_returns_all[s] = 0.0 # Dummy score to pass filter

                rs_ratings_map = {}
                if rs_returns_all:
                    _rs_ser = pd.Series(rs_returns_all)
                    rs_ratings_map = (_rs_ser.rank(pct=True) * 99).clip(0, 99).astype(int).to_dict()

                # Phase 2: เจาะลึกหุ้นที่เข้ารอบ
                results_to_process = [s for s in sym_list if rs_ratings_map.get(s, 0) >= 60]
                if not results_to_process:
                    results_to_process = sym_list[:20] # Safety fallback

                def _process_precision_scan(symbol):
                    try:
                        # ใช้ yf.Ticker().history() แทน yf.download() เพราะ yf.download() 
                        # มีบั๊ก Thread-safety กรองข้อมูลข้าม Symbol กันเมื่อรันใน ThreadPool 
                        ticker_obj = yf.Ticker(f"{symbol}.BK")
                        df = ticker_obj.history(start=scan_start_str, end=scan_end_str, interval="1d")

                        if df is None or df.empty:
                            try:
                                yq = YQTicker(f"{symbol}.BK")
                                df = yq.history(start=scan_start_str, end=scan_end_str, interval="1d")
                                if isinstance(df, pd.DataFrame) and not df.empty:
                                    df = df.reset_index()
                                    if 'date' in df.columns:
                                        df.set_index('date', inplace=True)
                                    if 'symbol' in df.columns:
                                        df.drop(columns=['symbol'], inplace=True)
                                    df.rename(columns={'open': 'Open', 'high': 'High', 'low': 'Low',
                                                       'close': 'Close', 'volume': 'Volume'}, inplace=True)
                            except Exception:
                                pass

                        if df is None or df.empty:
                            return None

                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = df.columns.droplevel(1)

                        df = df.dropna(subset=['Close', 'High'])
                        if len(df) < 200:
                            return None

                        # ====== Liquidity & Quality Filters (Institutional Grade) ======
                        avg_vol_20 = float(df['Volume'].tail(20).mean())
                        avg_close_20 = float(df['Close'].tail(20).mean())
                        avg_turnover_20 = avg_vol_20 * avg_close_20
                
                        import logging as _lg; _scan_log = _lg.getLogger('stocks.scan')
                        current_price = float(df['Close'].iloc[-1])

                        # 1. Turnover >= 10M THB
                        if avg_turnover_20 < 10_000_000:
                            _scan_log.info(f"[SCAN SKIP] {symbol}: Turnover ฿{avg_turnover_20/1e6:.1f}M < 10M")
                            return None

                        # 2. Minimum Price >= 1.00
                        if current_price < 1.00:
                            _scan_log.info(f"[SCAN SKIP] {symbol}: Price ฿{current_price} < 1.00")
                            return None

                        # 3. RS Rating >= 60 (None = ข้อมูลไม่พอ → ไม่ตัดออก)
                        rs_val = rs_ratings_map.get(symbol, None)
                        if rs_val is not None and rs_val < 60:
                            _scan_log.info(f"[SCAN SKIP] {symbol}: RS {rs_val} < 60")
                            return None

                        # ====== คำนวณ Indicators ======
                        df['EMA200'] = ta.ema(df['Close'], length=200)
                        df['EMA50'] = ta.ema(df['Close'], length=50)
                        df['RSI'] = ta.rsi(df['Close'], length=14)
                        adx_df = ta.adx(df['High'], df['Low'], df['Close'], length=14)
                        if adx_df is not None and not adx_df.empty:
                            df = pd.concat([df, adx_df], axis=1)
                        mfi_series = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                        df['MFI'] = mfi_series

                        last_row = df.iloc[-1]
                        current_price = float(last_row['Close'])
                        ema200 = float(df['EMA200'].iloc[-1]) if pd.notna(df['EMA200'].iloc[-1]) else current_price
                        year_high = float(df['High'].tail(252).max())

                        # ====== ADX Filter ======
                        adx_val = float(df['ADX_14'].iloc[-1]) if 'ADX_14' in df.columns and pd.notna(df['ADX_14'].iloc[-1]) else 0
                        if adx_val < 15:
                            _scan_log.info(f"[SCAN SKIP] {symbol}: ADX {adx_val:.1f} < 15")
                            return None

                        # ====== Trend Template Filter ======
                        near_high  = current_price >= year_high * 0.65
                        if not near_high:
                            _scan_log.info(f"[SCAN SKIP] {symbol}: Price ฿{current_price} < 65% of 52wH ฿{year_high} ({current_price/year_high*100:.0f}%)")
                            return None

                        import logging; logger = logging.getLogger('stocks')
                        logger.debug(f"[Precision] MATCH: {symbol} (ADX:{adx_val:.1f})")

                        # ====== Precision Technical Analysis (v3) ======
                        tech = analyze_momentum_technical_v2(df)
                        integrated_score = tech['score']
                        rvol         = tech['rvol']
                        rsi          = tech['rsi']
                        rvol_bullish = tech['rvol_bullish']
                        sd_zone      = tech['sd_zone']
                        ema20_aligned_flag = tech.get('ema20_aligned', False)
                        ema20_slope_val    = tech.get('ema20_slope', 0.0)
                        ema20_rising_flag  = tech.get('ema20_rising', False)
                        hh_hl_flag         = tech.get('hh_hl_structure', False)

                        mfi_val = float(df['MFI'].iloc[-1]) if 'MFI' in df.columns and pd.notna(df['MFI'].iloc[-1]) else 0

                        # ====== MACD (12,26,9) - histogram + bullish crossover detection ======
                        macd_hist_val  = None
                        macd_cross_val = False
                        try:
                            macd_df = ta.macd(df['Close'], fast=12, slow=26, signal=9)
                            if macd_df is not None and not macd_df.empty:
                                hist_col = [c for c in macd_df.columns if 'h' in c.lower() or 'hist' in c.lower()]
                                macd_col = [c for c in macd_df.columns if c.lower().startswith('macd_')]
                                sig_col  = [c for c in macd_df.columns if 'macds' in c.lower() or 'signal' in c.lower()]
                                if hist_col:
                                    macd_hist_val = float(macd_df[hist_col[0]].iloc[-1]) if pd.notna(macd_df[hist_col[0]].iloc[-1]) else None
                                # Bullish crossover = MACD line crosses above signal line in last 3 bars
                                if macd_col and sig_col:
                                    m_ser = macd_df[macd_col[0]].dropna()
                                    s_ser = macd_df[sig_col[0]].dropna()
                                    if len(m_ser) >= 4 and len(s_ser) >= 4:
                                        # Check if MACD crossed above signal in last 3 candles
                                        for i in range(-3, 0):
                                            if m_ser.iloc[i-1] <= s_ser.iloc[i-1] and m_ser.iloc[i] > s_ser.iloc[i]:
                                                macd_cross_val = True
                                                break
                        except Exception:
                            pass

                        # ====== Bollinger Bands Squeeze - bandwidth in bottom 20th pct (pending breakout) ======
                        bb_squeeze_flag = False
                        try:
                            bb_df = ta.bbands(df['Close'], length=20, std=2)
                            if bb_df is not None and not bb_df.empty:
                                upper_col = [c for c in bb_df.columns if 'BBU' in c or 'upper' in c.lower()]
                                lower_col = [c for c in bb_df.columns if 'BBL' in c or 'lower' in c.lower()]
                                mid_col   = [c for c in bb_df.columns if 'BBM' in c or 'mid' in c.lower()]
                                if upper_col and lower_col and mid_col:
                                    bbu = bb_df[upper_col[0]].dropna()
                                    bbl = bb_df[lower_col[0]].dropna()
                                    bbm = bb_df[mid_col[0]].dropna()
                                    if len(bbu) >= 20:
                                        bw = (bbu - bbl) / bbm  # bandwidth ratio
                                        pct20 = bw.quantile(0.20)
                                        if float(bw.iloc[-1]) <= float(pct20):
                                            bb_squeeze_flag = True
                        except Exception:
                            pass

                        # ====== Stage 2 (Weinstein): price > SMA150 AND SMA150 rising ======
                        # Stage 2 = markup phase - หุ้นที่ผ่าน filter นี้อยู่ในช่วงที่ดีที่สุดสำหรับการซื้อ
                        stage2_flag = False
                        try:
                            sma150 = ta.sma(df['Close'], length=150)
                            if sma150 is not None:
                                sma150_clean = sma150.dropna()
                                if len(sma150_clean) >= 20:
                                    sma150_cur = float(sma150_clean.iloc[-1])
                                    sma150_4w  = float(sma150_clean.iloc[-20])  # ~4 สัปดาห์ก่อน
                                    stage2_flag = (current_price > sma150_cur) and (sma150_cur > sma150_4w)
                        except Exception:
                            pass

                        # ====== Fundamental Data (bulk-fetched after all threads complete) ======
                        # ตัวแปรเหล่านี้ไม่ถูกใช้ใน thread - bulk enrichment เป็นตัวทำใน step 2

                        # ====== Supply & Demand Zone ======
                        entry_strat = ""
                        dz_start = None
                        dz_end = None
                        sz_start = None
                        sz_end = None
                        sl_price = None
                        rr_val = None
                        erc_vol_confirmed = False
                        zone_target_src = '52w'

                        if sd_zone:
                            entry_strat = sd_zone['type']
                            dz_start = sd_zone['start']
                            dz_end = sd_zone['end']
                            sz_start = sd_zone['target']
                            sz_end = sd_zone['target'] * 1.02
                            sl_price = sd_zone['stop_loss']
                            rr_val = sd_zone['rr_ratio']
                            erc_vol_confirmed = sd_zone.get('erc_volume_confirmed', False)
                            zone_target_src = sd_zone.get('zone_target_source', '52w')

                        prox_val = 999.0
                        if dz_start:
                            if current_price <= dz_start:
                                prox_val = 0.0
                            else:
                                prox_val = ((current_price - dz_start) / dz_start) * 100

                        gap_to_high = ((year_high - current_price) / current_price) * 100

                        # ====== Pocket Pivot (Morales & Kacher) ======
                        # Up-day volume > highest down-day volume in prior 10 sessions
                        pocket_pivot_flag = False
                        try:
                            if len(df) >= 14:
                                closes = df['Close'].values
                                volumes = df['Volume'].values
                                for _i in [-1, -2]:
                                    if float(closes[_i]) <= float(closes[_i - 1]):
                                        continue  # not an up day
                                    _start = len(volumes) + _i - 10
                                    _end   = len(volumes) + _i
                                    if _start < 1:
                                        continue
                                    _prior_c = closes[_start:_end]
                                    _prior_v = volumes[_start:_end]
                                    _prior_prev_c = closes[_start - 1:_end - 1]
                                    _down_mask = _prior_c < _prior_prev_c
                                    if not _down_mask.any():
                                        continue
                                    _max_down_vol = float(_prior_v[_down_mask].max())
                                    if float(volumes[_i]) > _max_down_vol and _max_down_vol > 0:
                                        pocket_pivot_flag = True
                                        break
                        except Exception:
                            pass

                        # ====== Volume Dry-Up (VDU): เงียบสะสม - volume ลด 3 วันติด + ต่ำกว่า avg 70% ======
                        vdu_flag = False
                        try:
                            if len(df) >= 4:
                                _vols = df['Volume'].tail(4).values.astype(float)
                                _avg20 = float(df['Volume'].tail(20).mean())
                                _declining = (_vols[-1] < _vols[-2]) and (_vols[-2] < _vols[-3])
                                _quiet     = _vols[-1] < _avg20 * 0.7
                                vdu_flag   = _declining and _quiet
                        except Exception:
                            pass

                        # ====== Ichimoku Cloud ======
                        ichimoku_above_kumo = False
                        ichimoku_tk_cross   = False
                        ichimoku_kumo_green = False
                        ichimoku_chikou_ok  = False
                        ichimoku_score_val  = 0
                        try:
                            if len(df) >= 52:
                                _h9  = df['High'].rolling(9).max()
                                _l9  = df['Low'].rolling(9).min()
                                _h26 = df['High'].rolling(26).max()
                                _l26 = df['Low'].rolling(26).min()
                                _h52 = df['High'].rolling(52).max()
                                _l52 = df['Low'].rolling(52).min()
                                _tenkan = (_h9  + _l9)  / 2
                                _kijun  = (_h26 + _l26) / 2
                                _span_a = ((_tenkan + _kijun) / 2).shift(26)
                                _span_b = ((_h52   + _l52)  / 2).shift(26)
                                _sa_cur = float(_span_a.iloc[-1]) if pd.notna(_span_a.iloc[-1]) else 0
                                _sb_cur = float(_span_b.iloc[-1]) if pd.notna(_span_b.iloc[-1]) else 0
                                # 1) Price above Kumo
                                ichimoku_above_kumo = current_price > max(_sa_cur, _sb_cur) > 0
                                # 2) TK Cross bullish in last 5 bars
                                for _i in range(-5, 0):
                                    if (_tenkan.iloc[_i-1] <= _kijun.iloc[_i-1]
                                            and _tenkan.iloc[_i] > _kijun.iloc[_i]):
                                        ichimoku_tk_cross = True
                                        break
                                # 3) Future Kumo green (SpanA > SpanB at current shifted position)
                                ichimoku_kumo_green = _sa_cur > _sb_cur and _sa_cur > 0
                                # 4) Chikou Span clear (current close > close 26 bars ago)
                                if len(df) >= 27:
                                    ichimoku_chikou_ok = float(df['Close'].iloc[-1]) > float(df['Close'].iloc[-27])
                                ichimoku_score_val = sum([ichimoku_above_kumo, ichimoku_tk_cross,
                                                         ichimoku_kumo_green, ichimoku_chikou_ok])
                        except Exception:
                            pass

                        # Price Pattern detection (ใช้ df ที่มีอยู่แล้ว)
                        pattern_result = detect_price_pattern(df)
                        pattern_name  = pattern_result['name']
                        pattern_score = pattern_result['score']

                        close_series = df['Close'].dropna()
                        rel_1m = rel_3m = 0.0
                        stock_3m_ret = 0.0
                        if len(close_series) >= 66:
                            stock_1m = float((close_series.iloc[-1] - close_series.iloc[-22]) / close_series.iloc[-22] * 100)
                            stock_3m = float((close_series.iloc[-1] - close_series.iloc[-66]) / close_series.iloc[-66] * 100)
                            stock_3m_ret = stock_3m
                            rel_1m = round(stock_1m - set_1m_return, 2)
                            rel_3m = round(stock_3m - set_3m_return, 2)
                        elif len(close_series) >= 22:
                            stock_1m = float((close_series.iloc[-1] - close_series.iloc[-22]) / close_series.iloc[-22] * 100)
                            rel_1m = round(stock_1m - set_1m_return, 2)

                        # Return dict instead of model to allow bulk fundamental enrichment and RS Ranking
                        return {
                            'symbol': symbol,
                            'price': round(current_price, 2),
                            'rsi': round(rsi, 2),
                            'adx': round(adx_val, 2),
                            'mfi': round(mfi_val, 2),
                            'rvol': round(rvol, 2),
                            'technical_score': int(integrated_score),
                            'avg_volume_20d': round(avg_vol_20, 0),
                            'rvol_bullish': rvol_bullish,
                            'erc_volume_confirmed': erc_vol_confirmed,
                            'zone_target_src': zone_target_src,
                            'entry_strat': entry_strat,
                            'dz_start': dz_start,
                            'dz_end': dz_end,
                            'sz_start': sz_start,
                            'sz_end': sz_end,
                            'sl_price': sl_price,
                            'rr_val': rr_val,
                            'year_high': round(year_high, 2),
                            'upside_to_high': round(gap_to_high, 2),
                            'prox_val': round(prox_val, 2),
                            'pattern_name': pattern_name,
                            'pattern_score': pattern_score,
                            'rel_1m': rel_1m,
                            'rel_3m': rel_3m,
                            'macd_histogram': round(macd_hist_val, 4) if macd_hist_val is not None else None,
                            'macd_crossover': macd_cross_val,
                            'bb_squeeze': bb_squeeze_flag,
                            'ema20_aligned': ema20_aligned_flag,
                            'ema20_slope': round(ema20_slope_val, 3),
                            'ema20_rising': ema20_rising_flag,
                            'hh_hl_structure': hh_hl_flag,
                            'stock_3m_ret': stock_3m_ret,
                            'rs_rating': rs_ratings_map.get(symbol, 0),
                            'stage2': stage2_flag,
                            'pocket_pivot': pocket_pivot_flag,
                            'vdu_near_zone': vdu_flag,
                            'cmf': tech.get('cmf', 0.0),
                            'is_52w_breakout': tech.get('is_52w_breakout', False),
                            'volume_surge': tech.get('volume_surge', 1.0),
                            'is_volume_surge': tech.get('is_volume_surge', False),
                            'ichimoku_score': ichimoku_score_val,
                            # ====== VCP Detection ======
                            'vcp': detect_vcp_pattern(df),
                            # ====== Launcher Data (v10) ======
                            'launcher_score': tech.get('launcher_score', 0),
                            'turtle_dist_pct': tech.get('turtle_dist_pct', 99.0),
                            'is_explosive': tech.get('is_explosive', False),
                            'tightness_idx': tech.get('tightness_idx', 99.0),
                        }

                    except Exception as e:
                        import logging
                        logging.getLogger('stocks').exception(f"[Precision] Error scanning {symbol}: {e}")
                        return None


                # ====== Phase 2: Deep Scan (only for candidates) ======
                _cache.set(ckey, {'state': 'running', 'progress': 25, 'total': 100, 'phase': f'สแกนละเอียด (จำนวน {len(results_to_process)} ตัว)...'}, timeout=900)
                
                results = []
                done_count = 0
                with concurrent.futures.ThreadPoolExecutor(max_workers=20) as executor:
                    futures = [executor.submit(_process_precision_scan, sym) for sym in results_to_process]
                    for future in concurrent.futures.as_completed(futures):
                        res = future.result()
                        if res:
                            results.append(res)
                        done_count += 1
                        _cache.set(ckey, {'state': 'running', 'progress': 25 + int((done_count/len(results_to_process))*70), 
                                          'total': 100, 'phase': f'สแกนละเอียด {done_count}/{len(results_to_process)}...'}, timeout=900)

                if results:
                    scan_df = pd.DataFrame(results)
                    if 'rs_rating' not in scan_df.columns:
                        scan_df['rs_rating'] = 0

                    _cache.set(ckey, {'state': 'running', 'progress': done_count, 'total': len(sym_list), 'phase': 'ดึงข้อมูล Fundamental…'}, timeout=900)

                    # ====== Bulk Fundamental Enrichment ======
                    matched_symbols = [r['symbol'] for r in results]
                    symbols_bk = [f"{s}.BK" for s in matched_symbols]
                    fund_data = {}
                    try:
                        yq_all = YQTicker(symbols_bk)
                        modules = yq_all.get_modules('financialData summaryProfile')
                        for sym_bk, data in modules.items():
                            if not isinstance(data, dict):
                                continue
                            clean_sym = sym_bk.replace('.BK', '')
                            profile  = data.get('summaryProfile', {})
                            fin_data = data.get('financialData', {})
                            sector   = (
                                profile.get('sector')
                                or data.get('assetProfile', {}).get('sector')
                                or 'Unknown'
                            )
                            eps_growth = float(fin_data.get('earningsQuarterlyGrowth', 0) or 0) * 100
                            rev_growth = float(fin_data.get('revenueGrowth', 0) or 0) * 100
                            fund_data[clean_sym] = {'sector': sector, 'eps_growth': eps_growth, 'rev_growth': rev_growth}
                    except Exception as e:
                        print(f"[Precision] Bulk Fundamental fetch failed: {e}")

                    # ====== Bulk Create ======
                    bulk_candidates = []
                    for r in scan_df.to_dict('records'):
                        sym = r['symbol']
                        f = fund_data.get(sym, {'sector': 'N/A', 'eps_growth': 0.0, 'rev_growth': 0.0})
                        bulk_candidates.append(PrecisionScanCandidate(
                            user=user,
                            market='SET',
                            scan_run=scan_run_time,
                            symbol=sym,
                            symbol_bk=f"{sym}.BK",
                            sector=f.get('sector') or 'Unknown',
                            price=r['price'],
                            rsi=r['rsi'],
                            adx=r['adx'],
                            mfi=r['mfi'],
                            rvol=r['rvol'],
                            eps_growth=round(f.get('eps_growth', 0), 2),
                            rev_growth=round(f.get('rev_growth', 0), 2),
                            technical_score=r['technical_score'],
                            rs_rating=r['rs_rating'],
                            avg_volume_20d=r['avg_volume_20d'],
                            rvol_bullish=r['rvol_bullish'],
                            erc_volume_confirmed=r['erc_volume_confirmed'],
                            zone_target_source=r['zone_target_src'],
                            is_new_entry=(sym not in prev_symbols),
                            entry_strategy=r['entry_strat'],
                            demand_zone_start=r['dz_start'],
                            demand_zone_end=r['dz_end'],
                            supply_zone_start=r['sz_start'],
                            supply_zone_end=r['sz_end'],
                            stop_loss=r['sl_price'],
                            risk_reward_ratio=r['rr_val'],
                            year_high=r['year_high'],
                            upside_to_high=r['upside_to_high'],
                            zone_proximity=r['prox_val'],
                            price_pattern=r['pattern_name'],
                            price_pattern_score=r['pattern_score'],
                            rel_momentum_1m=r['rel_1m'],
                            rel_momentum_3m=r['rel_3m'],
                            macd_histogram=r['macd_histogram'],
                            macd_crossover=r['macd_crossover'],
                            bb_squeeze=r['bb_squeeze'],
                            ema20_aligned=r['ema20_aligned'],
                            ema20_slope=r.get('ema20_slope', 0.0),
                            ema20_rising=r.get('ema20_rising', False),
                            hh_hl_structure=r.get('hh_hl_structure', False),
                            stage2=r.get('stage2', False),
                            pocket_pivot=r.get('pocket_pivot', False),
                            vdu_near_zone=r.get('vdu_near_zone', False),
                            cmf=r.get('cmf', None),
                            is_52w_breakout=r.get('is_52w_breakout', False),
                            volume_surge=r.get('volume_surge', 1.0),
                            is_volume_surge=r.get('is_volume_surge', False),
                            ichimoku_above_kumo=r.get('ichimoku_above_kumo', False),
                            ichimoku_tk_cross=r.get('ichimoku_tk_cross', False),
                            ichimoku_kumo_green=r.get('ichimoku_kumo_green', False),
                            ichimoku_chikou_ok=r.get('ichimoku_chikou_ok', False),
                            ichimoku_score=r.get('ichimoku_score', 0),
                            # VCP v9
                            vcp_setup=r.get('vcp', {}).get('setup', False),
                            vcp_contractions=r.get('vcp', {}).get('contractions', 0),
                            vcp_tightness=r.get('vcp', {}).get('tightness', 0.0),
                            vcp_vdu=r.get('vcp', {}).get('vdu_confirmed', False),
                            # Launcher v10
                            launcher_score=r.get('launcher_score', 0),
                            turtle_dist_pct=r.get('turtle_dist_pct', 99.0),
                            is_explosive=r.get('is_explosive', False),
                            tightness_idx=r.get('tightness_idx', 99.0),
                        ))

                    if bulk_candidates:
                        PrecisionScanCandidate.objects.bulk_create(bulk_candidates)

                # เก็บ 3 รอบล่าสุด
                distinct_runs = (
                    PrecisionScanCandidate.objects
                    .filter(user=user, market='SET')
                    .values_list('scan_run', flat=True)
                    .order_by('-scan_run')
                    .distinct()
                )
                runs_list = list(distinct_runs)
                if len(runs_list) > 3:
                    old_runs = runs_list[3:]
                    PrecisionScanCandidate.objects.filter(user=user, market='SET', scan_run__in=old_runs).delete()

                _cache.set(ckey, {'state': 'done', 'count': len(results)}, timeout=300)

            except Exception as _bg_err:
                import logging
                logging.getLogger('stocks').exception(f"[PrecisionBG] Error: {_bg_err}")
                from django.core.cache import cache as _ec
                _ec.set(ckey, {'state': 'idle'}, timeout=60)

        # เปิด background thread แล้ว return ทันที
        _t = threading.Thread(
            target=_run_precision_bg,
            args=(user_id, cache_key, scan_symbols),
            daemon=True
        )
        _t.start()
        
        # Redirect กลับหน้าที่ส่งมา (เช่น SEPA)
        next_url = request.POST.get('next_url')
        if next_url == 'sepa':
            return redirect('stocks:minervini_sepa_scanner')
            
        return redirect('stocks:precision_momentum_scanner')

    # ====== จัดเรียงผลลัพธ์ ======
    sort_by = request.GET.get('sort', 'score')
    valid_db_sorts = {
        'symbol': 'symbol',
        'score': '-technical_score',
        'price': '-price',
        'rsi': '-rsi',
        'rvol': '-rvol',
        'adx': '-adx',
        'prox': 'zone_proximity',
        'round_rr': '-risk_reward_ratio',
        'rs': '-rs_rating',          # RS Rating (Minervini Relative Strength)
        'launcher': '-launcher_score', # Explosive Launcher Score
    }
    use_db_sort = sort_by in valid_db_sorts
    order_field = valid_db_sorts.get(sort_by, '-technical_score')

    # รายชื่อ scan runs ทั้งหมด (index 0 = ล่าสุด)
    all_runs = list(
        PrecisionScanCandidate.objects
        .filter(user=request.user, market='SET')
        .values_list('scan_run', flat=True)
        .order_by('-scan_run')
        .distinct()
    )

    # เลือกรอบสแกนตาม ?run_idx= (0 = ล่าสุด, 1 = ก่อนหน้า, ...)
    try:
        run_idx = int(request.GET.get('run_idx', 0))
    except (ValueError, TypeError):
        run_idx = 0
    run_idx = max(0, min(run_idx, len(all_runs) - 1)) if all_runs else 0

    candidates = []
    scanned_at = None
    if all_runs:
        selected_run = all_runs[run_idx]
        qs = PrecisionScanCandidate.objects.filter(user=request.user, scan_run=selected_run, market='SET')
        if use_db_sort:
            qs = qs.order_by(order_field)
        candidates = list(qs)
        scanned_at = selected_run

        # ====== Live Price Fetch (ถ้าตลาดเปิด) - แสดงราคาปัจจุบันคู่กับราคา close เมื่อวาน ======
        # Indicator ยังคำนวณจาก settled close (คงที่), live price ใช้แสดงเท่านั้น
        import pytz as _lpytz
        from datetime import datetime as _ldt, time as _ldtime
        _lbkk = _lpytz.timezone('Asia/Bangkok')
        _lnow = _ldt.now(_lbkk)
        _lt = _lnow.time()
        # Live price: เฉพาะช่วงที่ตลาด SET ซื้อขายจริง (ไม่รวม midday break 12:30-14:30)
        _lmarket_open = (
            _lnow.weekday() < 5 and
            (_ldtime(10, 0) <= _lt <= _ldtime(12, 30) or
             _ldtime(14, 30) <= _lt <= _ldtime(16, 30))
        )
        live_prices = {}
        live_mcaps  = {}
        live_zones  = {}   # fresh zones - keyed by symbol
        if candidates:
            try:
                import concurrent.futures as _lcf
                from datetime import timedelta as _ltd, time as _ltime2
                # end date เหมือน entry_finder
                _prec_market_open = (
                    _lnow.weekday() < 5 and
                    (
                        _lt < _ltime2(10, 0) or
                        (_ltime2(10, 0) <= _lt <= _ltime2(12, 30)) or
                        (_ltime2(12, 30) < _lt < _ltime2(14, 30)) or
                        (_ltime2(14, 30) <= _lt <= _ltime2(16, 30))
                    )
                )
                _prec_end_date  = (_lnow.date() - _ltd(days=1)) if _prec_market_open else _lnow.date()
                _prec_end_str   = _prec_end_date.strftime('%Y-%m-%d')
                _prec_start_str = (_prec_end_date - _ltd(days=400)).strftime('%Y-%m-%d')

                def _get_live(sym):
                    try:
                        full_sym = f"{sym}.BK"
                        # Use fast_info for quick price/mcap without downloading full history
                        fi = yf.Ticker(full_sym).fast_info
                        p  = getattr(fi, 'last_price', None)
                        mc = getattr(fi, 'market_cap', None)
                        return sym, (float(p) if p else None), (round(float(mc)/1e9, 2) if mc else None), None
                    except Exception:
                        return sym, None, None, None
                with _lcf.ThreadPoolExecutor(max_workers=6) as _lex:
                    for _sym, _p, _mc, _fz in _lex.map(_get_live, [c.symbol for c in candidates]):
                        if _p:  live_prices[_sym] = _p
                        if _mc: live_mcaps[_sym]  = _mc
                        if _fz: live_zones[_sym]  = _fz
            except Exception:
                pass

        for c in candidates:
            lp = live_prices.get(c.symbol)
            c.live_price      = lp
            c.live_market_cap = live_mcaps.get(c.symbol)
            c.is_live         = _lmarket_open and lp is not None

            # อัปเดต zone ด้วย fresh data ถ้ามี
            fz = live_zones.get(c.symbol)
            if fz:
                c.demand_zone_start = fz.get('start') or c.demand_zone_start
                c.demand_zone_end   = fz.get('end')   or c.demand_zone_end
                c.supply_zone_start = fz.get('target') or c.supply_zone_start
                c.stop_loss         = fz.get('stop_loss') or c.stop_loss

            ref_price = lp if lp else float(c.price or 0)
            dz_top = float(c.demand_zone_start or 0)
            dz_bot = float(c.demand_zone_end   or 0)
            tp     = float(c.supply_zone_start or 0)

            if lp and dz_top > 0:
                c.live_zone_prox = 0.0 if lp <= dz_top else round(((lp - dz_top) / dz_top) * 100, 1)
            else:
                c.live_zone_prox = None
            if lp and c.price and c.price > 0:
                c.live_change_pct = round(((lp - float(c.price)) / float(c.price)) * 100, 2)
            else:
                c.live_change_pct = None

            # precompute zone status flags - ใช้ใน template แทนการคำนวณใน {% if %}
            c.live_at_tp      = tp > 0 and ref_price >= tp
            c.live_broke_zone = dz_bot > 0 and ref_price < dz_bot
            c.live_in_zone    = dz_top > 0 and dz_bot > 0 and dz_bot <= ref_price <= dz_top

            # upside_to_tp: % ของ range Entry→TP ที่ยังเหลืออยู่
            entry = dz_top
            total_range = tp - entry
            if tp > 0 and entry > 0 and total_range > 0 and ref_price > 0:
                remaining = tp - ref_price
                c.upside_to_tp = round((remaining / total_range) * 100, 1)
            else:
                c.upside_to_tp = 999

        # ====== คำนวณ BUY/SELL Score ด้วย _compute_signals() เดียวกับ Dashboard/Watchlist ======
        for c in candidates:
            sigs = _compute_signals(c)
            c.buy_score  = sigs['buy_score']
            c.sell_score = sigs['sell_score']
            c.exit_signal = sigs['exit_signal']

        # ====== BUY Score Delta เทียบกับรอบก่อนหน้า ======
        prev_buy_scores = {}
        if len(all_runs) > run_idx + 1:
            for _p in PrecisionScanCandidate.objects.filter(
                    user=request.user, scan_run=all_runs[run_idx + 1]):
                _ps = _compute_signals(_p)
                prev_buy_scores[_p.symbol] = _ps['buy_score']
        for c in candidates:
            _prev = prev_buy_scores.get(c.symbol)
            c.buy_score_delta = (c.buy_score - _prev) if _prev is not None else None

        # เรียงตาม BUY/SELL/RS score ด้วย Python (fallback ถ้าไม่ใช่ DB sort)
        if sort_by == 'buy':
            candidates.sort(key=lambda x: x.buy_score, reverse=True)
        elif sort_by == 'sell':
            candidates.sort(key=lambda x: x.sell_score, reverse=True)
        elif sort_by == 'rs':
            candidates.sort(key=lambda x: getattr(x, 'rs_rating', 0), reverse=True)

        # ====== Top 5 หุ้นแนะนำซื้อ (BUY score สูง) ======
        # เงื่อนไข: RVOL Bull ≥ 1.0x (มีแรงซื้อจริง) + RSI ไม่ overbought
        def _top5_filter(min_rvol, max_rsi=85):
            # max_rsi=85 สอดคล้องกับ _compute_signals ที่ยังให้ +2 กับ RSI 80-85
            return sorted(
                [c for c in candidates
                 if c.buy_score >= 50
                 and c.rvol_bullish
                 and c.rvol >= min_rvol
                 and c.rsi <= max_rsi
                 and not (c.demand_zone_end and float(getattr(c, 'live_price', None) or c.price or 0) < float(c.demand_zone_end))],
                key=lambda x: x.buy_score, reverse=True
            )[:5]

        top5_buy = _top5_filter(1.0)
        if len(top5_buy) < 5:
            top5_buy = _top5_filter(0.7)
        if len(top5_buy) < 3:
            top5_buy = _top5_filter(0.0)

        # ====== Top 5 หุ้นที่ "ผ่านเกณฑ์ครบทุกข้อ" ======
        # ผ่อนปรนเกณฑ์ ADX 20 (เดิม 25) และ RSI ขยายเพื่อให้มีตัวเลือกมากขึ้น 
        def _is_fully_qualified(c):
            rr = c.risk_reward_ratio or 0
            dz_start = float(c.demand_zone_start or 0)
            dz_end   = float(c.demand_zone_end   or 0)
            # ใช้ live price ถ้ามี เพราะ zone_proximity ใน DB เป็นค่า ณ เวลาสแกน
            price    = float(getattr(c, 'live_price', None) or c.price or 0)
            live_prox = getattr(c, 'live_zone_prox', None)
            effective_prox = live_prox if live_prox is not None else c.zone_proximity
            in_zone  = dz_start > 0 and dz_end > 0 and dz_end <= price <= dz_start
            above_zone = price > dz_start and effective_prox <= 30
            near_zone  = in_zone or above_zone
            # ตัดหุ้นที่ราคาหลุดต่ำกว่า demand zone (ทะลุ SL) - ไม่ใช่จุดซื้อแล้ว
            broke_zone = dz_end > 0 and price < dz_end
            # ตัดหุ้นที่ราคาวิ่งขึ้นเกือบถึง/เกิน target แล้ว
            target = c.supply_zone_start or 0
            upside_pct = ((target - price) / price * 100) if (target > 0 and price > 0) else 999
            price_near_target = target > 0 and upside_pct < 8
            return (
                c.buy_score >= 65
                and rr >= 1.5
                and c.adx >= 20
                and 45 <= c.rsi <= 82
                and c.rvol_bullish
                and c.rvol >= 0.8
                and near_zone
                and not broke_zone             # กรอง: ราคาหลุดต่ำกว่า zone (ทะลุ SL)
                and not price_near_target      # กรอง: ราคาใกล้ target แล้ว
                and (c.sell_score or 0) < 50
                and getattr(c, 'rs_rating', 0) >= 60
            )

        top5_qualified = sorted(
            [c for c in candidates if _is_fully_qualified(c)],
            key=lambda x: x.buy_score, reverse=True
        )

        for c in top5_buy:
            reasons = []
            in_zone = (c.demand_zone_start and c.demand_zone_end and
                       c.price <= c.demand_zone_start and c.price >= c.demand_zone_end)
            if in_zone:
                reasons.append("อยู่ใน Entry Zone แล้ว")
            elif c.zone_proximity <= 10:
                reasons.append(f"เหนือโซน {c.zone_proximity:.0f}% (รอย่อ)")
            elif c.zone_proximity <= 30:
                reasons.append(f"เหนือโซน {c.zone_proximity:.0f}%")

            if c.rvol_bullish and c.rvol >= 1.5:
                reasons.append(f"RVOL {c.rvol:.1f}x Bull แรง")
            elif c.rvol_bullish and c.rvol >= 1.0:
                reasons.append(f"RVOL {c.rvol:.1f}x Bull ยืนยัน")

            rr = c.risk_reward_ratio or 0
            if rr >= 3:
                reasons.append(f"RR 1:{rr:.1f} ดีเยี่ยม")
            elif rr >= 2:
                reasons.append(f"RR 1:{rr:.1f} ดี")

            if c.adx >= 30:
                reasons.append(f"ADX {c.adx:.0f} เทรนด์แข็ง")
            elif c.adx >= 25:
                reasons.append(f"ADX {c.adx:.0f} มีเทรนด์")

            if c.technical_score >= 85:
                reasons.append(f"Precision {c.technical_score} สูงมาก")
            elif c.technical_score >= 75:
                reasons.append(f"Precision {c.technical_score} ดี")

            if c.erc_volume_confirmed:
                reasons.append("ERC ยืนยันแล้ว")

            if 55 <= c.rsi <= 70:
                reasons.append(f"RSI {c.rsi:.0f} จุดหวาน")

            if c.price_pattern and c.price_pattern_score > 0:
                reasons.append(f"Pattern: {c.price_pattern}")

            rel = c.rel_momentum_3m if c.rel_momentum_3m != 0.0 else c.rel_momentum_1m
            if rel >= 8:
                reasons.append(f"ชนะ SET +{rel:.1f}% (3m)")
            elif rel >= 3:
                reasons.append(f"ชนะ SET +{rel:.1f}%")

            rs = getattr(c, 'rs_rating', 0)
            if rs >= 85:
                reasons.insert(1, f"RS {rs} - ผู้นำตลาด")   # สอดในตำแหน่ง 2 เสมอ
            elif rs >= 70:
                reasons.insert(1, f"RS {rs} - แข็งแกร่ง")

            c.top_reasons = reasons[:4]

        # เพิ่ม reasons ให้ top5_qualified ด้วย (บางตัวอาจซ้ำกับ top5_buy)
        for c in top5_qualified:
            if not hasattr(c, 'top_reasons'):
                reasons = []
                in_zone = (c.demand_zone_start and c.demand_zone_end and
                           c.price <= c.demand_zone_start and c.price >= c.demand_zone_end)
                if in_zone:
                    reasons.append("อยู่ใน Entry Zone แล้ว")
                elif c.zone_proximity <= 10:
                    reasons.append(f"เหนือโซน {c.zone_proximity:.0f}% (รอย่อ)")
                else:
                    reasons.append(f"เหนือโซน {c.zone_proximity:.0f}%")
                rr = c.risk_reward_ratio or 0
                reasons.append(f"RR 1:{rr:.1f} ✓")
                reasons.append(f"ADX {c.adx:.0f} ✓")
                rs = getattr(c, 'rs_rating', 0)
                if rs >= 85:
                    reasons.insert(1, f"RS {rs} - ผู้นำตลาด")
                elif rs >= 70:
                    reasons.insert(1, f"RS {rs} - แข็งแกร่ง")
                if c.rvol >= 1.5:
                    reasons.append(f"RVOL {c.rvol:.1f}x Bull ✓")
                c.top_reasons = reasons[:4]

        # ====== Leading Sectors Analysis (v3) ======
        # นับจำนวนหุ้นที่ 'แข็งแกร่ง' (buy_score >= 65) ในแต่ละ sector เพื่อหาผู้นำกลุ่ม
        sector_counts = {}
        for c in candidates:
            if c.buy_score >= 65:
                sec = c.sector or 'Unknown'
                sector_counts[sec] = sector_counts.get(sec, 0) + 1
        
        # เรียงลำดับกลุ่มที่แข็งแกร่งที่สุด 5 อันดับแรก
        top_sectors = sorted(
            [{'name': k, 'count': v} for k, v in sector_counts.items()],
            key=lambda x: x['count'], reverse=True
        )[:5]

        # ====== Automated AI Insights (คำอธิบายวิเคราะห์ผลสแกน) ======
        scan_insights = []
        if top5_qualified:
            best_setup = top5_qualified[0]
            rr_val = best_setup.risk_reward_ratio or 0
            rs_val = getattr(best_setup, "rs_rating", 0)
            target_val = best_setup.supply_zone_start or 0
            upside_left = ((target_val - best_setup.price) / best_setup.price * 100) if (target_val > 0 and best_setup.price > 0) else 0
            scan_insights.append({
                'icon': '🏆',
                'title': f'โฟกัสหลัก: {best_setup.symbol} (หน้าเทรดคุ้มค่า ปลอดภัยสูง)',
                'desc': f'ถือเป็นหน้าเทรด Swing Trade ที่สมบูรณ์แบบที่สุดในรอบนี้ เพราะมีความแข็งแกร่ง (RS {rs_val}) เทรนด์ทำมุมสวยงาม แต่ราคาปัจจุบันกลับอยู่ใกล้จุดเข้าซื้อ ทำให้มีความคุ้มค่ากับความเสี่ยงสูง (RR 1:{rr_val:.1f}) Upside เหลืออีก {upside_left:.1f}% ถึง Target ซื้อแล้วมีโอกาสชนะสูง'
            })
        
        # หาหุ้นซิ่งที่สุดแต่เสี่ยงสูง (RS > 90 แต่ RR < 1.0)
        high_risk_momentum = [c for c in top5_buy if getattr(c, 'rs_rating', 0) >= 90 and (c.risk_reward_ratio or 0) < 1.5]
        if high_risk_momentum:
            hm = high_risk_momentum[0]
            if not top5_qualified or hm.symbol != top5_qualified[0].symbol:
                scan_insights.append({
                    'icon': '🚀',
                    'title': f'สายซิ่ง (เสี่ยงสูง): {hm.symbol} (แรงทะลุกราฟ)',
                    'desc': f'นี่คือ "หุ้นโคตรโมเมนตัม" วิ่งแรงชนะตลาดถึง {getattr(hm, "rs_rating", 0)}% วอลุ่มเข้าหนัก แต่อัตราคุ้มทุน (RR) ณ ราคานี้ได้เพียง 1:{hm.risk_reward_ratio or 0:.1f} แปลว่าราคาลอยขึ้นมาพอสมควรแล้ว "อย่าเพิ่งไล่ราคา (Market Buy) แนะนำให้เอาเข้า Watchlist แล้วรอซื้อตอนย่อตัวพักฐาน"'
                })

        # หาหุ้นเพิ่งเริ่มกลับตัว (MACD crossover)
        reversal_stocks = [c for c in top5_buy if any('MACD' in str(r) for r in (c.top_reasons or []))]
        if reversal_stocks and len(scan_insights) < 3:
            rev = reversal_stocks[0]
            skip = False
            if top5_qualified and rev.symbol == top5_qualified[0].symbol: skip = True
            if high_risk_momentum and rev.symbol == high_risk_momentum[0].symbol: skip = True
            if not skip:
                scan_insights.append({
                    'icon': '🥈',
                    'title': f'สัญญาณกลับตัว: {rev.symbol} (ต้นเทรนด์)',
                    'desc': f'หุ้นตัวนี้มีสัญญาณเชิงบวกคือเกิด MACD Crossover (ตัดเส้น Signal ขึ้นมา) มักใช้ระบุจุดเริ่มต้นของรอบขาขึ้นชุดใหม่ น่าเก็บสะสมที่บริเวณแนวรับปัจจุบัน'
                })
        
        # ถ้าระบบยังไม่มีคำแนะนำเลย
        if not scan_insights and top5_buy:
            top = top5_buy[0]
            scan_insights.append({
                'icon': '💡',
                'title': f'น่าปั้นเทรนด์: {top.symbol}',
                'desc': f'ได้คะแนนเข้าซื้อสูงสุดในรอบนี้ มีโครงสร้างทางเทคนิคโดยรวมเฉลี่ยดีที่สุด เหมาะเป็นหุ้นน่าจับตามอง'
            })

    else:
        top5_buy = []
        top5_qualified = []
        top_sectors = []
        scan_insights = []

    # ====== Market Condition - ดึงข้อมูล SET Index สำหรับแสดงผล (GET + POST) ======
    market_condition = {'phase': 'UNKNOWN', 'label': 'ไม่มีข้อมูล', 'color': 'secondary', 'score': 0}
    try:
        from datetime import datetime as _mcdt, timedelta as _mctd
        import pytz as _mcpytz
        _mc_bkk   = _mcpytz.timezone('Asia/Bangkok')
        _mc_now   = _mcdt.now(_mc_bkk)
        _mc_end   = _mc_now.date().strftime('%Y-%m-%d')
        _mc_start = (_mc_now.date() - _mctd(days=430)).strftime('%Y-%m-%d')
        _mc_df = yf.download("^SET", start=_mc_start, end=_mc_end, interval="1d", progress=False)
        if _mc_df is not None and not _mc_df.empty:
            if isinstance(_mc_df.columns, pd.MultiIndex):
                _mc_df.columns = _mc_df.columns.droplevel(1)
            market_condition = _get_market_condition(_mc_df)
    except Exception:
        pass

    context = {
        'title': 'Precision Momentum Scanner - กรองคุณภาพ',
        'candidates': candidates,
        'scanned_at': scanned_at,
        'current_sort': sort_by,
        'all_runs': all_runs,
        'selected_run_idx': run_idx,
        'has_scanned': request.method == "POST" or request.GET.get('scan') == 'true' or bool(all_runs),
        'top5_buy': top5_buy,
        'top5_qualified': top5_qualified,
        'scan_total': len(scan_symbols),
        'scan_passed': len(candidates),
        'top_sectors': top_sectors,
        'scan_insights': scan_insights,
        'scan_data_date': None,  # คำนวณด้านล่าง
        'market_condition': market_condition,
    }
    # คำนวณ scan_data_date จาก scanned_at - ถ้า scan ทำหลัง 16:30 BKK ข้อมูลคือวันเดียวกัน
    # ถ้า scan ทำระหว่าง 10:00-16:30 (ตลาดเปิด) ข้อมูลจะเป็นวันก่อนหน้า
    if scanned_at:
        import pytz as _sddtz
        _bkk = _sddtz.timezone('Asia/Bangkok')
        _st = scanned_at.astimezone(_bkk) if hasattr(scanned_at, 'astimezone') else scanned_at
        from datetime import time as _t, timedelta as _tdd
        # scan_data_date: midday break ยังถือว่า candle ของวันนั้นยังไม่ settle
        _in_mkt = (
            _st.weekday() < 5 and (
                _t(10, 0) <= _st.time() <= _t(12, 30) or
                _t(12, 30) < _st.time() < _t(14, 30) or
                _t(14, 30) <= _st.time() <= _t(16, 30)
            )
        )
        context['scan_data_date'] = (_st.date() - _tdd(days=1)) if _in_mkt else _st.date()

    # Build AI scan JSON for Gemini analysis button
    import json as _scan_json
    def _ser_c(c):
        return {
            "symbol": c.symbol,
            "price": c.price,
            "buy_score": getattr(c, 'buy_score', 0),
            "rs_rating": getattr(c, 'rs_rating', 0),
            "rsi": round(c.rsi, 1),
            "adx": round(c.adx, 1),
            "rvol": round(c.rvol, 2),
            "rvol_bullish": c.rvol_bullish,
            "risk_reward_ratio": c.risk_reward_ratio,
            "zone_proximity": round(c.zone_proximity, 1) if c.zone_proximity else None,
            "macd_crossover": getattr(c, 'macd_crossover', False),
            "ema20_aligned": getattr(c, 'ema20_aligned', False),
            "ema20_rising": getattr(c, 'ema20_rising', False),
            "hh_hl_structure": getattr(c, 'hh_hl_structure', False),
            "bb_squeeze": getattr(c, 'bb_squeeze', False),
            "rel_momentum_3m": getattr(c, 'rel_momentum_3m', 0),
            "sector": c.sector,
            "exit_signal": getattr(c, 'exit_signal', ''),
            "top_reasons": getattr(c, 'top_reasons', []),
        }
    _ai_data = {
        "scan_date": str(context.get('scan_data_date', '')),
        "qualified_stocks": [_ser_c(c) for c in top5_qualified],
        "top_buy_stocks": [_ser_c(c) for c in top5_buy],
        "total_passed": len(candidates),
        "top_sectors": [{"name": s["name"], "count": s["count"]} for s in top_sectors],
    }
    context['ai_scan_json'] = _scan_json.dumps(_ai_data, ensure_ascii=False, default=str)

    # Watchlist symbols for toggle button state
    from .models import ScanWatchlistItem
    context['watchlist_symbols'] = set(
        ScanWatchlistItem.objects.filter(user=request.user).values_list('symbol', flat=True)
    )

    return render(request, 'stocks/precision_scan.html', context)


# ====== Portfolio Momentum Scan - สแกนเฉพาะหุ้นใน Portfolio ======

@login_required
def portfolio_scan(request):
    """
    สแกน Momentum เฉพาะหุ้นที่อยู่ใน Portfolio ของผู้ใช้
    ใช้ Logic เดียวกันกับ momentum_scanner() แต่เปลี่ยน Input จาก SET100+MAI เป็นหุ้นใน Portfolio
    """
    from .utils import analyze_momentum_technical, find_supply_demand_zones
    from types import SimpleNamespace

    portfolio_items = Portfolio.objects.filter(user=request.user, category='STOCK')

    candidates = []
    scanned_at = None

    if request.method == "POST" or request.GET.get('scan') == 'true':
        import pandas_ta as ta
        import datetime

        for item in portfolio_items:
            symbol = item.symbol.upper().replace('.BK', '')
            try:
                print(f"[PortfolioScan] Scanning {symbol}...")
                df = yf.download(f"{symbol}.BK", period="1y", interval="1d", progress=False)

                if df is None or df.empty:
                    try:
                        yq = YQTicker(f"{symbol}.BK")
                        df = yq.history(period="1y", interval="1d")
                        if isinstance(df, pd.DataFrame) and not df.empty:
                            df = df.reset_index()
                            if 'date' in df.columns:
                                df.set_index('date', inplace=True)
                            if 'symbol' in df.columns:
                                df.drop(columns=['symbol'], inplace=True)
                            df.rename(columns={'open': 'Open', 'high': 'High', 'low': 'Low',
                                               'close': 'Close', 'volume': 'Volume'}, inplace=True)
                    except Exception:
                        pass

                if df is None or df.empty:
                    continue

                if isinstance(df.columns, pd.MultiIndex):
                    df.columns = df.columns.droplevel(1)

                df = df.dropna(subset=['Close', 'High'])
                if len(df) < 150:
                    continue

                # ====== คำนวณ Technical Indicators (เหมือน momentum_scanner) ======
                df['EMA50'] = ta.ema(df['Close'], length=50)
                df['EMA150'] = ta.ema(df['Close'], length=150)
                df['EMA200'] = ta.ema(df['Close'], length=200)

                adx_df = ta.adx(df['High'], df['Low'], df['Close'], length=14)
                if adx_df is not None and not adx_df.empty:
                    df = pd.concat([df, adx_df], axis=1)

                mfi = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                df['MFI'] = mfi

                avg_vol_20 = df['Volume'].rolling(window=20).mean()
                df['RVOL'] = df['Volume'] / avg_vol_20

                tech = analyze_momentum_technical(df)

                current_price = float(df['Close'].iloc[-1])
                year_high = float(df['High'].tail(252).max())

                integrated_score = tech['score']
                rvol = tech['rvol']
                rsi = tech['rsi']
                ema200 = tech['ema200']

                mfi_val = float(df['MFI'].iloc[-1]) if 'MFI' in df.columns and pd.notna(df['MFI'].iloc[-1]) else 0
                adx = float(df['ADX_14'].iloc[-1]) if 'ADX_14' in df.columns and pd.notna(df['ADX_14'].iloc[-1]) else 0
                gap_to_high = ((year_high - current_price) / current_price) * 100

                # ====== เกณฑ์กรอง Trend Template (เหมือน momentum_scanner) ======
                is_uptrend = (current_price > ema200)
                near_high = (current_price >= year_high * 0.60)

                if is_uptrend and near_high:
                    sector = "Unknown"
                    eps_growth = 0.0
                    rev_growth = 0.0
                    fund_bonus = 0

                    try:
                        ticker_yf = yf.Ticker(f"{symbol}.BK")
                        info = ticker_yf.info
                        if isinstance(info, dict) and len(info) >= 5:
                            sector = info.get('sector', 'Other')
                            eps_growth = float(info.get('earningsQuarterlyGrowth', 0) or 0) * 100
                            rev_growth = float(info.get('revenueGrowth', 0) or 0) * 100
                        else:
                            sector = "N/A"
                    except Exception:
                        pass

                    if eps_growth >= 20:
                        fund_bonus += 10
                    if rev_growth >= 10:
                        fund_bonus += 10

                    sd_zone = find_supply_demand_zones(df)
                    dz_start = dz_end = sz_start = sz_end = sl_price = rr_val = None
                    entry_strat = ""
                    prox_val = 999.0

                    if sd_zone:
                        entry_strat = sd_zone['type']
                        dz_start = sd_zone['start']
                        dz_end = sd_zone['end']
                        sz_start = sd_zone['target']
                        sz_end = sd_zone['target'] * 1.02
                        sl_price = sd_zone['stop_loss']
                        rr_val = sd_zone['rr_ratio']

                    if dz_start:
                        prox_val = 0.0 if current_price <= dz_start else ((current_price - dz_start) / dz_start) * 100

                    candidates.append(SimpleNamespace(
                        symbol=symbol,
                        symbol_bk=f"{symbol}.BK",
                        sector=sector,
                        price=round(current_price, 2),
                        rsi=round(rsi, 2),
                        adx=round(adx, 2),
                        mfi=round(mfi_val, 2),
                        rvol=round(rvol, 2),
                        eps_growth=round(eps_growth, 2),
                        rev_growth=round(rev_growth, 2),
                        technical_score=int(integrated_score + fund_bonus),
                        entry_strategy=entry_strat,
                        demand_zone_start=dz_start,
                        demand_zone_end=dz_end,
                        supply_zone_start=sz_start,
                        supply_zone_end=sz_end,
                        stop_loss=sl_price,
                        risk_reward_ratio=rr_val,
                        year_high=round(year_high, 2),
                        upside_to_high=round(gap_to_high, 2),
                        zone_proximity=round(prox_val, 2),
                        portfolio_name=item.name,
                    ))

            except Exception as e:
                print(f"!!! [PortfolioScan] Error scanning {symbol}: {str(e)}")
                continue

        candidates.sort(key=lambda x: x.technical_score, reverse=True)
        scanned_at = datetime.datetime.now()

    context = {
        'title': 'Portfolio Momentum Scan',
        'candidates': candidates,
        'portfolio_count': portfolio_items.count(),
        'scanned_at': scanned_at,
        'has_scanned': request.method == "POST" or request.GET.get('scan') == 'true',
    }
    return render(request, 'stocks/portfolio_scan.html', context)


# ====== Entry Finder - กราฟ Sniper Entry พร้อม Supply & Demand Zone ======

@login_required
def entry_finder(request, symbol):
    """
    Detailed view for Supply & Demand / Sniper Entry zone for a specific symbol.
    แสดงกราฟ 120 วันพร้อม:
    - Demand Zone (โซนเข้าซื้อ)
    - Supply Zone / Target (โซนขาย)
    - Stop Loss Line
    - EMA50 และ EMA200
    """
    # ถ้า market=US ไม่ต้องเติม .BK (US stocks ใช้ ticker เดิม)
    market = request.GET.get('market')
    
    # หากไม่ได้ระบุ market มาใน URL ให้ลองตรวจสอบจากฐานข้อมูลก่อน
    if not market:
        from .models import PrecisionScanCandidate, MomentumCandidate
        # ลองหาจากรอบสแกนล่าสุด
        cand = PrecisionScanCandidate.objects.filter(symbol=symbol).order_by('-scan_run').first()
        if cand:
            market = cand.market
        else:
            # ลองหาจาก momentum
            mom = MomentumCandidate.objects.filter(symbol=symbol).first()
            if mom:
                market = mom.market
            else:
                market = 'SET' # Default
    
    if market == 'US':
        full_symbol = symbol
    else:
        full_symbol = f"{symbol}.BK" if not symbol.endswith(".BK") else symbol

    try:
        # ใช้ scan_end_date เดียวกับ Precision Scanner เพื่อให้ Zone ตรงกัน
        from datetime import datetime as _efdt, timedelta as _eftd, time as _efdtime
        import pytz as _efpytz
        _ef_bkk = _efpytz.timezone('Asia/Bangkok')
        _ef_now = _efdt.now(_ef_bkk)
        _ef_t   = _ef_now.time()
        _ef_market_day = (
            _ef_now.weekday() < 5 and
            (
                _ef_t < _efdtime(10, 0) or                                      # ก่อนเปิด
                (_efdtime(10, 0) <= _ef_t <= _efdtime(12, 30)) or               # เช้า
                (_efdtime(12, 30) < _ef_t < _efdtime(14, 30)) or               # พัก
                (_efdtime(14, 30) <= _ef_t <= _efdtime(16, 30))                 # บ่าย
            )
        )
        _ef_end_date  = (_ef_now.date() - _eftd(days=1)) if _ef_market_day else _ef_now.date()
        _ef_end_str   = _ef_end_date.strftime('%Y-%m-%d')
        _ef_start_str = (_ef_end_date - _eftd(days=600)).strftime('%Y-%m-%d')

        # Retry up to 4 times with exponential backoff (handles yfinance rate limits)
        import time as _eftime, random as _efrnd
        df = None
        for _ef_attempt in range(4):
            try:
                _ef_ticker = yf.Ticker(full_symbol)
                df = _ef_ticker.history(start=_ef_start_str, end=_ef_end_str, interval='1d')
                if df is not None and not df.empty:
                    break
            except Exception as _ef_exc:
                _ef_msg = str(_ef_exc).lower()
                if 'rate' in _ef_msg or 'too many' in _ef_msg or '429' in _ef_msg:
                    _wait = 2 ** _ef_attempt + _efrnd.uniform(0, 1)
                    _eftime.sleep(_wait)
                    continue
                raise
            if _ef_attempt < 3:
                _eftime.sleep(1.5 * (_ef_attempt + 1))

        if df is None or df.empty:
            messages.error(request, f"ไม่พบข้อมูลสำหรับ {symbol}")
            if market == 'US':
                return redirect('stocks:us_precision_scanner')
            return redirect('stocks:momentum_scanner')

        # แก้ไข MultiIndex columns (ถ้ามี)
        if isinstance(df.columns, pd.MultiIndex):
            df.columns = df.columns.droplevel(1)

        # คำนวณ Supply & Demand Zone ด้วย v2 (ใช้ข้อมูลชุดเดียวกับ Precision Scanner)
        sd_zone = find_supply_demand_zones_v2(df)

        # คำนวณ EMA บน df เต็ม (600 วัน) ก่อน เพื่อให้ EMA200 warm-up ครบ
        # ถ้าคำนวณบน subset 120 วัน → EMA200 จะเป็น NaN ทั้งหมด
        import pandas_ta as ta
        df['EMA50']  = ta.ema(df['Close'], length=50)
        df['EMA200'] = ta.ema(df['Close'], length=200)

        # เตรียมข้อมูลกราฟ 120 วันล่าสุด (slice หลังจากคำนวณ EMA แล้ว)
        history_subset = df.tail(120).copy()
        chart_labels = [d.strftime('%Y-%m-%d') for d in history_subset.index]
        chart_values = [round(float(v), 2) for v in history_subset['Close'].values]

        ema50_vals  = [round(float(v), 2) if pd.notna(v) else None for v in history_subset['EMA50'].values]
        ema200_vals = [round(float(v), 2) if pd.notna(v) else None for v in history_subset['EMA200'].values]

        # OHLCV สำหรับ candlestick chart
        ohlcv_data = [
            {
                'x': chart_labels[i],
                'o': round(float(history_subset['Open'].values[i]), 2),
                'h': round(float(history_subset['High'].values[i]), 2),
                'l': round(float(history_subset['Low'].values[i]), 2),
                'c': round(float(history_subset['Close'].values[i]), 2),
            }
            for i in range(len(chart_labels))
        ]

        # ====== RS Line (Relative Strength vs SET) ======
        rs_line_vals = []
        try:
            set_df = yf.download("^SET", start=_ef_start_str, end=_ef_end_str, interval='1d', progress=False)
            if set_df is not None and not set_df.empty:
                if isinstance(set_df.columns, pd.MultiIndex):
                    set_df.columns = set_df.columns.droplevel(1)
                
                # รวมข้อมูลเพื่อเฉลี่ย ratio (RS Line = Stock / Index)
                combined = pd.concat([df['Close'], set_df['Close']], axis=1, keys=['stock', 'set']).dropna()
                combined['rs'] = combined['stock'] / combined['set']
                
                rs_subset = combined['rs'].tail(len(history_subset))
                rs_line_vals = [float(v) for v in rs_subset.values]
        except Exception as e:
            import logging; logging.getLogger('stocks').error(f"RS Line Error: {e}")

        # แปลงข้อมูล chart เป็น JSON สำหรับ JavaScript
        chart_labels_json = json.dumps(chart_labels)
        chart_values_json = json.dumps(chart_values)
        ema50_vals_json = json.dumps(ema50_vals)
        ema200_vals_json = json.dumps(ema200_vals)
        rs_line_json = json.dumps(rs_line_vals)
        ohlcv_json = json.dumps(ohlcv_data)
        sd_zone_json = json.dumps(sd_zone)

        # ราคาปิดวันสุดท้ายของ historical data (ใช้คำนวณ zone เท่านั้น)
        hist_close = float(df['Close'].iloc[-1])

        # Live price จาก fast_info (ตรงกับ momentum card)
        try:
            _live_fi = yf.Ticker(full_symbol).fast_info
            _live_p  = getattr(_live_fi, 'last_price', None)
            curr_price = float(_live_p) if _live_p else hist_close
        except Exception:
            curr_price = hist_close

        # ถ้า live price ต่างจาก hist_close → เพิ่มเป็น data point สุดท้ายในกราฟ
        if abs(curr_price - hist_close) > 0.001:
            _today_str = _ef_now.strftime('%Y-%m-%d')
            chart_labels.append(_today_str)
            chart_values.append(round(curr_price, 2))
            ema50_vals.append(None)
            ema200_vals.append(None)
            ohlcv_data.append({
                'x': _today_str,
                'o': round(curr_price, 2),
                'h': round(curr_price, 2),
                'l': round(curr_price, 2),
                'c': round(curr_price, 2),
            })
            chart_labels_json = json.dumps(chart_labels)
            chart_values_json = json.dumps(chart_values)
            ema50_vals_json   = json.dumps(ema50_vals)
            ema200_vals_json  = json.dumps(ema200_vals)
            ohlcv_json        = json.dumps(ohlcv_data)

        def _set_tick(price):
            """คืน tick size ของหุ้น SET ตามช่วงราคา"""
            p = float(price or 0)
            if p < 2:    return 0.01
            if p < 5:    return 0.02
            if p < 10:   return 0.05
            if p < 25:   return 0.10
            if p < 100:  return 0.25
            if p < 200:  return 0.50
            return 1.00

        # คำนวณ zone_proximity และ zone_status
        ef_zone_prox   = None
        ef_zone_ticks  = None   # จำนวน ticks ที่ห่างจาก zone
        ef_zone_baht   = None   # ระยะห่างในหน่วยบาท
        ef_zone_status = 'above'   # 'at_tp' | 'in_zone' | 'broke' | 'above'
        if sd_zone and sd_zone.get('start'):
            dz_top  = float(sd_zone['start'])
            dz_bot  = float(sd_zone.get('end') or 0)
            target  = float(sd_zone.get('target') or 0)
            tick    = _set_tick(curr_price)
            if target > 0 and curr_price >= target:
                # ราคาถึงหรือเกิน target แล้ว - Take Profit zone
                gap = round(curr_price - target, 4)
                ef_zone_prox  = round((gap / target) * 100, 1)
                ef_zone_ticks = round(gap / tick)
                ef_zone_baht  = round(gap, 2)
                ef_zone_status = 'at_tp'
            elif dz_bot > 0 and curr_price < dz_bot:
                # ราคาหลุดต่ำกว่า zone (ทะลุ SL)
                gap = round(dz_bot - curr_price, 4)
                ef_zone_prox  = round((gap / dz_bot) * 100, 1)
                ef_zone_ticks = round(gap / tick)
                ef_zone_baht  = round(gap, 2)
                ef_zone_status = 'broke'
            elif curr_price <= dz_top:
                # ราคาอยู่ใน demand zone
                ef_zone_prox  = 0.0
                ef_zone_ticks = 0
                ef_zone_baht  = 0.0
                ef_zone_status = 'in_zone'
            else:
                # ราคาอยู่เหนือ zone
                gap = round(curr_price - dz_top, 4)
                ef_zone_prox  = round((gap / dz_top) * 100, 1)
                ef_zone_ticks = round(gap / tick)
                ef_zone_baht  = round(gap, 2)
                ef_zone_status = 'above'

        currency = '$' if market == 'US' else '฿'
        context = {
            'symbol': symbol,
            'full_symbol': full_symbol,
            'market': market,
            'currency': currency,
            'sd_zone': sd_zone,
            'sd_zone_json': sd_zone_json,
            'curr_price': round(curr_price, 2),
            'zone_proximity': ef_zone_prox,
            'zone_ticks': ef_zone_ticks,
            'zone_baht': ef_zone_baht,
            'zone_status': ef_zone_status,
            'scan_end_date': _ef_end_str,
            'chart_labels': chart_labels_json,
            'chart_values': chart_values_json,
            'ema50_vals': ema50_vals_json,
            'ema200_vals': ema200_vals_json,
            'rs_line_vals': rs_line_json,
            'ohlcv_data': ohlcv_json,
            'title': f"Sniper Entry: {symbol}"
        }
        return render(request, 'stocks/entry_finder.html', context)
    except Exception as e:
        messages.error(request, f"Error finding zones for {symbol}: {str(e)}")
        # Redirect back to the referring page, fallback to US or SET scanner by symbol suffix
        referer = request.META.get('HTTP_REFERER', '')
        if 'momentum/us' in referer:
            return redirect('stocks:us_momentum_scanner')
        if 'us-precision' in referer or 'us-sepa' in referer:
            return redirect('stocks:us_precision_scanner')
        if referer:
            from django.http import HttpResponseRedirect
            return HttpResponseRedirect(referer)
        # Fallback: US scanner if symbol has no .BK suffix, else SET
        if symbol.endswith('.BK') or '.' not in symbol and len(symbol) <= 5:
            return redirect('stocks:momentum_scanner')
        return redirect('stocks:us_momentum_scanner')

# ====== Signup - สมัครสมาชิกใหม่ ======

@login_required
@require_POST
def clear_scan_data(request):
    """
    ลบผลสแกนของ user ออกจาก database
    POST parameter: scanner = momentum_set | momentum_us | precision_set | precision_us |
                               multifactor | cup_handle | us_sepa | us_value
    """
    from .models import (
        MomentumCandidate, MultiFactorCandidate, PrecisionScanCandidate,
        ValueScanCandidate, CupHandleCandidate, USSepaCandidate,
    )
    from django.contrib import messages as _msg

    scanner = request.POST.get('scanner', '')
    next_url = request.POST.get('next', '/')

    _map = {
        'momentum_set':  (MomentumCandidate,       {'user': request.user, 'market': 'SET'}),
        'momentum_us':   (MomentumCandidate,        {'user': request.user, 'market': 'US'}),
        'precision_set': (PrecisionScanCandidate,   {'user': request.user, 'market': 'SET'}),
        'precision_us':  (PrecisionScanCandidate,   {'user': request.user, 'market': 'US'}),
        'multifactor':   (MultiFactorCandidate,     {'user': request.user, 'market': 'SET'}),
        'us_multifactor':(MultiFactorCandidate,     {'user': request.user, 'market': 'US'}),
        'cup_handle':    (CupHandleCandidate,        {'user': request.user, 'market': 'SET'}),
        'us_cup_handle': (CupHandleCandidate,        {'user': request.user, 'market': 'US'}),
        'us_sepa':       (USSepaCandidate,           {'user': request.user}),
        'us_value':      (ValueScanCandidate,        {'user': request.user}),
    }

    if scanner in _map:
        model_cls, filters = _map[scanner]
        deleted_count, _ = model_cls.objects.filter(**filters).delete()
        _msg.success(request, f'ลบข้อมูลสแกน {deleted_count} รายการเรียบร้อยแล้ว')
    else:
        _msg.error(request, 'ไม่พบประเภท scanner ที่ระบุ')

    return redirect(next_url)


def signup(request):
    """
    หน้าสมัครสมาชิก ใช้ Django built-in UserCreationForm
    เมื่อสมัครสำเร็จจะ login อัตโนมัติและ redirect ไปหน้า dashboard
    """
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Login อัตโนมัติหลังจากสมัครสำเร็จ
            login(request, user)
            messages.success(request, f"ยินดีต้อนรับคุณ {user.username}! ระบบของคุณพร้อมใช้งานแล้ว")
            return redirect('stocks:dashboard')
    else:
        form = UserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})


# ====== Multi-Factor Scoring Scanner ======

@login_required
def multi_factor_scanner(request):
    """
    สแกนหุ้นด้วยระบบ Multi-Factor Super Score
    รวม 4 ปัจจัย: Momentum(40) + Volume/Flow(30) + Sentiment AI(20) + Fundamental(10)
    """
    # ====== SCAN STATUS POLL (AJAX) ======
    if request.GET.get('scan_status') == '1':
        from django.core.cache import cache
        from django.http import JsonResponse as _JsonResponse
        key = f'multifactor_scan_{request.user.id}'
        status = cache.get(key, {'state': 'idle'})
        # เมื่อ done ถูก poll แล้ว ให้ reset เป็น idle ทันที ป้องกัน reload loop
        if status.get('state') == 'done':
            cache.delete(key)
        return _JsonResponse(status)

    # ====== SCAN (POST - ทำ background เพื่อไม่ให้ nginx timeout) ======
    if request.method == "POST" and request.POST.get('action') == 'scan':
        from django.core.cache import cache
        import threading

        user_id  = request.user.id
        cache_key = f'multifactor_scan_{user_id}'

        # ป้องกันการสแกนซ้ำถ้ายังรันอยู่
        current = cache.get(cache_key, {})
        if current.get('state') == 'running':
            return redirect('stocks:multi_factor_scanner')

        cache.set(cache_key, {'state': 'running', 'progress': 0, 'total': 0}, timeout=600)

        def _run_scan(user_id, cache_key):
            import django
            django.setup()
            import pandas_ta as ta
            from concurrent.futures import ThreadPoolExecutor, as_completed
            from django.contrib.auth import get_user_model
            from django.core.cache import cache as _cache

            User = get_user_model()
            user = User.objects.get(pk=user_id)

            try:
                scan_symbols = ScannableSymbol.objects.filter(
                    is_active=True, market='SET'
                ).values_list('symbol', flat=True).distinct()
                if not scan_symbols:
                    refresh_all_thai_symbols()
                    scan_symbols = ScannableSymbol.objects.filter(
                        is_active=True, market='SET'
                    ).values_list('symbol', flat=True).distinct()

                # Delete any existing SET records for these symbols
                MultiFactorCandidate.objects.filter(user=user, symbol__in=sym_list).delete()
                # deduplicate while preserving order
                seen = set()
                sym_list = [s for s in scan_symbols if not (s in seen or seen.add(s))]

                _cache.set(cache_key, {'state': 'running', 'progress': 0, 'total': len(sym_list)}, timeout=600)

                # โหลด sector cache จาก DB ครั้งเดียว
                sector_cache = {
                    s.symbol: s.sector
                    for s in ScannableSymbol.objects.filter(
                        is_active=True, market='SET'
                    ).only('symbol', 'sector')
                }

                # Phase 1: Batch download
                tickers_str = " ".join(f"{s}.BK" for s in sym_list)
                try:
                    batch_df = yf.download(
                        tickers_str, period="1y", interval="1d",
                        group_by='ticker', progress=False, threads=True,
                    )
                except Exception as e:
                    print(f"[MultiFactorScan] Batch download failed: {e}")
                    batch_df = None

                def _get_df(symbol):
                    sym_bk = f"{symbol}.BK"
                    df = None
                    if batch_df is not None and not batch_df.empty:
                        try:
                            df = batch_df[sym_bk].copy() if len(sym_list) > 1 else batch_df.copy()
                        except Exception:
                            df = None
                    if df is None or (hasattr(df, 'empty') and df.empty):
                        try:
                            df = yf.download(sym_bk, period="1y", interval="1d", progress=False)
                        except Exception:
                            return None
                    return df

                def process_one(symbol):
                    try:
                        df = _get_df(symbol)
                        if df is None or df.empty:
                            return None
                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = df.columns.droplevel(1)
                        df = df.dropna(subset=['Close', 'High'])
                        if len(df) < 60:
                            return None
                        df = df.copy()
                        df['EMA50']  = ta.ema(df['Close'], length=50)
                        df['EMA200'] = ta.ema(df['Close'], length=200)
                        df['RSI']    = ta.rsi(df['Close'], length=14)
                        adx_df = ta.adx(df['High'], df['Low'], df['Close'], length=14)
                        if adx_df is not None and not adx_df.empty:
                            df = pd.concat([df, adx_df], axis=1)
                        df['MFI']  = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                        df['RVOL'] = df['Volume'] / df['Volume'].rolling(20).mean()
                        last    = df.iloc[-1]
                        price   = float(df['Close'].iloc[-1])
                        rsi     = float(last.get('RSI')    or 0) if pd.notna(last.get('RSI'))    else 0
                        adx_val = float(last.get('ADX_14') or 0) if 'ADX_14' in df.columns and pd.notna(last.get('ADX_14')) else 0
                        mfi_val = float(last.get('MFI')    or 0) if pd.notna(last.get('MFI'))    else 0
                        rvol    = float(last.get('RVOL')   or 1) if pd.notna(last.get('RVOL'))   else 1.0
                        ema50   = float(last.get('EMA50')  or 0) if pd.notna(last.get('EMA50'))  else 0
                        ema200  = float(last.get('EMA200') or 0) if pd.notna(last.get('EMA200')) else 0
                        above_ema200 = bool(price > ema200) if ema200 else False
                        above_ema50  = bool(price > ema50)  if ema50  else False
                        mom = 0
                        if above_ema200:                        mom += 15
                        if above_ema50:                         mom += 5
                        if 55 <= rsi <= 72:                     mom += 15
                        elif 45 <= rsi < 55 or 72 < rsi <= 80: mom += 7
                        if adx_val >= 30:                       mom += 5
                        vol = 0
                        if rvol >= 3.0:     vol += 15
                        elif rvol >= 2.0:   vol += 12
                        elif rvol >= 1.5:   vol += 8
                        elif rvol >= 1.0:   vol += 4
                        if mfi_val >= 70:   vol += 15
                        elif mfi_val >= 60: vol += 10
                        elif mfi_val >= 50: vol += 5
                        sector = sector_cache.get(symbol, 'Unknown')
                        vol_score = min(vol, 30)
                        return dict(
                            symbol=symbol, sector=sector, price=round(price, 2),
                            market='SET',
                            momentum_score=mom, volume_score=vol_score,
                            sentiment_score=0, fundamental_score=0,
                            super_score=mom + vol_score,
                            rsi=round(rsi, 2), adx=round(adx_val, 2),
                            mfi=round(mfi_val, 2), rvol=round(rvol, 2),
                            eps_growth=0.0, rev_growth=0.0,
                            above_ema200=above_ema200, above_ema50=above_ema50,
                        )
                    except Exception as e:
                        print(f"[MultiFactorScan] {symbol}: {e}")
                        return None

                # Phase 2: รันขนาน + update progress
                raw_results = []
                seen_symbols = set()
                done = 0
                with ThreadPoolExecutor(max_workers=12) as executor:
                    futures = {executor.submit(process_one, s): s for s in sym_list}
                    for future in as_completed(futures):
                        r = future.result()
                        if r and r['symbol'] not in seen_symbols:
                            raw_results.append(r)
                            seen_symbols.add(r['symbol'])
                        done += 1
                        _cache.set(cache_key, {'state': 'running', 'progress': done, 'total': len(sym_list)}, timeout=600)

                # Phase 3: Atomic delete+create to guarantee no duplicates
                from django.db import transaction
                with transaction.atomic():
                    MultiFactorCandidate.objects.filter(
                        user=user, symbol__in=[r['symbol'] for r in raw_results]
                    ).delete()
                    MultiFactorCandidate.objects.bulk_create([
                        MultiFactorCandidate(user=user, **r) for r in raw_results
                    ])
                _cache.set(cache_key, {'state': 'done', 'count': len(raw_results)}, timeout=300)
            except Exception as e:
                import traceback
                print(f"[MultiFactorScan] CRITICAL ERROR: {e}\n{traceback.format_exc()}")
                _cache.set(cache_key, {'state': 'done', 'error': str(e)}, timeout=300)

        # เปิด background thread แล้ว return ทันที - ไม่ block nginx
        t = threading.Thread(target=_run_scan, args=(user_id, cache_key), daemon=True)
        t.start()
        return redirect('stocks:multi_factor_scanner')

    # ====== AI SENTIMENT (batch) ======
    if request.GET.get('sentiment') == 'true':
        candidates_qs = MultiFactorCandidate.objects.filter(user=request.user, market='SET').order_by('-super_score')[:30]
        symbols_list  = [c.symbol for c in candidates_qs]
        if symbols_list:
            try:
                client     = genai.Client(api_key=settings.GEMINI_API_KEY)
                prompt = f"""วิเคราะห์ข่าวและ Sentiment ล่าสุดสำหรับหุ้นไทยเหล่านี้ในตลาด SET:
{', '.join(symbols_list)}

ตอบเป็น JSON array เท่านั้น ไม่มีข้อความอื่นนอก array:
[{{"symbol":"PTT","score":15,"label":"บวก","reason":"สรุปเหตุผล 1 ประโยค"}}]

score: 0-20 (20=บวกมาก, 10=กลาง, 0=ลบมาก)
label: "บวก" หรือ "กลาง" หรือ "ลบ"
reason: ภาษาไทย ไม่เกิน 60 ตัวอักษร"""
                resp = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                import json, re
                raw = resp.text.strip()
                raw = re.sub(r'^```(?:json)?\s*', '', raw)
                raw = re.sub(r'\s*```$', '', raw)
                data = json.loads(raw)
                for item in data:
                    sym   = item.get('symbol','').upper()
                    score = int(item.get('score', 0))
                    label = item.get('label', '')
                    reason= item.get('reason', '')
                    MultiFactorCandidate.objects.filter(
                        user=request.user, symbol=sym, market='SET'
                    ).update(
                        sentiment_score=score,
                        sentiment_label=label,
                        sentiment_reason=reason,
                    )
                # recalculate super_score for updated records
                for c in MultiFactorCandidate.objects.filter(user=request.user, market='SET'):
                    c.super_score = c.momentum_score + c.volume_score + c.sentiment_score + c.fundamental_score
                    c.save(update_fields=['super_score'])
            except Exception as e:
                messages.warning(request, f"AI Sentiment Error: {e}")
        return redirect('stocks:multi_factor_scanner')

    # ====== Sort & Render ======
    sort_by = request.GET.get('sort', 'super')
    valid_sorts = {
        'super': '-super_score',
        'momentum': '-momentum_score',
        'volume': '-volume_score',
        'sentiment': '-sentiment_score',
        'fundamental': '-fundamental_score',
        'rsi': '-rsi',
        'rvol': '-rvol',
        'symbol': 'symbol',
    }
    order_field = valid_sorts.get(sort_by, '-super_score')
    candidates  = MultiFactorCandidate.objects.filter(user=request.user, market='SET').order_by(order_field)
    last_scan   = candidates.first()

    context = {
        'candidates':    candidates,
        'current_sort':  sort_by,
        'has_scanned':   candidates.exists(),
        'scanned_at':    last_scan.scanned_at if last_scan else None,
        'has_sentiment': candidates.filter(sentiment_score__gt=0).exists(),
    }
    return render(request, 'stocks/multi_factor.html', context)


@login_required
def us_multi_factor_scanner(request):
    """
    US Multi-Factor Super Score Scanner
    Same 4-factor logic as SET scanner but for US stocks (Nasdaq/S&P 500 universe)
    """
    # ====== SCAN STATUS POLL (AJAX) ======
    if request.GET.get('scan_status') == '1':
        from django.core.cache import cache
        from django.http import JsonResponse as _JsonResponse
        key = f'us_multifactor_scan_{request.user.id}'
        status = cache.get(key, {'state': 'idle'})
        if status.get('state') == 'done':
            cache.delete(key)
        return _JsonResponse(status)

    # ====== SCAN (POST) ======
    if request.method == "POST" and request.POST.get('action') == 'scan':
        from django.core.cache import cache
        import threading

        user_id   = request.user.id
        cache_key = f'us_multifactor_scan_{user_id}'

        current = cache.get(cache_key, {})
        if current.get('state') == 'running':
            return redirect('stocks:us_multi_factor_scanner')

        cache.set(cache_key, {'state': 'running', 'progress': 0, 'total': 0}, timeout=600)

        def _run_scan(user_id, cache_key):
            import django
            django.setup()
            import pandas_ta as ta
            from concurrent.futures import ThreadPoolExecutor, as_completed
            from django.contrib.auth import get_user_model
            from django.core.cache import cache as _cache

            User = get_user_model()
            user = User.objects.get(pk=user_id)

            try:
                # deduplicate while preserving order
                _excl = {'SPY', 'QQQ', 'IWM'}
                _seen = set()
                sym_list = [s for s in _US_MOMENTUM_SYMBOLS
                            if s not in _excl and not (s in _seen or _seen.add(s))]
                # Delete any existing US records (market='US') AND any stale records
                # with wrong market value (e.g. market='SET' from before migration)
                MultiFactorCandidate.objects.filter(user=user, symbol__in=sym_list).delete()

                _cache.set(cache_key, {'state': 'running', 'progress': 0, 'total': len(sym_list)}, timeout=600)

                # Phase 1: Batch download
                tickers_str = " ".join(sym_list)
                try:
                    batch_df = yf.download(
                        tickers_str, period="1y", interval="1d",
                        group_by='ticker', progress=False, threads=True,
                    )
                except Exception as e:
                    print(f"[USMultiFactorScan] Batch download failed: {e}")
                    batch_df = None

                def _get_df(symbol):
                    df = None
                    if batch_df is not None and not batch_df.empty:
                        try:
                            df = batch_df[symbol].copy() if len(sym_list) > 1 else batch_df.copy()
                        except Exception:
                            df = None
                    if df is None or (hasattr(df, 'empty') and df.empty):
                        try:
                            df = yf.download(symbol, period="1y", interval="1d", progress=False)
                        except Exception:
                            return None
                    return df

                def process_one(symbol):
                    try:
                        df = _get_df(symbol)
                        if df is None or df.empty:
                            return None
                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = df.columns.droplevel(1)
                        df = df.dropna(subset=['Close', 'High'])
                        if len(df) < 60:
                            return None
                        df = df.copy()
                        df['EMA50']  = ta.ema(df['Close'], length=50)
                        df['EMA200'] = ta.ema(df['Close'], length=200)
                        df['RSI']    = ta.rsi(df['Close'], length=14)
                        adx_df = ta.adx(df['High'], df['Low'], df['Close'], length=14)
                        if adx_df is not None and not adx_df.empty:
                            df = pd.concat([df, adx_df], axis=1)
                        df['MFI']  = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                        df['RVOL'] = df['Volume'] / df['Volume'].rolling(20).mean()
                        last    = df.iloc[-1]
                        price   = float(df['Close'].iloc[-1])
                        rsi     = float(last.get('RSI')    or 0) if pd.notna(last.get('RSI'))    else 0
                        adx_val = float(last.get('ADX_14') or 0) if 'ADX_14' in df.columns and pd.notna(last.get('ADX_14')) else 0
                        mfi_val = float(last.get('MFI')    or 0) if pd.notna(last.get('MFI'))    else 0
                        rvol    = float(last.get('RVOL')   or 1) if pd.notna(last.get('RVOL'))   else 1.0
                        ema50   = float(last.get('EMA50')  or 0) if pd.notna(last.get('EMA50'))  else 0
                        ema200  = float(last.get('EMA200') or 0) if pd.notna(last.get('EMA200')) else 0
                        above_ema200 = bool(price > ema200) if ema200 else False
                        above_ema50  = bool(price > ema50)  if ema50  else False
                        # Use static sector map - avoid per-symbol API call (193 calls = very slow)
                        sector = _US_SECTOR_MAP.get(symbol, 'Unknown')
                        mom = 0
                        if above_ema200:                        mom += 15
                        if above_ema50:                         mom += 5
                        if 55 <= rsi <= 72:                     mom += 15
                        elif 45 <= rsi < 55 or 72 < rsi <= 80: mom += 7
                        if adx_val >= 30:                       mom += 5
                        vol = 0
                        if rvol >= 3.0:     vol += 15
                        elif rvol >= 2.0:   vol += 12
                        elif rvol >= 1.5:   vol += 8
                        elif rvol >= 1.0:   vol += 4
                        if mfi_val >= 70:   vol += 15
                        elif mfi_val >= 60: vol += 10
                        elif mfi_val >= 50: vol += 5
                        vol_score = min(vol, 30)
                        return dict(
                            symbol=symbol, sector=sector, price=round(price, 2),
                            market='US',
                            momentum_score=mom, volume_score=vol_score,
                            sentiment_score=0, fundamental_score=0,
                            super_score=mom + vol_score,
                            rsi=round(rsi, 2), adx=round(adx_val, 2),
                            mfi=round(mfi_val, 2), rvol=round(rvol, 2),
                            eps_growth=0.0, rev_growth=0.0,
                            above_ema200=above_ema200, above_ema50=above_ema50,
                        )
                    except Exception as e:
                        print(f"[USMultiFactorScan] {symbol}: {e}")
                        return None

                raw_results = []
                seen_symbols = set()
                done = 0
                with ThreadPoolExecutor(max_workers=10) as executor:
                    futures = {executor.submit(process_one, s): s for s in sym_list}
                    for future in as_completed(futures):
                        r = future.result()
                        if r and r['symbol'] not in seen_symbols:
                            raw_results.append(r)
                            seen_symbols.add(r['symbol'])
                        done += 1
                        _cache.set(cache_key, {'state': 'running', 'progress': done, 'total': len(sym_list)}, timeout=600)

                # Final atomic delete+create to guarantee no duplicates
                from django.db import transaction
                with transaction.atomic():
                    MultiFactorCandidate.objects.filter(
                        user=user, symbol__in=[r['symbol'] for r in raw_results]
                    ).delete()
                    MultiFactorCandidate.objects.bulk_create([
                        MultiFactorCandidate(user=user, **r) for r in raw_results
                    ])
                _cache.set(cache_key, {'state': 'done', 'count': len(raw_results)}, timeout=300)
            except Exception as e:
                import traceback
                print(f"[USMultiFactorScan] CRITICAL ERROR: {e}\n{traceback.format_exc()}")
                _cache.set(cache_key, {'state': 'done', 'error': str(e)}, timeout=300)

        t = threading.Thread(target=_run_scan, args=(user_id, cache_key), daemon=True)
        t.start()
        return redirect('stocks:us_multi_factor_scanner')

    # ====== AI SENTIMENT (batch) ======
    if request.GET.get('sentiment') == 'true':
        candidates_qs = MultiFactorCandidate.objects.filter(user=request.user, market='US').order_by('-super_score')[:30]
        symbols_list  = [c.symbol for c in candidates_qs]
        if symbols_list:
            try:
                client = genai.Client(api_key=settings.GEMINI_API_KEY)
                prompt = f"""Analyze the latest news sentiment for these US stocks:
{', '.join(symbols_list)}

Reply with a JSON array only, no other text:
[{{"symbol":"AAPL","score":15,"label":"Positive","reason":"One-sentence summary"}}]

score: 0-20 (20=very positive, 10=neutral, 0=very negative)
label: "Positive" or "Neutral" or "Negative"
reason: English, max 80 characters"""
                resp = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                import json, re
                raw = resp.text.strip()
                raw = re.sub(r'^```(?:json)?\s*', '', raw)
                raw = re.sub(r'\s*```$', '', raw)
                data = json.loads(raw)
                for item in data:
                    sym   = item.get('symbol','').upper()
                    score = int(item.get('score', 0))
                    label = item.get('label', '')
                    reason= item.get('reason', '')
                    MultiFactorCandidate.objects.filter(
                        user=request.user, symbol=sym, market='US'
                    ).update(
                        sentiment_score=score,
                        sentiment_label=label,
                        sentiment_reason=reason,
                    )
                for c in MultiFactorCandidate.objects.filter(user=request.user, market='US'):
                    c.super_score = c.momentum_score + c.volume_score + c.sentiment_score + c.fundamental_score
                    c.save(update_fields=['super_score'])
            except Exception as e:
                messages.warning(request, f"AI Sentiment Error: {e}")
        return redirect('stocks:us_multi_factor_scanner')

    # ====== Sort & Render ======
    sort_by = request.GET.get('sort', 'super')
    valid_sorts = {
        'super': '-super_score',
        'momentum': '-momentum_score',
        'volume': '-volume_score',
        'sentiment': '-sentiment_score',
        'fundamental': '-fundamental_score',
        'rsi': '-rsi',
        'rvol': '-rvol',
        'symbol': 'symbol',
    }
    order_field = valid_sorts.get(sort_by, '-super_score')
    candidates  = MultiFactorCandidate.objects.filter(user=request.user, market='US').order_by(order_field)
    last_scan   = candidates.first()

    context = {
        'candidates':    candidates,
        'current_sort':  sort_by,
        'has_scanned':   candidates.exists(),
        'scanned_at':    last_scan.scanned_at if last_scan else None,
        'has_sentiment': candidates.filter(sentiment_score__gt=0).exists(),
        'total_symbols': len([s for s in _US_MOMENTUM_SYMBOLS if s not in ('SPY', 'QQQ', 'IWM')]),
    }
    return render(request, 'stocks/us_multi_factor.html', context)

@login_required
def realized_pl_report(request):
    """
    รายงานกำไรขาดทุนสะสมที่เกิดขึ้นจริง (Realized P/L)
    พร้อมตัวกรอง รายวัน รายเดือน รายปี และช่วงเวลา
    แปลงกำไรหุ้น US → เงินบาท สำหรับใช้คำนวณทศางค์
    """
    import json
    from collections import defaultdict

    usd_thb = _get_usd_thb()

    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    group_by = request.GET.get('group_by', 'month')

    sold_stocks = list(SoldStock.objects.filter(user=request.user).order_by('sold_at'))

    if start_date:
        from datetime import datetime as _dt
        sold_stocks = [s for s in sold_stocks if s.sold_at.date() >= _dt.strptime(start_date, '%Y-%m-%d').date()]
    if end_date:
        from datetime import datetime as _dt2
        sold_stocks = [s for s in sold_stocks if s.sold_at.date() <= _dt2.strptime(end_date, '%Y-%m-%d').date()]

    # Annotate each record with is_us + pl_thb
    # ใช้ s.market ก่อน (บันทึกตอนขาย) - ถ้าไม่มีหรือ default SET ให้ fallback _is_us_symbol
    us_set = _build_us_symbol_set(request.user)
    for s in sold_stocks:
        if s.market and s.market != MarketType.SET:
            s.is_us = s.market == MarketType.US
        else:
            s.is_us = _is_us_symbol(s.symbol, us_set)
        
        # ลอจิกใหม่: ถ้ามี profit_loss_thb (ที่บันทึกตอนขาย) ให้ใช้ค่านั้นเลย
        # ถ้าเป็น 0 หรือเป็นข้อมูลเก่า ให้คำนวณจาก usd_thb ปัจจุบัน (fallback)
        if hasattr(s, 'profit_loss_thb') and s.profit_loss_thb != 0:
            s.pl_thb = float(s.profit_loss_thb)
        else:
            s.pl_thb = float(s.profit_loss) * usd_thb if s.is_us else float(s.profit_loss)

    summary_dict = defaultdict(lambda: {'items': [], 'total_pl': 0, 'total_pl_thb': 0})
    chart_labels = []
    chart_data = []
    running_pl = 0

    for s in sold_stocks:
        if group_by == 'day':
            key = s.sold_at.strftime('%Y-%m-%d')
        elif group_by == 'year':
            key = s.sold_at.strftime('%Y')
        else:
            key = s.sold_at.strftime('%Y-%m')

        summary_dict[key]['items'].append(s)
        summary_dict[key]['total_pl'] += float(s.profit_loss)
        summary_dict[key]['total_pl_thb'] += s.pl_thb

        running_pl += s.pl_thb
        chart_labels.append(s.sold_at.strftime('%Y-%m-%d %H:%M'))
        chart_data.append(round(running_pl, 2))

    summary_list = []
    for k in sorted(summary_dict.keys(), reverse=True):
        summary_list.append({
            'period': k,
            'items': summary_dict[k]['items'],
            'total_pl': summary_dict[k]['total_pl'],
            'total_pl_thb': summary_dict[k]['total_pl_thb'],
            'count': len(summary_dict[k]['items'])
        })

    context = {
        'summary_list': summary_list,
        'sold_stocks': sold_stocks,
        'chart_labels': json.dumps(chart_labels),
        'chart_data': json.dumps(chart_data),
        'start_date': start_date,
        'end_date': end_date,
        'group_by': group_by,
        'title': 'Realized P/L Report',
        'usd_thb': round(usd_thb, 2),
        'total_pl_thb': sum(s.pl_thb for s in sold_stocks),
    }
    return render(request, 'stocks/realized_pl_report.html', context)


# ======================================================================
# PRECISION SCAN AI ANALYSIS
# ======================================================================

@login_required
def precision_scan_ai_analysis(request):
    import json as _json_lib
    from django.http import JsonResponse
    if request.method != "POST":
        return JsonResponse({"error": "POST only"}, status=405)
    try:
        body = _json_lib.loads(request.body)
        scan_data = body.get("scan_data", {})
    except Exception:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    api_key = getattr(settings, "GEMINI_API_KEY", None)
    if not api_key:
        return JsonResponse({"error": "ไม่พบ GEMINI_API_KEY"}, status=500)

    qualified = scan_data.get("qualified_stocks", [])
    top_buy = scan_data.get("top_buy_stocks", [])
    scan_date = scan_data.get("scan_date", "ไม่ระบุ")
    total_passed = scan_data.get("total_passed", 0)
    top_sectors = scan_data.get("top_sectors", [])

    def _fmt_stock(s):
        lines = [
            "  - {sym}: ราคา {price} | BUY {buy} | RS {rs}".format(
                sym=s["symbol"], price=s["price"], buy=s["buy_score"], rs=s["rs_rating"]),
            "    RSI {rsi} | ADX {adx} | RVOL {rvol}x {dir} | RR 1:{rr}".format(
                rsi=s["rsi"], adx=s["adx"], rvol=s["rvol"],
                dir="Bull ▲" if s.get("rvol_bullish") else "Bear ▼",
                rr=s.get("risk_reward_ratio", "-")),
            "    Zone {z}% | {sec} | RelMom3m {rm}%".format(
                z=s.get("zone_proximity", "-"), sec=s.get("sector", "-"),
                rm=s.get("rel_momentum_3m", 0)),
            "    Signals: {m}{e}{h}{b}{a}".format(
                m="MACD✕ " if s.get("macd_crossover") else "",
                e="EMA↑ " if s.get("ema20_rising") else "",
                h="HH/HL " if s.get("hh_hl_structure") else "",
                b="BB Squeeze " if s.get("bb_squeeze") else "",
                a="EMA20 Aligned " if s.get("ema20_aligned") else ""),
        ]
        if s.get("top_reasons"):
            lines.append("    เหตุผล: {r}".format(r=", ".join(s["top_reasons"])))
        return "\n".join(lines)

    q_text = "\n".join([_fmt_stock(s) for s in qualified]) if qualified else "  (ไม่มีหุ้นผ่านเกณฑ์ครบ)"
    b_text = "\n".join([_fmt_stock(s) for s in top_buy]) if top_buy else "  (ไม่มีข้อมูล)"
    sec_text = ", ".join(["{n}({c})".format(n=s["name"], c=s["count"]) for s in top_sectors]) if top_sectors else "ไม่ระบุ"

    prompt = (
        "คุณคือผู้เชี่ยวชาญด้านหุ้น SET"
        " ที่ใช้ Precision Momentum (สไตล์ Mark Minervini + William O'Neil)\n"
        "เชี่ยวชาญ Stage Analysis, RS Rating, Supply/Demand Zone,"
        " RVOL Bull/Bear, MACD Crossover, EMA Alignment, Trend Following (HH/HL)\n\n"
        "ผล Precision Scan วันที่ {sd} (ผ่านเกณฑ์ {tp} ตัว)"
        " | Leading Sectors: {sec}\n\n"
        "=== หุ้นผ่านเกณฑ์ครบทุกข้อ (Fully Qualified) ===\n{q}\n\n"
        "=== Top หุ้นแนะนำซื้อ (BUY Score) ===\n{b}\n\n"
        "วิเคราะห์ในหัวข้อต่อไปนี้:\n\n"
        "## 1. ✨ ภาพรวมตลาดวันนี้\n"
        "- บรรยากาศตลาด (Bullish/Mixed/Bearish), Sector นำตลาด\n\n"
        "## 2. ✅ วิเคราะห์หุ้นผ่านเกณฑ์ครบทุกข้อ\n"
        "- แต่ละตัว: จุดเด่น, ความเสี่ยง, โอกาสวิ่ง, ลำดับน่าสนใจ\n\n"
        "## 3. 🏆 วิเคราะห์ Top BUY Score\n"
        "- ตัวที่น่าสนใจที่สุด และทำไม | ระวังอะไร\n\n"
        "## 4. ⚡ กลยุทธ์แนะนำ\n"
        "- ลำดับซื้อ: ตัวไหนก่อน หลัง หรือรอ | Entry Zone ที่เหมาะสม\n\n"
        "## 5. ⚠ ข้อควรระวัง\n"
        "- RSI/RVOL/Zone น่าเป็นห่วง, Stop Loss discipline\n\n"
        "ตอบภาษาไทย กระชับ เป็นมืออาชีพ"
        " จัดรูปแบบ Markdown"
        " เน้น Actionable Insights"
    ).format(sd=scan_date, tp=total_passed, sec=sec_text, q=q_text, b=b_text)

    try:
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        if not response.text:
            return JsonResponse({"error": "AI ไม่ตอบกลับ"}, status=500)
        return JsonResponse({"status": "success", "analysis": response.text})
    except Exception as e:
        err = str(e)
        if "API_KEY_INVALID" in err:
            return JsonResponse({"error": "GEMINI_API_KEY ไม่ถูกต้อง"}, status=500)
        return JsonResponse({"error": "Gemini error: {}".format(err)}, status=500)


# ====== tithe_report - รายงานทศางค์ 10% จากกำไรหุ้นรายเดือน ======

@login_required
def tithe_report(request):
    """
    แสดงกำไร/ขาดทุนรายเดือนจากการขายหุ้น
    คำนวณทศางค์ 10% จากเดือนที่มีกำไร พร้อม track การจ่าย
    แปลงกำไรหุ้น US → เงินบาท ด้วยอัตราแลกเปลี่ยนปัจจุบัน
    """
    from collections import defaultdict
    import calendar

    # ── Exchange rate ──
    usd_thb = _get_usd_thb()
    usd_thb_d = Decimal(str(round(usd_thb, 4)))

    # ── Aggregate per month with USD→THB conversion ──
    sold_stocks = SoldStock.objects.filter(user=request.user).order_by('sold_at')
    us_set = _build_us_symbol_set(request.user)
    monthly_raw = defaultdict(Decimal)
    for s in sold_stocks:
        if s.market and s.market != MarketType.SET:
            is_us = s.market == MarketType.US
        else:
            is_us = _is_us_symbol(s.symbol, us_set)
        
        # ลอจิกใหม่: ใช้ profit_loss_thb ที่บันทึกไว้ ณ วันที่ขาย (ถ้ามี)
        if hasattr(s, 'profit_loss_thb') and s.profit_loss_thb != 0:
            pl_thb = Decimal(str(s.profit_loss_thb))
        else:
            pl_raw = Decimal(str(s.profit_loss or 0))
            pl_thb = pl_raw * usd_thb_d if is_us else pl_raw
            
        key = (s.sold_at.year, s.sold_at.month)
        monthly_raw[key] += pl_thb

    tithe_map = {
        (t.year, t.month): t
        for t in TitheRecord.objects.filter(user=request.user)
    }

    months = []
    total_profit = Decimal('0')
    total_tithe_owed = Decimal('0')
    total_tithe_paid = Decimal('0')

    for (yr, mo) in sorted(monthly_raw.keys(), reverse=True):
        pl = monthly_raw[(yr, mo)].quantize(Decimal('0.01'))
        tithe = (pl * Decimal('0.10')).quantize(Decimal('0.01')) if pl > 0 else Decimal('0')

        rec = tithe_map.get((yr, mo))
        is_paid = rec.is_paid if rec else False
        paid_at = rec.paid_at if rec else None

        if pl > 0:
            total_profit += pl
            total_tithe_owed += tithe
            if is_paid:
                total_tithe_paid += tithe

        months.append({
            'year': yr,
            'month': mo,
            'month_name': calendar.month_abbr[mo],
            'pl': pl,
            'tithe': tithe,
            'is_paid': is_paid,
            'paid_at': paid_at,
        })

    # Build chart data (chronological order = reversed from newest-first list)
    chart_months = list(reversed(months))
    chart_data = json.dumps({
        'labels':  [f"{m['month_name']} {m['year']}" for m in chart_months],
        'profit':  [float(m['pl'])    for m in chart_months],
        'tithe':   [float(m['tithe']) for m in chart_months],
    }, ensure_ascii=False)

    context = {
        'months': months,
        'total_profit': total_profit,
        'total_tithe_owed': total_tithe_owed,
        'total_tithe_paid': total_tithe_paid,
        'total_tithe_remaining': total_tithe_owed - total_tithe_paid,
        'chart_json': chart_data,
        'usd_thb': round(usd_thb, 2),
    }
    return render(request, 'stocks/tithe_report.html', context)


@login_required
def tithe_mark_paid(request):
    """Toggle paid/unpaid status for a tithe month via AJAX POST."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    from django.utils import timezone

    try:
        yr = int(request.POST.get('year', 0))
        mo = int(request.POST.get('month', 0))
    except (TypeError, ValueError):
        return JsonResponse({'error': 'Invalid year/month'}, status=400)

    rec, _ = TitheRecord.objects.get_or_create(
        user=request.user, year=yr, month=mo,
        defaults={'is_paid': False}
    )
    rec.is_paid = not rec.is_paid
    rec.paid_at = timezone.now() if rec.is_paid else None
    rec.save()

    return JsonResponse({
        'is_paid': rec.is_paid,
        'paid_at': rec.paid_at.strftime('%d %b %Y %H:%M') if rec.paid_at else None,
    })


# ======================================================================
# US MOMENTUM SCANNER (no-DB) - same logic as SET scanner, US stocks
# ======================================================================

# Shared US symbol universe - ~220 symbols (Nasdaq + S&P 500 leaders)
_US_SECTOR_MAP = {
    # Technology
    "AAPL":"Technology","MSFT":"Technology","NVDA":"Technology","GOOGL":"Technology",
    "GOOG":"Technology","AMZN":"Technology","META":"Technology","TSLA":"Technology",
    "AVGO":"Technology","ORCL":"Technology","AMD":"Technology","ARM":"Technology",
    "DELL":"Technology","HPE":"Technology","WDC":"Technology","TSM":"Technology",
    "QCOM":"Technology","INTC":"Technology","MU":"Technology","AMAT":"Technology",
    "LRCX":"Technology","KLAC":"Technology","MRVL":"Technology","ON":"Technology",
    "TXN":"Technology","SMCI":"Technology","ASML":"Technology","NXPI":"Technology",
    "MPWR":"Technology","WOLF":"Technology","CRM":"Technology","NOW":"Technology",
    "SNOW":"Technology","PLTR":"Technology","PANW":"Technology","CRWD":"Technology",
    "ZS":"Technology","NET":"Technology","DDOG":"Technology","MDB":"Technology",
    "ADBE":"Technology","INTU":"Technology","ANSS":"Technology","CDNS":"Technology",
    "SNPS":"Technology","FTNT":"Technology","OKTA":"Technology","HUBS":"Technology",
    "TWLO":"Technology","TTD":"Technology","BILL":"Technology","GTLB":"Technology",
    "DOCN":"Technology","ZM":"Technology",
    # Financial Services
    "JPM":"Financial Services","BAC":"Financial Services","WFC":"Financial Services",
    "GS":"Financial Services","MS":"Financial Services","BLK":"Financial Services",
    "SCHW":"Financial Services","AXP":"Financial Services","V":"Financial Services",
    "MA":"Financial Services","COF":"Financial Services","DFS":"Financial Services",
    "SYF":"Financial Services","USB":"Financial Services","TFC":"Financial Services",
    "KEY":"Financial Services","RF":"Financial Services","FITB":"Financial Services",
    "COIN":"Financial Services","SQ":"Financial Services","PYPL":"Financial Services",
    # Insurance
    "CB":"Financial Services","PGR":"Financial Services","ALL":"Financial Services",
    "TRV":"Financial Services","MET":"Financial Services","PRU":"Financial Services",
    # Healthcare
    "UNH":"Healthcare","LLY":"Healthcare","JNJ":"Healthcare","ABBV":"Healthcare",
    "MRK":"Healthcare","PFE":"Healthcare","ABT":"Healthcare","TMO":"Healthcare",
    "AMGN":"Healthcare","ISRG":"Healthcare","DXCM":"Healthcare","IDXX":"Healthcare",
    "ILMN":"Healthcare","MRNA":"Healthcare","REGN":"Healthcare","VRTX":"Healthcare",
    "BIIB":"Healthcare","GILD":"Healthcare","BMY":"Healthcare","CVS":"Healthcare",
    "CI":"Healthcare","HUM":"Healthcare","MDT":"Healthcare","SYK":"Healthcare",
    "BSX":"Healthcare","EW":"Healthcare",
    # Consumer Staples
    "COST":"Consumer Staples","WMT":"Consumer Staples","TGT":"Consumer Staples",
    "KR":"Consumer Staples","PG":"Consumer Staples","KO":"Consumer Staples",
    "PEP":"Consumer Staples","CL":"Consumer Staples","MDLZ":"Consumer Staples","MO":"Consumer Staples",
    # Consumer Discretionary
    "HD":"Consumer Discretionary","LOW":"Consumer Discretionary","NKE":"Consumer Discretionary",
    "LULU":"Consumer Discretionary","DECK":"Consumer Discretionary","ONON":"Consumer Discretionary",
    "RH":"Consumer Discretionary","SBUX":"Consumer Discretionary","MCD":"Consumer Discretionary",
    "YUM":"Consumer Discretionary","CMG":"Consumer Discretionary","DPZ":"Consumer Discretionary",
    "NFLX":"Consumer Discretionary","ABNB":"Consumer Discretionary","UBER":"Consumer Discretionary",
    "DASH":"Consumer Discretionary","ETSY":"Consumer Discretionary","EBAY":"Consumer Discretionary",
    "BABA":"Consumer Discretionary","JD":"Consumer Discretionary","PDD":"Consumer Discretionary",
    "SPOT":"Consumer Discretionary","RBLX":"Consumer Discretionary",
    # Energy
    "XOM":"Energy","CVX":"Energy","COP":"Energy","EOG":"Energy","SLB":"Energy",
    "PSX":"Energy","MPC":"Energy","VLO":"Energy","OXY":"Energy","HAL":"Energy",
    "DVN":"Energy","FANG":"Energy","APA":"Energy","MRO":"Energy","WMB":"Energy","KMI":"Energy",
    # Industrials
    "CAT":"Industrials","DE":"Industrials","HON":"Industrials","GE":"Industrials",
    "RTX":"Industrials","LMT":"Industrials","BA":"Industrials","UPS":"Industrials",
    "FDX":"Industrials","CSX":"Industrials","ITW":"Industrials","EMR":"Industrials",
    "ETN":"Industrials","PH":"Industrials","ROK":"Industrials","AME":"Industrials",
    "TT":"Industrials","DHR":"Industrials","NOC":"Industrials","GD":"Industrials",
    # Real Estate
    "AMT":"Real Estate","PLD":"Real Estate","CCI":"Real Estate","EQIX":"Real Estate",
    "O":"Real Estate","WELL":"Real Estate","VICI":"Real Estate","PSA":"Real Estate",
    # Utilities
    "NEE":"Utilities","DUK":"Utilities","SO":"Utilities","AEP":"Utilities","EXC":"Utilities",
    # Communication
    "PINS":"Communication","SNAP":"Communication",
    # Conglomerate
    "BRK-B":"Conglomerate","MSTR":"Technology",
}

_US_MOMENTUM_SYMBOLS = [
    # ── Mega-cap Tech ──────────────────────────────────────────────────
    "AAPL", "MSFT", "NVDA", "GOOGL", "GOOG", "AMZN", "META", "TSLA", "AVGO", "ORCL",
    "AMD", "ARM", "DELL", "HPE", "WDC",
    # ── Semiconductor ──────────────────────────────────────────────────
    "TSM", "QCOM", "INTC", "MU", "AMAT", "LRCX", "KLAC", "MRVL", "ON", "TXN",
    "SMCI", "ASML", "NXPI", "MPWR", "WOLF",
    # ── Cloud / Software ───────────────────────────────────────────────
    "CRM", "NOW", "SNOW", "PLTR", "PANW", "CRWD", "ZS", "NET", "DDOG", "MDB",
    "ADBE", "INTU", "ANSS", "CDNS", "SNPS", "FTNT", "OKTA", "HUBS", "TWLO",
    "TTD", "BILL", "GTLB", "DOCN", "ZM",
    # ── Financials ─────────────────────────────────────────────────────
    "JPM", "BAC", "WFC", "GS", "MS", "BLK", "SCHW", "AXP", "V", "MA",
    "COF", "DFS", "SYF", "USB", "TFC", "KEY", "RF", "FITB",
    "CB", "PGR", "ALL", "TRV", "MET", "PRU",
    # ── Healthcare / Biotech ───────────────────────────────────────────
    "UNH", "LLY", "JNJ", "ABBV", "MRK", "PFE", "ABT", "TMO", "AMGN", "ISRG",
    "DXCM", "IDXX", "ILMN", "MRNA", "REGN", "VRTX", "BIIB", "GILD", "BMY",
    "CVS", "CI", "HUM", "MDT", "SYK", "BSX", "EW",
    # ── Consumer Staples ───────────────────────────────────────────────
    "COST", "WMT", "TGT", "KR", "PG", "KO", "PEP", "CL", "MDLZ", "MO",
    # ── Consumer Discretionary ─────────────────────────────────────────
    "HD", "LOW", "NKE", "LULU", "DECK", "ONON", "RH",
    "SBUX", "MCD", "YUM", "CMG", "DPZ",
    "NFLX", "ABNB", "UBER", "DASH", "ETSY", "EBAY",
    "BABA", "JD", "PDD",
    # ── Energy ─────────────────────────────────────────────────────────
    "XOM", "CVX", "COP", "EOG", "SLB", "PSX", "MPC", "VLO", "OXY", "HAL",
    "DVN", "FANG", "APA", "MRO", "WMB", "KMI",
    # ── Industrials / Aerospace ────────────────────────────────────────
    "CAT", "DE", "HON", "GE", "RTX", "LMT", "BA", "UPS", "FDX", "CSX",
    "ITW", "EMR", "ETN", "PH", "ROK", "AME", "TT", "DHR", "NOC", "GD",
    # ── FinTech / Payments / Crypto ────────────────────────────────────
    "SPOT", "RBLX", "COIN", "SQ", "PYPL", "MSTR",
    # ── REIT / Utilities ───────────────────────────────────────────────
    "AMT", "PLD", "CCI", "EQIX", "O", "WELL", "VICI", "PSA",
    "NEE", "DUK", "SO", "AEP", "EXC",
    # ── Conglomerate / Other ───────────────────────────────────────────
    "BRK-B", "PINS", "SNAP",
    # ── Benchmarks (RS calc only - filtered from results) ──────────────
    "SPY", "QQQ", "IWM",
]

@login_required
def us_momentum_scanner(request):
    """
    US Momentum Scanner - scans ~200 Nasdaq/S&P 500 stocks using Minervini Trend Template.
    Results saved to MomentumCandidate (market='US') - same as SET scanner.
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'us_momentum_scan_{user_id}'

    # ── AJAX scan progress poll ───────────────────────────────────────
    if request.GET.get('scan_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        if st.get('state') == 'done':
            _cp.delete(cache_key)
        return _JR(st)

    # ── Trigger background scan ───────────────────────────────────────
    if request.GET.get('scan') == 'true' or request.method == 'POST':
        cur = _cp.get(cache_key, {})
        if cur.get('state') != 'running':
            scan_syms = [s for s in _US_MOMENTUM_SYMBOLS if s not in ('SPY', 'QQQ', 'IWM')]
            _cp.set(cache_key, {
                'state': 'running', 'progress': 0,
                'total': len(scan_syms), 'phase': 'เตรียมข้อมูล…'
            }, timeout=900)

            def _run_us_bg(uid, ckey, sym_list):
                try:
                    import pandas_ta as _ta
                    import pandas as _pd
                    import concurrent.futures as _cf
                    from datetime import timedelta as _td
                    from django.core.cache import cache as _c
                    from django.utils import timezone as _tz
                    from django.contrib.auth import get_user_model
                    from .utils import find_supply_demand_zones_v2
                    from .models import MomentumCandidate as _MCM
                    _User = get_user_model()
                    _user = _User.objects.get(pk=uid)

                    import pytz as _pytz
                    from datetime import datetime as _dt
                    _ny_tz      = _pytz.timezone('America/New_York')
                    _now_ny     = _dt.now(_ny_tz)
                    _end_date   = _now_ny.date()
                    _end_str    = (_end_date + _td(days=1)).strftime('%Y-%m-%d')
                    _start_str  = (_end_date - _td(days=600)).strftime('%Y-%m-%d')
                    _spy_start  = (_end_date - _td(days=400)).strftime('%Y-%m-%d')

                    total = len(sym_list)

                    # ── Step 1: SPY benchmark returns ─────────────────
                    _c.set(ckey, {'state': 'running', 'progress': 0, 'total': total, 'phase': 'โหลด SPY Benchmark…'}, timeout=900)
                    spy_1m = spy_3m = 0.0
                    try:
                        spy_df = yf.download("SPY", start=_spy_start, end=_end_str, interval="1d", progress=False)
                        if spy_df is not None and not spy_df.empty:
                            if isinstance(spy_df.columns, _pd.MultiIndex):
                                spy_df.columns = spy_df.columns.droplevel(1)
                            sc = spy_df['Close'].dropna()
                            if len(sc) >= 66:
                                spy_1m = float((sc.iloc[-1] - sc.iloc[-22]) / sc.iloc[-22] * 100)
                                spy_3m = float((sc.iloc[-1] - sc.iloc[-66]) / sc.iloc[-66] * 100)
                    except Exception:
                        pass

                    # ── Step 2: RS Rating (4-quarter weighted) ────────
                    _c.set(ckey, {'state': 'running', 'progress': 0, 'total': total, 'phase': 'คำนวณ RS Rating…'}, timeout=900)

                    def _fetch_rs(sym):
                        import time as _t, random as _rnd
                        _t.sleep(_rnd.uniform(0.05, 0.3))
                        for _att in range(3):
                            try:
                                d = yf.Ticker(sym).history(start=_start_str, end=_end_str, interval="1d")
                                if d is None or d.empty:
                                    return sym, None
                                if isinstance(d.columns, _pd.MultiIndex):
                                    d.columns = d.columns.droplevel(1)
                                cl = d['Close'].dropna()
                                if len(cl) >= 252:
                                    r = float(
                                        (cl.iloc[-1] - cl.iloc[-64]) / abs(cl.iloc[-64]) * 0.4 +
                                        (cl.iloc[-64] - cl.iloc[-127]) / abs(cl.iloc[-127]) * 0.2 +
                                        (cl.iloc[-127] - cl.iloc[-190]) / abs(cl.iloc[-190]) * 0.2 +
                                        (cl.iloc[-190] - cl.iloc[-253]) / abs(cl.iloc[-253]) * 0.2
                                    ) * 100
                                    return sym, r
                                return sym, None
                            except Exception as _e:
                                _emsg = str(_e).lower()
                                if 'rate' in _emsg or 'too many' in _emsg or '429' in _emsg:
                                    _t.sleep(2 ** _att + _rnd.uniform(0, 1))
                                    continue
                                return sym, None
                        return sym, None

                    rs_raw = {}
                    _rs_done = 0
                    with _cf.ThreadPoolExecutor(max_workers=6) as ex:
                        futs = {ex.submit(_fetch_rs, s): s for s in sym_list}
                        for f in _cf.as_completed(futs, timeout=180):
                            try:
                                s, r = f.result()
                                if r is not None:
                                    rs_raw[s] = r
                            except Exception:
                                pass
                            _rs_done += 1
                            if _rs_done % 5 == 0 or _rs_done == total:
                                _c.set(ckey, {
                                    'state': 'running',
                                    'progress': _rs_done,
                                    'total': total,
                                    'phase': f'RS Ratings ({_rs_done}/{total})…',
                                }, timeout=900)

                    rs_map = {}
                    if rs_raw:
                        ser = _pd.Series(rs_raw)
                        rs_map = (ser.rank(pct=True) * 99).clip(0, 99).astype(int).to_dict()

                    # ── Step 3: Technical scan ─────────────────────────
                    _c.set(ckey, {'state': 'running', 'progress': 0, 'total': total, 'phase': 'Technical Scan…'}, timeout=900)

                    results    = []
                    lock       = _th.Lock()
                    done_count = [0]

                    def _scan_one(symbol):
                        import time as _time
                        import random as _random
                        # Small random jitter to spread requests
                        _time.sleep(_random.uniform(0.1, 0.6))
                        try:
                            # Retry up to 3 times with exponential backoff on rate-limit
                            df = _pd.DataFrame()
                            for _attempt in range(3):
                                try:
                                    ticker = yf.Ticker(symbol)
                                    df = ticker.history(start=_start_str, end=_end_str, interval="1d")
                                    if df is not None and not df.empty:
                                        break
                                except Exception as _e:
                                    _emsg = str(_e).lower()
                                    if 'rate' in _emsg or 'too many' in _emsg or '429' in _emsg:
                                        _time.sleep(2 ** _attempt + _random.uniform(0, 1))
                                        continue
                                    break
                            if df is None or df.empty:
                                return None
                            if isinstance(df.columns, _pd.MultiIndex):
                                df.columns = df.columns.droplevel(1)
                            df = df.dropna(subset=['Close', 'High'])
                            if len(df) < 150:
                                return None

                            # Liquidity filter - ≥500K avg daily volume
                            av20 = float(df['Volume'].tail(20).mean())
                            if av20 < 500_000:
                                return None

                            # Indicators
                            df['EMA50']  = _ta.ema(df['Close'], length=50)
                            df['EMA150'] = _ta.ema(df['Close'], length=150)
                            df['EMA200'] = _ta.ema(df['Close'], length=200)
                            df['RSI']    = _ta.rsi(df['Close'], length=14)
                            adx_df = _ta.adx(df['High'], df['Low'], df['Close'], length=14)
                            if adx_df is not None and not adx_df.empty:
                                df = _pd.concat([df, adx_df], axis=1)
                            mfi_s = _ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                            if mfi_s is not None:
                                df['MFI'] = mfi_s
                            vol_avg = df['Volume'].rolling(20).mean()
                            df['RVOL'] = df['Volume'] / vol_avg

                            last         = df.iloc[-1]
                            current_p    = float(last['Close'])
                            year_high    = float(df['High'].tail(252).max())
                            ema200       = float(last.get('EMA200', 0) or 0)
                            ema50        = float(last.get('EMA50', 0) or 0)
                            rsi_val      = float(last.get('RSI', 50) or 50)
                            adx_val      = float(last['ADX_14']) if 'ADX_14' in df.columns and _pd.notna(last.get('ADX_14')) else 0
                            mfi_val      = float(last['MFI']) if 'MFI' in df.columns and _pd.notna(last.get('MFI')) else 0
                            rvol_val     = float(last['RVOL']) if _pd.notna(last.get('RVOL')) else 1.0

                            # ── Minervini Trend Template filters ──────
                            if not (current_p > ema200 and current_p >= year_high * 0.65):
                                return None
                            if adx_val < 15:
                                return None

                            rs_v = rs_map.get(symbol, 0)

                            # ── Score (0–100) ─────────────────────────
                            score = 0
                            # Trend alignment (max 35)
                            if current_p > ema200:               score += 15
                            if current_p > ema50:                score += 10
                            if ema50 > ema200:                   score += 10
                            # Momentum (max 20)
                            if 55 < rsi_val < 80:                score += 15
                            elif 45 < rsi_val <= 55:             score += 5
                            # Volume/RVOL (max 20)
                            last_close  = float(df['Close'].iloc[-1])
                            prev_close  = float(df['Close'].iloc[-2]) if len(df) > 1 else last_close
                            is_bullish  = last_close >= prev_close
                            if is_bullish and rvol_val >= 1.5:   score += 20
                            elif is_bullish and rvol_val >= 1.0: score += 12
                            elif rvol_val >= 1.0:                score += 5
                            # Price strength (max 10)
                            if current_p >= year_high * 0.90:    score += 10
                            elif current_p >= year_high * 0.80:  score += 5
                            # RS Rating bonus (max 15)
                            if rs_v >= 85:                       score += 15
                            elif rs_v >= 70:                     score += 8
                            elif rs_v >= 60:                     score += 3
                            score = min(score, 100)

                            # ── MACD crossover ────────────────────────
                            m_cross = False
                            try:
                                md = _ta.macd(df['Close'])
                                if md is not None and not md.empty:
                                    mc = md.columns[0]
                                    ms = next((c for c in md.columns if 'MACDs' in c), None)
                                    if mc and ms:
                                        for i in range(-3, 0):
                                            if md[mc].iloc[i-1] <= md[ms].iloc[i-1] and md[mc].iloc[i] > md[ms].iloc[i]:
                                                m_cross = True
                                                break
                            except Exception:
                                pass

                            # ── BB Squeeze ────────────────────────────
                            bb_sqz = False
                            try:
                                bb = _ta.bbands(df['Close'])
                                if bb is not None and not bb.empty:
                                    bu = next((c for c in bb.columns if 'BBU' in c), None)
                                    bl = next((c for c in bb.columns if 'BBL' in c), None)
                                    bm = next((c for c in bb.columns if 'BBM' in c), None)
                                    if bu and bl and bm:
                                        bw = (bb[bu] - bb[bl]) / bb[bm]
                                        if float(bw.iloc[-1]) <= float(bw.quantile(0.2)):
                                            bb_sqz = True
                            except Exception:
                                pass

                            # ── Stage 2 ───────────────────────────────
                            stage2 = False
                            try:
                                s150 = _ta.sma(df['Close'], length=150)
                                if s150 is not None and not s150.empty:
                                    stage2 = (current_p > float(s150.iloc[-1])) and (float(s150.iloc[-1]) > float(s150.iloc[-20]))
                            except Exception:
                                pass

                            # ── Supply / Demand Zone ──────────────────
                            sd_zone = find_supply_demand_zones_v2(df)
                            dz_s = dz_e = sz_s = rr_v = None
                            prox = 999.0
                            if sd_zone:
                                dz_s = sd_zone.get('start')
                                dz_e = sd_zone.get('end')
                                sz_s = sd_zone.get('target')
                                rr_v = sd_zone.get('rr_ratio')
                                if dz_s:
                                    prox = 0.0 if current_p <= dz_s else round((current_p - dz_s) / dz_s * 100, 2)

                            # ── Relative returns ──────────────────────
                            rel_1m = rel_3m = 0.0
                            cl = df['Close'].dropna()
                            if len(cl) >= 22:
                                rel_1m = round(float((cl.iloc[-1] - cl.iloc[-22]) / cl.iloc[-22] * 100) - spy_1m, 2)
                            if len(cl) >= 66:
                                rel_3m = round(float((cl.iloc[-1] - cl.iloc[-66]) / cl.iloc[-66] * 100) - spy_3m, 2)

                            # ── Sector ────────────────────────────────
                            sector = 'Unknown'
                            try:
                                info = ticker.info or {}
                                sector = info.get('sector', 'Unknown') or 'Unknown'
                            except Exception:
                                pass

                            return {
                                'symbol':            symbol,
                                'price':             round(current_p, 2),
                                'technical_score':   score,
                                'rs_rating':         rs_v,
                                'rsi':               round(rsi_val, 2),
                                'adx':               round(adx_val, 2),
                                'mfi':               round(mfi_val, 2),
                                'rvol':              round(rvol_val, 2),
                                'rvol_bullish':      is_bullish and rvol_val >= 1.0,
                                'demand_zone_start': round(dz_s, 2) if dz_s else None,
                                'demand_zone_end':   round(dz_e, 2) if dz_e else None,
                                'supply_zone_start': round(sz_s, 2) if sz_s else None,
                                'risk_reward_ratio': round(rr_v, 2) if rr_v else None,
                                'zone_proximity':    round(prox, 2),
                                'year_high':         round(year_high, 2),
                                'upside_to_high':    round((year_high - current_p) / current_p * 100, 2),
                                'sector':            sector,
                                'stage2':            stage2,
                                'macd_crossover':    m_cross,
                                'bb_squeeze':        bb_sqz,
                                'rel_1m':            rel_1m,
                                'rel_3m':            rel_3m,
                                # live fields (filled on display)
                                'live_price':       None,
                                'live_change_pct':  None,
                                'live_in_zone':     False,
                                'live_broke_zone':  False,
                                'live_above_tp':    False,
                                'live_near_tp':     False,
                                'live_zone_prox':   999,
                            }
                        except Exception:
                            return None

                    with _cf.ThreadPoolExecutor(max_workers=4) as ex:
                        futs = {ex.submit(_scan_one, s): s for s in sym_list}
                        for idx, f in enumerate(_cf.as_completed(futs, timeout=900)):
                            try:
                                r = f.result()
                                if r:
                                    with lock:
                                        results.append(r)
                            except Exception:
                                pass
                            done_count[0] += 1
                            if done_count[0] % 5 == 0:
                                _c.set(ckey, {
                                    'state': 'running',
                                    'progress': done_count[0],
                                    'total': total,
                                    'phase': f'Scanning… ({done_count[0]}/{total})',
                                }, timeout=900)

                    # ── Save to DB (delete old, bulk create new) ──────
                    _MCM.objects.filter(user=_user, market='US').delete()
                    _bulk = []
                    for r in results:
                        _bulk.append(_MCM(
                            user=_user,
                            market='US',
                            symbol=r['symbol'],
                            symbol_bk=r['symbol'],   # US has no .BK suffix
                            sector=r.get('sector', 'Unknown'),
                            price=r['price'],
                            rsi=r['rsi'],
                            adx=r['adx'],
                            mfi=r['mfi'],
                            rvol=r['rvol'],
                            rvol_bullish=r.get('rvol_bullish', False),
                            eps_growth=0.0,
                            rev_growth=0.0,
                            technical_score=r['technical_score'],
                            rs_rating=r.get('rs_rating', 0),
                            stage2=r.get('stage2', False),
                            macd_crossover=r.get('macd_crossover', False),
                            bb_squeeze=r.get('bb_squeeze', False),
                            rel_1m=r.get('rel_1m', 0.0),
                            rel_3m=r.get('rel_3m', 0.0),
                            demand_zone_start=r.get('demand_zone_start'),
                            demand_zone_end=r.get('demand_zone_end'),
                            supply_zone_start=r.get('supply_zone_start'),
                            supply_zone_end=None,
                            stop_loss=None,
                            risk_reward_ratio=r.get('risk_reward_ratio'),
                            year_high=r.get('year_high', 0.0),
                            upside_to_high=r.get('upside_to_high', 0.0),
                            zone_proximity=r.get('zone_proximity', 999.0),
                        ))
                    if _bulk:
                        _MCM.objects.bulk_create(_bulk, ignore_conflicts=True)

                    _c.set(ckey, {'state': 'done'}, timeout=300)

                except Exception as exc:
                    from django.core.cache import cache as _c2
                    _c2.set(ckey, {'state': 'done'}, timeout=300)

            _th.Thread(
                target=_run_us_bg,
                args=(user_id, cache_key, scan_syms),
                daemon=True,
            ).start()

        from django.shortcuts import redirect as _redir
        return _redir('stocks:us_momentum_scanner')

    # ── Display results from DB ───────────────────────────────────────
    sort_by = request.GET.get('sort', 'score')
    valid_sorts = {
        'symbol':    'symbol',
        'score':     '-technical_score',
        'rs':        '-rs_rating',
        'rsi':       '-rsi',
        'rvol':      '-rvol',
        'adx':       '-adx',
        'mfi':       '-mfi',
        'price':     '-price',
        'gap':       'upside_to_high',
        'prox':      'zone_proximity',
        'round_rr':  '-risk_reward_ratio',
        'rel1m':     '-rel_1m',
    }
    order_field = valid_sorts.get(sort_by, '-technical_score')

    _scan_state = _cp.get(cache_key, {})
    is_scanning = _scan_state.get('state') == 'running'

    db_candidates = (
        MomentumCandidate.objects.filter(user=request.user, market='US')
        .order_by(order_field)
        if not is_scanning else MomentumCandidate.objects.none()
    )

    # Attach live price + zone status (same pattern as SET scanner)
    candidate_list = list(db_candidates)
    if candidate_list:
        try:
            import concurrent.futures as _mcf

            def _live_us(sym):
                try:
                    fi = yf.Ticker(sym).fast_info
                    p  = getattr(fi, 'last_price', None)
                    return sym, float(p) if p else None
                except Exception:
                    return sym, None

            live_map = {}
            with _mcf.ThreadPoolExecutor(max_workers=6) as ex:
                for sym, lp in ex.map(_live_us, [c.symbol for c in candidate_list]):
                    if lp:
                        live_map[sym] = lp
        except Exception:
            live_map = {}

        for c in candidate_list:
            lp  = live_map.get(c.symbol)
            c.live_price = lp
            ref  = lp if lp else float(c.price or 0)
            dz_s = float(c.demand_zone_start or 0)
            dz_e = float(c.demand_zone_end   or 0)
            sz_s = float(c.supply_zone_start or 0)
            c.live_in_zone    = dz_s > 0 and dz_e > 0 and dz_e <= ref <= dz_s
            c.live_broke_zone = dz_e > 0 and ref < dz_e
            c.live_above_tp   = sz_s > 0 and ref >= sz_s
            c.live_near_tp    = (
                not c.live_above_tp and sz_s > 0 and dz_s > 0 and
                (sz_s - dz_s) > 0 and (sz_s - ref) / (sz_s - dz_s) * 100 <= 15
            )
            c.live_zone_prox  = (
                0.0 if ref <= dz_s else
                round((ref - dz_s) / dz_s * 100, 1) if dz_s > 0 else 999
            )
            c.live_change_pct = (
                round((lp - float(c.price)) / float(c.price) * 100, 2)
                if lp and float(c.price or 0) > 0 else None
            )

    # ── AI Superperformance filter ────────────────────────────────────
    ai_analysis = None
    if candidate_list and request.GET.get('analyze') == 'true':
        syms = [c.symbol for c in candidate_list[:30]]
        try:
            client = genai.Client(api_key=settings.GEMINI_API_KEY)
            prompt = f"""You are a top US momentum stock analyst using Minervini/O'Neil methodology.

From this list of US stocks that passed the Trend Template filter:
{', '.join(syms)}

Analyze current news, earnings momentum, sector rotation, and market sentiment to identify
the top 5-7 "Superperformance" candidates most likely to make significant moves in the next 2-6 weeks.

Write in Thai language, markdown format:
- For each pick: symbol, why it's the leader (RS Rating, Stage 2, Catalyst), key risk
- Focus on Relative Strength leaders and stocks near breakout pivots
- Note any Earnings dates or Fed events to watch
- No intro, no outro"""
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            ai_analysis = response.text
            if ai_analysis and ai_analysis.startswith("```"):
                ai_analysis = ai_analysis.split('\n', 1)[-1].rsplit('```', 1)[0].strip()
        except Exception as e:
            ai_analysis = f"AI Error: {str(e)}"

    last_scan = MomentumCandidate.objects.filter(user=request.user, market='US').order_by('-scanned_at').first()
    scanned_at = last_scan.scanned_at if last_scan else None

    return render(request, 'stocks/us_momentum.html', {
        'title':        'US Momentum Scanner',
        'candidates':   candidate_list,
        'ai_analysis':  ai_analysis,
        'scanned_at':   scanned_at,
        'current_sort': sort_by,
        'is_scanning':  is_scanning,
        'has_scanned':  db_candidates.exists() if not is_scanning else True,
    })


@login_required
def us_momentum_quick_analysis(request, symbol):
    """
    Quick CrewAI multi-agent analysis for a US momentum stock.
    AJAX endpoint - returns JSON, displayed in a modal.
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'us_mq_analysis_{user_id}_{symbol}'

    # Poll
    if request.GET.get('mq_status') == '1':
        return _JR(_cp.get(cache_key, {'state': 'idle'}))

    cached = _cp.get(cache_key)
    if cached:
        if cached.get('state') == 'running':
            return _JR({'state': 'running'})
        if cached.get('state') == 'done':
            _cp.delete(cache_key)
            return _JR({'state': 'done', 'result': cached.get('result', '')})

    # Get scan data from cache results
    scan_data = {}
    # Load scan data from DB (market='US')
    try:
        cand = MomentumCandidate.objects.filter(user=request.user, symbol=symbol, market='US').first()
        if cand:
            scan_data = {
                'symbol':            cand.symbol,
                'price':             float(cand.price),
                'technical_score':   cand.technical_score,
                'rs_rating':         cand.rs_rating,
                'rsi':               float(cand.rsi),
                'adx':               float(cand.adx),
                'mfi':               float(cand.mfi),
                'rvol':              float(cand.rvol),
                'rvol_bullish':      cand.rvol_bullish,
                'demand_zone_start': float(cand.demand_zone_start) if cand.demand_zone_start else None,
                'demand_zone_end':   float(cand.demand_zone_end) if cand.demand_zone_end else None,
                'supply_zone_start': float(cand.supply_zone_start) if cand.supply_zone_start else None,
                'risk_reward_ratio': float(cand.risk_reward_ratio) if cand.risk_reward_ratio else None,
                'zone_proximity':    float(cand.zone_proximity),
                'sector':            cand.sector or 'Unknown',
                'year_high':         float(cand.year_high),
                'upside_to_high':    float(cand.upside_to_high),
                'stage2':            cand.stage2,
                'macd_crossover':    cand.macd_crossover,
                'bb_squeeze':        cand.bb_squeeze,
                'rel_1m':            float(cand.rel_1m),
                'rel_3m':            float(cand.rel_3m),
            }
    except Exception:
        pass

    # Background analysis
    _cp.set(cache_key, {'state': 'running'}, timeout=600)

    def _run_us_crew(ckey, sym, sd):
        from django.core.cache import cache as _c
        try:
            _c.set(ckey, {'state': 'running', 'phase': 'กำลังวิเคราะห์ด้วย 3 US Expert Agents…'}, timeout=600)
            from .crew_analysis import USMomentumShortTermCrew as _USC
            import concurrent.futures as _cf
            crew = _USC(sym, scan_data=sd)
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(crew.run_analysis)
                try:
                    result = fut.result(timeout=180)
                except _cf.TimeoutError:
                    result = '## Timeout\n\nกรุณาลองใหม่อีกครั้ง'
            _c.set(ckey, {'state': 'done', 'result': result}, timeout=900)
        except Exception as exc:
            from django.core.cache import cache as _c2
            _c2.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=60)

    _th.Thread(target=_run_us_crew, args=(cache_key, symbol, scan_data), daemon=True).start()
    return _JR({'state': 'running'})


@login_required
def us_momentum_crew_page(request, symbol):
    """
    Standalone full-page CrewAI analysis for a US momentum stock.
    Opens in a new tab - does not block the scanner page.
    GET  → render the page (triggers analysis in background)
    GET ?mq_status=1 → AJAX poll (same logic as us_momentum_quick_analysis)
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    # share the same cache key so existing in-progress analysis is reused
    cache_key = f'us_mq_analysis_{user_id}_{symbol}'

    # AJAX poll
    if request.GET.get('mq_status') == '1':
        return _JR(_cp.get(cache_key, {'state': 'idle'}))

    # Load scan data from DB
    scan_data = {}
    try:
        cand = MomentumCandidate.objects.filter(user=request.user, symbol=symbol, market='US').first()
        if cand:
            scan_data = {
                'symbol':            cand.symbol,
                'price':             float(cand.price),
                'technical_score':   cand.technical_score,
                'rs_rating':         cand.rs_rating,
                'rsi':               float(cand.rsi),
                'adx':               float(cand.adx),
                'mfi':               float(cand.mfi),
                'rvol':              float(cand.rvol),
                'rvol_bullish':      cand.rvol_bullish,
                'demand_zone_start': float(cand.demand_zone_start) if cand.demand_zone_start else None,
                'demand_zone_end':   float(cand.demand_zone_end) if cand.demand_zone_end else None,
                'supply_zone_start': float(cand.supply_zone_start) if cand.supply_zone_start else None,
                'risk_reward_ratio': float(cand.risk_reward_ratio) if cand.risk_reward_ratio else None,
                'zone_proximity':    float(cand.zone_proximity),
                'sector':            cand.sector or 'Unknown',
                'year_high':         float(cand.year_high),
                'upside_to_high':    float(cand.upside_to_high),
                'stage2':            cand.stage2,
                'macd_crossover':    cand.macd_crossover,
                'bb_squeeze':        cand.bb_squeeze,
                'rel_1m':            float(cand.rel_1m),
                'rel_3m':            float(cand.rel_3m),
            }
    except Exception:
        pass

    # Trigger background analysis (skip if already running/done)
    existing = _cp.get(cache_key)
    if not existing or existing.get('state') not in ('running', 'done'):
        _cp.set(cache_key, {'state': 'running'}, timeout=600)

        def _run(ckey, sym, sd):
            from django.core.cache import cache as _c
            try:
                _c.set(ckey, {'state': 'running', 'phase': 'กำลังวิเคราะห์ด้วย 3 US Expert Agents…'}, timeout=600)
                from .crew_analysis import USMomentumShortTermCrew as _USC
                import concurrent.futures as _cf
                crew = _USC(sym, scan_data=sd)
                with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                    fut = ex.submit(crew.run_analysis)
                    try:
                        result = fut.result(timeout=180)
                    except _cf.TimeoutError:
                        result = '## Timeout\n\nกรุณาลองใหม่อีกครั้ง'
                _c.set(ckey, {'state': 'done', 'result': result}, timeout=900)
            except Exception as exc:
                from django.core.cache import cache as _c2
                _c2.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=60)

        _th.Thread(target=_run, args=(cache_key, symbol, scan_data), daemon=True).start()

    return render(request, 'stocks/us_momentum_crew_page.html', {
        'symbol':    symbol,
        'scan_data': scan_data,
    })


# ======================================================================
# US PRECISION MOMENTUM SCANNER - Nasdaq & S&P 500
# ======================================================================

def _seed_us_symbols():
    """Seed curated US stock universe (~220 symbols, Nasdaq & S&P 500) into ScannableSymbol."""
    US_SYMBOLS = [
        # ── Mega-cap Tech ──────────────────────────────────────────────
        "AAPL", "MSFT", "NVDA", "GOOGL", "GOOG", "AMZN", "META", "TSLA", "AVGO", "ORCL",
        "AMD", "ARM", "DELL", "HPE", "WDC",
        # ── Semiconductor ──────────────────────────────────────────────
        "TSM", "QCOM", "INTC", "MU", "AMAT", "LRCX", "KLAC", "MRVL", "ON", "TXN",
        "SMCI", "ASML", "NXPI", "MPWR", "WOLF",
        # ── Cloud / Software ──────────────────────────────────────────
        "CRM", "NOW", "SNOW", "PLTR", "PANW", "CRWD", "ZS", "NET", "DDOG", "MDB",
        "ADBE", "INTU", "ANSS", "CDNS", "SNPS", "FTNT", "OKTA", "HUBS", "TWLO",
        "TTD", "BILL", "GTLB", "DOCN", "ZM",
        # ── Financials ────────────────────────────────────────────────
        "JPM", "BAC", "WFC", "GS", "MS", "BLK", "SCHW", "AXP", "V", "MA",
        "COF", "DFS", "SYF", "USB", "TFC", "KEY", "RF", "FITB",
        "CB", "PGR", "ALL", "TRV", "MET", "PRU",
        # ── Healthcare / Biotech ──────────────────────────────────────
        "UNH", "LLY", "JNJ", "ABBV", "MRK", "PFE", "ABT", "TMO", "AMGN", "ISRG",
        "DXCM", "IDXX", "ILMN", "MRNA", "REGN", "VRTX", "BIIB", "GILD", "BMY",
        "CVS", "CI", "HUM", "MDT", "SYK", "BSX", "EW",
        # ── Consumer Staples ──────────────────────────────────────────
        "COST", "WMT", "TGT", "KR", "PG", "KO", "PEP", "CL", "MDLZ", "MO",
        # ── Consumer Discretionary ────────────────────────────────────
        "HD", "LOW", "NKE", "LULU", "DECK", "ONON", "RH",
        "SBUX", "MCD", "YUM", "CMG", "DPZ",
        "NFLX", "ABNB", "UBER", "DASH", "LYFT", "ETSY", "EBAY",
        "BABA", "JD", "PDD",
        # ── Energy ────────────────────────────────────────────────────
        "XOM", "CVX", "COP", "EOG", "SLB", "PSX", "MPC", "VLO", "OXY", "HAL",
        "DVN", "FANG", "APA", "MRO", "WMB", "KMI",
        # ── Industrials / Aerospace ───────────────────────────────────
        "CAT", "DE", "HON", "GE", "RTX", "LMT", "BA", "UPS", "FDX", "CSX",
        "ITW", "EMR", "ETN", "PH", "ROK", "AME", "TT", "DHR", "NOC", "GD",
        # ── FinTech / Payments ────────────────────────────────────────
        "SPOT", "RBLX", "COIN", "SQ", "PYPL", "MSTR",
        # ── REIT / Utilities ──────────────────────────────────────────
        "AMT", "PLD", "CCI", "EQIX", "O", "WELL", "VICI", "PSA",
        "NEE", "DUK", "SO", "AEP", "EXC",
        # ── Conglomerate / Other ──────────────────────────────────────
        "BRK-B", "PINS", "SNAP",
        # ── Benchmarks ────────────────────────────────────────────────
        "SPY", "QQQ", "IWM",
    ]
    for sym in US_SYMBOLS:
        ScannableSymbol.objects.update_or_create(
            symbol=sym, market='US',
            defaults={'index_name': 'Nasdaq+S&P500', 'is_active': True},
        )
    return US_SYMBOLS


@login_required
def us_precision_scanner(request):
    """
    US Precision Momentum Scanner - Nasdaq & S&P 500
    - market='US' filter on all DB queries
    - Background scanning to prevent timeouts
    """
    from .models import PrecisionScanCandidate
    from .utils import analyze_momentum_technical_v2
    from django.utils import timezone as tz
    from yahooquery import Ticker as YQTicker

    # AJAX status polling
    if request.GET.get('scan_status') == '1':
        from django.core.cache import cache as _cp
        from django.http import JsonResponse as _JR
        _key = f'us_precision_scan_{request.user.id}'
        _st = _cp.get(_key, {'state': 'idle'})
        if _st.get('state') == 'done':
            _cp.delete(_key)
        return _JR(_st)

    scan_symbols = list(
        ScannableSymbol.objects.filter(is_active=True, market='US').values_list('symbol', flat=True)
    )
    # If symbols are missing or deactivated (e.g. by Thai refresh bug), re-seed them
    if len(scan_symbols) < 100:
        _seed_us_symbols()
        scan_symbols = list(
            ScannableSymbol.objects.filter(is_active=True, market='US').values_list('symbol', flat=True)
        )

    if request.method == "POST" or request.GET.get('scan') == 'true':
        from django.core.cache import cache as _cache_bg
        import threading
        user_id = request.user.id
        cache_key = f'us_precision_scan_{user_id}'

        _cur = _cache_bg.get(cache_key, {})
        if _cur.get('state') == 'running':
            return redirect('stocks:us_precision_scanner')

        _cache_bg.set(cache_key, {'state': 'running', 'progress': 0, 'total': 0, 'phase': 'เตรียมข้อมูล…'}, timeout=1200)

        def _run_us_scan_bg(uid, ckey, sym_list):
            try:
                import django
                django.setup()
                import yfinance as yf
                import pandas as pd
                import pandas_ta as ta
                import requests
                from datetime import datetime as _dt, timedelta as _td, time as _dtime
                import pytz as _pytz
                import concurrent.futures
                from django.contrib.auth import get_user_model
                from django.core.cache import cache as _cache_inner

                # Create a session with timeout to prevent hanging
                _session = requests.Session()
                _session.mount("https://", requests.adapters.HTTPAdapter(max_retries=2))
                _session.mount("http://", requests.adapters.HTTPAdapter(max_retries=2))
                # Patch yfinance to use this session with a timeout if needed, 
                # but usually passing session to Ticker is enough.
                
                User = get_user_model()
                user = User.objects.get(pk=uid)
                scan_run_time = tz.now()

                _ny_tz = _pytz.timezone('America/New_York')
                _now_ny = _dt.now(_ny_tz)
                
                # yfinance end date is exclusive. To include today's data, we must set end to tomorrow.
                scan_end_date  = _now_ny.date()
                scan_end_str   = (scan_end_date + _td(days=1)).strftime('%Y-%m-%d')
                scan_start_str = (scan_end_date - _td(days=600)).strftime('%Y-%m-%d')
                spy_start_str  = (scan_end_date - _td(days=400)).strftime('%Y-%m-%d')

                _cache_inner.set(ckey, {'state': 'running', 'progress': 0, 'total': len(sym_list), 'phase': 'Benchmarks…'}, timeout=1200)

                # Previous symbols
                prev_run = PrecisionScanCandidate.objects.filter(user=user, market='US').order_by('-scan_run').values_list('scan_run', flat=True).distinct().first()
                prev_symbols = set(PrecisionScanCandidate.objects.filter(user=user, market='US', scan_run=prev_run).values_list('symbol', flat=True)) if prev_run else set()

                # SPY
                spy_1m = spy_3m = 0.0
                try:
                    spy_df = yf.download("SPY", start=spy_start_str, end=scan_end_str, interval="1d", progress=False)
                    if spy_df is not None and not spy_df.empty:
                        if isinstance(spy_df.columns, pd.MultiIndex): spy_df.columns = spy_df.columns.droplevel(1)
                        c = spy_df['Close'].dropna()
                        if len(c) >= 66:
                            spy_1m = float((c.iloc[-1] - c.iloc[-22])/c.iloc[-22]*100)
                            spy_3m = float((c.iloc[-1] - c.iloc[-66])/c.iloc[-66]*100)
                except: pass

                # RS Rating (Bulk Fetch for speed)
                total_syms = len(sym_list)
                _cache_inner.set(ckey, {'state': 'running', 'progress': 0, 'total': total_syms, 'phase': 'Fetching RS Data (Bulk)…'}, timeout=1200)
                rs_returns_all = {}
                try:
                    from yahooquery import Ticker as _TQ
                    chunk_size = 80
                    for i in range(0, len(sym_list), chunk_size):
                        chunk = sym_list[i : i + chunk_size]
                        _cache_inner.set(ckey, {'state': 'running', 'progress': 5 + int((i/total_syms)*15), 'total': total_syms, 'phase': f'Fetching RS Data ({i//chunk_size + 1})…'}, timeout=900)
                        try:
                            tq = _TQ(chunk)
                            tq_hist = tq.history(start=scan_start_str, end=scan_end_str, interval="1d")
                            if tq_hist is not None and not tq_hist.empty:
                                if isinstance(tq_hist.index, pd.MultiIndex):
                                    for symbol in chunk:
                                        try:
                                            if symbol in tq_hist.index.get_level_values(0):
                                                _close = tq_hist.loc[symbol]['adjclose'].dropna()
                                                if len(_close) >= 66:
                                                    ret = float((_close.iloc[-1] - _close.iloc[-66]) / abs(_close.iloc[-66]) * 100)
                                                    rs_returns_all[symbol] = ret
                                        except Exception: continue
                        except Exception: pass
                except Exception as e:
                    print(f"Bulk RS Error: {e}")

                # FAILSAFE: If results are empty or too small, force evaluation of a subset
                if len(rs_returns_all) < 10:
                    for s in sym_list[:100]:
                        if s not in rs_returns_all: rs_returns_all[s] = 0.0

                rs_map = {}
                if rs_returns_all:
                    ser = pd.Series(rs_returns_all)
                    rs_map = (ser.rank(pct=True)*99).clip(0,99).astype(int).to_dict()

                # Main Scan
                results_to_process = [s for s in sym_list if rs_map.get(s, 0) >= 60]
                _cache_inner.set(ckey, {'state': 'running', 'progress': 0, 'total': len(results_to_process), 'phase': 'Technical Scan…'}, timeout=1200)
                
                results = []
                def _scan_one(symbol):
                    try:
                        rs_v = rs_map.get(symbol, 0)
                        if rs_v < 60: return None

                        try:
                            ticker_obj = yf.Ticker(symbol)
                            df = ticker_obj.history(start=scan_start_str, end=scan_end_str, interval="1d")

                            if df is None or df.empty:
                                try:
                                    yq = YQTicker(symbol)
                                    df = yq.history(start=scan_start_str, end=scan_end_str, interval="1d")
                                    if isinstance(df, pd.DataFrame) and not df.empty:
                                        df = df.reset_index()
                                        if 'date' in df.columns:
                                            df.set_index('date', inplace=True)
                                        if 'symbol' in df.columns:
                                            df.drop(columns=['symbol'], inplace=True)
                                        df.rename(columns={'open': 'Open', 'high': 'High', 'low': 'Low',
                                                           'close': 'Close', 'volume': 'Volume'}, inplace=True)
                                except Exception:
                                    pass
                        except Exception:
                            df = None

                        if df is None or df.empty: return None
                        if isinstance(df.columns, pd.MultiIndex): df.columns = df.columns.droplevel(1)
                        df = df.dropna(subset=['Close','High'])
                        if len(df) < 200: return None
                        
                        av20 = float(df['Volume'].tail(20).mean())
                        if av20 < 1_000_000: return None
                        
                        rs_v = rs_map.get(symbol, 0)
                        if rs_v < 60: return None
                        
                        # Indicators
                        df['EMA200'] = ta.ema(df['Close'], length=200)
                        df['EMA50']  = ta.ema(df['Close'], length=50)
                        df['RSI']    = ta.rsi(df['Close'], length=14)
                        adx_d = ta.adx(df['High'], df['Low'], df['Close'], length=14)
                        if adx_d is not None and not adx_d.empty: df = pd.concat([df, adx_d], axis=1)
                        df['MFI'] = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)
                        
                        current_p = float(df['Close'].iloc[-1])
                        year_h = float(df['High'].tail(252).max())
                        
                        if current_p < year_h * 0.65: return None
                        adx_v = float(df['ADX_14'].iloc[-1]) if 'ADX_14' in df.columns and pd.notna(df['ADX_14'].iloc[-1]) else 0
                        if adx_v < 15: return None

                        tech = analyze_momentum_technical_v2(df)

                        mfi_val = float(df['MFI'].iloc[-1]) if 'MFI' in df.columns and pd.notna(df['MFI'].iloc[-1]) else 0

                        # MACD (12,26,9)
                        macd_hist_val = None
                        macd_cross_val = False
                        try:
                            macd_df = ta.macd(df['Close'], fast=12, slow=26, signal=9)
                            if macd_df is not None and not macd_df.empty:
                                hist_col = [c for c in macd_df.columns if 'h' in c.lower() or 'hist' in c.lower()]
                                macd_col = [c for c in macd_df.columns if c.lower().startswith('macd_')]
                                sig_col  = [c for c in macd_df.columns if 'macds' in c.lower() or 'signal' in c.lower()]
                                if hist_col: macd_hist_val = float(macd_df[hist_col[0]].iloc[-1]) if pd.notna(macd_df[hist_col[0]].iloc[-1]) else None
                                if macd_col and sig_col:
                                    m_ser = macd_df[macd_col[0]].dropna()
                                    s_ser = macd_df[sig_col[0]].dropna()
                                    if len(m_ser) >= 4 and len(s_ser) >= 4:
                                        for i in range(-3, 0):
                                            if m_ser.iloc[i-1] <= s_ser.iloc[i-1] and m_ser.iloc[i] > s_ser.iloc[i]:
                                                macd_cross_val = True; break
                        except: pass

                        # Bollinger Bands Squeeze
                        bb_squeeze_flag = False
                        try:
                            bb_df = ta.bbands(df['Close'], length=20, std=2)
                            if bb_df is not None and not bb_df.empty:
                                upper_col = [c for c in bb_df.columns if 'BBU' in c or 'upper' in c.lower()]
                                lower_col = [c for c in bb_df.columns if 'BBL' in c or 'lower' in c.lower()]
                                mid_col   = [c for c in bb_df.columns if 'BBM' in c or 'mid' in c.lower()]
                                if upper_col and lower_col and mid_col:
                                    bbu = bb_df[upper_col[0]].dropna()
                                    bbl = bb_df[lower_col[0]].dropna()
                                    bbm = bb_df[mid_col[0]].dropna()
                                    if len(bbu) >= 20:
                                        bw = (bbu - bbl) / bbm
                                        pct20 = bw.quantile(0.20)
                                        if float(bw.iloc[-1]) <= float(pct20): bb_squeeze_flag = True
                        except: pass

                        # Stage 2 (Weinstein)
                        stage2_flag = False
                        try:
                            sma150 = ta.sma(df['Close'], length=150)
                            if sma150 is not None:
                                sma150_clean = sma150.dropna()
                                if len(sma150_clean) >= 20:
                                    sma150_cur = float(sma150_clean.iloc[-1])
                                    sma150_4w  = float(sma150_clean.iloc[-20])
                                    stage2_flag = (current_p > sma150_cur) and (sma150_cur > sma150_4w)
                        except: pass

                        # Pocket Pivot
                        pocket_pivot_flag = False
                        try:
                            if len(df) >= 14:
                                closes = df['Close'].values
                                volumes = df['Volume'].values
                                for _i in [-1, -2]:
                                    if float(closes[_i]) <= float(closes[_i - 1]): continue
                                    _start = len(volumes) + _i - 10
                                    _end   = len(volumes) + _i
                                    if _start < 1: continue
                                    _prior_c = closes[_start:_end]
                                    _prior_v = volumes[_start:_end]
                                    _prior_prev_c = closes[_start - 1:_end - 1]
                                    _down_mask = _prior_c < _prior_prev_c
                                    if not _down_mask.any(): continue
                                    _max_down_vol = float(_prior_v[_down_mask].max())
                                    if float(volumes[_i]) > _max_down_vol and _max_down_vol > 0:
                                        pocket_pivot_flag = True; break
                        except: pass

                        # Volume Dry-Up (VDU)
                        vdu_flag = False
                        try:
                            if len(df) >= 4:
                                _vols = df['Volume'].tail(4).values.astype(float)
                                _avg20 = float(df['Volume'].tail(20).mean())
                                _declining = (_vols[-1] < _vols[-2]) and (_vols[-2] < _vols[-3])
                                _quiet     = _vols[-1] < _avg20 * 0.7
                                vdu_flag   = _declining and _quiet
                        except: pass

                        # Ichimoku Score
                        ichimoku_score_val = 0
                        ichimoku_above_kumo = False
                        ichimoku_tk_cross = False
                        ichimoku_kumo_green = False
                        ichimoku_chikou_ok = False
                        try:
                            if len(df) >= 52:
                                _h9 = df['High'].rolling(9).max(); _l9 = df['Low'].rolling(9).min()
                                _h26 = df['High'].rolling(26).max(); _l26 = df['Low'].rolling(26).min()
                                _h52 = df['High'].rolling(52).max(); _l52 = df['Low'].rolling(52).min()
                                _tenkan = (_h9 + _l9) / 2; _kijun = (_h26 + _l26) / 2
                                _span_a = ((_tenkan + _kijun) / 2).shift(26); _span_b = ((_h52 + _l52) / 2).shift(26)
                                _sa_cur = float(_span_a.iloc[-1]); _sb_cur = float(_span_b.iloc[-1])
                                ichimoku_above_kumo = current_p > max(_sa_cur, _sb_cur) > 0
                                for _i in range(-5, 0):
                                    if _tenkan.iloc[_i-1] <= _kijun.iloc[_i-1] and _tenkan.iloc[_i] > _kijun.iloc[_i]:
                                        ichimoku_tk_cross = True; break
                                ichimoku_kumo_green = _sa_cur > _sb_cur > 0
                                ichimoku_chikou_ok = current_p > float(df['Close'].iloc[-27]) if len(df) >= 27 else False
                                ichimoku_score_val = sum([ichimoku_above_kumo, ichimoku_tk_cross, ichimoku_kumo_green, ichimoku_chikou_ok])
                        except: pass

                        # Price Pattern detection
                        try:
                            pattern_result = detect_price_pattern(df)
                            pattern_name = pattern_result['name']
                            pattern_score = pattern_result['score']
                        except:
                            pattern_name = "None"
                            pattern_score = 0

                        return {
                            'symbol': symbol, 'price': round(current_p, 2),
                            'rsi': round(tech['rsi'], 2), 'adx': round(adx_v, 2), 'mfi': round(mfi_val, 2),
                            'rvol': round(tech['rvol'], 2), 'technical_score': int(tech['score']),
                            'avg_volume_20d': round(av20, 0), 'rvol_bullish': tech['rvol_bullish'],
                            'erc_volume_confirmed': tech.get('erc_volume_confirmed', False),
                            'zone_target_src': tech.get('zone_target_source', '52w'),
                            'entry_strat': tech['sd_zone']['type'] if tech['sd_zone'] else '',
                            'dz_start': tech['sd_zone']['start'] if tech['sd_zone'] else None,
                            'dz_end': tech['sd_zone']['end'] if tech['sd_zone'] else None,
                            'sz_start': tech['sd_zone']['target'] if tech['sd_zone'] else None,
                            'sz_end': (tech['sd_zone']['target']*1.02) if tech['sd_zone'] else None,
                            'sl_price': tech['sd_zone']['stop_loss'] if tech['sd_zone'] else None,
                            'rr_val': tech['sd_zone']['rr_ratio'] if tech['sd_zone'] else None,
                            'year_high': round(year_h, 2), 'upside_to_high': round((year_h-current_p)/current_p*100, 2),
                            'prox_val': round(0.0 if current_p <= (tech['sd_zone']['start'] or 0) else ((current_p - tech['sd_zone']['start']) / tech['sd_zone']['start']) * 100, 2) if tech['sd_zone'] else 999,
                            'rel_1m': round(float((df['Close'].iloc[-1]-df['Close'].iloc[-22])/df['Close'].iloc[-22]*100) - spy_1m, 2) if len(df)>=22 else 0,
                            'rel_3m': round(float((df['Close'].iloc[-1]-df['Close'].iloc[-66])/df['Close'].iloc[-66]*100) - spy_3m, 2) if len(df)>=66 else 0,
                            'macd_histogram': round(macd_hist_val, 4) if macd_hist_val is not None else None,
                            'macd_crossover': macd_cross_val, 'bb_squeeze': bb_squeeze_flag, 'stage2': stage2_flag, 'rs_rating': rs_v,
                            'ema20_aligned': tech.get('ema20_aligned', False), 'ema20_rising': tech.get('ema20_rising', False),
                            'ema20_slope': tech.get('ema20_slope', 0.0),
                            'hh_hl_structure': tech.get('hh_hl_structure', False),
                            'pocket_pivot': pocket_pivot_flag,
                            'vdu_near_zone': vdu_flag,
                            'cmf': tech.get('cmf', 0.0),
                            'is_52w_breakout': tech.get('is_52w_breakout', False),
                            'volume_surge': tech.get('volume_surge', 1.0),
                            'is_volume_surge': tech.get('is_volume_surge', False),
                            'ichimoku_above_kumo': ichimoku_above_kumo, 'ichimoku_tk_cross': ichimoku_tk_cross,
                            'ichimoku_kumo_green': ichimoku_kumo_green, 'ichimoku_chikou_ok': ichimoku_chikou_ok,
                            'ichimoku_score': ichimoku_score_val,
                            'price_pattern': pattern_name, 'price_pattern_score': pattern_score,
                            'vcp': detect_vcp_pattern(df),
                            'launcher_score': tech.get('launcher_score', 0),
                            'turtle_dist_pct': tech.get('turtle_dist_pct', 99.0),
                            'is_explosive': tech.get('is_explosive', False),
                            'tightness_idx': tech.get('tightness_idx', 99.0),
                        }
                    except Exception as e:
                        import logging
                        logging.getLogger('stocks').exception(f"[US Precision] Error scanning {symbol}: {e}")
                        return None

                count = 0
                with concurrent.futures.ThreadPoolExecutor(max_workers=10) as ex:
                    futs = {ex.submit(_scan_one, s): s for s in results_to_process}
                    for f in concurrent.futures.as_completed(futs):
                        try:
                            r = f.result()
                            if r: results.append(r)
                        except: pass
                        count += 1
                        if count % 10 == 0 or count == len(results_to_process):
                            _cache_inner.set(ckey, {'state': 'running', 'progress': count, 'total': len(results_to_process), 'phase': f'Technical Scan ({count}/{len(results_to_process)})…'}, timeout=1200)

                if results:
                    _cache_inner.set(ckey, {'state': 'running', 'progress': count, 'total': len(results_to_process), 'phase': f'Enriching Fundamentals ({len(results)} matches)…'}, timeout=1200)
                    matched = [r['symbol'] for r in results]
                    fund = {}
                    try:
                        yqa = YQTicker(matched)
                        mods = yqa.get_modules('financialData summaryProfile')
                        for k, d in mods.items():
                            if isinstance(d, dict):
                                p = d.get('summaryProfile', {}); fd = d.get('financialData', {})
                                fund[k.upper()] = {
                                    'sector': p.get('sector') or 'Unknown',
                                    'eps': float(fd.get('earningsQuarterlyGrowth',0) or 0)*100,
                                    'rev': float(fd.get('revenueGrowth',0) or 0)*100
                                }
                    except: pass
                    
                    bulk = []
                    for r in results:
                        f = fund.get(r['symbol'].upper(), {'sector':'N/A','eps':0,'rev':0})
                        bulk.append(PrecisionScanCandidate(
                            user=user, market='US', scan_run=scan_run_time, symbol=r['symbol'], symbol_bk=r['symbol'],
                            sector=f['sector'], price=r['price'], rsi=r['rsi'], adx=r['adx'], mfi=r['mfi'], rvol=r['rvol'],
                            eps_growth=round(f['eps'], 2), rev_growth=round(f['rev'], 2),
                            technical_score=r['technical_score'], rs_rating=r['rs_rating'],
                            avg_volume_20d=r['avg_volume_20d'], rvol_bullish=r['rvol_bullish'],
                            erc_volume_confirmed=r['erc_volume_confirmed'], zone_target_source=r['zone_target_src'],
                            is_new_entry=(r['symbol'] not in prev_symbols), entry_strategy=r['entry_strat'],
                            demand_zone_start=r['dz_start'], demand_zone_end=r['dz_end'],
                            supply_zone_start=r['sz_start'], supply_zone_end=r['sz_end'],
                            stop_loss=r['sl_price'], risk_reward_ratio=r['rr_val'],
                            year_high=r['year_high'], upside_to_high=r['upside_to_high'],
                            zone_proximity=r['prox_val'], rel_momentum_1m=r['rel_1m'], rel_momentum_3m=r['rel_3m'],
                            price_pattern=r.get('price_pattern', 'None'),
                            price_pattern_score=r.get('price_pattern_score', 0),
                            macd_histogram=r.get('macd_histogram'),
                            macd_crossover=r['macd_crossover'], bb_squeeze=r['bb_squeeze'],
                            ema20_aligned=r['ema20_aligned'], ema20_slope=r.get('ema20_slope', 0.0), ema20_rising=r['ema20_rising'],
                            hh_hl_structure=r['hh_hl_structure'], stage2=r['stage2'],
                            pocket_pivot=r.get('pocket_pivot', False),
                            vdu_near_zone=r.get('vdu_near_zone', False),
                            cmf=r.get('cmf', None),
                            is_52w_breakout=r.get('is_52w_breakout', False),
                            volume_surge=r.get('volume_surge', 1.0),
                            is_volume_surge=r.get('is_volume_surge', False),
                            ichimoku_above_kumo=r.get('ichimoku_above_kumo', False),
                            ichimoku_tk_cross=r.get('ichimoku_tk_cross', False),
                            ichimoku_kumo_green=r.get('ichimoku_kumo_green', False),
                            ichimoku_chikou_ok=r.get('ichimoku_chikou_ok', False),
                            ichimoku_score=r.get('ichimoku_score', 0),
                            vcp_setup=r.get('vcp', {}).get('setup', False),
                            vcp_contractions=r.get('vcp', {}).get('contractions', 0),
                            vcp_tightness=r.get('vcp', {}).get('tightness', 0.0),
                            vcp_vdu=r.get('vcp', {}).get('vdu_confirmed', False),
                            launcher_score=r.get('launcher_score', 0),
                            turtle_dist_pct=r.get('turtle_dist_pct', 99.0),
                            is_explosive=r.get('is_explosive', False),
                            tightness_idx=r.get('tightness_idx', 99.0),
                        ))
                    PrecisionScanCandidate.objects.bulk_create(bulk)
                    
                    runs = list(PrecisionScanCandidate.objects.filter(user=user, market='US').values_list('scan_run', flat=True).order_by('-scan_run').distinct())
                    if len(runs) > 3: PrecisionScanCandidate.objects.filter(user=user, market='US', scan_run__in=runs[3:]).delete()

                _cache_inner.set(ckey, {'state': 'done'}, timeout=300)
            except Exception as e:
                import traceback
                with open('us_scan_error.txt', 'w') as err_file:
                    err_file.write(traceback.format_exc())
                _cache_inner.set(ckey, {'state': 'done', 'error': str(e)}, timeout=300)

        t = threading.Thread(target=_run_us_scan_bg, args=(user_id, cache_key, scan_symbols), daemon=True)
        t.start()
        return redirect('stocks:us_precision_scanner')

    # Selection & Rendering
    sort_by = request.GET.get('sort', 'score')
    valid_sorts = {
        'symbol':'symbol','score':'-technical_score','price':'-price',
        'rsi':'-rsi','rvol':'-rvol','adx':'-adx','prox':'zone_proximity','rs':'-rs_rating',
        'launcher': '-launcher_score',
    }
    order = valid_sorts.get(sort_by, '-technical_score')
    
    all_runs = list(PrecisionScanCandidate.objects.filter(user=request.user, market='US').values_list('scan_run', flat=True).order_by('-scan_run').distinct())
    try: run_idx = int(request.GET.get('run_idx', 0))
    except: run_idx = 0
    run_idx = max(0, min(run_idx, len(all_runs)-1)) if all_runs else 0
    
    candidates = []
    scanned_at = None
    top5_buy = []
    top5_qualified = []
    top_sectors = []
    scan_insights = []

    if all_runs:
        sel_run = all_runs[run_idx]
        candidates = list(PrecisionScanCandidate.objects.filter(user=request.user, market='US', scan_run=sel_run).order_by(order))
        scanned_at = sel_run
        
        # Live prices
        import pytz as _lpytz
        from datetime import datetime as _ldt, time as _ldtime
        _lny = _lpytz.timezone('America/New_York')
        _lnow = _ldt.now(_lny)
        _lt = _lnow.time()
        _lmarket_open = _lnow.weekday()<5 and _ldtime(9,30)<=_lt<=_ldtime(16,0)
        
        lp_map = {}
        if candidates:
            try:
                import concurrent.futures as lcf
                def _glp(s):
                    try: 
                        fi = yf.Ticker(s).fast_info
                        return s, float(fi.last_price) if fi.last_price else None
                    except: return s, None
                with lcf.ThreadPoolExecutor(max_workers=10) as lex:
                    for s, p in lex.map(_glp, [c.symbol for c in candidates]):
                        if p: lp_map[s] = p
            except: pass
        
        for c in candidates:
            lp = lp_map.get(c.symbol)
            c.live_price = lp; c.is_live = _lmarket_open and lp is not None
            if lp and c.demand_zone_start: c.live_zone_prox = 0.0 if lp <= c.demand_zone_start else round((lp-c.demand_zone_start)/c.demand_zone_start*100, 1)
            if lp and c.price: c.live_change_pct = round((lp-c.price)/c.price*100, 2)
            
            sigs = _compute_signals(c)
            c.buy_score = sigs['buy_score']; c.sell_score = sigs['sell_score']; c.exit_signal = sigs['exit_signal']

            # Reasons logic
            reasons = []
            if c.demand_zone_start and c.demand_zone_end and c.price <= c.demand_zone_start and c.price >= c.demand_zone_end: reasons.append("In Entry Zone")
            elif (c.zone_proximity or 999) <= 10: reasons.append(f"Near Zone {c.zone_proximity:.0f}%")
            if c.rvol_bullish and c.rvol >= 1.0: reasons.append(f"RVOL {c.rvol:.1f}x Bull")
            if (c.risk_reward_ratio or 0) >= 2: reasons.append(f"RR 1:{c.risk_reward_ratio:.1f} Good")
            if c.adx >= 25: reasons.append(f"ADX {c.adx:.0f} Trendy")
            if c.rs_rating >= 85: reasons.insert(0, f"RS {c.rs_rating} Leader")
            c.top_reasons = reasons[:4]

        # Top 5 Buy
        top5_buy = sorted([c for c in candidates if c.buy_score >= 50], key=lambda x: x.buy_score, reverse=True)[:5]
        # Top 5 Qualified
        top5_qualified = sorted([c for c in candidates if c.buy_score >= 65 and (c.risk_reward_ratio or 0) >= 1.5], key=lambda x: x.buy_score, reverse=True)[:5]
        
        # Sectors
        sec_counts = {}
        for c in candidates:
            if c.buy_score >= 65:
                s = c.sector or 'Unknown'
                sec_counts[s] = sec_counts.get(s, 0) + 1
        top_sectors = sorted([{'name': k, 'count': v} for k, v in sec_counts.items()], key=lambda x: x['count'], reverse=True)[:5]

        # Insights
        if top5_qualified:
            b = top5_qualified[0]
            scan_insights.append({'icon': '🏆', 'title': f'Best Setup: {b.symbol}', 'desc': f'RS {b.rs_rating}, RR 1:{b.risk_reward_ratio or 0:.1f}. Strong technical structure in entry zone.'})
        elif top5_buy:
            scan_insights.append({'icon': '💡', 'title': f'Watchlist: {top5_buy[0].symbol}', 'desc': 'High momentum buy score. Watch for price consolidation.'})

    from .models import ScanWatchlistItem
    watchlist_symbols = set(ScanWatchlistItem.objects.filter(user=request.user).values_list('symbol', flat=True))

    scan_data_date = None
    if scanned_at:
        import pytz as _sddtz
        _ny = _sddtz.timezone('America/New_York')
        _st = scanned_at.astimezone(_ny) if hasattr(scanned_at, 'astimezone') else scanned_at
        from datetime import time as _t, timedelta as _tdd
        _in_mkt = (_st.weekday() < 5 and _t(9, 30) <= _st.time() <= _t(16, 0))
        scan_data_date = (_st.date() - _tdd(days=1)) if _in_mkt else _st.date()

    import json as _scan_json
    def _ser_c_us(c):
        return {
            "symbol": c.symbol, "price": c.price, "buy_score": c.buy_score, "rs_rating": c.rs_rating,
            "rsi": round(c.rsi, 1), "adx": round(c.adx, 1), "rvol": round(c.rvol, 2),
            "rvol_bullish": c.rvol_bullish, "risk_reward_ratio": c.risk_reward_ratio,
            "sector": c.sector, "exit_signal": c.exit_signal, "top_reasons": getattr(c, 'top_reasons', []),
        }

    ai_scan_json = _scan_json.dumps({
        "scan_date": str(scan_data_date),
        "qualified_stocks": [_ser_c_us(c) for c in top5_qualified],
        "top_buy_stocks": [_ser_c_us(c) for c in top5_buy],
        "total_passed": len(candidates),
        "top_sectors": top_sectors,
    }, ensure_ascii=False, default=str)

    return render(request, 'stocks/us_precision_scan.html', {
        'title': 'US Precision Momentum Scanner - Nasdaq & S&P 500',
        'candidates': candidates, 'scanned_at': scanned_at, 'current_sort': sort_by,
        'all_runs': all_runs, 'selected_run_idx': run_idx,
        'has_scanned': bool(all_runs), 'top5_buy': top5_buy, 'top5_qualified': top5_qualified,
        'scan_total': len(scan_symbols), 'scan_passed': len(candidates),
        'top_sectors': top_sectors, 'scan_insights': scan_insights,
        'scan_data_date': scan_data_date, 'watchlist_symbols': watchlist_symbols,
        'ai_scan_json': ai_scan_json,
    })


# ======================================================================
# US PRECISION SCAN AI ANALYSIS
# ======================================================================


@login_required
def us_precision_scan_ai_analysis(request):
    import json as _json_lib
    from django.http import JsonResponse
    if request.method != "POST":
        return JsonResponse({"error": "POST only"}, status=405)
    try:
        body = _json_lib.loads(request.body)
        scan_data = body.get("scan_data", {})
    except Exception:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    api_key = getattr(settings, "GEMINI_API_KEY", None)
    if not api_key:
        return JsonResponse({"error": "No GEMINI_API_KEY configured"}, status=500)

    qualified    = scan_data.get("qualified_stocks", [])
    top_buy      = scan_data.get("top_buy_stocks", [])
    scan_date    = scan_data.get("scan_date", "N/A")
    total_passed = scan_data.get("total_passed", 0)
    top_sectors  = scan_data.get("top_sectors", [])

    def _fmt_stock(s):
        lines = [
            "  - {sym}: Price ${price} | BUY {buy} | RS {rs}".format(
                sym=s["symbol"], price=s["price"], buy=s["buy_score"], rs=s["rs_rating"]),
            "    RSI {rsi} | ADX {adx} | RVOL {rvol}x {dir} | RR 1:{rr}".format(
                rsi=s["rsi"], adx=s["adx"], rvol=s["rvol"],
                dir="Bull ▲" if s.get("rvol_bullish") else "Bear ▼",
                rr=s.get("risk_reward_ratio", "-")),
            "    Zone {z}% | {sec} | RelMom3m vs SPY {rm}%".format(
                z=s.get("zone_proximity", "-"), sec=s.get("sector", "-"),
                rm=s.get("rel_momentum_3m", 0)),
            "    Signals: {m}{e}{h}{b}{a}".format(
                m="MACD✕ " if s.get("macd_crossover") else "",
                e="EMA↑ " if s.get("ema20_rising") else "",
                h="HH/HL " if s.get("hh_hl_structure") else "",
                b="BB Squeeze " if s.get("bb_squeeze") else "",
                a="EMA20 Aligned " if s.get("ema20_aligned") else ""),
        ]
        if s.get("top_reasons"):
            lines.append("    Reasons: {r}".format(r=", ".join(s["top_reasons"])))
        return "\n".join(lines)

    q_text   = "\n".join([_fmt_stock(s) for s in qualified]) if qualified else "  (No fully qualified stocks)"
    b_text   = "\n".join([_fmt_stock(s) for s in top_buy]) if top_buy else "  (No data)"
    sec_text = ", ".join(["{n}({c})".format(n=s["name"], c=s["count"]) for s in top_sectors]) if top_sectors else "N/A"

    prompt = (
        "You are an expert US equity analyst specializing in Precision Momentum trading"
        " (Mark Minervini SEPA + William O'Neil CAN SLIM methodology).\n"
        "Expert in Stage Analysis, RS Rating, Supply/Demand Zone,"
        " RVOL Bull/Bear, MACD Crossover, EMA Alignment, Trend Following (HH/HL).\n\n"
        "US Precision Scan Date: {sd} | Stocks Passed Filter: {tp}"
        " | Leading Sectors: {sec}\n\n"
        "=== Fully Qualified Stocks (All Criteria Met) ===\n{q}\n\n"
        "=== Top BUY Score Stocks ===\n{b}\n\n"
        "Analyze the following:\n\n"
        "## 1. 📊 US Market Overview\n"
        "- Market sentiment (Bullish/Mixed/Bearish), Leading sectors, SPY context\n\n"
        "## 2. ✅ Fully Qualified Setups\n"
        "- Each stock: Key strengths, risks, upside potential, priority ranking\n\n"
        "## 3. 🏆 Top BUY Score Analysis\n"
        "- Most attractive setup and why | What to watch out for\n\n"
        "## 4. ⚡ Trading Strategy\n"
        "- Entry order: Which first, which to wait | Optimal entry zones\n\n"
        "## 5. ⚠ Risk Warnings\n"
        "- RSI/RVOL/Zone concerns | Stop Loss discipline | Macro risks\n\n"
        "Reply in English. Be concise and professional."
        " Format as Markdown. Focus on Actionable Insights."
    ).format(sd=scan_date, tp=total_passed, sec=sec_text, q=q_text, b=b_text)

    try:
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        if not response.text:
            return JsonResponse({"error": "AI did not respond"}, status=500)
        return JsonResponse({"status": "success", "analysis": response.text})
    except Exception as e:
        err = str(e)
        if "API_KEY_INVALID" in err:
            return JsonResponse({"error": "GEMINI_API_KEY is invalid"}, status=500)
        return JsonResponse({"error": "Gemini error: {}".format(err)}, status=500)




# ============================================================
# US VALUE SCANNER
# ============================================================

def _seed_value_symbols():
    """~200 US value-oriented stocks across all sectors."""
    return [
        # Financials
        'JPM','BAC','WFC','C','GS','MS','BRK-B','BLK','AXP','USB',
        'TFC','MET','PRU','AFL','CB','ALL','SCHW','COF','DFS','SYF',
        'PNC','HBAN','MTB','CFG','ZION',
        # Healthcare
        'JNJ','PFE','MRK','BMY','ABBV','AMGN','GILD','CVS','UNH',
        'MDT','ABT','SYK','BSX','BIIB','CI','HUM','ELV',
        # Energy
        'XOM','CVX','COP','EOG','SLB','HAL','WMB','KMI','DVN',
        'FANG','MRO','OXY','PSX','VLO','MPC',
        # Consumer Staples
        'KO','PEP','WMT','PG','CL','PM','MO','KMB','GIS','K',
        'CPB','SJM','MKC','HSY','CAG','ADM','BG',
        # Utilities
        'NEE','DUK','SO','D','SRE','AEP','EXC','ED','XEL',
        'WEC','ES','ETR','NI','CMS','AES',
        # Industrials
        'GE','HON','MMM','EMR','ETN','ROK','ITW','PH','GD',
        'LMT','RTX','NOC','UPS','FDX','CSX','NSC','UNP','CAT','DE',
        'CMI','PCAR','IR','AME','ROP',
        # Materials
        'LIN','APD','NEM','FCX','NUE','X','AA','MOS','CF','ALB',
        # Real Estate
        'AMT','PLD','O','VICI','CCI','DLR','EXR','WPC','SPG','PSA',
        # Tech (value-priced)
        'CSCO','IBM','INTC','HPQ','HPE','ORCL','QCOM','TXN',
        'AVGO','AMAT','KLAC','GOOGL','META','MSFT','CTSH','CDW',
        # Consumer Discretionary
        'TGT','KR','GM','F','LEN','PHM','DHI','NVR','TOL',
    ]


def _score_value_candidate(info, df):
    """
    Score a stock on value criteria. Returns (val_score, qual_score, price_score, total).
    """
    import pandas_ta as ta

    val_score = 0
    qual_score = 0
    price_score = 0

    # ── Extract fundamentals ──────────────────────────────
    pe  = info.get('trailingPE') or info.get('forwardPE')
    fpe = info.get('forwardPE')
    pb  = info.get('priceToBook')
    peg = info.get('pegRatio')
    ps  = info.get('priceToSalesTrailing12Months')
    div = (info.get('dividendYield') or 0) * 100          # fraction → %
    roe = (info.get('returnOnEquity') or 0) * 100          # fraction → %
    margin = (info.get('profitMargins') or 0) * 100        # fraction → %
    de_raw = info.get('debtToEquity')
    de  = (de_raw / 100) if de_raw is not None else None   # yf sends %, convert to ratio
    cr  = info.get('currentRatio')
    rev_g = (info.get('revenueGrowth') or 0) * 100
    mkt_cap = (info.get('marketCap') or 0) / 1e9           # → USD billions
    fcf_raw = info.get('freeCashflow') or 0
    fcf_yield = (fcf_raw / (info.get('marketCap') or 1)) * 100 if fcf_raw and mkt_cap > 0 else 0

    # ── Valuation Score (max 40) ──────────────────────────
    if pe and pe > 0:
        if pe < 10:    val_score += 15
        elif pe < 15:  val_score += 12
        elif pe < 20:  val_score += 8
        elif pe < 25:  val_score += 4

    if pb and pb > 0:
        if pb < 1:     val_score += 10
        elif pb < 1.5: val_score += 7
        elif pb < 2.5: val_score += 4

    if div > 0:
        if div >= 4:   val_score += 10
        elif div >= 3: val_score += 7
        elif div >= 2: val_score += 4
        elif div >= 1: val_score += 2

    if peg and peg > 0:
        if peg < 1:    val_score += 5
        elif peg < 1.5: val_score += 3

    # ── Quality Score (max 35) ────────────────────────────
    if roe > 0:
        if roe > 25:   qual_score += 15
        elif roe > 20: qual_score += 12
        elif roe > 15: qual_score += 8
        elif roe > 10: qual_score += 5

    if margin > 0:
        if margin > 25:  qual_score += 10
        elif margin > 15: qual_score += 7
        elif margin > 10: qual_score += 5
        elif margin > 5:  qual_score += 3

    if de is not None:
        if de < 0.3:   qual_score += 10
        elif de < 0.5: qual_score += 7
        elif de < 1.0: qual_score += 4
        elif de < 1.5: qual_score += 2

    # ── Price Action Score (max 25) ───────────────────────
    if df is not None and len(df) >= 50:
        try:
            close = df['Close']
            ema200 = ta.ema(close, length=200)
            rsi14  = ta.rsi(close, length=14)
            last_close = float(close.iloc[-1])
            y_high = float(df['High'].max())
            y_low  = float(df['Low'].min())
            pct_from_high = ((y_high - last_close) / y_high * 100) if y_high > 0 else 0

            # EMA200 trend
            if ema200 is not None and not ema200.dropna().empty:
                ema200_val = float(ema200.dropna().iloc[-1])
                if last_close > ema200_val:
                    price_score += 10

            # RSI - underowned zone
            if rsi14 is not None and not rsi14.dropna().empty:
                rsi_val = float(rsi14.dropna().iloc[-1])
                if 30 <= rsi_val <= 50:  price_score += 7
                elif 50 < rsi_val <= 60: price_score += 4

            # Distance from 52w high
            if pct_from_high > 20:   price_score += 5
            elif pct_from_high > 10: price_score += 3

        except Exception:
            pass

    # FCF yield bonus
    if fcf_yield > 5:   price_score += 3
    elif fcf_yield > 3: price_score += 1

    total = min(val_score + qual_score + price_score, 100)
    return val_score, qual_score, price_score, total


@login_required
def us_value_scanner(request):
    """
    US Value Stock Scanner - fundamental quality + cheap valuation.
    P/E < 25 across all sectors (Financials, Energy, Healthcare, Tech, etc.)
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from datetime import datetime as _dt, timezone as _tz
    import pandas_ta as ta

    run_scan = request.GET.get('scan') == 'true'
    current_sort = request.GET.get('sort', 'score')

    sort_map = {
        'score': '-total_score', 'val': '-valuation_score',
        'qual': '-quality_score', 'pe': 'pe_ratio',
        'pb': 'pb_ratio',        'div': '-dividend_yield',
        'roe': '-roe',           'symbol': 'symbol',
        'price': '-price',       'rsi': 'rsi',
    }

    # ── Load from DB (display mode) ───────────────────────
    if not run_scan:
        all_runs = list(
            ValueScanCandidate.objects
            .filter(user=request.user)
            .values_list('scan_run', flat=True)
            .order_by('-scan_run').distinct()
        )
        has_scanned = bool(all_runs)
        candidates = []
        scanned_at = None

        try:
            run_idx = int(request.GET.get('run_idx', 0))
        except (ValueError, TypeError):
            run_idx = 0
        run_idx = max(0, min(run_idx, len(all_runs) - 1)) if all_runs else 0

        if has_scanned:
            selected_run = all_runs[run_idx]
            scanned_at   = selected_run
            qs = (ValueScanCandidate.objects
                  .filter(user=request.user, scan_run=selected_run)
                  .order_by(sort_map.get(current_sort, '-total_score')))
            candidates = list(qs)

            # Live prices (fast_info)
            live_prices = {}
            try:
                import concurrent.futures as _lcf
                def _get_live_val(sym):
                    try:
                        fi = yf.Ticker(sym).fast_info
                        p = getattr(fi, 'last_price', None)
                        return sym, round(float(p), 2) if p else None
                    except Exception:
                        return sym, None
                with _lcf.ThreadPoolExecutor(max_workers=12) as ex:
                    for sym, p in ex.map(_get_live_val, [c.symbol for c in candidates]):
                        if p:
                            live_prices[sym] = p
            except Exception:
                pass
            for c in candidates:
                c.live_price = live_prices.get(c.symbol)

        return render(request, 'stocks/us_value_scan.html', {
            'candidates':       candidates,
            'has_scanned':      has_scanned,
            'scanned_at':       scanned_at,
            'all_runs':         all_runs,
            'selected_run_idx': run_idx,
            'current_sort':     current_sort,
        })

    # ── RUN SCAN ──────────────────────────────────────────
    symbols = _seed_value_symbols()
    scan_time = _dt.now(_tz.utc)
    results = []

    def _process_value_symbol(sym):
        try:
            ticker = yf.Ticker(sym)
            info   = ticker.info or {}

            price = info.get('regularMarketPrice') or info.get('currentPrice') or 0
            if not price:
                fi = getattr(ticker, 'fast_info', None)
                price = getattr(fi, 'last_price', 0) or 0
            if not price or price <= 0:
                return None

            # P/E filter - skip pure growth stocks (P/E > 30)
            pe = info.get('trailingPE') or info.get('forwardPE')
            if pe and pe > 30:
                return None

            # Market cap filter - at least $2B (mid/large cap)
            mkt_cap = (info.get('marketCap') or 0) / 1e9
            if mkt_cap < 2:
                return None

            # Download 1-year price history for technical indicators
            df = yf.download(sym, period='1y', progress=False, auto_adjust=True)
            if df is None or len(df) < 50:
                return None
            if isinstance(df.columns, pd.MultiIndex):
                df.columns = df.columns.droplevel(1)

            val_score, qual_score, price_score, total = _score_value_candidate(info, df)

            # Minimum quality threshold
            if total < 20:
                return None

            # Compute additional price stats
            close = df['Close']
            rsi14  = ta.rsi(close, length=14)
            rsi_val = float(rsi14.dropna().iloc[-1]) if rsi14 is not None and not rsi14.dropna().empty else 50
            y_high = float(df['High'].max())
            y_low  = float(df['Low'].min())
            pct_from_high = ((y_high - float(close.iloc[-1])) / y_high * 100) if y_high > 0 else 0

            ema200 = ta.ema(close, length=200)
            above_ema200 = False
            if ema200 is not None and not ema200.dropna().empty:
                above_ema200 = float(close.iloc[-1]) > float(ema200.dropna().iloc[-1])

            div = (info.get('dividendYield') or 0) * 100
            roe = (info.get('returnOnEquity') or 0) * 100
            margin = (info.get('profitMargins') or 0) * 100
            de_raw = info.get('debtToEquity')
            de  = (de_raw / 100) if de_raw is not None else None
            fcf_raw = info.get('freeCashflow') or 0
            fcf_yield = (fcf_raw / (info.get('marketCap') or 1)) * 100 if fcf_raw and mkt_cap > 0 else 0

            return {
                'symbol':       sym,
                'name':         info.get('longName') or info.get('shortName') or sym,
                'sector':       info.get('sector') or 'Unknown',
                'price':        round(float(price), 2),
                'market_cap':   round(mkt_cap, 2),
                'pe_ratio':     round(float(pe), 2) if pe and pe > 0 else None,
                'forward_pe':   round(float(info.get('forwardPE') or 0), 2) or None,
                'pb_ratio':     round(float(info.get('priceToBook') or 0), 2) or None,
                'peg_ratio':    round(float(info.get('pegRatio') or 0), 2) or None,
                'ps_ratio':     round(float(info.get('priceToSalesTrailing12Months') or 0), 2) or None,
                'dividend_yield': round(div, 2),
                'roe':          round(roe, 1) if roe else None,
                'profit_margin': round(margin, 1) if margin else None,
                'debt_equity':  round(de, 2) if de is not None else None,
                'current_ratio': round(float(info.get('currentRatio') or 0), 2) or None,
                'revenue_growth': round((info.get('revenueGrowth') or 0) * 100, 1),
                'fcf_yield':    round(fcf_yield, 1),
                'rsi':          round(rsi_val, 1),
                'year_high':    round(y_high, 2),
                'year_low':     round(y_low, 2),
                'pct_from_high': round(pct_from_high, 1),
                'above_ema200': above_ema200,
                'valuation_score':    val_score,
                'quality_score':      qual_score,
                'price_action_score': price_score,
                'total_score':        total,
            }
        except Exception:
            return None

    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = {pool.submit(_process_value_symbol, sym): sym for sym in symbols}
        for future in as_completed(futures):
            r = future.result()
            if r:
                results.append(r)

    results.sort(key=lambda x: -x['total_score'])

    # Fetch previous symbols for new-entry flag
    prev_run_q = (ValueScanCandidate.objects
                  .filter(user=request.user)
                  .order_by('-scan_run')
                  .values_list('scan_run', flat=True)
                  .distinct()[:1])
    prev_symbols = set()
    if prev_run_q:
        prev_symbols = set(
            ValueScanCandidate.objects.filter(user=request.user, scan_run=prev_run_q[0])
            .values_list('symbol', flat=True)
        )

    # Save new results FIRST
    to_create = []
    for r in results:
        to_create.append(ValueScanCandidate(
            user=request.user, scan_run=scan_time,
            symbol=r['symbol'], name=r['name'], sector=r['sector'],
            price=r['price'], market_cap=r['market_cap'],
            pe_ratio=r['pe_ratio'], forward_pe=r['forward_pe'],
            pb_ratio=r['pb_ratio'], peg_ratio=r['peg_ratio'],
            ps_ratio=r['ps_ratio'], dividend_yield=r['dividend_yield'],
            roe=r['roe'], profit_margin=r['profit_margin'],
            debt_equity=r['debt_equity'], current_ratio=r['current_ratio'],
            revenue_growth=r['revenue_growth'], fcf_yield=r['fcf_yield'],
            rsi=r['rsi'], year_high=r['year_high'], year_low=r['year_low'],
            pct_from_high=r['pct_from_high'], above_ema200=r['above_ema200'],
            valuation_score=r['valuation_score'],
            quality_score=r['quality_score'],
            price_action_score=r['price_action_score'],
            total_score=r['total_score'],
            is_new_entry=(r['symbol'] not in prev_symbols),
        ))
    ValueScanCandidate.objects.bulk_create(to_create)

    # THEN delete old runs (keep last 3)
    distinct_runs = list(
        ValueScanCandidate.objects
        .filter(user=request.user)
        .values_list('scan_run', flat=True)
        .order_by('-scan_run').distinct()
    )
    if len(distinct_runs) > 3:
        ValueScanCandidate.objects.filter(
            user=request.user, scan_run__in=distinct_runs[3:]
        ).delete()

    return redirect(f'/stocks/value/us-value/?sort={current_sort}&run_idx=0')


# ======================================================================
# US SEPA SCANNER - แยกต่างหากจาก US Precision Scanner
# Model: USSepaCandidate (ไม่แตะ PrecisionScanCandidate เลย)
# ======================================================================

@login_required
def us_sepa_scanner(request):
    """
    US SEPA Scanner - Stage 2 + VCP + RS ≥70 สำหรับหุ้น Nasdaq/S&P500
    ใช้ USSepaCandidate (แยกต่างหากจาก PrecisionScanCandidate อย่างสมบูรณ์)
    """
    import yfinance as yf
    import pandas as pd
    from .models import USSepaCandidate, ScannableSymbol, ScanWatchlistItem

    # AJAX scan status
    if request.GET.get('scan_status') == '1':
        from django.core.cache import cache as _cp
        from django.http import JsonResponse as _JR
        _key = f'us_sepa_scan_{request.user.id}'
        st = _cp.get(_key, {'state': 'idle'})
        if st.get('state') == 'done':
            _cp.delete(_key)
        return _JR(st)

    # ── Trigger background scan ────────────────────────────────────
    if request.method == 'POST' or request.GET.get('scan') == 'true':
        from django.core.cache import cache as _cache_bg
        import threading
        user_id   = request.user.id
        cache_key = f'us_sepa_scan_{user_id}'

        if _cache_bg.get(cache_key, {}).get('state') == 'running':
            return redirect('stocks:us_sepa_scanner')

        sym_list = list(ScannableSymbol.objects.filter(is_active=True, market='US').values_list('symbol', flat=True))
        if len(sym_list) < 100:
            _seed_us_symbols()
            sym_list = list(ScannableSymbol.objects.filter(is_active=True, market='US').values_list('symbol', flat=True))

        _cache_bg.set(cache_key, {'state': 'running', 'progress': 0, 'total': len(sym_list), 'phase': 'Starting US SEPA scan…'}, timeout=1200)

        def _run_bg(uid, ckey, syms):
            try:
                import django; django.setup()
                import pandas_ta as ta
                import concurrent.futures
                import pytz as _pytz
                from datetime import datetime as _dt, timedelta as _td
                from django.contrib.auth import get_user_model
                from django.core.cache import cache as _c
                from django.utils import timezone as tz
                from .models import USSepaCandidate as _USC, ScannableSymbol
                from .utils import detect_vcp_pattern

                User = get_user_model()
                user = User.objects.get(pk=uid)
                _ny  = _pytz.timezone('America/New_York')
                now  = _dt.now(_ny)
                end_str   = (now.date() + _td(days=1)).strftime('%Y-%m-%d')
                start_str = (now.date() - _td(days=600)).strftime('%Y-%m-%d')
                scan_run  = tz.now()

                # Keep only 3 latest runs
                old_runs = list(_USC.objects.filter(user=user).values_list('scan_run', flat=True).distinct().order_by('-scan_run')[3:])
                if old_runs:
                    _USC.objects.filter(user=user, scan_run__in=old_runs).delete()

                # ── Step 1: Compute RS Rating ─────────────────────────────
                _c.set(ckey, {'state': 'running', 'progress': 0, 'total': len(syms), 'phase': 'Computing RS Ratings…'}, timeout=1200)

                def _fetch_rs(s):
                    try:
                        d = yf.Ticker(s).history(start=start_str, end=end_str, interval='1d')
                        if d is None or d.empty: return s, None
                        if isinstance(d.columns, pd.MultiIndex): d.columns = d.columns.droplevel(1)
                        cl = d['Close'].dropna()
                        if len(cl) < 252: return s, None
                        r = float(
                            (cl.iloc[-1]-cl.iloc[-64])/abs(cl.iloc[-64])*0.4 +
                            (cl.iloc[-64]-cl.iloc[-127])/abs(cl.iloc[-127])*0.2 +
                            (cl.iloc[-127]-cl.iloc[-190])/abs(cl.iloc[-190])*0.2 +
                            (cl.iloc[-190]-cl.iloc[-253])/abs(cl.iloc[-253])*0.2
                        ) * 100
                        return s, r
                    except: return s, None

                rs_raw = {}
                with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
                    for s, r in ex.map(_fetch_rs, syms):
                        if r is not None: rs_raw[s] = r
                rs_map = {}
                if rs_raw:
                    ser = pd.Series(rs_raw)
                    rs_map = (ser.rank(pct=True) * 99).clip(0, 99).astype(int).to_dict()

                # ── Step 2: SEPA Technical Scan ───────────────────────────
                _c.set(ckey, {'state': 'running', 'progress': 0, 'total': len(syms), 'phase': 'SEPA Technical Scan…'}, timeout=1200)

                results = []
                def _scan_one(symbol):
                    try:
                        rs_v = rs_map.get(symbol, 0)
                        if rs_v < 60: return None  # pre-filter; display shows ≥70

                        df = yf.Ticker(symbol).history(start=start_str, end=end_str, interval='1d')
                        if df is None or df.empty: return None
                        if isinstance(df.columns, pd.MultiIndex): df.columns = df.columns.droplevel(1)
                        df = df.dropna(subset=['Close', 'High', 'Low'])
                        if len(df) < 200: return None

                        # Liquidity: avg daily volume ≥ 1M shares
                        if float(df['Volume'].tail(20).mean()) < 1_000_000: return None

                        curr = float(df['Close'].iloc[-1])
                        year_h = float(df['High'].tail(252).max())

                        # Stage 2: price > SMA150 AND SMA150 trending up
                        s2 = False
                        try:
                            s150 = ta.sma(df['Close'], 150)
                            if s150 is not None and pd.notna(s150.iloc[-1]) and pd.notna(s150.iloc[-20]):
                                s2 = (curr > float(s150.iloc[-1])) and (float(s150.iloc[-1]) > float(s150.iloc[-20]))
                        except: pass
                        if not s2: return None

                        # ADX
                        adx_v = 0.0
                        try:
                            adx_df = ta.adx(df['High'], df['Low'], df['Close'], 14)
                            if adx_df is not None:
                                col = [c for c in adx_df.columns if c.startswith('ADX_')]
                                if col and pd.notna(adx_df[col[0]].iloc[-1]):
                                    adx_v = float(adx_df[col[0]].iloc[-1])
                        except: pass

                        # RSI
                        rsi_v = 50.0
                        try:
                            r = ta.rsi(df['Close'], 14)
                            if r is not None and pd.notna(r.iloc[-1]):
                                rsi_v = float(r.iloc[-1])
                        except: pass

                        # RVOL
                        rvol_v = 1.0
                        try:
                            avg20 = float(df['Volume'].tail(20).mean())
                            if avg20 > 0:
                                rvol_v = round(float(df['Volume'].iloc[-1]) / avg20, 2)
                        except: pass

                        # VCP
                        vcp = detect_vcp_pattern(df)

                        # VDU (Volume Dry-Up near zone)
                        vdu_near = False
                        try:
                            rv5 = float(df['Volume'].tail(5).mean())
                            rv50 = float(df['Volume'].tail(50).mean())
                            vdu_near = (rv5 < rv50 * 0.70) and (curr >= year_h * 0.88)
                        except: pass

                        # Pocket Pivot
                        pp = False
                        try:
                            vols = df['Volume'].values
                            closes = df['Close'].values
                            if len(vols) >= 12:
                                today_vol = vols[-1]
                                today_up  = closes[-1] > closes[-2]
                                dn_vols   = [vols[-(i+2)] for i in range(10) if closes[-(i+2)] < closes[-(i+3)]]
                                if today_up and dn_vols and today_vol > max(dn_vols):
                                    pp = True
                        except: pass

                        return {
                            'symbol': symbol,
                            'price': round(curr, 2),
                            'stage2': True,
                            'rs_rating': rs_v,
                            'vcp_setup': vcp.get('setup', False),
                            'vcp_contractions': vcp.get('contractions', 0),
                            'vcp_tightness': vcp.get('tightness', 0.0),
                            'vcp_vdu': vcp.get('vdu_confirmed', False),
                            'pocket_pivot': pp,
                            'vdu_near_zone': vdu_near,
                            'adx': round(adx_v, 1),
                            'rsi': round(rsi_v, 1),
                            'rvol': rvol_v,
                            'year_high': round(year_h, 2),
                            'upside_to_high': round((year_h - curr) / curr * 100, 2),
                        }
                    except: return None

                done = 0
                with concurrent.futures.ThreadPoolExecutor(max_workers=12) as ex:
                    futs = {ex.submit(_scan_one, s): s for s in syms}
                    for fut in concurrent.futures.as_completed(futs):
                        done += 1
                        if done % 15 == 0:
                            _c.set(ckey, {'state': 'running', 'progress': done, 'total': len(syms), 'phase': f'Scanning {done}/{len(syms)}…'}, timeout=1200)
                        r = fut.result()
                        if r: results.append(r)

                # ── Step 3: Enrich sector names ───────────────────────────
                if results:
                    _c.set(ckey, {'state': 'running', 'progress': done, 'total': len(syms), 'phase': 'Fetching sector data…'}, timeout=1200)
                    from yahooquery import Ticker as YQT
                    sector_map = {}
                    name_map   = {}
                    try:
                        yqa = YQT([r['symbol'] for r in results])
                        mods = yqa.get_modules('summaryProfile quoteType')
                        for k, d in mods.items():
                            if isinstance(d, dict):
                                sp = d.get('summaryProfile', {})
                                qt = d.get('quoteType', {})
                                sector_map[k.upper()] = sp.get('sector', 'Unknown') or 'Unknown'
                                name_map[k.upper()]   = qt.get('shortName', '') or ''
                    except: pass

                    bulk = [_USC(
                        user=user, scan_run=scan_run,
                        symbol=r['symbol'],
                        name=name_map.get(r['symbol'], ''),
                        sector=sector_map.get(r['symbol'], 'Unknown'),
                        price=r['price'],
                        stage2=r['stage2'],
                        rs_rating=r['rs_rating'],
                        vcp_setup=r['vcp_setup'],
                        vcp_contractions=r['vcp_contractions'],
                        vcp_tightness=r['vcp_tightness'],
                        vcp_vdu=r['vcp_vdu'],
                        pocket_pivot=r['pocket_pivot'],
                        vdu_near_zone=r['vdu_near_zone'],
                        adx=r['adx'],
                        rsi=r['rsi'],
                        rvol=r['rvol'],
                        year_high=r['year_high'],
                        upside_to_high=r['upside_to_high'],
                    ) for r in results]
                    _USC.objects.bulk_create(bulk)

                _c.set(ckey, {'state': 'done'}, timeout=300)
            except Exception as exc:
                import logging
                logging.getLogger('stocks').exception(f'[US SEPA] bg scan error: {exc}')
                from django.core.cache import cache as _c2
                _c2.set(ckey, {'state': 'done'}, timeout=300)

        import threading as _thr
        _thr.Thread(target=_run_bg, args=(request.user.id, cache_key, sym_list), daemon=True).start()
        return redirect('stocks:us_sepa_scanner')

    # ── Display ────────────────────────────────────────────────────
    all_runs = list(
        USSepaCandidate.objects.filter(user=request.user)
        .values_list('scan_run', flat=True).distinct().order_by('-scan_run')
    )
    try:
        run_idx = max(0, min(int(request.GET.get('run_idx', 0)), len(all_runs) - 1)) if all_runs else 0
    except (ValueError, TypeError):
        run_idx = 0

    candidates  = []
    last_updated = None
    if all_runs:
        run_time    = all_runs[run_idx]
        candidates  = list(USSepaCandidate.objects.filter(user=request.user, scan_run=run_time).order_by('-vcp_setup', '-rs_rating'))
        last_updated = run_time

    # Filters
    vcp_only     = request.GET.get('vcp_only') == '1'
    hide_at_tp   = request.GET.get('hide_at_tp', '1') == '1'

    if vcp_only:
        candidates = [c for c in candidates if c.vcp_setup]
    if hide_at_tp:
        candidates = [c for c in candidates if c.upside_to_high >= 5.0]

    # RS filter: enforce ≥70 (scan saves down to 60 for flexibility)
    candidates = [c for c in candidates if c.rs_rating >= 70]

    # Computed display fields + SEPA Score
    for c in candidates:
        c.dist_from_pivot = round(c.upside_to_high, 1)
        if c.upside_to_high < 5:
            c.tp_status = 'at_tp'
        elif c.upside_to_high < 10:
            c.tp_status = 'near_tp'
        else:
            c.tp_status = None

        # ── SEPA Score (0-200) ─────────────────────────────────
        # VCP Setup quality (0-65 pts)
        sc = 0
        if c.vcp_setup:
            sc += 30
            sc += int(max(0, (10 - min(c.vcp_tightness, 10)) * 2))  # tighter = better (max 20)
            sc += min(c.vcp_contractions, 5) * 3                     # more contractions = better (max 15)
        # Volume Dry-Up confirmed (20 pts)
        if c.vcp_vdu or c.vdu_near_zone:
            sc += 20
        # Pocket Pivot (10 pts)
        if c.pocket_pivot:
            sc += 10
        # RS Rating (weight 0.7, max ~69 pts)
        sc += int(c.rs_rating * 0.7)
        # ADX strength (0-10 pts)
        if c.adx >= 25:
            sc += 10
        elif c.adx >= 15:
            sc += 5
        # Proximity to pivot bonus (only for VCP stocks)
        if c.vcp_setup:
            dist = c.dist_from_pivot
            if dist <= 5:
                sc += 10
            elif dist <= 10:
                sc += 5
            elif dist > 15:
                sc -= 5
        c.sepa_score = sc

    # Sort by SEPA Score descending
    candidates.sort(key=lambda c: c.sepa_score, reverse=True)

    # Assign rank
    for i, c in enumerate(candidates, 1):
        c.sepa_rank = i

    watchlist_symbols = set(ScanWatchlistItem.objects.filter(user=request.user).values_list('symbol', flat=True))

    context = {
        'candidates':    candidates,
        'last_updated':  last_updated,
        'all_runs':      all_runs,
        'selected_run_idx': run_idx,
        'vcp_only':      vcp_only,
        'hide_at_tp':    hide_at_tp,
        'watchlist_symbols': watchlist_symbols,
    }
    return render(request, 'stocks/us_sepa_scanner.html', context)


@login_required
def cup_handle_scanner(request):
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'cup_handle_scan_{user_id}'

    if request.GET.get('scan_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        if st.get('state') == 'done':
            _cp.delete(cache_key)
        return _JR(st)

    if request.GET.get('scan') == 'true' or request.method == 'POST':
        from .utils import refresh_all_thai_symbols, get_top_ranked_symbols
        # ใช้ Top 300 หุ้นใหญ่เท่านั้นเพื่อความเร็วและแม่นยำ (Cup & Handle ต้องการสภาพคล่อง)
        scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)
        
        if not scan_symbols:
            refresh_all_thai_symbols()
            scan_symbols = get_top_ranked_symbols(market='SET', limit=200, auto_refresh=True)

        already = _cp.get(cache_key, {})
        if already.get('state') != 'running':
            _cp.set(cache_key, {'state': 'running', 'progress': 0, 'total': len(scan_symbols), 'phase': 'เริ่มสแกน Cup & Handle...'}, timeout=900)

            def _run_cup_handle_bg(uid, ckey, sym_list):
                try:
                    import pandas as _pd
                    import pandas_ta as _ta
                    import yfinance as _yf
                    import concurrent.futures as _cf
                    from django.core.cache import cache as _c
                    from django.contrib.auth import get_user_model
                    from django.utils import timezone as tz
                    from datetime import datetime as _dt, timedelta as _td
                    import pytz as _pytz
                    from .models import CupHandleCandidate as _CHC
                    from .utils import detect_cup_and_handle, get_top_ranked_symbols as _GTRS
                    from yahooquery import Ticker as _TQ
                    sym_list = _GTRS(market='SET', limit=200, auto_refresh=True)

                    User      = get_user_model()
                    user      = User.objects.get(pk=uid)
                    _bkk      = _pytz.timezone('Asia/Bangkok')
                    _now      = _dt.now(_bkk)
                    _end_str  = (_now.date() + _td(days=1)).strftime('%Y-%m-%d')
                    _start    = (_now.date() - _td(days=600)).strftime('%Y-%m-%d')
                    _scan_run = tz.now()

                    # 1. ลบข้อมูลเก่าที่เกิน 3 รอบ
                    old_runs = list(_CHC.objects.filter(user=user).values_list('scan_run', flat=True).distinct().order_by('-scan_run'))
                    if len(old_runs) >= 3:
                        _CHC.objects.filter(user=user, scan_run__in=old_runs[2:]).delete()

                    _CHC.objects.filter(user=user, scan_run=_scan_run).delete()

                    total = len(sym_list)
                    results = []

                    # --- STAGE 1: Bulk Screening (Fast) ---
                    _c.set(ckey, {'state': 'running', 'progress': 5, 'total': total, 'phase': 'Stage 1: 🔎 กรองสุขภาพคล่อง (Bulk)...'}, timeout=900)
                    chunk_size = 100
                    candidates = []
                    for i in range(0, total, chunk_size):
                        chunk = sym_list[i : i + chunk_size]
                        chunk_bk = [f"{s}.BK" for s in chunk]
                        try:
                            tq = _TQ(chunk_bk, timeout=60)
                            # summary_detail has 'averageVolume' (3-month) which is more stable for screening
                            details = tq.summary_detail
                            for symbol in chunk:
                                s_bk = f"{symbol}.BK"
                                if isinstance(details, dict) and s_bk in details:
                                    d_data = details[s_bk]
                                    if isinstance(d_data, dict):
                                        vol_avg = d_data.get('averageVolume', 0) or 0
                                        price   = d_data.get('previousClose', 0) or 0
                                        # Liquidity filter: Value > 1,000,000 THB/day
                                        if (vol_avg * price) >= 1_000_000:
                                            candidates.append(symbol)
                        except Exception:
                            candidates.extend(chunk)

                    # --- STAGE 2: Pattern Analysis (Threaded) ---
                    total_cand = len(candidates)
                    _c.set(ckey, {'state': 'running', 'progress': 20, 'total': total_cand, 'phase': f'Stage 2: ☕ วิเคราะห์รูปแบบ {total_cand} ตัว...'}, timeout=900)

                    def _scan_one(symbol):
                        try:
                            s_bk = f'{symbol}.BK'
                            ticker_obj = _yf.Ticker(s_bk)
                            df = ticker_obj.history(start=_start, end=_end_str, interval="1d")
                            if df is None or df.empty or len(df) < 80:
                                return None
                            if isinstance(df.columns, _pd.MultiIndex):
                                df.columns = df.columns.droplevel(1)
                            df = df.dropna(subset=['Close', 'High', 'Low'])
                            
                            pat = detect_cup_and_handle(df)
                            if pat is None: return None

                            # RS return (3M)
                            close_s = df['Close'].dropna()
                            rs_return = None
                            if len(close_s) >= 66:
                                rs_return = float((close_s.iloc[-1] - close_s.iloc[-66]) / abs(close_s.iloc[-66]) * 100)

                            # ADX & RSI
                            adx_val, rsi_val = 0.0, 50.0
                            try:
                                adx_df = _ta.adx(df['High'], df['Low'], df['Close'], length=14)
                                if adx_df is not None and not adx_df.empty:
                                    col = [c for c in adx_df.columns if c.startswith('ADX_')]
                                    if col: adx_val = float(adx_df[col[0]].iloc[-1])
                                rsi_s = _ta.rsi(df['Close'], length=14)
                                if rsi_s is not None: rsi_val = float(rsi_s.iloc[-1])
                            except Exception: pass

                            return {'symbol': symbol, 'pat': pat, 'rs_return': rs_return,
                                    'adx': adx_val, 'rsi': rsi_val, 'avg_vol': float(df['Volume'].tail(20).mean())}
                        except Exception: return None

                    with _cf.ThreadPoolExecutor(max_workers=15) as ex:
                        futs = {ex.submit(_scan_one, s): s for s in candidates}
                        done = 0
                        for fut in _cf.as_completed(futs):
                            done += 1
                            if done % 5 == 0:
                                _c.set(ckey, {'state': 'running', 'progress': 20 + int((done/total_cand)*75), 'total': total_cand,
                                              'phase': f'วิเคราะห์ {done}/{total_cand}...'}, timeout=900)
                            try:
                                res = fut.result()
                                if res: results.append(res)
                            except Exception: pass
                    
                    # RS Percentile
                    rs_map = {}
                    rs_vals = {r['symbol']: r['rs_return'] for r in results if r['rs_return'] is not None}
                    if rs_vals:
                        rs_ser = _pd.Series(rs_vals)
                        rs_map = (rs_ser.rank(pct=True) * 99).clip(0, 99).astype(int).to_dict()

                    # Save to DB
                    objs = []
                    for r in results:
                        pat = dict(r['pat'])
                        price_val = pat.pop('current_price', 0.0)
                        objs.append(_CHC(
                            user=user, scan_run=_scan_run, symbol=r['symbol'],
                            price=price_val,
                            market='SET',
                            rs_rating=rs_map.get(r['symbol'], 0),
                            adx=r['adx'], rsi=r['rsi'], avg_vol_20d=r['avg_vol'],
                            **pat
                        ))
                    _CHC.objects.bulk_create(objs)
                    _c.set(ckey, {'state': 'done', 'progress': 100}, timeout=300)
                except Exception as exc:
                    import logging
                    logging.getLogger('stocks').exception(f'[Cup&Handle] bg scan error: {exc}')
                    from django.core.cache import cache as _c2
                    _c2.set(ckey, {'state': 'done'}, timeout=300)

            _th.Thread(target=_run_cup_handle_bg, args=(user_id, cache_key, scan_symbols), daemon=True).start()

        from django.shortcuts import redirect as _redir
        return _redir('stocks:cup_handle_scanner')

    # ── Display results ───────────────────────────────────────────
    from .models import CupHandleCandidate as _CHC, ScanWatchlistItem as _SWI

    all_runs = list(
        _CHC.objects.filter(user=request.user, market='SET')
        .values_list('scan_run', flat=True)
        .distinct().order_by('-scan_run')
    )

    try:
        run_idx = int(request.GET.get('run_idx', 0))
    except (ValueError, TypeError):
        run_idx = 0
    run_idx = max(0, min(run_idx, len(all_runs) - 1)) if all_runs else 0

    candidates = []
    scanned_at = None
    if all_runs:
        selected_run = all_runs[run_idx]
        candidates   = list(_CHC.objects.filter(user=request.user, scan_run=selected_run, market='SET'))
        scanned_at   = selected_run

    # เรียงตาม Stage Priority → Confidence เสมอ
    stage_order = {'breakout': 0, 'ready': 1, 'handle': 2, 'forming': 3}
    candidates.sort(key=lambda c: (stage_order.get(c.stage, 9), -c.confidence_score))

    # Stage counts (ก่อน filter เพื่อแสดงใน summary cards)
    stage_counts = {
        'breakout': sum(1 for c in candidates if c.stage == 'breakout'),
        'ready':    sum(1 for c in candidates if c.stage == 'ready'),
        'handle':   sum(1 for c in candidates if c.stage == 'handle'),
        'forming':  sum(1 for c in candidates if c.stage == 'forming'),
    }

    # ── Filters ───────────────────────────────────────────────────
    stage_filter = request.GET.get('stage', 'all')
    try:
        min_conf = int(request.GET.get('min_conf', 0))
    except (ValueError, TypeError):
        min_conf = 0
    rs_only = request.GET.get('rs_only') == '1'

    if stage_filter != 'all':
        candidates = [c for c in candidates if c.stage == stage_filter]
    if min_conf > 0:
        candidates = [c for c in candidates if c.confidence_score >= min_conf]
    if rs_only:
        candidates = [c for c in candidates if c.rs_rating >= 70]

    # ── Computed fields ───────────────────────────────────────────
    for c in candidates:
        # % ห่างจาก Breakout Price
        if c.breakout_price > 0 and c.price > 0:
            c.pct_to_breakout = round((c.breakout_price - c.price) / c.price * 100, 1)
            c.pct_to_breakout = max(0.0, c.pct_to_breakout)
        else:
            c.pct_to_breakout = 0.0
        # Recovery % (ใช้แสดง progress bar สำหรับ Forming)
        if c.cup_high > c.cup_low:
            raw = (c.price - c.cup_low) / (c.cup_high - c.cup_low) * 100
            c.recovery_pct = round(min(100.0, max(0.0, raw)), 1)
        else:
            c.recovery_pct = 0.0

    # หุ้น Forming ที่ฟื้นตัวมากที่สุด (สำหรับ smart summary)
    forming_list = [c for c in candidates if c.stage == 'forming']
    closest_forming = max(forming_list, key=lambda c: c.recovery_pct, default=None)

    # Watchlist symbols ของ user
    watchlist_symbols = set(
        _SWI.objects.filter(user=request.user, market='SET').values_list('symbol', flat=True)
    )

    context = {
        'candidates':       candidates,
        'has_scanned':      bool(all_runs),
        'scanned_at':       scanned_at,
        'all_runs':         all_runs,
        'selected_run_idx': run_idx,
        'current_sort':     request.GET.get('sort', 'stage'),
        'stage_counts':     stage_counts,
        'stage_filter':     stage_filter,
        'min_conf':         min_conf,
        'rs_only':          rs_only,
        'closest_forming':  closest_forming,
        'watchlist_symbols': watchlist_symbols,
        'stage_labels': {
            'breakout': ('Breakout',       '#16a34a'),
            'ready':    ('Ready to Break', '#2563eb'),
            'handle':   ('Handle Forming', '#d97706'),
            'forming':  ('Cup Forming',    '#64748b'),
        },
    }
    return render(request, 'stocks/cup_handle_scan.html', context)


# ====== US Cup & Handle Scanner ======

@login_required
def us_cup_handle_scanner(request):
    """
    US Cup & Handle Scanner - สแกน หุ้น US จาก _US_MOMENTUM_SYMBOLS universe
    ใช้ logic เดียวกับ SET scanner แต่:
    - ไม่เติม .BK suffix
    - liquidity filter เป็น USD (avg_vol * avg_price >= 1,000,000 USD)
    - บันทึก market='US' แยกต่างหาก
    - ตรวจ breakout_vol_ok (volume ≥1.5x avg on breakout bar) - O'Neil rule
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'us_cup_handle_scan_{user_id}'

    if request.GET.get('scan_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        if st.get('state') == 'done':
            _cp.delete(cache_key)
        return _JR(st)

    if request.GET.get('scan') == 'true' or request.method == 'POST':
        scan_symbols = [s for s in _US_MOMENTUM_SYMBOLS if s not in ('SPY', 'QQQ', 'IWM')]

        already = _cp.get(cache_key, {})
        if already.get('state') != 'running':
            _cp.set(cache_key, {
                'state': 'running', 'progress': 0,
                'total': len(scan_symbols), 'phase': 'เริ่มสแกน US Cup & Handle...'
            }, timeout=900)

            def _run_us_cup_handle_bg(uid, ckey, sym_list):
                try:
                    import pandas as _pd
                    import pandas_ta as _ta
                    import yfinance as _yf
                    import concurrent.futures as _cf
                    from django.core.cache import cache as _c
                    from django.contrib.auth import get_user_model
                    from datetime import datetime as _dt, timedelta as _td
                    import pytz as _pytz
                    from .models import CupHandleCandidate as _CHC
                    from .utils import detect_cup_and_handle

                    User      = get_user_model()
                    user      = User.objects.get(pk=uid)
                    _now      = _dt.now(_pytz.utc)
                    _end_str  = _now.date().strftime('%Y-%m-%d')
                    _start    = (_now.date() - _td(days=600)).strftime('%Y-%m-%d')
                    _scan_run = _dt.now(_pytz.utc)

                    # เก็บ 3 รอบล่าสุด ลบเก่า
                    old_runs = list(
                        _CHC.objects.filter(user=user, market='US')
                        .values_list('scan_run', flat=True)
                        .distinct().order_by('-scan_run')[2:]
                    )
                    if old_runs:
                        _CHC.objects.filter(user=user, market='US', scan_run__in=old_runs).delete()

                    total   = len(sym_list)
                    results = []

                    def _scan_one(symbol):
                        try:
                            df = _yf.Ticker(symbol).history(start=_start, end=_end_str, interval='1d')
                            if df is None or df.empty or len(df) < 60:
                                return None
                            if isinstance(df.columns, _pd.MultiIndex):
                                df.columns = df.columns.droplevel(1)
                            df = df.dropna(subset=['Close', 'High', 'Low'])

                            # Liquidity filter - USD (≥$1M daily turnover)
                            avg_vol   = float(df['Volume'].tail(20).mean())
                            avg_price = float(df['Close'].tail(20).mean())
                            if avg_vol * avg_price < 1_000_000:
                                return None

                            pat = detect_cup_and_handle(df)
                            if pat is None:
                                return None

                            # Breakout volume confirmation (O'Neil rule: ≥1.5x avg on breakout day)
                            breakout_vol_ok = False
                            if pat['stage'] == 'breakout':
                                last_vol = float(df['Volume'].iloc[-1])
                                breakout_vol_ok = last_vol >= avg_vol * 1.5

                            # RS return (3M)
                            close_s   = df['Close'].dropna()
                            rs_return = None
                            if len(close_s) >= 66:
                                rs_return = float(
                                    (close_s.iloc[-1] - close_s.iloc[-66]) / abs(close_s.iloc[-66]) * 100
                                )

                            # ADX & RSI
                            adx_val = 0.0
                            rsi_val = 50.0
                            try:
                                adx_df = _ta.adx(df['High'], df['Low'], df['Close'], length=14)
                                if adx_df is not None and not adx_df.empty:
                                    col = [c for c in adx_df.columns if c.startswith('ADX_')]
                                    if col and _pd.notna(adx_df[col[0]].iloc[-1]):
                                        adx_val = float(adx_df[col[0]].iloc[-1])
                                rsi_s = _ta.rsi(df['Close'], length=14)
                                if rsi_s is not None and _pd.notna(rsi_s.iloc[-1]):
                                    rsi_val = float(rsi_s.iloc[-1])
                            except Exception:
                                pass

                            return {
                                'symbol': symbol, 'pat': pat,
                                'rs_return': rs_return, 'adx': adx_val,
                                'rsi': rsi_val, 'avg_vol': avg_vol,
                                'breakout_vol_ok': breakout_vol_ok,
                            }
                        except Exception:
                            return None

                    with _cf.ThreadPoolExecutor(max_workers=20) as ex:
                        futs = {ex.submit(_scan_one, s): s for s in sym_list}
                        done = 0
                        for fut in _cf.as_completed(futs):
                            done += 1
                            _c.set(ckey, {
                                'state': 'running', 'progress': done,
                                'total': total, 'phase': f'สแกน {done}/{total}...'
                            }, timeout=900)
                            res = fut.result()
                            if res:
                                results.append(res)

                    # RS percentile rank
                    rs_vals = {r['symbol']: r['rs_return'] for r in results if r['rs_return'] is not None}
                    rs_map  = {}
                    if rs_vals:
                        rs_ser = _pd.Series(rs_vals)
                        rs_map = (rs_ser.rank(pct=True) * 99).clip(0, 99).astype(int).to_dict()

                    # Save to DB
                    for res in results:
                        pat = res['pat']
                        _CHC.objects.create(
                            user=user, scan_run=_scan_run, market='US',
                            symbol=res['symbol'],
                            price=pat['current_price'],
                            cup_high=pat['cup_high'],
                            cup_low=pat['cup_low'],
                            cup_depth_pct=pat['cup_depth_pct'],
                            cup_length_days=pat['cup_length_days'],
                            cup_start_date=pat['cup_start_date'],
                            cup_end_date=pat['cup_end_date'],
                            handle_high=pat['handle_high'],
                            handle_low=pat['handle_low'],
                            handle_depth_pct=pat['handle_depth_pct'],
                            handle_length_days=pat['handle_length_days'],
                            handle_start_date=pat['handle_start_date'],
                            breakout_price=pat['breakout_price'],
                            target_price=pat['target_price'],
                            stop_loss=pat['stop_loss'],
                            risk_reward=pat['risk_reward'],
                            avg_vol_20d=res['avg_vol'],
                            cup_vol_confirmed=pat['cup_vol_confirmed'],
                            handle_vol_dry=pat['handle_vol_dry'],
                            breakout_vol_ok=res['breakout_vol_ok'],
                            stage=pat['stage'],
                            confidence_score=pat['confidence_score'],
                            rs_rating=rs_map.get(res['symbol'], 0),
                            adx=res['adx'],
                            rsi=res['rsi'],
                        )

                    _c.set(ckey, {'state': 'done'}, timeout=300)
                except Exception as exc:
                    import logging
                    logging.getLogger('stocks').exception(f'[US Cup&Handle] bg scan error: {exc}')
                    from django.core.cache import cache as _c2
                    _c2.set(ckey, {'state': 'done'}, timeout=300)

            _th.Thread(
                target=_run_us_cup_handle_bg,
                args=(user_id, cache_key, scan_symbols),
                daemon=True
            ).start()

        from django.shortcuts import redirect as _redir
        return _redir('stocks:us_cup_handle_scanner')

    # ── Display results ───────────────────────────────────────────
    from .models import CupHandleCandidate as _CHC

    all_runs = list(
        _CHC.objects.filter(user=request.user, market='US')
        .values_list('scan_run', flat=True)
        .distinct().order_by('-scan_run')
    )

    try:
        run_idx = int(request.GET.get('run_idx', 0))
    except (ValueError, TypeError):
        run_idx = 0
    run_idx = max(0, min(run_idx, len(all_runs) - 1)) if all_runs else 0

    candidates = []
    scanned_at = None
    if all_runs:
        selected_run = all_runs[run_idx]
        candidates   = list(_CHC.objects.filter(user=request.user, market='US', scan_run=selected_run))
        scanned_at   = selected_run

    stage_order = {'breakout': 0, 'ready': 1, 'handle': 2, 'forming': 3}
    candidates.sort(key=lambda c: (stage_order.get(c.stage, 9), -c.confidence_score))

    stage_counts = {
        'breakout': sum(1 for c in candidates if c.stage == 'breakout'),
        'ready':    sum(1 for c in candidates if c.stage == 'ready'),
        'handle':   sum(1 for c in candidates if c.stage == 'handle'),
        'forming':  sum(1 for c in candidates if c.stage == 'forming'),
    }

    stage_filter = request.GET.get('stage', 'all')
    try:
        min_conf = int(request.GET.get('min_conf', 0))
    except (ValueError, TypeError):
        min_conf = 0
    rs_only = request.GET.get('rs_only') == '1'

    if stage_filter != 'all':
        candidates = [c for c in candidates if c.stage == stage_filter]
    if min_conf > 0:
        candidates = [c for c in candidates if c.confidence_score >= min_conf]
    if rs_only:
        candidates = [c for c in candidates if c.rs_rating >= 70]

    for c in candidates:
        if c.breakout_price > 0 and c.price > 0:
            c.pct_to_breakout = round((c.breakout_price - c.price) / c.price * 100, 1)
            c.pct_to_breakout = max(0.0, c.pct_to_breakout)
        else:
            c.pct_to_breakout = 0.0
        if c.cup_high > c.cup_low:
            raw = (c.price - c.cup_low) / (c.cup_high - c.cup_low) * 100
            c.recovery_pct = round(min(100.0, max(0.0, raw)), 1)
        else:
            c.recovery_pct = 0.0

    forming_list    = [c for c in candidates if c.stage == 'forming']
    closest_forming = max(forming_list, key=lambda c: c.recovery_pct, default=None)

    context = {
        'candidates':        candidates,
        'has_scanned':       bool(all_runs),
        'scanned_at':        scanned_at,
        'all_runs':          all_runs,
        'selected_run_idx':  run_idx,
        'current_sort':      request.GET.get('sort', 'stage'),
        'stage_counts':      stage_counts,
        'stage_filter':      stage_filter,
        'min_conf':          min_conf,
        'rs_only':           rs_only,
        'closest_forming':   closest_forming,
        'stage_labels': {
            'breakout': ('Breakout',       '#16a34a'),
            'ready':    ('Ready to Break', '#2563eb'),
            'handle':   ('Handle Forming', '#d97706'),
            'forming':  ('Cup Forming',    '#64748b'),
        },
        'total_symbols': len(_US_MOMENTUM_SYMBOLS),
    }
    return render(request, 'stocks/us_cup_handle_scan.html', context)


@login_required
def scanner_guide(request):
    """
    แสดงคู่มือการใช้งานและการทำงานของ Scanner ทั้ง 3 รูปแบบ
    (Precision Momentum, Cup & Handle, Standard Momentum)
    """
    return render(request, 'stocks/scanner_guide.html')


@login_required
def crypto_hub(request):
    """
    หน้าต่างศูนย์กลางการวิเคราะห์ Cryptocurrency
    ดึงข้อมูลจาก Alternative.me (Fear & Greed Index) และข้อมูลราคา Real-time เบื้องต้น
    """
    import urllib.request
    import json
    
    # ── Fetch Fear & Greed Index ──
    fng_value = 50
    fng_value_classification = "Neutral"
    try:
        url = "https://api.alternative.me/fng/?limit=1"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode())
            if data and "data" in data and len(data["data"]) > 0:
                fng_value = int(data["data"][0]["value"])
                fng_value_classification = data["data"][0]["value_classification"]
    except Exception as e:
        print(f"Error fetching Fear and Greed Index: {e}")

    # กำหนดเหรียญสำคัญ
    crypto_symbols = [
        {'symbol': 'BTC-USD', 'name': 'Bitcoin'},
        {'symbol': 'ETH-USD', 'name': 'Ethereum'},
        {'symbol': 'SOL-USD', 'name': 'Solana'}
    ]

    context = {
        'fng_value': fng_value,
        'fng_classification': fng_value_classification,
        'crypto_symbols': crypto_symbols
    }
    
    return render(request, 'stocks/crypto_hub.html', context)


@login_required
def macro_playbook_view(request):
    """
    หน้าแสดงรายงาน Daily Mastermind Briefing (Playbook)
    โครงสร้างแบบ AJAX Loading เหมือนหน้า Crew Analysis อื่นๆ
    """
    return render(request, 'stocks/macro_playbook.html')


@login_required
def macro_playbook_run_ajax(request):
    """
    รัน CrewAI 5 Agents เบื้องหลังและเคลียร์ Cache เมื่อให้ผลลัพธ์
    พร้อมวิเคราะห์ Portfolio ปัจจุบันของผู้ใช้งาน
    """
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR
    from .models import Portfolio
    
    user_id = request.user.id
    cache_key = f'macro_playbook_{user_id}'

    # Check status
    if request.GET.get('status_check') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        return _JR(st)

    # Fetch user portfolio data before threading to avoid DB context issues in thread
    try:
        user_portfolio_list = list(Portfolio.objects.filter(user=request.user).values('symbol', 'quantity', 'entry_price', 'market'))
    except Exception:
        user_portfolio_list = []

    def _run_bg(portfolio_data):
        from .crew_analysis import MacroPlaybookCrew
        try:
            _cp.set(cache_key, {'state': 'running', 'phase': 'Agents กำลังประชุมและดึงราคาตลาด...'}, timeout=600)
            crew = MacroPlaybookCrew(portfolio_data=portfolio_data)
            result = crew.run_analysis()
            _cp.set(cache_key, {'state': 'done', 'result': result}, timeout=600)
        except Exception as exc:
            _cp.set(cache_key, {'state': 'done', 'result': f'## Error\nเกิดข้อผิดพลาดในการรัน Agents: {exc}'}, timeout=600)

    # Start if idle
    current_state = _cp.get(cache_key, {}).get('state', 'idle')
    if current_state == 'idle' or request.GET.get('force') == '1':
        _cp.set(cache_key, {'state': 'running', 'phase': 'กำลังเรียกทีมผู้เชี่ยวชาญ...'}, timeout=600)
        _th.Thread(target=_run_bg, args=(user_portfolio_list,), daemon=True).start()
        
    return _JR({'status': 'started'})

# ======================================================================
# Turtle Trader Scanner (Mechanical Breakout System)
# ======================================================================

@login_required
def turtle_scanner(request):
    """
    หน้าแสดงผลการสแกนด้วยระบบ Turtle Trading
    - System 1: Breakout 20-day high (Exit: 10-day low)
    - System 2: Breakout 55-day high (Exit: 20-day low)
    """
    from .models import TurtleScanCandidate, PrecisionScanCandidate
    
    market = request.GET.get('market', 'SET')
    candidates_qs = TurtleScanCandidate.objects.filter(user=request.user, market=market)
    
    candidates = []
    if candidates_qs.exists():
        latest_run = candidates_qs.order_by('-scan_run').values_list('scan_run', flat=True).first()
        candidates = list(candidates_qs.filter(scan_run=latest_run).order_by('symbol'))
        last_updated = latest_run
        
        prec_qs = PrecisionScanCandidate.objects.filter(user=request.user, market=market)
        latest_prec_run = prec_qs.values_list('scan_run', flat=True).order_by('-scan_run').first()
        if latest_prec_run:
            prec_dict = {p.symbol: p for p in prec_qs.filter(scan_run=latest_prec_run)}
            for c in candidates:
                p_match = prec_dict.get(c.symbol)
                if p_match:
                    c.technical_score = p_match.technical_score
                    c.rs_rating = p_match.rs_rating
                    c.launcher_score = p_match.launcher_score
                else:
                    c.technical_score = None
                    c.rs_rating = None
                    c.launcher_score = None
    else:
        last_updated = None

    context = {
        'candidates': candidates,
        'last_updated': last_updated,
        'selected_market': market,
        'title': "Turtle Trader Scanner"
    }
    return render(request, 'stocks/turtle_scanner.html', context)


@login_required
def turtle_scanner_run_ajax(request):
    """
    รัน Turtle Scanner เบื้องหลัง 
    (ดึงข้อมูลย้อนหลัง 3-6 เดือนเพื่อหาสถิติ 20-day, 55-day High/Low)
    """
    from .models import ScannableSymbol, TurtleScanCandidate
    import concurrent.futures as _cf
    import pandas as _pd
    import yfinance as _yf
    from django.utils import timezone as _tz
    import threading as _th
    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR
    import time as _time
    import random as _random

    user_id = request.user.id
    market_param = request.GET.get('market', 'SET')
    ckey = f'turtle_scan_{user_id}_{market_param}'

    c_state = _cp.get(ckey, {'state': 'idle'})
    if request.GET.get('status_check') == '1':
        return _JR(c_state)

    if c_state.get('state') == 'running' and request.GET.get('force') != '1':
        return _JR({'status': 'started'})

    # Get symbols
    precision_only = request.GET.get('precision_only') == 'true'
    
    if precision_only:
        from .models import PrecisionScanCandidate
        # Get latest precision run
        latest_prec = PrecisionScanCandidate.objects.filter(user=request.user, market=market_param).order_by('-scan_run').first()
        if not latest_prec:
            return _JR({'state': 'done', 'error': 'ไม่พบข้อมุลจากหน้า Precision Scan กรุณาสแกน Precision ก่อนเปิดโหมดกรองคุณภาพ'})
        
        sym_list = list(PrecisionScanCandidate.objects.filter(
            user=request.user, 
            market=market_param, 
            scan_run=latest_prec.scan_run,
            technical_score__gte=75
        ).values_list('symbol', flat=True))
        
        if not sym_list:
            return _JR({'state': 'done', 'error': 'ไม่พบหุ้นที่มีคะแนน > 75 ในฐานข้อมูล Precision ล่าสุด'})
    else:
        from .utils import get_top_ranked_symbols
        sym_list = get_top_ranked_symbols(market=market_param, limit=200)
        
    if not sym_list:
        return _JR({'state': 'done', 'error': f'ไม่พบหุ้นใน watchlist สำหรับตลาด {market_param}'})

    def _bg_task(syms, market):
        from django.contrib.auth import get_user_model as _GUM
        from .models import PrecisionScanCandidate # ย้ายมาไว้ตรงนี้เพื่อให้ใน Thread มองเห็น
        user = _GUM().objects.get(pk=user_id)

        # Auto-refresh market cap rankings daily (SET only)
        if market == 'SET':
            try:
                from .utils import get_top_ranked_symbols as _GTRS
                new_syms = _GTRS(market='SET', limit=200, auto_refresh=True)
                if new_syms: # ป้องกันกรณี refresh แล้วได้ค่าว่าง
                    syms = new_syms
            except Exception:
                pass 

        scan_time = _tz.now()
        total_syms = len(syms)
        _cp.set(ckey, {'state': 'running', 'progress': 0, 'total': total_syms}, timeout=3600)
        
        results = []
        processed = 0

        # --- STAGE 1: Systematic Analysis (Detailed Scan) ---
        # ข้าม Stage 1 (YahooQuery) เพื่อความเร็วและป้องกันการค้าง
        candidates = [{'symbol': s} for s in syms]
        total_cand = len(candidates)
        _cp.set(ckey, {'state': 'running', 'progress': 0, 'total': total_cand, 'phase': f'🐢 กำลังเริ่มวิเคราะห์หุ้น {total_cand} ตัว...'}, timeout=3600)
        # กำหนดขนาดกลุ่มข้อมูลตามตลาด
        c_size = 20 if market == 'US' else 50
        for i in range(0, len(candidates), c_size):
            chunk = candidates[i : i + c_size]
            chunk_syms = [c['symbol'] for c in chunk]
            chunk_bk = [f"{s}.BK" if market == 'SET' and '.' not in s else s for s in chunk_syms]
            
            _cp.set(ckey, {'state': 'running', 'progress': i, 'total': total_cand, 'phase': f'กำลังวิเคราะห์กลุ่มที่ {i//c_size + 1}...'}, timeout=3600)
            
            try:
                # Use yfinance with auto_adjust
                data = yf.download(chunk_bk, period="1y", interval="1d", progress=False, group_by='ticker', threads=True, timeout=30, auto_adjust=True)
                
                for symbol in chunk_syms:
                    try:
                        s_bk = f"{symbol}.BK" if market == 'SET' and '.' not in symbol else symbol
                        df = None
                        if s_bk in data and not data[s_bk].empty:
                            df = data[s_bk].dropna(subset=['Close'])
                        
                        # Fallback: ถ้าดึงแบบกลุ่มไม่ได้ ให้ดึงรายตัว (สำหรับ US ที่มักโดนบล็อก)
                        if df is None or df.empty:
                            try:
                                t_obj = yf.Ticker(s_bk)
                                df = t_obj.history(period="1y", interval="1d")
                                if df is not None and not df.empty:
                                    df = df.dropna(subset=['Close'])
                            except Exception: pass

                        if df is None or df.empty or len(df) < 55: continue
                        
                        # Debug: เห็นกันชัดๆ ว่าหุ้นตัวไหนกำลังถูกตรวจ
                        # print(f"--- Analyzing {symbol} ---")
                        
                        avg_vol = float(df['Volume'].tail(20).mean())
                        avg_val = (df['Close'] * df['Volume']).tail(20).mean()
                        
                        # ยกเลิก Liquidity filter ชั่วคราว หรือปรับให้ต่ำสุดๆ เพื่อเช็คผล
                        if market == 'SET' and avg_val < 100_000: continue 
                        
                        # Calculate Turtle
                        df['High_20'] = df['High'].rolling(20).max().shift(1)
                        df['Low_10'] = df['Low'].rolling(10).min().shift(1)
                        df['High_55'] = df['High'].rolling(55).max().shift(1)
                        df['Low_20'] = df['Low'].rolling(20).min().shift(1)
                        
                        # -- Stage 2 Analysis (v2) --
                        df['SMA150'] = df['Close'].rolling(150).mean()
                        df['SMA200'] = df['Close'].rolling(200).mean()
                        
                        # -- ADX Calculation --
                        df['H_L'] = df['High'] - df['Low']
                        df['H_PC'] = abs(df['High'] - df['Close'].shift(1))
                        df['L_PC'] = abs(df['Low'] - df['Close'].shift(1))
                        df['TR'] = df[['H_L', 'H_PC', 'L_PC']].max(axis=1)
                        
                        # ATR
                        df['ATR_20'] = df['TR'].ewm(span=20, adjust=False).mean()

                        # Weinsteins Stage 2
                        last_row = df.iloc[-1]
                        current_close = float(last_row['Close'])
                        h20 = float(last_row.get('High_20', 0) or 0)
                        h55 = float(last_row.get('High_55', 0) or 0)
                        atr = float(last_row.get('ATR_20', 0) or 0)
                        l10 = float(last_row.get('Low_10', 0) or 0)
                        l20 = float(last_row.get('Low_20', 0) or 0)
                        sma150 = float(last_row.get('SMA150', 0) or 0)
                        sma200 = float(last_row.get('SMA200', 0) or 0)
                        is_stage2 = current_close > sma150 and sma150 > sma200

                        # --- Just Broke (Expanded window to 10 days) ---
                        window = df.tail(10)
                        sys1_days_ago = None
                        sys2_days_ago = None
                        for d_ago, (_, row) in enumerate(window.iloc[::-1].iterrows()):
                            rh20 = float(row.get('High_20', 0) or 0)
                            rh55 = float(row.get('High_55', 0) or 0)
                            rc   = float(row['Close'])
                            if sys1_days_ago is None and rh20 > 0 and rc >= rh20:
                                sys1_days_ago = d_ago
                            if sys2_days_ago is None and rh55 > 0 and rc >= rh55:
                                sys2_days_ago = d_ago

                        sys1 = sys1_days_ago is not None
                        sys2 = sys2_days_ago is not None
                        sys1_near = (not sys1) and h20 > 0 and current_close >= h20 * 0.97
                        sys2_near = (not sys2) and h55 > 0 and current_close >= h55 * 0.97

                        pct_to_20d = round((current_close - h20) / h20 * 100, 2) if h20 > 0 else None
                        pct_to_55d = round((current_close - h55) / h55 * 100, 2) if h55 > 0 else None

                        if sys1 or sys2 or sys1_near or sys2_near:
                            p_score = None
                            rs_rat = None
                            p_match = PrecisionScanCandidate.objects.filter(user=user, symbol=symbol, market=market).order_by('-scan_run').first()
                            if p_match:
                                p_score = p_match.technical_score
                                rs_rat = p_match.rs_rating
                            
                            # -- ปรับเกณฑ์ Elite ให้เข้มงวดขึ้น (โดยเฉพาะ US) --
                            ps_val = p_score if p_score is not None else 0
                            rs_val = rs_rat if rs_rat is not None else 0
                            
                            if market == 'US':
                                # US: ต้องผ่านเกณฑ์ทั้ง RS >= 85 และ Score >= 80 (AND)
                                is_elite = is_stage2 and (rs_val >= 85 and ps_val >= 80)
                            else:
                                # SET: ต้องผ่านเกณฑ์ทั้ง RS >= 80 และ Score >= 75 (AND)
                                is_elite = is_stage2 and (rs_val >= 80 and ps_val >= 75)

                            results.append(TurtleScanCandidate(
                                user=user, scan_run=scan_time, symbol=symbol, market=market,
                                price=current_close,
                                sys1_breakout=sys1, sys1_days_ago=sys1_days_ago,
                                high_20d=round(h20, 2), low_10d=round(l10, 2),
                                sys2_breakout=sys2, sys2_days_ago=sys2_days_ago,
                                high_55d=round(h55, 2), low_20d=round(l20, 2),
                                sys1_near=sys1_near, sys2_near=sys2_near,
                                pct_to_20d=pct_to_20d, pct_to_55d=pct_to_55d,
                                avg_vol_20d=avg_vol, atr_20d=round(atr, 4),
                                adx=0.0, # Simplified
                                stage2=is_stage2,
                                technical_score=p_score, rs_rating=rs_rat,
                                is_elite=is_elite
                            ))
                    except Exception as e: 
                        print(f"Error analyzing {symbol}: {e}")
                        continue
            except Exception as e: 
                print(f"Chunk Error: {e}")
                continue

        # Save to DB — always delete old + save new so scan_run timestamp always updates
        TurtleScanCandidate.objects.filter(user=user, market=market).delete()
        if results:
            TurtleScanCandidate.objects.bulk_create(results)

        _cp.set(ckey, {'state': 'done', 'found': len(results)}, timeout=3600)

    _cp.set(ckey, {'state': 'running', 'progress': 0, 'total': len(sym_list)}, timeout=3600)
    _th.Thread(target=_bg_task, args=(sym_list, market_param), daemon=True).start()

    return _JR({'status': 'started'})


# ---------------------------------------------------------------------------
# Stock Chart View - Turtle Breakout + Momentum
# ---------------------------------------------------------------------------

@login_required
def stock_chart(request, symbol):
    market = request.GET.get('market', 'SET')
    context = {
        'symbol': symbol.upper(),
        'market': market,
    }
    return render(request, 'stocks/stock_chart.html', context)


@login_required
def stock_chart_data(request, symbol):
    import json as _json
    import yfinance as _yf
    import pandas as _pd
    import numpy as _np
    from django.http import JsonResponse as _JR

    symbol = symbol.upper()
    market = request.GET.get('market', 'SET')
    period = request.GET.get('period', '1y')

    # Append .BK for SET stocks
    yf_symbol = symbol + '.BK' if market == 'SET' else symbol
    
    # Padding: ดึงข้อมูลเผื่อล่วงหน้าเพื่อให้เส้น Donchian 55 และ RSI มีค่าเพียงพอ
    download_period = '2y' if period in ('1y', '2y') else '1y' 

    def _safe_val(val, default=0.0):
        try:
            v = float(val)
            if _np.isnan(v) or _np.isinf(v):
                return default
            return v
        except:
            return default

    try:
        df = _yf.download(yf_symbol, period=download_period, auto_adjust=True,
                          progress=False, group_by='column')
        
        if df is None or df.empty:
            return _JR({'error': f'ไม่พบข้อมูลสำหรับ {yf_symbol} (yfinance returned empty)'}, status=404)

        # 1. จัดการ MultiIndex (yfinance 0.2.x คืน (Price,Ticker) หรือ (Ticker,Price))
        if isinstance(df.columns, _pd.MultiIndex):
            _pf = {'Open', 'High', 'Low', 'Close', 'Volume', 'Adj Close'}
            if any(v in _pf for v in df.columns.get_level_values(0)):
                df.columns = df.columns.get_level_values(0)
            else:
                df.columns = df.columns.get_level_values(1)

        # 2. ปรับชื่อคอลัมน์ให้เป็นมาตรฐาน (Standardize Capitalization)
        df.columns = [str(c).capitalize() for c in df.columns]
        
        # คืนค่าคอลัมน์ที่จำเป็น (Handle Adj Close/Close redundancy)
        if 'Adj close' in df.columns:
            df.rename(columns={'Adj close': 'Close'}, inplace=True)
        elif 'Adj close' not in df.columns and 'Close' not in df.columns:
            # บางกรณี yfinance คืนค่า 'Regular market price' หรืออื่นๆ
            potential_close = [c for c in df.columns if 'close' in c.lower()]
            if potential_close:
                df.rename(columns={potential_close[0]: 'Close'}, inplace=True)

        required = ['Open', 'High', 'Low', 'Close', 'Volume']
        missing = [c for c in required if c not in df.columns]
        if missing:
            return _JR({'error': f'ไม่สามารถดึงข้อมูลที่จำเป็นได้: {missing}'}, status=500)

        df = df[required].copy()
        
        # 3. จัดการข้อมูลว่างเฉพาะจุด (หลีกเลี่ยงการ dropna ล้างทั้งแถว)
        # เติมค่าว่างด้วยวิธี ffill สำหรับราคาส่วน volume เป็น 0
        df[['Open', 'High', 'Low', 'Close']] = df[['Open', 'High', 'Low', 'Close']].ffill()
        df['Volume'] = df['Volume'].fillna(0)
        
        df.index = _pd.to_datetime(df.index)
        df.sort_index(inplace=True)

        # Donchian Channel 20 & 55
        df['dc10_upper'] = df['High'].rolling(10).max()
        df['dc10_lower'] = df['Low'].rolling(10).min()
        df['dc20_upper'] = df['High'].rolling(20).max()
        df['dc20_lower'] = df['Low'].rolling(20).min()
        df['dc55_upper'] = df['High'].rolling(55).max()
        df['dc55_lower'] = df['Low'].rolling(55).min()

        # RSI 14
        delta = df['Close'].diff()
        gain = delta.clip(lower=0).rolling(14).mean()
        loss = (-delta.clip(upper=0)).rolling(14).mean()
        rs = gain / loss.replace(0, _np.nan)
        df['rsi'] = 100 - (100 / (1 + rs))

        # --- Trend Following: EMA ---
        df['ema9']   = df['Close'].ewm(span=9, adjust=False).mean()
        df['ema20']  = df['Close'].ewm(span=20, adjust=False).mean()
        df['ema50']  = df['Close'].ewm(span=50, adjust=False).mean()
        df['ema200'] = df['Close'].ewm(span=200, adjust=False).mean()

        # --- Momentum: MACD ---
        exp12 = df['Close'].ewm(span=12, adjust=False).mean()
        exp26 = df['Close'].ewm(span=26, adjust=False).mean()
        df['macd']    = exp12 - exp26
        df['macd_sig'] = df['macd'].ewm(span=9, adjust=False).mean()
        df['macd_hist'] = df['macd'] - df['macd_sig']

        # --- Volatility: Bollinger Bands ---
        ma20 = df['Close'].rolling(window=20).mean()
        std20 = df['Close'].rolling(window=20).std()
        df['bb_upper'] = ma20 + (std20 * 2)
        df['bb_lower'] = ma20 - (std20 * 2)

        # --- Tactical Analysis (N and Next Units) ---
        # Calculate ATR (Wilder's method)
        df['h_l'] = df['High'] - df['Low']
        df['h_pc'] = (df['High'] - df['Close'].shift(1)).abs()
        df['l_pc'] = (df['Low'] - df['Close'].shift(1)).abs()
        df['tr'] = df[['h_l', 'h_pc', 'l_pc']].max(axis=1)
        df['atr_20'] = df['tr'].ewm(alpha=1/20, adjust=False).mean() # N for 20 days
        
        last_row = df.iloc[-1]
        n_val = round(float(last_row['atr_20']), 4)
        curr_price = float(last_row['Close'])
        
        # 🚥 Strategic Signal Logic
        def get_signal(buy_cond, sell_cond):
            if buy_cond: return 'BUY'
            if sell_cond: return 'SELL'
            return 'WAIT'

        # Short: DC10 + RSI
        short_buy = curr_price >= float(last_row['dc10_upper']) and float(last_row['rsi']) > 50
        short_sell = curr_price <= float(last_row['dc10_lower'])
        
        # ⏱️ Breakout Age Analysis (User Suggestion)
        breakout_age = 0
        dist_from_breakout = 0
        if short_buy:
            # Find how many candles ago it first broke DC10 upper
            past_df = df.iloc[:-1].iloc[::-1] # All except last, reversed
            for i, (idx, row) in enumerate(past_df.iterrows()):
                if row['Close'] >= row['dc10_upper']:
                    breakout_age += 1
                else:
                    break
            dist_from_breakout = round(((curr_price - float(last_row['dc10_upper'])) / float(last_row['dc10_upper'])) * 100, 2)

        # Medium: DC20 + EMA200
        med_buy = curr_price >= float(last_row['dc20_upper']) and curr_price > float(last_row['ema200'])
        med_sell = curr_price <= float(last_row['dc10_lower'])
        
        # Long: DC55
        long_buy = curr_price >= float(last_row['dc55_upper'])
        long_sell = curr_price <= float(last_row['dc20_lower'])

        tactical = {
            'price': round(_safe_val(curr_price), 2),
            'n': _safe_val(n_val),
            'signals': {
                'short': get_signal(short_buy, short_sell),
                'medium': get_signal(med_buy, med_sell),
                'long': get_signal(long_buy, long_sell)
            },
            'breakout_age': breakout_age,
            'breakout_dist': _safe_val(dist_from_breakout),
            'short_term_high': round(_safe_val(last_row['dc10_upper']), 2),
            'high_20d': round(_safe_val(last_row['dc20_upper']), 2),
            'high_55d': round(_safe_val(last_row['dc55_upper']), 2),
            'rsi': round(_safe_val(last_row['rsi']), 2),
            'ema9': round(_safe_val(last_row['ema9']), 2),
            'ema200': round(_safe_val(last_row['ema200']), 2),
            # Added missing fields for Tactical Command Center
            'exit_10d_low': round(_safe_val(last_row['dc10_lower']), 2),
            'exit_20d_low': round(_safe_val(last_row['dc20_lower']), 2),
            'next_unit': round(_safe_val(curr_price + (0.5 * n_val)), 2),
        }

        # Turtle breakout signals (compare close vs previous day's channel)
        df['sys1_signal'] = df['Close'] >= df['dc20_upper'].shift(1)
        df['sys2_signal'] = df['Close'] >= df['dc55_upper'].shift(1)
        df['sys1_exit']   = df['Close'] <= df['dc10_lower'].shift(1)

        def datestr(dt):
            return _pd.Timestamp(dt).strftime('%Y-%m-%d')

        candles, vol, rsi_data = [], [], []
        dc20u, dc20l, dc55u, dc55l = [], [], [], []
        ema20, ema50, ema200 = [], [], []
        bbu, bbl = [], []
        macd, macd_sig, macd_hist = [], [], []
        signals = []

        for dt, row in df.iterrows():
            t = datestr(dt)
            o = _safe_val(row['Open'])
            h = _safe_val(row['High'])
            l = _safe_val(row['Low'])
            c = _safe_val(row['Close'])
            candles.append({'time': t, 'open': round(o, 2), 'high': round(h, 2),
                            'low': round(l, 2), 'close': round(c, 2)})
            vol.append({'time': t, 'value': int(_safe_val(row['Volume'])),
                        'color': '#26a69a' if c >= o else '#ef5350'})

            if _pd.notna(row['dc20_upper']):
                dc20u.append({'time': t, 'value': round(float(row['dc20_upper']), 2)})
                dc20l.append({'time': t, 'value': round(float(row['dc20_lower']), 2)})
            if _pd.notna(row['dc55_upper']):
                dc55u.append({'time': t, 'value': round(float(row['dc55_upper']), 2)})
                dc55l.append({'time': t, 'value': round(float(row['dc55_lower']), 2)})
            
            if _pd.notna(row['rsi']):
                rsi_data.append({'time': t, 'value': round(float(row['rsi']), 2)})
            
            if _pd.notna(row['ema20']):
                ema20.append({'time': t, 'value': round(float(row['ema20']), 2)})
            if _pd.notna(row['ema50']):
                ema50.append({'time': t, 'value': round(float(row['ema50']), 2)})
            if _pd.notna(row['ema200']):
                ema200.append({'time': t, 'value': round(float(row['ema200']), 2)})
            
            if _pd.notna(row['bb_upper']):
                bbu.append({'time': t, 'value': round(float(row['bb_upper']), 2)})
                bbl.append({'time': t, 'value': round(float(row['bb_lower']), 2)})
            
            if _pd.notna(row['macd']):
                macd.append({'time': t, 'value': round(float(row['macd']), 2)})
                macd_sig.append({'time': t, 'value': round(float(row['macd_sig']), 2)})
                macd_hist.append({'time': t, 'value': round(float(row['macd_hist']), 2)})

            if row['sys1_signal']:
                signals.append({'time': t, 'type': 'sys1_buy', 'price': round(float(row['Close']), 2)})
            if row['sys2_signal']:
                signals.append({'time': t, 'type': 'sys2_buy', 'price': round(float(row['Close']), 2)})
            if row['sys1_exit']:
                signals.append({'time': t, 'type': 'sys1_exit', 'price': round(float(row['Close']), 2)})

        return JsonResponse({
            'symbol': symbol, 'market': market, 'tactical': tactical,
            'candles': candles, 'volume': vol, 'rsi': rsi_data,
            'dc20_upper': dc20u, 'dc20_lower': dc20l,
            'dc55_upper': dc55u, 'dc55_lower': dc55l,
            'ema20': ema20, 'ema50': ema50, 'ema200': ema200,
            'bb_upper': bbu, 'bb_lower': bbl,
            'macd': macd, 'macd_sig': macd_sig, 'macd_hist': macd_hist,
            'signals': signals
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def trading_accounts_view(request):
    """
    หน้าจอจัดการบัญชีเทรด (List & Add)
    """
    from .models import TradingAccount, BrokerType
    
    if request.method == 'POST':
        broker = request.POST.get('broker')
        acc_id = request.POST.get('account_id')
        key    = request.POST.get('api_key', '')
        secret = request.POST.get('api_secret', '')
        
        TradingAccount.objects.create(
            user=request.user,
            broker=broker,
            account_id=acc_id,
            api_key=key,
            api_secret=secret
        )
        return redirect('stocks:trading_accounts')

    accounts = TradingAccount.objects.filter(user=request.user)
    return render(request, 'stocks/trading_account_list.html', {
        'accounts': accounts,
        'broker_types': BrokerType.choices
    })

@login_required
def delete_trading_account_view(request, pk):
    from .models import TradingAccount
    acc = get_object_or_404(TradingAccount, pk=pk, user=request.user)
    acc.delete()
    return redirect('stocks:trading_accounts')

@csrf_exempt
@login_required
def sync_trading_account_ajax(request, pk):
    """
    ดึงยอดเงินล่าสุดจาก Broker มาอัปเดต (AJAX)
    """
    from .models import TradingAccount
    from .trading_bridge import RobotBridge
    
    acc = get_object_or_404(TradingAccount, pk=pk, user=request.user)
    bridge = RobotBridge(user=request.user, account=acc)
    
    try:
        success = bridge.sync_account_balance()
        if success:
            return JsonResponse({
                'success': True,
                'balance': float(acc.balance),
                'equity': float(acc.equity),
                'currency': acc.currency
            })
        else:
            return JsonResponse({'success': False, 'error': 'API returned failure. Check if account is connected in MetaApi.'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def get_gold_positions_ajax(request):
    """
    ดึงรายการออเดอร์ทองที่เปิดอยู่
    """
    from .trading_bridge import RobotBridge
    bridge = RobotBridge(user=request.user)
    try:
        positions = bridge.get_open_positions()
        return JsonResponse({'success': True, 'positions': positions})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@login_required
def close_all_gold_positions_ajax(request):
    """
    สั่งปิดออเดอร์ทองทั้งหมด
    """
    from .trading_bridge import RobotBridge
    bridge = RobotBridge(user=request.user)
    try:
        success = bridge.close_all_positions()
        return JsonResponse({'success': success})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@login_required
def execute_gold_trade_ajax(request):
    """
    รับคำสั่งจากปุ่มเทรดในหน้า Gold Command Center
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Post required'}, status=400)
    
    import json
    from .trading_bridge import RobotBridge
    
    try:
        data = json.loads(request.body)
        symbol = data.get('symbol', 'GC=F')
        side   = data.get('side', 'BUY')
        price  = data.get('price')
        sl     = data.get('sl')
        tp     = data.get('tp')
        volume = data.get('volume', 0.01)
        strategy = data.get('strategy', 'Manual')

        # เรียกใช้ RobotBridge
        bridge = RobotBridge(user=request.user)
        order = bridge.execute_trade(
            symbol=symbol,
            side=side,
            volume=volume,
            price=price,
            sl=sl,
            tp=tp,
            strategy=strategy
        )

        return JsonResponse({
            'success': True,
            'order_id': order.order_id,
            'message': f'ส่งคำสั่ง {side} เรียบร้อยแล้ว (ID: {order.order_id})'
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def refresh_market_caps_view(request):
    """
    Manual trigger to refresh market caps for all SET symbols.
    """
    from .utils import refresh_market_caps
    from .models import ScannableSymbol
    from django.contrib import messages
    from django.shortcuts import redirect
    from django.utils import timezone
    
    # ⚡ High-speed check: ดึงแค่ timestamp ล่าสุดมาดูค่าเดียว + แปลงเป็น BKK timezone ก่อนเช็ค
    last_update_dt = ScannableSymbol.objects.filter(is_active=True, market='SET', market_cap__gt=0)\
                                          .values_list('last_cap_update', flat=True).first()
    today = timezone.localtime(timezone.now()).date()
    
    # ต้องแปลง UTC → Bangkok timezone ก่อนเปรียบเทียบวันที่ มิฉะนั้นจะ off ได้ถึง 7 ชั่วโมง
    already_today = (
        last_update_dt is not None and
        timezone.localtime(last_update_dt).date() == today
    )

    if already_today:
        messages.info(request, "ข้อมูลอันดับ Market Cap ของวันนี้อัปเดตเรียบร้อยแล้วครับ สแกนต่อได้ทันที!")
    else:
        count = refresh_market_caps()
        messages.success(request, f"สำเร็จ! อัปเดตข้อมูล Market Cap หุ้นไทยแล้ว {count} ตัว ระบบพร้อมจัดอันดับ Top 200 เพื่อสแกนแล้วครับ")
    
    next_url = request.GET.get('next') or 'stocks:momentum_scanner'
    return redirect(next_url)

@login_required
def gold_trading(request):
    """
    Gold Trading & Robot Command Center (XAU/USD).
    """
    from .models import TradingAccount
    account = TradingAccount.objects.filter(user=request.user, is_active=True).first()
    capital = float(account.equity or account.balance) if account else 100.0
    symbol = "GC=F"
    return render(request, 'stocks/gold_trading.html', {
        'symbol': symbol,
        'title': 'Gold Robot Command Center',
        'market': 'US',
        'capital': capital,
    })


@login_required
def debug_scan_symbol(request, symbol):
    """
    Debug endpoint — fetch data for one symbol and return full breakdown.
    Call: /stocks/debug-scan/CK/
    Shows exactly what the scanner computes (RSI, RVOL, ADX, score breakdown).
    Staff-only for safety.
    """
    from django.http import JsonResponse
    if not request.user.is_staff:
        return JsonResponse({'error': 'staff only'}, status=403)

    import yfinance as yf
    import pandas_ta as ta
    import pandas as pd
    from .utils import find_supply_demand_zones, find_supply_demand_zones_v2, analyze_momentum_technical_v2

    sym_bk = f"{symbol.upper()}.BK"
    result = {'symbol': symbol, 'sym_bk': sym_bk}

    try:
        df = yf.download(sym_bk, period='1y', interval='1d', progress=False, timeout=20)
        result['rows_fetched'] = len(df)
        result['is_multindex'] = isinstance(df.columns, pd.MultiIndex)

        if isinstance(df.columns, pd.MultiIndex):
            df.columns = df.columns.droplevel(1)

        if df.empty or len(df) < 55:
            result['error'] = f'Not enough data: {len(df)} rows'
            return JsonResponse(result)

        # Raw last bar
        result['last_close'] = float(df['Close'].iloc[-1])
        result['last_open']  = float(df['Open'].iloc[-1])
        result['last_vol']   = float(df['Volume'].iloc[-1])
        result['avg_vol_20'] = float(df['Volume'].tail(20).mean())
        result['rvol_raw']   = round(result['last_vol'] / result['avg_vol_20'], 3) if result['avg_vol_20'] > 0 else 0

        # Compute indicators
        df['EMA50']  = ta.ema(df['Close'], length=50)
        df['EMA200'] = ta.ema(df['Close'], length=200)
        df['EMA20']  = ta.ema(df['Close'], length=20)
        df['RSI']    = ta.rsi(df['Close'], length=14)
        adx_df = ta.adx(df['High'], df['Low'], df['Close'], length=14)
        if adx_df is not None:
            df = pd.concat([df, adx_df], axis=1)
        df['MFI'] = ta.mfi(df['High'], df['Low'], df['Close'], df['Volume'], length=14)

        result['rsi']    = round(float(df['RSI'].iloc[-1]), 2) if pd.notna(df['RSI'].iloc[-1]) else None
        result['ema20']  = round(float(df['EMA20'].iloc[-1]), 2) if pd.notna(df['EMA20'].iloc[-1]) else None
        result['ema50']  = round(float(df['EMA50'].iloc[-1]), 2) if pd.notna(df['EMA50'].iloc[-1]) else None
        result['ema200'] = round(float(df['EMA200'].iloc[-1]), 2) if pd.notna(df['EMA200'].iloc[-1]) else None
        result['adx']    = round(float(df['ADX_14'].iloc[-1]), 2) if 'ADX_14' in df.columns and pd.notna(df['ADX_14'].iloc[-1]) else None
        result['mfi']    = round(float(df['MFI'].iloc[-1]), 2) if pd.notna(df['MFI'].iloc[-1]) else None
        result['year_high'] = round(float(df['High'].tail(252).max()), 2)

        # Score conditions breakdown
        p = result['last_close']
        result['cond_above_ema200']   = bool(p > result['ema200']) if result['ema200'] else False
        result['cond_above_ema50']    = bool(p > result['ema50']) if result['ema50'] else False
        result['cond_golden_cross']   = bool(result['ema50'] > result['ema200']) if (result['ema50'] and result['ema200']) else False
        result['cond_ema20_aligned']  = bool(p > result['ema20'] > result['ema50'] > result['ema200']) if all([result['ema20'], result['ema50'], result['ema200']]) else False
        result['cond_rsi_ok']         = bool(55 <= (result['rsi'] or 0) <= 75)
        result['cond_near_52w_high']  = bool(p >= result['year_high'] * 0.85)
        result['cond_52w_breakout']   = bool(p >= result['year_high'] * 0.99)
        result['rvol_bullish']        = bool(p >= result['last_open'])

        # Full v2 score
        tech = analyze_momentum_technical_v2(df)
        result['v2_score'] = tech.get('score')
        result['v2_rvol']  = tech.get('rvol')
        result['v2_rsi']   = tech.get('rsi')

    except Exception as e:
        result['exception'] = str(e)

    return JsonResponse(result, json_dumps_params={'indent': 2})


@login_required
@require_POST
def portfolio_refresh_prices(request):
    """
    Lightweight price refresh — fetches current price via fast_info (parallel)
    and updates highest_price in Portfolio if price has risen.
    Returns JSON: { updated: [...], skipped: [...], errors: [...] }
    """
    from django.http import JsonResponse
    import concurrent.futures

    items = list(Portfolio.objects.filter(user=request.user, category='STOCK'))
    if not items:
        return JsonResponse({'updated': [], 'skipped': [], 'errors': []})

    def _fetch_price(item):
        symbol = item.symbol.upper()
        market = item.market  # 'SET', 'US', 'CRYPTO', 'OTHER'

        if market == 'SET':
            sym_yf = symbol if symbol.endswith('.BK') else f"{symbol}.BK"
        elif market == 'CRYPTO':
            # BTC → BTC-USD, BTC-USD → BTC-USD (ใช้ตามที่บันทึกไว้)
            sym_yf = symbol if '-' in symbol else f"{symbol}-USD"
        else:
            # US + OTHER: ใช้ symbol ตรงๆ (DELL, KEY, AAPL ฯลฯ)
            sym_yf = symbol.replace('.BK', '')

        try:
            fi = yf.Ticker(sym_yf).fast_info
            price = float(getattr(fi, 'last_price', None) or 0)
            return item, price
        except Exception:
            return item, 0.0

    updated, skipped, errors = [], [], []

    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as ex:
        futures = {ex.submit(_fetch_price, item): item for item in items}
        for fut in concurrent.futures.as_completed(futures):
            orig_item = futures[fut]
            try:
                item, price = fut.result(timeout=15)
                if price <= 0:
                    errors.append(item.symbol)
                    continue
                old_high = float(item.highest_price or 0)
                if price > old_high:
                    item.highest_price = price
                    item.save(update_fields=['highest_price'])
                    updated.append({'symbol': item.symbol, 'price': price, 'prev_high': old_high})
                else:
                    skipped.append({'symbol': item.symbol, 'price': price, 'highest': old_high})
            except Exception:
                errors.append(orig_item.symbol)

    return JsonResponse({'updated': updated, 'skipped': skipped, 'errors': errors})

@login_required
def get_bot_status_ajax(request):
    """
    ดึงสถานะล่าสุดของบอทที่รันบน Server มาโชว์ที่หน้าจอ UI
    โดยเช็คทั้งจาก Database และเช็ค Process จริงในระบบ
    """
    from .models import BotActivity
    from django.utils import timezone
    import datetime
    
    # 1. เช็คว่ามีไฟล์ PID และ Process ยังรันอยู่ไหม
    is_process_alive = False
    if os.path.exists(PID_FILE):
        try:
            with open(PID_FILE, 'r') as f:
                pid = int(f.read().strip())
            if os.name == 'nt':
                import ctypes
                handle = ctypes.windll.kernel32.OpenProcess(0x0400, False, pid)
                if handle:
                    is_process_alive = True
                    ctypes.windll.kernel32.CloseHandle(handle)
            else:
                os.kill(pid, 0)
                is_process_alive = True
        except:
            is_process_alive = False

    # 2. ดึงข้อมูลจากฐานข้อมูล
    try:
        activity = BotActivity.objects.get(bot_name="Gold Server Bot")
        diff = timezone.now() - activity.last_heartbeat
        # ถ้า DB บอกว่า Active และเวลาไม่เก่าเกินไป หรือ Process ในเครื่องยังรันอยู่
        is_active = (activity.status == "ACTIVE" and diff.total_seconds() < 300) or is_process_alive
        
        return JsonResponse({
            'status': "ACTIVE" if is_active else "OFFLINE",
            'last_heartbeat': activity.last_heartbeat.strftime('%H:%M:%S'),
            'message': activity.message,
            'is_alive': is_active,
            'process_running': is_process_alive,
            'debug_diff': diff.total_seconds()
        })
    except BotActivity.DoesNotExist:
        return JsonResponse({
            'status': "ACTIVE" if is_process_alive else "OFFLINE",
            'is_alive': is_process_alive,
            'process_running': is_process_alive
        })

import subprocess
import os
import signal

PID_FILE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'gold_bot.pid')

@login_required
def start_gold_bot_ajax(request):
    """สั่งเริ่มการทำงานของบอทใน Background"""
    if os.path.exists(PID_FILE):
        try:
            with open(PID_FILE, 'r') as f:
                pid = int(f.read().strip())
            
            # เช็คว่า Process ยังอยู่จริงไหม
            if os.name == 'nt':
                import ctypes
                PROCESS_QUERY_INFORMATION = 0x0400
                handle = ctypes.windll.kernel32.OpenProcess(PROCESS_QUERY_INFORMATION, False, pid)
                if handle:
                    ctypes.windll.kernel32.CloseHandle(handle)
                    return JsonResponse({'success': False, 'error': 'Bot is already running (Live Process)'})
            else:
                os.kill(pid, 0) # ส่งสัญญาณ 0 เพื่อเช็คว่า Process มีอยู่จริงไหม
                return JsonResponse({'success': False, 'error': 'Bot is already running (Live Process)'})
        except (ValueError, ProcessLookupError, OSError):
            # ถ้ามาถึงตรงนี้แปลว่า PID เก่าตายไปแล้ว ให้ลบไฟล์ทิ้งและรันใหม่
            if os.path.exists(PID_FILE):
                os.remove(PID_FILE)
    
    try:
        # สั่งรันบอทใน Background
        import sys
        python_exe = sys.executable
        log_dir = os.path.dirname(os.path.dirname(__file__))
        stdout_log = open(os.path.join(log_dir, 'bot_stdout.log'), 'a')
        stderr_log = open(os.path.join(log_dir, 'bot_stderr.log'), 'a')
        
        if os.name == 'nt':
            process = subprocess.Popen(
                [python_exe, 'manage.py', 'run_gold_bot'],
                creationflags=subprocess.CREATE_NEW_CONSOLE
            )
        else:
            # สำหรับ Linux (Ubuntu): บันทึก Log ลงไฟล์เพื่อตรวจสอบ
            process = subprocess.Popen(
                [python_exe, 'manage.py', 'run_gold_bot'],
                stdout=stdout_log,
                stderr=stderr_log,
                start_new_session=True,
                env=os.environ.copy()
            )
        
        # บันทึก PID ไว้สำหรับสั่งปิด
        with open(PID_FILE, 'w') as f:
            f.write(str(process.pid))
            
        return JsonResponse({'success': True, 'pid': process.pid})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def stop_gold_bot_ajax(request):
    """สั่งหยุดบอทโดยอ้างอิงจาก PID"""
    if not os.path.exists(PID_FILE):
        return JsonResponse({'success': False, 'error': 'No running bot found'})
    
    try:
        with open(PID_FILE, 'r') as f:
            pid = int(f.read())
        
        if os.name == 'nt':
            # Windows: ใช้ taskkill เพื่อปิดทั้ง process tree
            subprocess.run(['taskkill', '/F', '/T', '/PID', str(pid)])
        else:
            os.kill(pid, signal.SIGTERM)
            
        os.remove(PID_FILE)
        
        # อัปเดตสถานะใน DB ด้วย
        from .models import BotActivity
        BotActivity.objects.filter(bot_name="Gold Server Bot").update(status="STOPPED", message="Bot stopped by user via UI")
        
        return JsonResponse({'success': True})
    except Exception as e:
        if os.path.exists(PID_FILE): os.remove(PID_FILE)
        return JsonResponse({'success': False, 'error': str(e)})

# ====== Investment Dashboard (Premium Insights) ======

@login_required
def investment_dashboard(request):
    """
    หน้า Dashboard หลักสำหรับนักลงทุน Premium
    แสดง Top 5 SET และ US พร้อมบทวิเคราะห์ AI แบบถาวร
    """
    insight_id = request.GET.get('id')
    insights = InvestmentDashboardInsight.objects.filter(user=request.user).order_by('-created_at')[:3]
    
    if insight_id:
        latest_insight = InvestmentDashboardInsight.objects.filter(user=request.user, id=insight_id).first()
    else:
        latest_insight = insights[0] if insights else None

    # Fetch user's watchlist for toggle state
    from .models import ScanWatchlistItem
    watchlist_set = set(ScanWatchlistItem.objects.filter(user=request.user).values_list('symbol', flat=True))

    return render(request, 'stocks/investment_dashboard.html', {
        'insight': latest_insight,
        'insights_history': insights,
        'watchlist_set': watchlist_set,
    })

@login_required
@require_POST
def investment_dashboard_refresh(request):
    """
    ระบบคัดกรองหุ้นแบบ Multi-Scanner Funnel:
    Cup & Handle (Setup) -> Precision Momentum (Power) -> Minervini SEPA (Quality) -> Turtle Breakout (Trigger)
    """
    import json
    from django.db.models import Max
    from .models import (
        CupHandleCandidate, PrecisionScanCandidate, USSepaCandidate, TurtleScanCandidate, InvestmentDashboardInsight
    )
    from google import genai
    from django.conf import settings
    from django.contrib import messages
    from django.shortcuts import redirect
    
    def get_consensus_top_10(market):
        # 1. หา scan_run ล่าสุดของแต่ละระบบ
        latest_prec   = PrecisionScanCandidate.objects.filter(market=market).aggregate(Max('scan_run'))['scan_run__max']
        latest_ch     = CupHandleCandidate.objects.filter(market=market).aggregate(Max('scan_run'))['scan_run__max']
        latest_turtle = TurtleScanCandidate.objects.filter(market=market).aggregate(Max('scan_run'))['scan_run__max']
        latest_sepa   = USSepaCandidate.objects.aggregate(Max('scan_run'))['scan_run__max'] if market == 'US' else None

        if not latest_prec and not latest_ch: return []

        # 2. Batch-fetch ข้อมูลทุกระบบพร้อมกัน (4 queries เท่านั้น แทน N*4)
        prec_map   = {c.symbol: c for c in PrecisionScanCandidate.objects.filter(market=market, scan_run=latest_prec)} if latest_prec else {}
        ch_map     = {c.symbol: c for c in CupHandleCandidate.objects.filter(market=market, scan_run=latest_ch)} if latest_ch else {}
        turtle_map = {c.symbol: c for c in TurtleScanCandidate.objects.filter(market=market, scan_run=latest_turtle)} if latest_turtle else {}
        sepa_map   = {c.symbol: c for c in USSepaCandidate.objects.filter(scan_run=latest_sepa)} if latest_sepa else {}

        # 3. รวม universe แล้ว sort ตัวอักษร → iteration order คงที่ทุกครั้ง
        all_symbols = sorted(set(prec_map) | set(ch_map) | set(turtle_map) | set(sepa_map))

        ch_lane    = []
        other_lane = []

        for sym in all_symbols:
            prec   = prec_map.get(sym)
            ch     = ch_map.get(sym)
            turtle = turtle_map.get(sym)
            sepa   = sepa_map.get(sym)

            score  = 0
            badges = []

            # Step 1: The Radar (Cup & Handle) (+40)
            if ch:
                score += 40
                if ch.stage in ['ready', 'handle']: score += 10 # Bonus for late-stage setup
                badges.append('RADAR')

            # Step 2: The Power (Precision Momentum) (+30)
            if prec:
                score += 30
                if prec.rs_rating >= 80: score += 15 # The Power bonus (RS > 80)
                badges.append('POWER')

            # Step 3: The Quality (SEPA / Stage 2) (+20)
            is_quality = bool(sepa) if market == 'US' else bool(prec and prec.stage2)
            if is_quality:
                score += 20
                badges.append('QUALITY')

            # Step 4: The Trigger (Turtle Breakout) (+20)
            if turtle:
                if turtle.sys1_breakout or turtle.sys2_breakout:
                    score += 20
                    badges.append('TRIGGER')
                elif turtle.sys1_near or turtle.sys2_near:
                    score += 5
                    badges.append('NEAR')

            if score < 25: continue

            entry = {
                'symbol':         sym,
                'price':          float((prec or ch or turtle).price),
                'total_score':    score,
                'badges':         badges,
                'sector':         (prec.sector if prec else (ch.sector if ch else (getattr(turtle, 'sector', None) or 'Unknown'))),
                'technical_score': prec.technical_score if prec else (ch.confidence_score if ch else 0),
                'rs_rating':      prec.rs_rating if prec else (ch.rs_rating if ch else 0),
                'cup_stage':      ch.stage if ch else "None",
                'turtle_breakout': "YES" if (turtle and (turtle.sys1_breakout or turtle.sys2_breakout)) else "No",
                'vdu':  prec.vdu_near_zone if prec else False,
                'vcp':  bool((prec and prec.vcp_setup) or (ch and ch.handle_vol_dry)),
                'is_explosive': bool(prec and prec.is_explosive),
            }

            if ch:
                ch_lane.append(entry)
            elif score >= 45:
                other_lane.append(entry)

        # 4. Sort deterministic: score ↓, technical_score ↓, rs_rating ↓, symbol ↑
        def sort_key(x):
            return (-x['total_score'], -x['technical_score'], -x['rs_rating'], x['symbol'])

        ch_lane.sort(key=sort_key)
        other_lane.sort(key=sort_key)

        return (ch_lane + other_lane)[:10]

    set_top = get_consensus_top_10('SET')
    us_top = get_consensus_top_10('US')
    
    if not set_top and not us_top:
        messages.warning(request, "ไม่พบข้อมูลสแกนที่สอดคล้องกับ Funnel ในขณะนี้ กรุณารัน Scanner ให้ครบถ้วน")
        return redirect('stocks:investment_dashboard')

    set_summary = json.dumps(set_top, indent=2)
    us_summary = json.dumps(us_top, indent=2)

    prompt = f"""
คุณคือ Senior Quantitative Strategist วิเคราะห์หุ้นด้วยระบบ Funnel หลายระบบ
กฎการวิเคราะห์: Cup & Handle (Radar) -> Precision (Power) -> SEPA (Quality) -> Turtle (Trigger)

ข้อมูลหุ้น TOP 10 ที่ผ่านการคัดกรอง Confluence สูงสุด:
[SET]: {set_summary}
[US]: {us_summary}

เขียนรายงานวิเคราะห์เป็นภาษาไทย โดยใช้โครงสร้าง Markdown ดังนี้:

## 🔍 Funnel Consensus
สรุปว่ารอบนี้หุ้นส่วนใหญ่ติดสแกนในขั้นตอนไหนมากที่สุด และ Confluence ของตลาดเป็นอย่างไร

---

## 💎 High Conviction Picks
| ตลาด | หุ้น | เหตุผล (ระบุระบบที่ติด) | ความเชื่อมั่น |
|------|------|-----------------------|------------|
| SET | ... | ... | สูง/ปานกลาง |
| US | ... | ... | สูง/ปานกลาง |

---

## 🎯 Risk Management
- **Stop Loss:** ระบุแนวทางตาม ATR/Low ของฐาน
- **Trailing Stop:** ใช้ระบบ Turtle (10D/20D Low) หรือ Minervini

---

## 💡 Strategist's View
มุมมองรวมของตลาด ควรรุกหรือรับ (2-3 ประโยค)
"""
    
    ai_strategy = "ระบบ AI ไม่สามารถประมวลผลได้ในขณะนี้"
    market_outlook = "Neutral"
    try:
        api_key = getattr(settings, "GEMINI_API_KEY", None)
        if api_key:
            client = genai.Client(api_key=api_key)
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            ai_strategy = response.text or ai_strategy
            _txt = ai_strategy.lower()
            if any(w in _txt for w in ["bullish", "ขาขึ้น", "รุก", "ซื้อ"]): market_outlook = "Bullish"
            elif any(w in _txt for w in ["bearish", "ขาลง", "ระวัง", "ขาย"]): market_outlook = "Bearish"
            else: market_outlook = "Sideways/Neutral"
    except Exception as e:
        ai_strategy = f"เกิดข้อผิดพลาด: {str(e)}"
        
    # เก็บเฉพาะ 3 ครั้งล่าสุดเพื่อประหยัดพื้นที่
    InvestmentDashboardInsight.objects.create(
        user=request.user, set_top_stocks=set_top, us_top_stocks=us_top,
        ai_strategy=ai_strategy, market_outlook=market_outlook, is_active=True
    )
    
    # ลบตัวที่เก่ากว่า 3 อันดับแรก
    all_ids = list(InvestmentDashboardInsight.objects.filter(user=request.user).order_by('-created_at').values_list('id', flat=True))
    if len(all_ids) > 3:
        InvestmentDashboardInsight.objects.filter(id__in=all_ids[3:]).delete()

    messages.success(request, "วิเคราะห์ข้อมูลแบบ Multi-System Funnel เรียบร้อยแล้ว")
    return redirect('stocks:investment_dashboard')
