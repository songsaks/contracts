from .base import * 

from .base import (
    _get_usd_thb, _compute_signals, _get_market_condition, _get_precision_scan_data,
    _US_SECTOR_MAP, _US_MOMENTUM_SYMBOLS, _build_us_symbol_set, _is_us_symbol,
    _seed_us_symbols, _seed_value_symbols, _score_value_candidate, _check_rate_limit,
    _is_commodity, _fetch_commodity_macro, _score_commodity_signal
)

@login_required
def dashboard(request):
    """
    แสดงรายการ Watchlist ของผู้ใช้พร้อมราคาปัจจุบัน, % เปลี่ยนแปลง
    และค่า RSI 14 วัน คำนวณแบบ Real-time ผ่าน yfinance
    """
    watchlist = Watchlist.objects.filter(user=request.user)
    items = []
    import pandas_ta as ta

    from stocks.utils import analyze_momentum_technical
    for item in watchlist:
        try:
            t = yf.Ticker(item.symbol)
            hist = t.history(period="1y")

            # Fallback: try alternate symbol if empty
            if hist.empty:
                alt_sym = f"{item.symbol}.BK" if ".BK" not in item.symbol else item.symbol.replace(".BK", "")
                t = yf.Ticker(alt_sym)
                hist = t.history(period="1y")

            current = None
            change = 0
            if not hist.empty:
                if isinstance(hist.columns, pd.MultiIndex):
                    hist.columns = [col[0] for col in hist.columns]
                hist = hist.loc[:, ~hist.columns.duplicated()]
                current = float(hist['Close'].iloc[-1])
                if len(hist) >= 2:
                    prev = float(hist['Close'].iloc[-2])
                    change = ((current - prev) / prev * 100) if prev else 0
            if not current:
                try:
                    info = t.info
                    if isinstance(info, dict):
                        current = info.get('currentPrice') or info.get('regularMarketPrice') or info.get('previousClose')
                        change = info.get('regularMarketChangePercent', 0)
                except:
                    pass

            rsi_val = None
            rsi_status = "Neutral"
            if not hist.empty and len(hist) >= 14:
                rsi_series = ta.rsi(hist['Close'], length=14)
                if not rsi_series.empty:
                    rsi_val = rsi_series.iloc[-1]
                    if rsi_val < 30: rsi_status = "Oversold"
                    elif rsi_val > 70: rsi_status = "Overbought"

            # ดึงข้อมูลจาก PrecisionScanCandidate ก่อน (ตรงกับ Scanner ทุกค่า)
            clean_symbol = item.symbol.split('.')[0].upper()
            from stocks.models import PrecisionScanCandidate
            prec_data = (PrecisionScanCandidate.objects
                         .filter(user=request.user, symbol=clean_symbol)
                         .order_by('-scan_run').first())

            if prec_data:
                class QuickMom: pass
                mom_data = QuickMom()
                mom_data.technical_score      = prec_data.technical_score
                mom_data.rvol                 = prec_data.rvol
                mom_data.rvol_bullish         = prec_data.rvol_bullish
                mom_data.adx                  = prec_data.adx
                mom_data.rsi                  = prec_data.rsi
                mom_data.erc_volume_confirmed = prec_data.erc_volume_confirmed
                mom_data.risk_reward_ratio    = prec_data.risk_reward_ratio
                mom_data.demand_zone_start    = prec_data.demand_zone_start
                mom_data.demand_zone_end      = prec_data.demand_zone_end
                mom_data.supply_zone_start    = prec_data.supply_zone_start
                mom_data.stop_loss            = prec_data.stop_loss
                mom_data.zone_proximity       = prec_data.zone_proximity
                mom_data.year_high            = prec_data.year_high
                mom_data.price_pattern        = prec_data.price_pattern
                mom_data.price_pattern_score  = prec_data.price_pattern_score
                mom_data.rel_momentum_1m      = prec_data.rel_momentum_1m
                mom_data.rel_momentum_3m      = prec_data.rel_momentum_3m
            else:
                # Fallback: คำนวณ on-the-fly ด้วย v2
                from stocks.utils import analyze_momentum_technical_v2
                mom_data = None
                if not hist.empty:
                    tech_analysis = analyze_momentum_technical_v2(hist)
                    if tech_analysis and tech_analysis.get('score', 0) > 0:
                        class QuickMom: pass
                        mom_data = QuickMom()
                        mom_data.technical_score      = tech_analysis['score']
                        mom_data.rvol                 = tech_analysis.get('rvol', 1.0)
                        mom_data.rvol_bullish         = tech_analysis.get('rvol_bullish', True)
                        mom_data.adx                  = 0
                        mom_data.rsi                  = float(rsi_val or 0)
                        mom_data.erc_volume_confirmed = False
                        mom_data.year_high            = 0
                        mom_data.price_pattern        = ''
                        mom_data.price_pattern_score  = 0
                        mom_data.rel_momentum_1m      = 0.0
                        mom_data.rel_momentum_3m      = 0.0
                        sd = tech_analysis.get('sd_zone')
                        if sd and sd.get('start') and sd['start'] > 0:
                            mom_data.risk_reward_ratio = sd.get('rr_ratio', 0)
                            mom_data.demand_zone_start = sd['start']
                            mom_data.demand_zone_end   = sd.get('end', 0)
                            mom_data.supply_zone_start = sd.get('target', 0)
                            mom_data.stop_loss         = sd.get('stop_loss', None)
                            mom_data.zone_proximity    = 0 if (current and current <= sd['start']) else ((float(current or 0) - sd['start']) / sd['start']) * 100
                        else:
                            mom_data.risk_reward_ratio = 0
                            mom_data.demand_zone_start = 0
                            mom_data.demand_zone_end   = 0
                            mom_data.supply_zone_start = 0
                            mom_data.stop_loss         = None
                            mom_data.zone_proximity    = 999

            signals = _compute_signals(mom_data, current) if mom_data else {'buy_score': 0, 'sell_score': 0, 'exit_signal': ''}

            # Heuristic market detection
            mkt = 'SET'
            if '.BK' in item.symbol: mkt = 'SET'
            elif item.category == AssetCategory.CRYPTO: mkt = 'CRYPTO'
            elif '-' in item.symbol and item.category != AssetCategory.CRYPTO: mkt = 'US'
            elif '.' not in item.symbol and '=' not in item.symbol and '-' not in item.symbol:
                mkt = 'US'
            elif '=' in item.symbol: mkt = 'OTHER' # Commodities usually use =F

            items.append({
                'obj': item,
                'price': current,
                'change': change,
                'rsi': rsi_val,
                'rsi_status': rsi_status,
                'mom_data': mom_data,
                'buy_score': signals['buy_score'],
                'sell_score': signals['sell_score'],
                'exit_signal': signals['exit_signal'],
                'symbol_base': item.symbol.split('.')[0],
                'market': mkt,
            })

        except:
            items.append({'obj': item, 'price': 'Error', 'change': 0, 'rsi': None, 'rsi_status': 'Error', 'mom_data': None})

    # --- สรุปข้อมูลสำหรับ Real Dashboard ---
    from stocks.models import Portfolio, SoldStock
    owned_assets = Portfolio.objects.filter(user=request.user)
    sold_assets = SoldStock.objects.filter(user=request.user)
    
    usd_thb = _get_usd_thb()
    total_val_thb = 0
    total_cost_thb = 0
    
    # คำนวณ Realized P/L (กำไรที่ขายไปแล้ว) ตาม Logic เดียวกับหน้า Report
    us_set = _build_us_symbol_set(request.user)
    total_realized_pl = 0
    for s in sold_assets:
        is_us = (s.market == MarketType.US) if s.market else _is_us_symbol(s.symbol, us_set)
        
        if hasattr(s, 'profit_loss_thb') and s.profit_loss_thb is not None:
            pl_thb = float(s.profit_loss_thb)
        else:
            pl_thb = float(s.profit_loss or 0) * usd_thb if is_us else float(s.profit_loss or 0)
        total_realized_pl += pl_thb
    
    set_val = 0
    us_val = 0
    crypto_val = 0
    
    # สร้าง map ของราคาเพื่อความรวดเร็ว
    price_map = {x['obj'].symbol: x['price'] for x in items if x.get('price') is not None and isinstance(x.get('price'), (int, float))}
    
    for p in owned_assets:
        curr_p = price_map.get(p.symbol, 0)
        if curr_p == 0:
            try:
                import yfinance as yf
                t = yf.Ticker(p.symbol)
                curr_p = t.fast_info['lastPrice']
            except: curr_p = float(p.entry_price or 0)
            
        # Ensure values are not None before conversion
        q_val = float(p.quantity or 0)
        p_val = float(curr_p or 0)
        
        val = q_val * p_val
        cost = q_val * float(p.entry_price or 0)
        
        if p.market == 'US' or p.market == 'CRYPTO':
            total_val_thb += val * usd_thb
            total_cost_thb += cost * usd_thb
            if p.market == 'US': us_val += val * usd_thb
            else: crypto_val += val * usd_thb
        else:
            total_val_thb += val
            total_cost_thb += cost
            set_val += val
            
    total_unrealized_pl = total_val_thb - total_cost_thb
    unrealized_pl_pct = (total_unrealized_pl / total_cost_thb * 100) if total_cost_thb else 0
    
    # ข้อมูลสำหรับกราฟ Performance (Realized P/L สะสม)
    recent_transactions = sold_assets.order_by('-sold_at')[:5]
    
    sorted_items = sorted([x for x in items if isinstance(x['change'], (int, float))], key=lambda x: x['change'], reverse=True)
    top_gainers = sorted_items[:3]
    top_losers = sorted_items[-3:][::-1] if len(sorted_items) > 3 else []

    context = {
        'items': items[:6], 
        'total_items': len(items),
        'categories': AssetCategory.choices,
        'market_types': MarketType.choices,
        'summary': {
            'total_value': total_val_thb,
            'unrealized_pl': total_unrealized_pl,
            'realized_pl': total_realized_pl,
            'total_net_pl': total_unrealized_pl + total_realized_pl,
            'pl_pct': unrealized_pl_pct,
            'asset_alloc': {
                'SET': set_val,
                'US': us_val,
                'CRYPTO': crypto_val,
                'total_count': owned_assets.count()
            }
        },
        'recent_transactions': recent_transactions,
        'top_gainers': top_gainers,
        'top_losers': top_losers,
        'usd_thb': usd_thb
    }
    return render(request, 'stocks/dashboard.html', context)

# ====== Analyze - วิเคราะห์หุ้นรายตัวด้วย AI (Gemini) ======

@login_required
def analyze(request, symbol):
    """
    ดึงข้อมูลหุ้นจาก yfinance + yahooquery และส่งให้ Gemini AI วิเคราะห์
    ผลการวิเคราะห์จะถูกแคชไว้ใน AnalysisCache เพื่อใช้ซ้ำได้
    แสดงกราฟราคา 90 วัน, ข่าวล่าสุด, และข้อมูลพื้นฐาน
    """
    # ── market=SET guard: redirect M → M.BK so yfinance fetches Thai stock ──
    market_param = request.GET.get('market', '')
    if market_param == 'SET' and not symbol.upper().endswith('.BK'):
        from django.http import HttpResponseRedirect
        qs = request.GET.urlencode()
        return HttpResponseRedirect(f"/stocks/analyze/{symbol.upper()}.BK/?{qs}")

    # ====== 1. Check Cache First (to prevent 504 Time-out) ======
    cache_timeout_hours = 12
    cached_analysis = AnalysisCache.objects.filter(user=request.user, symbol=symbol).first()
    
    if cached_analysis:
        from django.utils import timezone
        now = timezone.now()
        # ถ้าแคชยังไม่เกิน 12 ชม. ให้ลองใช้ของเดิม (ถ้าไม่ได้กด Force Refresh)
        if (now - cached_analysis.last_updated).total_seconds() < (cache_timeout_hours * 3600):
            # ดึงข้อมูลหุ้นแบบเร็วเพื่อแสดงกราฟและราคาปัจจุบัน
            try:
                data = get_stock_data(symbol)
                history = data.get('history', pd.DataFrame())
                
                # ถ้าดึงข้อมูลสำเร็จ ให้แสดงหน้าเว็บด้วยข้อมูลจาก Cache ได้เลย
                context = {
                    'symbol': symbol,
                    'analysis': cached_analysis.analysis_data,
                    'data': data,
                    'history': history,
                    # ... (ข้อมูลอื่นๆ ที่จำเป็นต้องใช้ใน template)
                }
                # หมายเหตุ: เพื่อความปลอดภัย ผมจะให้มันรันต่อไปด้านล่างก่อนถ้าข้อมูลพื้นฐาน (data) จำเป็นต้องใช้ 
                # แต่จะข้ามเฉพาะส่วนการเรียก AI ที่หนักที่สุด
            except: pass

    try:
        # ดึงข้อมูลหุ้นทั้งหมดผ่าน utility function
        data = get_stock_data(symbol)
        # ====== Fetch Extra Context from Cache (Value Stock data) ======
        extra_ctx = ""
        cached_lists = ['THAI_REC_ALL', 'US_REC_ALL']
        for list_sym in cached_lists:
            c = AnalysisCache.objects.filter(user=request.user, symbol=list_sym).first()
            if c:
                try:
                    s_list = json.loads(c.analysis_data)
                    match = next((s for s in s_list if s['symbol'] == symbol), None)
                    if match:
                        extra_ctx = f"Legendary Score: {match.get('value_score')}\n"
                        extra_ctx += f"Pillars: {match.get('legendary')}\n"
                        extra_ctx += f"Fair Value (Estimated): {match.get('fair_value')}\n"
                        extra_ctx += f"Upside (%): {match.get('upside')}\n"
                        break
                except: pass

        # ====== Fetch Extra Context from Momentum Scanner (Technical data) ======
        mom = MomentumCandidate.objects.filter(user=request.user, symbol_bk=symbol).first()
        if not mom:
            # ลองค้นหาด้วย symbol แบบไม่มี .BK
            clean_sym = symbol.replace('.BK', '')
            mom = MomentumCandidate.objects.filter(user=request.user, symbol=clean_sym).first()

        if mom:
            extra_ctx += f"\n[Technical Momentum Analysis]:\n"
            extra_ctx += f"Momentum Score: {mom.technical_score}/100\n"
            extra_ctx += f"RVOL: {mom.rvol}x, ADX: {mom.adx}, MFI: {mom.mfi}\n"
            if mom.entry_strategy:
                extra_ctx += f"Entry Strategy: {mom.entry_strategy}\n"
                extra_ctx += f"Demand Zone: {mom.demand_zone_start} - {mom.demand_zone_end}\n"
                extra_ctx += f"Target (Supply): {mom.supply_zone_start}\n"
                extra_ctx += f"Stop Loss: {mom.stop_loss}\n"
                extra_ctx += f"RR Ratio: {mom.risk_reward_ratio}\n"

        # ====== Fetch Mean Reversion context (when arriving from MR Scanner) ======
        mr_context = None
        from_mr = request.GET.get('from_mr') == '1'
        if from_mr:
            try:
                from stocks.models import MeanReversionCandidate as _MRC_A
                _clean = symbol.replace('.BK', '')
                _mr = _MRC_A.objects.filter(user=request.user, symbol=_clean).order_by('-scan_run').first()
                if _mr:
                    mr_context = {
                        'direction':   _mr.direction,
                        'rsi':         _mr.rsi,
                        'adx':         _mr.adx,
                        'rvol':        _mr.rvol,
                        'pattern':     _mr.pattern,
                        'support':     _mr.support_level,
                        'resistance':  _mr.resistance_level,
                        'mean_target': _mr.mean_target,
                        'dist_sup':    _mr.dist_to_support_pct,
                        'upside_pct':  _mr.upside_pct,
                        'r_score':     _mr.r_score,
                        'rs_rating':   _mr.rs_rating,
                    }
            except Exception:
                pass

        # ====== Pre-fetch macro data & compute signal for commodities ======
        # Done BEFORE AI call so we can pass signal into the prompt
        history = data.get('history', pd.DataFrame())
        is_commodity = _is_commodity(symbol)
        macro_data   = {}
        macro_signal = None
        if is_commodity:
            macro_data = _fetch_commodity_macro()
            _ema200 = float(history['EMA_200'].iloc[-1]) if 'EMA_200' in history.columns and not history.empty else None
            _rsi    = history['RSI'].iloc[-1]             if 'RSI'     in history.columns and not history.empty else None
            _price  = float(history['Close'].iloc[-1])    if not history.empty else 0
            macro_signal = _score_commodity_signal(symbol, _price, _ema200, _rsi, macro_data)
            # Stash raw macro inside signal so AI function can re-use without double-fetching
            macro_signal['_raw_macro'] = macro_data

        # ====== 2. AI Analysis Logic (Skip if cached) ======
        analysis_text = None
        # ถ้ามาจาก MR Scanner และมี MR context ใหม่ → บังคับ regenerate เพื่อให้ AI วิเคราะห์ MR ด้วย
        _force_regen = from_mr and mr_context is not None
        if cached_analysis and not _force_regen:
            from django.utils import timezone
            if (timezone.now() - cached_analysis.last_updated).total_seconds() < (cache_timeout_hours * 3600):
                analysis_text = cached_analysis.analysis_data

        if not analysis_text:
            # ส่งข้อมูลให้ AI วิเคราะห์และรับผลเป็น Markdown (เฉพาะกรณีไม่มีแคชหรือแคชเก่า)
            analysis_text = analyze_with_ai(symbol, data, extra_context=extra_ctx,
                                            macro_signal=macro_signal, mr_context=mr_context)

        # ====== เตรียมข้อมูลกราฟราคาและวอลลุ่ม ======
        # Prepare Chart Data (Price & Volume)
        chart_labels = []
        chart_values = []
        chart_volumes = []
        if not history.empty:
            # แสดงเฉพาะ 90 วันล่าสุด
            history_subset = history.tail(90)
            chart_labels = [d.strftime('%Y-%m-%d') for d in history_subset.index]
            chart_values = [round(float(v), 2) for v in history_subset['Close'].values]
            chart_volumes = [int(v) for v in history_subset['Volume'].values]

        # ====== เตรียมข้อมูลข่าว - แปลง timestamp ให้อ่านได้ ======
        # Prepare News Data (Convert timestamp to readable)
        from datetime import datetime
        news_list = data.get('news', [])
        for n in news_list:
            if 'providerPublishTime' in n:
                try:
                    # รองรับทั้ง string (ISO format) และ int (Unix timestamp)
                    if isinstance(n['providerPublishTime'], str):
                        n['display_time'] = datetime.fromisoformat(n['providerPublishTime'].replace('Z', '+00:00'))
                    else:
                        n['display_time'] = datetime.fromtimestamp(n['providerPublishTime'])
                except Exception:
                    n['display_time'] = n['providerPublishTime']

        # ====== บันทึกผลวิเคราะห์ลงใน cache ของแต่ละ user ======
        # Cache it per user
        AnalysisCache.objects.update_or_create(
            user=request.user,
            symbol=symbol,
            defaults={'analysis_data': analysis_text}
        )

        # ====== คำนวณค่า RSI ล่าสุดเพื่อแสดงใน header ======
        # Prepare RSI
        current_rsi = history['RSI'].iloc[-1] if 'RSI' in history.columns and not history.empty else None
        rsi_status = "Neutral"
        if current_rsi:
            if current_rsi < 30: rsi_status = "Oversold"
            elif current_rsi > 70: rsi_status = "Overbought"

        # ====== เตรียมข้อมูลพื้นฐานการเงิน ======
        info = data.get('info') or {}
        if not isinstance(info, dict):
            info = {}

        # ทำ copy เพื่อไม่แก้ไข dict ต้นฉบับ
        info = info.copy() # Safe copy
        # แปลงค่าสัดส่วนพื้นฐาน (ROE, Dividend Yield, NPM) จากทศนิยมเป็นเปอร์เซ็นต์ (เฉพาะถ้าเป็นทศนิยม < 1)
        for key in ['returnOnEquity', 'dividendYield', 'profitMargins']:
            val = info.get(key)
            if isinstance(val, (int, float)):
                if abs(val) < 1.0:
                    info[key] = val * 100
                else:
                    info[key] = val

        # ====== คำนวณ D/E Ratio จาก Balance Sheet ======
        bs = data.get('balance_sheet')
        de_calculated = None
        if bs is not None and not bs.empty:
            try:
                col = bs.columns[0]
                # พยายามดึง Total Liabilities จากหลายชื่อ field ที่ yfinance อาจใช้
                tot_liab = bs.loc['Total Liabilities Net Minority Interest', col] if 'Total Liabilities Net Minority Interest' in bs.index else bs.loc['Total Liabilities', col]
                tot_eq = bs.loc['Stockholders Equity', col] if 'Stockholders Equity' in bs.index else bs.loc['Total Equity Gross Minority Interest', col]
                de_calculated = tot_liab / tot_eq
            except Exception:
                pass

        # อัปเดต D/E ใน info dict (ใช้ค่าจาก balance sheet ถ้ามี หรือ fallback เป็นค่าจาก yfinance)
        if de_calculated is not None:
            info['debtToEquity'] = de_calculated
        elif isinstance(info.get('debtToEquity'), (int, float)) if isinstance(info, dict) else False:
            # yfinance ส่งค่า D/E เป็น % (คูณ 100 มาแล้ว) ต้องหาร 100 กลับ
            info['debtToEquity'] = info['debtToEquity'] / 100

        # ====== คำนวณ Breakout / Resistance Levels ======
        fifty_two_week_high = history['High'].max() if not history.empty and 'High' in history.columns else None
        # แนวต้านย่อยในช่วง 20 วันล่าสุด
        recent_resistance = history['High'].tail(20).max() if not history.empty and 'High' in history.columns else None
        # แนวรับ/จุดตัดขาดทุนในช่วง 20 วันล่าสุด
        recent_support = history['Low'].tail(20).min() if not history.empty and 'Low' in history.columns else None
        curr_price = history['Close'].iloc[-1] if not history.empty and 'Close' in history.columns else info.get('currentPrice', 0)
        # ตรวจสอบว่าราคาปัจจุบันทะลุ 52-Week High หรือไม่ (Breakout Signal)
        is_breakout = (curr_price >= fifty_two_week_high) if (fifty_two_week_high and curr_price) else False

        context = {
            'symbol': symbol,
            'info': info,
            'analysis': analysis_text,
            'chart_labels': chart_labels,
            'chart_values': chart_values,
            'chart_volumes': chart_volumes,
            'fifty_two_week_high': fifty_two_week_high,
            'recent_resistance': recent_resistance,
            'recent_support': recent_support,
            'is_breakout': is_breakout,
            'current_rsi': current_rsi,
            'rsi_status': rsi_status,
            'news': news_list,
            'title': f"AI Analysis: {symbol}",
            'is_commodity': is_commodity,
            'macro_data': macro_data,
            'macro_signal': macro_signal,
        }
        return render(request, 'stocks/analysis.html', context)
    except Exception as e:
        messages.error(request, f"Error analyzing {symbol}: {str(e)}")
        return redirect('stocks:dashboard')

@login_required
def crew_analyze(request, symbol):
    """
    CrewAI Multi-Agent Deep Analysis - runs in background thread,
    progress polled via AJAX so the browser never times out.

    Optional GET params (from Portfolio page):
      ?entry_price=2.40&quantity=10000&gain_loss_pct=5.2&gain_loss=1200&market_value=25200
    """
    import threading as _th

    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'crew_analysis_{user_id}_{symbol}'

    # ── Portfolio context from GET params ────────────────────────────
    def _safe_float(val, default=0.0):
        try:
            return float(val) if val else default
        except (ValueError, TypeError):
            return default

    portfolio_context = {}
    if request.GET.get('entry_price'):
        portfolio_context = {
            'entry_price':   _safe_float(request.GET.get('entry_price')),
            'quantity':      _safe_float(request.GET.get('quantity')),
            'gain_loss_pct': _safe_float(request.GET.get('gain_loss_pct')),
            'gain_loss':     _safe_float(request.GET.get('gain_loss')),
            'market_value':  _safe_float(request.GET.get('market_value')),
        }
        ep_key = str(int(_safe_float(request.GET.get('entry_price')) * 100))
        cache_key = f'crew_analysis_{user_id}_{symbol}_p{ep_key}'

    # ── AJAX status poll - use ck param to find correct cache key ────
    if request.GET.get('crew_status') == '1':
        ck = request.GET.get('ck')
        poll_key = ck if ck else cache_key
        st = _cp.get(poll_key, {'state': 'idle'})
        return _JR(st)

    strategy_param = request.GET.get('strategy')
    market_param = request.GET.get('market')
    
    # ── Auto-Detect Market if not provided ──────────────────────────
    if not market_param:
        us_set = _build_us_symbol_set(request.user)
        if _is_us_symbol(symbol, us_set):
            market_param = 'US'
        else:
            market_param = 'SET'

    # ── Result ready (page reload after done) ────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'done':
        result = cached.get('result', '')
        _cp.delete(cache_key)
        data = get_stock_data(symbol)
        return render(request, 'stocks/crew_result.html', {
            'symbol':            symbol,
            'crew_result':       result,
            'info':              data.get('info', {}),
            'title':             f'CrewAI Deep Analysis: {symbol}',
            'loading':           False,
            'portfolio_context': portfolio_context,
        })

    # ── Background worker ────────────────────────────────────────────
    def _run_crew_bg(ckey, sym, pctx, strat, mkt):
        import concurrent.futures as _cf

        from django.core.cache import cache as _c

        from stocks.crew_analysis import MomentumCrew as _MC

        try:
            phase = 'วิเคราะห์ Portfolio + Technical…' if pctx else 'กำลังวิเคราะห์…'
            _c.set(ckey, {'state': 'running', 'phase': phase}, timeout=600)

            mc = _MC(sym, portfolio_context=pctx, strategy=strat, market=mkt)

            # Hard timeout: kill entire analysis after 90 seconds
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(mc.run_analysis)
                try:
                    result = future.result(timeout=120)
                except _cf.TimeoutError:
                    result = '## วิเคราะห์ไม่สำเร็จ\n\nหมดเวลา 90 วินาที กรุณาลองใหม่อีกครั้ง'

            _c.set(ckey, {'state': 'done', 'result': result}, timeout=600)
        except Exception as exc:
            _c.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=600)

    # Start only if not already running
    if not cached or cached.get('state') == 'idle':
        _cp.set(cache_key, {'state': 'running', 'phase': 'เริ่มต้น Multi-Agent…'}, timeout=600)
        _th.Thread(target=_run_crew_bg, args=(cache_key, symbol, portfolio_context, strategy_param, market_param), daemon=True).start()

    # ── Show loading page (fast - no heavy data fetch) ───────────────
    try:
        import yfinance as _yf
        _t = _yf.Ticker(symbol)
        _fi = _t.fast_info
        info = {
            'longName': getattr(_fi, 'display_name', None) or symbol,
            'currentPrice': getattr(_fi, 'last_price', None),
        }
    except Exception:
        info = {}

    return render(request, 'stocks/crew_result.html', {
        'symbol':            symbol,
        'info':              info,
        'title':             f'CrewAI Deep Analysis: {symbol}',
        'loading':           True,
        'portfolio_context': portfolio_context,
        'cache_key':         cache_key,   # pass to template for correct poll URL
    })

@login_required
def core_analyze(request, symbol):
    """
    "The Core Project" - Renaissance-inspired Multi-Agent Deep Analysis.
    Uses TheCoreCrew with Anomaly Hunter, Backtest Engineer, and Execution Decider.
    """
    import threading as _th

    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'core_analysis_{user_id}_{symbol}'

    # ── AJAX status poll ─────────────────────────────────────────────
    if request.GET.get('core_status') == '1':
        ck = request.GET.get('ck')
        poll_key = ck if ck else cache_key
        st = _cp.get(poll_key, {'state': 'idle'})
        return _JR(st)

    strategy_param = request.GET.get('strategy')
    market_param = request.GET.get('market', 'SET')

    # ── Result ready ─────────────────────────────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'done':
        result = cached.get('result', '')
        _cp.delete(cache_key)
        data = get_stock_data(symbol)
        return render(request, 'stocks/crew_result.html', {
            'symbol':      symbol,
            'crew_result': result,
            'info':        data.get('info', {}),
            'title':       f'The Core Project: {symbol}',
            'loading':     False,
        })

    # ── Background worker ────────────────────────────────────────────
    def _run_core_bg(ckey, sym, mkt):
        import concurrent.futures as _cf

        from django.core.cache import cache as _c

        from stocks.crew_analysis import TheCoreCrew as _TCC

        try:
            _c.set(ckey, {'state': 'running', 'phase': 'Anomaly Hunter กำลังสแกน WACC/PEGY…'}, timeout=600)
            
            # รันการวิเคราะห์เชิงลึก
            crew = _TCC(sym, market=mkt)
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(crew.run_analysis)
                try:
                    # Anomaly Hunter + Backtest ใช้เวลาพอสมควร (timeout 2 นาที)
                    result = future.result(timeout=120) 
                except _cf.TimeoutError:
                    result = '## วิเคราะห์ "The Core" ไม่สำเร็จ\n\nการวิเคราะห์เชิงลึกใช้เวลาเกินกำหนด กรุณาลองใหม่อีกครั้ง'

            _c.set(ckey, {'state': 'done', 'result': result}, timeout=600)
        except Exception as exc:
            _c.set(ckey, {'state': 'done', 'result': f'## Error\n\n{exc}'}, timeout=600)

    # Start Worker
    if not cached or cached.get('state') == 'idle':
        _cp.set(cache_key, {'state': 'running', 'phase': 'เริ่มต้นโครงวิเคราะห์ The Core…'}, timeout=600)
        _th.Thread(target=_run_core_bg, args=(cache_key, symbol, market_param), daemon=True).start()

    # Show loading
    return render(request, 'stocks/crew_result.html', {
        'symbol':    symbol,
        'title':     f'The Core Analysis: {symbol}',
        'loading':   True,
        'cache_key': cache_key,
        # Special polling parameter for the client JS to use Correct URL
        'is_core':   True, 
    })

# ====== Momentum Quick CrewAI Analysis (AJAX modal) ======

@login_required
def momentum_quick_analysis(request, symbol):
    """
    Short-term CrewAI multi-agent analysis for a momentum candidate.
    Returns JSON - designed to be called from a modal (no page reload).

    Flow:
      1. POST/GET → start background analysis → return {'state': 'running'}
      2. Poll ?mq_status=1 until {'state': 'done', 'result': '...'}
      3. Render markdown in modal via marked.js
    """
    import threading as _th

    from django.core.cache import cache as _cp
    from django.http import JsonResponse as _JR

    user_id   = request.user.id
    cache_key = f'mq_analysis_{user_id}_{symbol}'

    # ── Poll ─────────────────────────────────────────────────────────
    if request.GET.get('mq_status') == '1':
        st = _cp.get(cache_key, {'state': 'idle'})
        return _JR(st)

    # ── Already running ───────────────────────────────────────────────
    cached = _cp.get(cache_key)
    if cached and cached.get('state') == 'running':
        return _JR({'state': 'running'})

    # ── Return cached done result ─────────────────────────────────────
    if cached and cached.get('state') == 'done':
        _cp.delete(cache_key)
        return _JR({'state': 'done', 'result': cached.get('result', '')})

    # ── Collect scan data from Candidate models ───────────────
    scan_data = {}
    try:
        def _sf(val):
            try:
                return float(val) if val is not None else None
            except (TypeError, ValueError):
                return None

        cand = None
        from stocks.models import PrecisionScanCandidate as _PCM
        from stocks.models import USSepaCandidate as _USC
        from stocks.models import MomentumCandidate as _MCM
        
        p_cand = _PCM.objects.filter(user=request.user, symbol=symbol).order_by('-scan_run').first()
        if p_cand:
            cand = p_cand
        else:
            u_cand = _USC.objects.filter(user=request.user, symbol=symbol).order_by('-scan_run').first()
            if u_cand:
                cand = u_cand
            else:
                cand = _MCM.objects.filter(user=request.user, symbol=symbol).order_by('-id').first()
        
        if cand:
            scan_data = {
                'price':             _sf(cand.price),
                'technical_score':   cand.technical_score,
                'rsi':               _sf(cand.rsi),
                'adx':               _sf(cand.adx),
                'mfi':               _sf(cand.mfi),
                'rvol':              _sf(cand.rvol),
                'demand_zone_start': _sf(cand.demand_zone_start),
                'demand_zone_end':   _sf(cand.demand_zone_end),
                'supply_zone_start': _sf(cand.supply_zone_start),
                'supply_zone_end':   _sf(cand.supply_zone_end),
                'risk_reward_ratio': _sf(cand.risk_reward_ratio),
                'zone_proximity':    _sf(cand.zone_proximity),
                'eps_growth':        _sf(cand.eps_growth) or 0,
                'rev_growth':        _sf(cand.rev_growth) or 0,
                'sector':            cand.sector or 'N/A',
                'year_high':         _sf(cand.year_high),
                'upside_to_high':    _sf(cand.upside_to_high),
                'cmf':               _sf(getattr(cand, 'cmf', None)),
                'volume_surge':      _sf(getattr(cand, 'volume_surge', None)),
                'pocket_pivot':      bool(getattr(cand, 'pocket_pivot', False)),
                'vcp_setup':         bool(getattr(cand, 'vcp_setup', False)),
                'vcp_contractions':  int(getattr(cand, 'vcp_contractions', 0)),
                'vcp_tightness':     _sf(getattr(cand, 'vcp_tightness', 0.0)),
                'vcp_vdu':           bool(getattr(cand, 'vcp_vdu', False)),
            }
    except Exception:
        pass

    # ── Background worker ─────────────────────────────────────────────
    def _run_bg(ckey, sym, sd, mkt, u_id):
        from django.core.cache import cache as _c
        try:
            _c.set(ckey, {'state': 'running', 'phase': 'กำลังวิเคราะห์ด้วย 3 Expert Agents…'}, timeout=600)

            from stocks.crew_analysis import run_single_call_analysis as _rsc
            import concurrent.futures as _cf
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(_rsc, sym, sd, mkt)
                try:
                    result = future.result(timeout=60)   # 1 call = 60s max
                except _cf.TimeoutError:
                    result = '## หมดเวลาวิเคราะห์\n\nกรุณาลองใหม่อีกครั้ง'
            _c.set(ckey, {'state': 'done', 'result': result}, timeout=900)
            
            # Save to Database for persistence
            try:
                from django.contrib.auth.models import User
                from stocks.models import AnalysisCache
                usr = User.objects.get(id=u_id)
                AnalysisCache.objects.update_or_create(
                    user=usr, symbol=f'crewai_{sym}',
                    defaults={'analysis_data': result}
                )
            except Exception as e:
                pass
                
        except Exception as exc:
            from django.core.cache import cache as _c2
            _c2.set(ckey, {'state': 'done', 'result': f'## เกิดข้อผิดพลาด\n\n{exc}'}, timeout=60)

    market = cand.market if cand else 'SET'
    _cp.set(cache_key, {'state': 'running'}, timeout=600)
    _th.Thread(target=_run_bg, args=(cache_key, symbol, scan_data, market, user_id), daemon=True).start()
    return _JR({'state': 'running', 'cache_key': cache_key})


# ====== CrewAI Export - Word / PDF ======

@login_required
def crew_export_docx(request, symbol):
    """Export CrewAI analysis as a formatted Word document (.docx)"""
    if request.method != 'POST':
        from django.shortcuts import redirect
        return redirect('stocks:crew_analyze', symbol=symbol)

    import re
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    markdown_text = request.POST.get('markdown_content', '')
    company_name  = request.POST.get('company_name', symbol)

    # ── Document setup ──────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Theme colors ────────────────────────────────────────────────
    COLOR_TITLE   = RGBColor(0x0f, 0x17, 0x2a)   # navy
    COLOR_H1      = RGBColor(0x1e, 0x29, 0x3b)
    COLOR_H2      = RGBColor(0x1d, 0x40, 0xaf)   # blue-800
    COLOR_H3      = RGBColor(0x07, 0x89, 0x16)   # green
    COLOR_BODY    = RGBColor(0x1e, 0x29, 0x3b)
    COLOR_MUTED   = RGBColor(0x64, 0x74, 0x8b)

    def _set_font(run, size, bold=False, color=None, italic=False):
        run.font.name      = 'Sarabun'
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.italic    = italic
        if color:
            run.font.color.rgb = color

    def _para_spacing(para, before=0, after=6, line=None):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(before * 20))
        spacing.set(qn('w:after'),  str(after  * 20))
        if line:
            spacing.set(qn('w:line'),     str(line * 20))
            spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)

    def _add_border_bottom(para, color='1d40af', size=12):
        """Add bottom border to paragraph (used for h1/h2)"""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    str(size))
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _apply_inline(run_text, para, default_size=11, default_color=None):
        """Parse **bold**, *italic*, `code` inline markers into runs"""
        segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', run_text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                run = para.add_run(seg[2:-2])
                _set_font(run, default_size, bold=True, color=default_color or COLOR_BODY)
            elif seg.startswith('*') and seg.endswith('*'):
                run = para.add_run(seg[1:-1])
                _set_font(run, default_size, italic=True, color=default_color or COLOR_BODY)
            elif seg.startswith('`') and seg.endswith('`'):
                run = para.add_run(seg[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(default_size - 1)
                run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            else:
                run = para.add_run(seg)
                _set_font(run, default_size, color=default_color or COLOR_BODY)

    # ── Cover header ────────────────────────────────────────────────
    # Top accent bar (using shaded paragraph)
    accent = doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(0)
    run_accent = accent.add_run('▬' * 60)
    run_accent.font.color.rgb = COLOR_H2
    run_accent.font.size      = Pt(6)

    # Title
    title_p = doc.add_paragraph()
    _para_spacing(title_p, before=8, after=2)
    r1 = title_p.add_run(f'🤖  CrewAI Multi-Agent Analysis')
    _set_font(r1, 22, bold=True, color=COLOR_TITLE)

    # Symbol + company
    sym_p = doc.add_paragraph()
    _para_spacing(sym_p, before=0, after=2)
    r2 = sym_p.add_run(f'{symbol}')
    _set_font(r2, 28, bold=True, color=COLOR_H2)
    r3 = sym_p.add_run(f'  ·  {company_name}')
    _set_font(r3, 16, color=COLOR_MUTED)

    # Date line
    from django.utils import timezone as _tz
    date_p = doc.add_paragraph()
    _para_spacing(date_p, before=0, after=12)
    r4 = date_p.add_run(f'Generated: {_tz.now().strftime("%d %B %Y  %H:%M")}')
    _set_font(r4, 10, italic=True, color=COLOR_MUTED)

    # Divider
    div = doc.add_paragraph()
    _add_border_bottom(div, color='1d40af', size=16)
    _para_spacing(div, before=0, after=12)

    # ── Parse markdown ──────────────────────────────────────────────
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # H1
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            _add_border_bottom(p, color='0f172a', size=12)
            _para_spacing(p, before=14, after=4)
            run = p.add_run(text)
            _set_font(run, 18, bold=True, color=COLOR_H1)
            i += 1
            continue

        # H2
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            _add_border_bottom(p, color='1d40af', size=8)
            _para_spacing(p, before=12, after=3)
            run = p.add_run(text)
            _set_font(run, 14, bold=True, color=COLOR_H2)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=10, after=2)
            run = p.add_run(text)
            _set_font(run, 12, bold=True, color=COLOR_H3)
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=8, after=2)
            run = p.add_run(text)
            _set_font(run, 11, bold=True, color=COLOR_BODY)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', line.strip()):
            hr = doc.add_paragraph()
            _add_border_bottom(hr, color='cbd5e1', size=6)
            _para_spacing(hr, before=8, after=8)
            i += 1
            continue

        # Unordered list
        if re.match(r'^[-*+]\s', line):
            text = re.sub(r'^[-*+]\s', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            _para_spacing(p, before=1, after=1)
            _apply_inline(text, p)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            _para_spacing(p, before=1, after=1)
            _apply_inline(text, p)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            _para_spacing(p, before=4, after=4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left_bdr = OxmlElement('w:left')
            left_bdr.set(qn('w:val'),   'single')
            left_bdr.set(qn('w:sz'),    '16')
            left_bdr.set(qn('w:space'), '8')
            left_bdr.set(qn('w:color'), '7c3aed')
            pBdr.append(left_bdr)
            pPr.append(pBdr)
            run = p.add_run(text)
            _set_font(run, 11, italic=True, color=COLOR_MUTED)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _para_spacing(p, before=2, after=4)
        _apply_inline(line, p)
        i += 1

    # ── Footer disclaimer ────────────────────────────────────────────
    doc.add_paragraph()
    disc_outer = doc.add_paragraph()
    _add_border_bottom(disc_outer, color='f59e0b', size=16)
    _para_spacing(disc_outer, before=16, after=4)
    dr = disc_outer.add_run('⚠  หมายเหตุ / Disclaimer')
    _set_font(dr, 10, bold=True, color=RGBColor(0xd9, 0x77, 0x06))

    disc_p = doc.add_paragraph()
    _para_spacing(disc_p, before=2, after=2)
    disc_r = disc_p.add_run(
        'รายงานนี้สร้างโดย AI Multi-Agent (CrewAI + Gemini) อิงจากข้อมูลตลาดจริงและข่าวล่าสุด '
        'ใช้เป็นข้อมูลประกอบการตัดสินใจเท่านั้น - ไม่ใช่คำแนะนำการลงทุน'
    )
    _set_font(disc_r, 9, italic=True, color=COLOR_MUTED)

    # ── Return file ──────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from django.http import HttpResponse
    resp = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    safe_symbol = re.sub(r'[^\w\-]', '_', symbol)
    resp['Content-Disposition'] = f'attachment; filename="CrewAI_{safe_symbol}_Analysis.docx"'
    return resp


# ====== Watchlist Management - เพิ่ม/ลบ รายการ Watchlist ======

def signup(request):
    """
    หน้าสมัครสมาชิก ใช้ Django built-in UserCreationForm
    เมื่อสมัครสำเร็จจะ login อัตโนมัติและ redirect ไปหน้า dashboard
    """
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Login อัตโนมัติหลังจากสมัครสำเร็จ
            login(request, user)
            messages.success(request, f"ยินดีต้อนรับคุณ {user.username}! ระบบของคุณพร้อมใช้งานแล้ว")
            return redirect('stocks:dashboard')
    else:
        form = UserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})


# ====== Multi-Factor Scoring Scanner ======

@login_required
def trading_accounts_view(request):
    """
    หน้าจอจัดการบัญชีเทรด (List & Add)
    """
    from stocks.models import BrokerType, TradingAccount
    
    if request.method == 'POST':
        broker = request.POST.get('broker')
        acc_id = request.POST.get('account_id')
        key    = request.POST.get('api_key', '')
        secret = request.POST.get('api_secret', '')
        
        TradingAccount.objects.create(
            user=request.user,
            broker=broker,
            account_id=acc_id,
            api_key=key,
            api_secret=secret
        )
        return redirect('stocks:trading_accounts')

    accounts = TradingAccount.objects.filter(user=request.user)
    return render(request, 'stocks/trading_account_list.html', {
        'accounts': accounts,
        'broker_types': BrokerType.choices
    })

@login_required
def delete_trading_account_view(request, pk):
    from stocks.models import TradingAccount
    acc = get_object_or_404(TradingAccount, pk=pk, user=request.user)
    acc.delete()
    return redirect('stocks:trading_accounts')

@csrf_exempt
@login_required
def sync_trading_account_ajax(request, pk):
    """
    ดึงยอดเงินล่าสุดจาก Broker มาอัปเดต (AJAX)
    """
    from stocks.models import TradingAccount
    from .trading_bridge import RobotBridge
    
    acc = get_object_or_404(TradingAccount, pk=pk, user=request.user)
    bridge = RobotBridge(user=request.user, account=acc)
    
    try:
        success = bridge.sync_account_balance()
        if success:
            return JsonResponse({
                'success': True,
                'balance': float(acc.balance),
                'equity': float(acc.equity),
                'currency': acc.currency
            })
        else:
            return JsonResponse({'success': False, 'error': 'API returned failure. Check if account is connected in MetaApi.'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def refresh_market_caps_view(request):
    """
    Manual trigger to refresh market caps for all SET symbols.
    """
    from django.contrib import messages
    from django.shortcuts import redirect
    from django.utils import timezone

    from stocks.models import ScannableSymbol
    from stocks.utils import refresh_market_caps
    
    # ⚡ High-speed check: ดึงแค่ timestamp ล่าสุดมาดูค่าเดียว + แปลงเป็น BKK timezone ก่อนเช็ค
    last_update_dt = ScannableSymbol.objects.filter(is_active=True, market='SET', market_cap__gt=0)\
                                          .values_list('last_cap_update', flat=True).first()
    today = timezone.localtime(timezone.now()).date()
    
    # ต้องแปลง UTC → Bangkok timezone ก่อนเปรียบเทียบวันที่ มิฉะนั้นจะ off ได้ถึง 7 ชั่วโมง
    already_today = (
        last_update_dt is not None and
        timezone.localtime(last_update_dt).date() == today
    )

    if already_today:
        messages.info(request, "ข้อมูลอันดับ Market Cap ของวันนี้อัปเดตเรียบร้อยแล้วครับ สแกนต่อได้ทันที!")
    else:
        count = refresh_market_caps()
        messages.success(request, f"สำเร็จ! อัปเดตข้อมูล Market Cap หุ้นไทยแล้ว {count} ตัว ระบบพร้อมจัดอันดับเพื่อสแกนแล้วครับ")
    
    next_url = request.GET.get('next') or 'stocks:momentum_scanner'
    return redirect(next_url)

